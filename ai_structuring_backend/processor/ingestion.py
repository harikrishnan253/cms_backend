"""
STAGE 1: Document Ingestion
- DOCX Parser (python-docx)
- Structure Extractor
- Document Serializer
- Context Zone Detection (Front Matter, Body, Back Matter, Table, Box)

Extracts all paragraphs with formatting metadata and assigns unique IDs.
Detects document context zones to provide better classification context.

CONTEXT ZONES:
- FRONT_MATTER: Metadata, author info → PMI only
- BODY: Main chapter content → Full style range
- BACK_MATTER: References, index → REF-N, EOC-*, IX-*
- TABLE: Word table content → T, T2, T4, TBL-*, TFN, TSN
- BOX: Pedagogical boxes → NBX-*, BX1-*, BX2-*, BX3-*
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from typing import Optional
import re
import logging
from .zone_styles import ZONE_VALID_STYLES as ZONE_VALID_STYLES_SET

logger = logging.getLogger(__name__)


# =============================================================================
# ZONE DETECTION PATTERNS
# =============================================================================

# Patterns that indicate FRONT MATTER (metadata, author info, etc.)
FRONT_MATTER_PATTERNS = [
    r'^book\s*title:',
    r'^chapter\s*#:',
    r'^chapter\s*title:',
    r'^corresponding\s*author:',
    r'^orcid\s*identifier:',
    r'^<metadata>',
    r'^key\s*words:',
    r'^abstract:',
    r'^section:$',
    r'^phone:\s*\(\d{3}\)',
    r'@.*\.(org|com|edu|net)$',  # Email patterns
    r'^\d{5}$',  # ZIP codes
]

# Patterns that indicate BACK MATTER (references, index, etc.)
BACK_MATTER_PATTERNS = [
    r'^<ref>references',
    r'^references\s*$',
    r'^bibliography\s*$',
    r'^index\s*$',
    r'^appendix\s+[a-z]',
    r'^glossary\s*$',
    r'^suggested\s+reading',
    r'^further\s+reading',
    r'^answer\s+key',
]

# Patterns that indicate CHAPTER OPENER (CN/CT - start of front matter)
CHAPTER_OPENER_PATTERNS = [
    r'^<cn>chapter\s+\d+',
    r'^chapter\s+\d+\s*$',
    r'^chapter\s+\d+[:\s]',
    r'^\d+\s*$',  # Just a number (chapter number)
]

# Patterns that indicate BODY START (first H1 heading - main content begins)
BODY_START_PATTERNS = [
    r'^<h1>',  # Explicit H1 marker
    r'^<h1\s',  # H1 with attributes
    r'^\s*introduction\s*$',  # Common first H1
    r'^\s*overview\s*$',
    r'^\s*background\s*$',
    r'^\s*getting\s+started\s*$',
]

# Patterns that indicate METADATA (pure PMI - before chapter opener)
METADATA_PATTERNS = [
    r'^book\s*title:',
    r'^chapter\s*#:',
    r'^chapter\s*title:',
    r'^corresponding\s*author:',
    r'^orcid\s*identifier:',
    r'^<metadata>',
    r'^key\s*words:',
    r'^abstract:',
    r'^section:\s*$',
    r'^phone:\s*\(\d{3}\)',
    r'@.*\.(org|com|edu|net)$',  # Email patterns
    r'^\d{5}$',  # ZIP codes
]

# Patterns that indicate BOX START
BOX_START_PATTERNS = [
    r'^<note>',
    r'^<clinical\s*pearl>',
    r'^<red\s*flag>',
    r'^<box>',
    r'^<tip>',
    r'^<example>',
    r'^<warning>',
    r'^<alert>',
    r'^<case\s*study>',
    r'^<reflection>',
    r'^<discussion>',
    r'^<practice>',
    r'^<key\s*point>',
    r'^<important>',
    r'^<remember>',
    r'^<unnumbered\s*box>',
]

# Patterns that indicate BOX END
BOX_END_PATTERNS = [
    r'^</note>',
    r'^</clinical\s*pearl>',
    r'^</red\s*flag>',
    r'^</box>',
    r'^</tip>',
    r'^</example>',
    r'^</warning>',
    r'^</alert>',
    r'^</case\s*study>',
    r'^</reflection>',
    r'^</discussion>',
    r'^</practice>',
    r'^</key\s*point>',
    r'^</important>',
    r'^</remember>',
    r'^</unnumbered\s*box>',
]

# Box type to style prefix mapping
BOX_TYPE_MAPPING = {
    'note': 'NBX',
    'tip': 'NBX',
    'clinical pearl': 'BX1',
    'example': 'BX1',
    'red flag': 'BX2',
    'warning': 'BX2',
    'alert': 'BX2',
    'reflection': 'BX3',
    'discussion': 'BX3',
    'case study': 'BX4',
    'practice': 'EXER',
    'key point': 'NBX',
    'important': 'NBX',
    'remember': 'NBX',
    'unnumbered box': 'NBX',
    'box': 'NBX',
}

# =============================================================================
# ZONE-SPECIFIC VALID STYLES
# Canonical source: backend/processor/zone_styles.py
#
# ingestion metadata stores valid_styles as list/None (JSON-friendly),
# while zone_styles.py stores set/None for fast membership checks.
# =============================================================================

ZONE_VALID_STYLES = {
    zone: (None if styles is None else sorted(styles))
    for zone, styles in ZONE_VALID_STYLES_SET.items()
}


def validate_style_for_zone(style: str, zone: str) -> bool:
    """
    Check if a style is valid for a given zone.
    
    Args:
        style: The style tag to validate
        zone: The context zone (FRONT_MATTER, BODY, TABLE, BOX_*, BACK_MATTER, EXERCISE)
    
    Returns:
        True if style is valid for zone, False otherwise
    """
    valid_styles = ZONE_VALID_STYLES.get(zone)
    
    # BODY has no restrictions
    if valid_styles is None:
        return True
    
    # Check exact match
    if style in valid_styles:
        return True
    
    # Check if style starts with any valid prefix (for variants like BL-MID0, TXT1, etc.)
    # Only for styles that might have numeric suffixes
    base_style = style.rstrip('0123456789')
    if base_style != style and base_style in valid_styles:
        return True
    
    return False


def get_valid_styles_for_zone(zone: str) -> list:
    """
    Get list of valid styles for a zone.
    
    Args:
        zone: The context zone
    
    Returns:
        List of valid style names, or empty list if no restrictions
    """
    styles = ZONE_VALID_STYLES.get(zone)
    return styles if styles else []


def get_zone_style_summary(zone: str) -> str:
    """
    Get a human-readable summary of valid styles for a zone.
    
    Args:
        zone: The context zone
    
    Returns:
        String summary of valid style prefixes
    """
    styles = ZONE_VALID_STYLES.get(zone)
    
    if styles is None:
        return "Full style range (no restrictions)"
    
    # Group by prefix for readability
    prefixes = set()
    for s in styles:
        if '-' in s:
            prefixes.add(s.split('-')[0] + '-*')
        else:
            prefixes.add(s)
    
    # Limit to most common prefixes
    common = sorted(prefixes)[:15]
    if len(prefixes) > 15:
        return ', '.join(common) + f' (+{len(prefixes)-15} more)'
    return ', '.join(common)


class DocumentIngestion:
    """
    Extract paragraphs and structure from DOCX files.
    Detects document context zones for better classification.
    
    Zones: FRONT_MATTER, BODY, BACK_MATTER, TABLE, BOX_*
    """
    
    def __init__(self, max_text_length: int = 200):
        """
        Initialize the document ingestion module.
        
        Args:
            max_text_length: Max characters for truncated display
        """
        self.max_text_length = max_text_length
        self._current_box_type = None  # Track current box context
    
    def _detect_box_start(self, text: str) -> Optional[str]:
        """
        Check if text starts a box and return box type.
        
        Returns:
            Box type (e.g., 'note', 'clinical pearl') or None
        """
        text_lower = text.lower().strip()
        
        for pattern in BOX_START_PATTERNS:
            if re.match(pattern, text_lower, re.IGNORECASE):
                # Extract box type from pattern
                for box_type in BOX_TYPE_MAPPING.keys():
                    if box_type.replace(' ', r'\s*') in pattern or box_type in text_lower:
                        return box_type
                # Default to 'box' if we can't determine type
                return 'box'
        return None
    
    def _detect_box_end(self, text: str) -> bool:
        """Check if text ends a box."""
        text_lower = text.lower().strip()
        
        for pattern in BOX_END_PATTERNS:
            if re.match(pattern, text_lower, re.IGNORECASE):
                return True
        return False
    
    def _get_box_zone(self, box_type: str) -> str:
        """Get the zone identifier for a box type."""
        if box_type is None:
            return 'BODY'
        
        prefix = BOX_TYPE_MAPPING.get(box_type, 'NBX')
        return f'BOX_{prefix}'
    
    def _detect_zone(self, text: str, current_zone: str, is_table: bool = False) -> tuple[str, Optional[str]]:
        """
        Detect which context zone this paragraph belongs to.
        
        Zone flow: METADATA → FRONT_MATTER → BODY → BACK_MATTER
        
        - METADATA: Pure metadata (book title, ORCID, etc.) - before chapter opener
        - FRONT_MATTER: Chapter opener through objectives (CN, CT, CAU, OBJ-*)
        - BODY: Main content starting at first H1
        - BACK_MATTER: References, index, etc.
        
        Args:
            text: Paragraph text
            current_zone: Current zone state
            is_table: Whether this is table content
            
        Returns:
            Tuple of (zone_identifier, box_type or None)
        """
        text_lower = text.lower().strip()
        
        # Table content always gets TABLE zone (unless in a box within table)
        if is_table:
            # Check if we're in a box context
            if self._current_box_type:
                # Check for box end
                if self._detect_box_end(text):
                    box_type = self._current_box_type
                    self._current_box_type = None
                    return (self._get_box_zone(box_type), box_type)
                return (self._get_box_zone(self._current_box_type), self._current_box_type)
            
            # Check for box start within table
            box_type = self._detect_box_start(text)
            if box_type:
                self._current_box_type = box_type
                return (self._get_box_zone(box_type), box_type)
            
            return ('TABLE', None)
        
        # Check for box end first (if we're in a box)
        if self._current_box_type:
            if self._detect_box_end(text):
                box_type = self._current_box_type
                self._current_box_type = None
                return (self._get_box_zone(box_type), box_type)
            return (self._get_box_zone(self._current_box_type), self._current_box_type)
        
        # Check for box start
        box_type = self._detect_box_start(text)
        if box_type:
            self._current_box_type = box_type
            return (self._get_box_zone(box_type), box_type)
        
        # If already in BODY, check for back matter
        if current_zone == 'BODY':
            for pattern in BACK_MATTER_PATTERNS:
                if re.match(pattern, text_lower, re.IGNORECASE):
                    return ('BACK_MATTER', None)
            return ('BODY', None)
        
        # If in BACK_MATTER, stay there
        if current_zone == 'BACK_MATTER':
            return ('BACK_MATTER', None)
        
        # Check for BODY start (first H1 heading)
        for pattern in BODY_START_PATTERNS:
            if re.match(pattern, text_lower, re.IGNORECASE):
                return ('BODY', None)
        
        # Check if this is a metadata section (pure PMI)
        if current_zone == 'METADATA':
            # Check for chapter opener - transitions to FRONT_MATTER
            for pattern in CHAPTER_OPENER_PATTERNS:
                if re.match(pattern, text_lower, re.IGNORECASE):
                    return ('FRONT_MATTER', None)
            
            # Check for end of metadata
            if '</metadata>' in text_lower:
                return ('METADATA', None)  # Still metadata, next will transition
            
            # Check if still metadata
            for pattern in METADATA_PATTERNS:
                if re.search(pattern, text_lower, re.IGNORECASE):
                    return ('METADATA', None)
            
            # If no metadata pattern but not chapter opener, stay in metadata
            return ('METADATA', None)
        
        # Check for chapter opener patterns - start of FRONT_MATTER
        for pattern in CHAPTER_OPENER_PATTERNS:
            if re.match(pattern, text_lower, re.IGNORECASE):
                return ('FRONT_MATTER', None)
        
        # If in FRONT_MATTER, stay there until H1
        if current_zone == 'FRONT_MATTER':
            return ('FRONT_MATTER', None)
        
        # Default: stay in current zone
        return (current_zone, None)
    
    # Keep old method for backward compatibility
    def _detect_section(self, text: str, current_section: str, prev_texts: list) -> str:
        """Legacy method - wraps _detect_zone for backward compatibility."""
        zone, _ = self._detect_zone(text, current_section, is_table=False)
        # Convert BOX zones back to BODY for old interface
        if zone.startswith('BOX_'):
            return 'BODY'
        if zone == 'METADATA':
            return 'FRONT_MATTER'
        return zone
    
    @staticmethod
    def _build_sdt_para_set(doc) -> set[int]:
        """Return a set of ``id(p_elem)`` for paragraphs inside body-level SDTs.

        Word Structured Document Tags (``<w:sdt>``) are content-control
        wrappers that appear as direct children of ``<w:body>``.  python-docx's
        ``doc.paragraphs`` flattens them into the normal paragraph stream, so
        they silently inherit whatever zone state is active.  This helper lets
        ``extract_paragraphs`` identify them and reset their zone to BODY.
        """
        try:
            from docx.oxml.ns import qn
        except ImportError:
            return set()

        sdt_paras: set[int] = set()
        for child in doc.element.body:
            if child.tag == qn('w:sdt'):
                sdt_content = child.find(qn('w:sdtContent'))
                if sdt_content is not None:
                    for p_elem in sdt_content.findall(qn('w:p')):
                        sdt_paras.add(id(p_elem))
        return sdt_paras

    def extract_paragraphs(self, docx_path: str | Path) -> list[dict]:
        """
        Extract all paragraphs from a DOCX file with context zone detection.

        Zone flow: METADATA → FRONT_MATTER → BODY → BACK_MATTER

        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            List of paragraph dictionaries with id, text, metadata including:
            - context_zone: METADATA, FRONT_MATTER, BODY, BACK_MATTER, TABLE, BOX_*
            - box_type: If in a box, the type (note, clinical pearl, etc.)
            - valid_styles: List of styles valid for this zone (None = all allowed)
        """
        doc = Document(docx_path)
        paragraphs = []
        para_id = 1
        self._current_box_type = None  # Reset box tracking

        # Pre-compute which paragraphs are inside body-level SDT (content
        # control) elements so their zone can be reset to BODY rather than
        # inheriting the surrounding box/zone state.
        sdt_para_ids = self._build_sdt_para_set(doc)
        
        # Determine starting zone based on first paragraph
        all_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Check what the document starts with
        if all_texts:
            first_text_lower = all_texts[0].lower()
            
            # Check for metadata patterns (pure PMI stuff)
            is_metadata_start = any(
                re.search(p, first_text_lower, re.IGNORECASE) 
                for p in METADATA_PATTERNS
            )
            
            # Check for chapter opener (CN/CT)
            is_chapter_start = any(
                re.match(p, first_text_lower, re.IGNORECASE) 
                for p in CHAPTER_OPENER_PATTERNS
            )
            
            # Check for H1/body start
            is_body_start = any(
                re.match(p, first_text_lower, re.IGNORECASE) 
                for p in BODY_START_PATTERNS
            )
            
            if is_metadata_start:
                current_zone = 'METADATA'
            elif is_chapter_start:
                current_zone = 'FRONT_MATTER'
            elif is_body_start:
                current_zone = 'BODY'
            else:
                current_zone = 'FRONT_MATTER'  # Default to front matter
        else:
            current_zone = 'FRONT_MATTER'
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # Skip completely empty paragraphs
            if not text:
                continue
            
            # Detect zone (includes box tracking)
            current_zone, box_type = self._detect_zone(text, current_zone, is_table=False)

            # SDT (content control) paragraphs must reset to BODY.
            # They are structural elements at the body level that should
            # not inherit the surrounding box zone.
            is_sdt_para = id(para._p) in sdt_para_ids
            if is_sdt_para and current_zone.startswith('BOX_'):
                current_zone = 'BODY'
                box_type = None

            # Extract formatting metadata
            metadata = self._extract_formatting(para)

            # Record SDT origin for downstream traceability
            if is_sdt_para:
                metadata['is_sdt'] = True

            # Add zone information
            metadata['context_zone'] = current_zone
            metadata['box_type'] = box_type
            metadata['valid_styles'] = ZONE_VALID_STYLES.get(current_zone)
            
            # Keep backward compatible field
            if current_zone.startswith('BOX_'):
                metadata['document_section'] = 'BODY'
            else:
                metadata['document_section'] = current_zone
            
            paragraphs.append({
                'id': para_id,
                'text': text,
                'text_truncated': self._truncate(text),
                'metadata': metadata
            })
            
            para_id += 1
        
        # Reset box tracking before processing tables
        self._current_box_type = None
        
        # Also extract table content with zone awareness
        table_paragraphs = self._extract_tables(doc, para_id, current_zone)
        paragraphs.extend(table_paragraphs)
        
        # Log zone breakdown
        zone_counts = {}
        for p in paragraphs:
            zone = p['metadata'].get('context_zone', 'UNKNOWN')
            zone_counts[zone] = zone_counts.get(zone, 0) + 1
        
        logger.info(f"Extracted {len(paragraphs)} paragraphs from {docx_path}")
        logger.info(f"Zone breakdown: {zone_counts}")
        
        return paragraphs
    
    def _extract_formatting(self, para) -> dict:
        """
        Extract formatting metadata from a paragraph.
        
        Args:
            para: python-docx Paragraph object
            
        Returns:
            Dictionary of formatting properties
        """
        metadata = {
            'style_name': para.style.name if para.style else 'Normal',
            'alignment': None,
            'is_bold': False,
            'is_italic': False,
            'is_all_caps': False,
            'font_size': None,
            'has_numbering': False,
            'has_bullet': False,
            'indent_level': 0,
        }
        
        # Check alignment
        if para.alignment:
            alignment_map = {
                WD_ALIGN_PARAGRAPH.LEFT: 'left',
                WD_ALIGN_PARAGRAPH.CENTER: 'center',
                WD_ALIGN_PARAGRAPH.RIGHT: 'right',
                WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
            }
            metadata['alignment'] = alignment_map.get(para.alignment, 'left')
        
        # Check run formatting (first run)
        if para.runs:
            run = para.runs[0]
            metadata['is_bold'] = run.bold or False
            metadata['is_italic'] = run.italic or False
            if run.font.size:
                metadata['font_size'] = run.font.size.pt
            if run.font.all_caps:
                metadata['is_all_caps'] = True
        
        # Check for numbering/bullets in text (manual lists)
        text = para.text.strip()
        if re.match(r'^\d+\.?\s', text):
            metadata['has_numbering'] = True
        if re.match(r'^[•\-\*]\s', text):
            metadata['has_bullet'] = True
            
        # Check for XML numbering (automatic Word lists)
        # This catches lists that don't have bullets/numbers in the text property
        try:
            if hasattr(para, '_p') and para._p.pPr is not None and para._p.pPr.numPr is not None:
                # It has a numbering property -> it is a list
                # Extract full numbering properties from Word XML
                numPr = para._p.pPr.numPr

                # Extract ilvl (indentation level)
                if hasattr(numPr, 'ilvl') and numPr.ilvl is not None:
                    try:
                        metadata['xml_list_level'] = int(numPr.ilvl.val)
                    except (AttributeError, ValueError, TypeError):
                        pass  # Fail gracefully if ilvl.val is malformed

                # Extract numId (identifies which list sequence)
                if hasattr(numPr, 'numId') and numPr.numId is not None:
                    try:
                        metadata['xml_num_id'] = int(numPr.numId.val)
                    except (AttributeError, ValueError, TypeError):
                        pass  # Fail gracefully if numId.val is malformed

                # Try to guess type from style name
                style_name = para.style.name.lower() if para.style else ""

                if 'bullet' in style_name:
                    metadata['has_bullet'] = True
                elif 'number' in style_name:
                    metadata['has_numbering'] = True
                else:
                    # Ambiguous list - mark as generic list
                    # DO NOT default to numbering as it causes false positives (BL becoming NL)
                    metadata['has_xml_list'] = True
        except Exception:
            pass  # Fail gracefully if XML access fails

        # Style-name-based list inference for paragraphs that use named list
        # styles WITHOUT XML numbering properties (numPr absent).  Catches
        # publisher templates that apply "List Bullet", "BulletList*",
        # "List Number", "NumberList*" etc. without adding numPr elements.
        #
        # NOTE: The same bullet/number keyword check already runs inside the
        # numPr block above (lines 634-637) for the numPr-present case; this
        # block is its mirror for the no-numPr case.  The guard ensures we
        # never overwrite a flag already set by text-regex or numPr detection.
        if not metadata.get('has_bullet') and not metadata.get('has_numbering') \
                and not metadata.get('has_xml_list'):
            _sn = metadata.get('style_name', '').lower()
            if 'bullet' in _sn:
                metadata['has_bullet'] = True
            elif 'number' in _sn:
                metadata['has_numbering'] = True

        # Check paragraph format for indentation
        if para.paragraph_format.left_indent:
            indent_inches = para.paragraph_format.left_indent.inches
            metadata['indent_level'] = int(indent_inches / 0.25)  # Estimate level
            # Also store raw twips for the hierarchy detector (720 twips = 0.5 inch = typical Level 1)
            metadata['indent_twips'] = int(indent_inches * 1440)

        # Canonical alias: expose xml_list_level as ooxml_ilvl so the
        # list_hierarchy_detector can find it without key-name mismatch.
        if 'xml_list_level' in metadata:
            metadata['ooxml_ilvl'] = metadata['xml_list_level']

        return metadata
    
    def _extract_tables(self, doc, start_id: int, current_zone: str = 'BODY') -> list[dict]:
        """
        Extract text content from tables with zone awareness.
        
        Args:
            doc: python-docx Document object
            start_id: Starting paragraph ID
            current_zone: Current document zone context
            
        Returns:
            List of paragraph dictionaries from tables
        """
        paragraphs = []
        para_id = start_id
        
        for table_idx, table in enumerate(doc.tables):
            num_rows = len(table.rows)
            num_cols = len(table.columns) if table.rows else 0
            
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Process each paragraph in the cell (cells can have multiple paragraphs)
                    for para_idx, para in enumerate(cell.paragraphs):
                        text = para.text.strip()
                        if not text:
                            continue
                        
                        # Get actual style from cell paragraph
                        cell_style = para.style.name if para.style else None
                        
                        # Determine position context
                        is_header = row_idx == 0
                        is_first_col = cell_idx == 0
                        is_last_row = row_idx == num_rows - 1
                        
                        # Check for box markers within table cells
                        cell_zone, box_type = self._detect_zone(text, current_zone, is_table=True)
                        
                        # Infer appropriate table style based on position
                        inferred_style = self._infer_table_style(
                            cell_style, 
                            text, 
                            is_header, 
                            is_first_col,
                            is_last_row,
                            para_idx
                        )
                        
                        paragraphs.append({
                            'id': para_id,
                            'text': text,
                            'text_truncated': self._truncate(text),
                            'metadata': {
                                'style_name': cell_style or 'TableCell',
                                'inferred_style': inferred_style,
                                'is_table': True,
                                'table_index': table_idx,
                                'table_id': table_idx,
                                'row_index': row_idx,
                                'table_row_idx': row_idx,
                                'cell_index': cell_idx,
                                'table_col_idx': cell_idx,
                                'para_in_cell': para_idx,
                                'is_header_row': is_header,
                                'is_first_column': is_first_col,
                                'is_stub_col': is_first_col,
                                'table_size': f"{num_rows}x{num_cols}",
                                'document_section': current_zone if not current_zone.startswith('BOX_') else 'BODY',
                                'context_zone': 'TABLE',
                                'table_cell_zone': cell_zone,
                                'box_type': box_type,
                                'valid_styles': ZONE_VALID_STYLES.get(cell_zone),
                            }
                        })
                        para_id += 1
        
        return paragraphs
    
    def _infer_table_style(
        self, 
        cell_style: str, 
        text: str, 
        is_header: bool, 
        is_first_col: bool,
        is_last_row: bool,
        para_idx: int
    ) -> str:
        """
        Infer appropriate WK template table style based on cell position and content.
        """
        # If cell has a recognized style, use it
        if cell_style:
            style_upper = cell_style.upper()
            # Map common table styles
            if style_upper in ['T', 'TABLEBODY', 'GT']:
                return 'T'  # Body cell (classifier decides T vs T4 from content)
            elif style_upper in ['T2', 'TABLECOLUMNHEAD1', 'TABLEHEADER']:
                return 'T2'  # Column header
            elif 'TBL' in style_upper or 'BULLET' in style_upper:
                return 'TBL-MID'  # Table bullet list
            elif style_upper in ['TFN', 'TABLEFOOTNOTE']:
                return 'TFN'
            elif style_upper.startswith('UNT'):
                if is_header:
                    return 'T2'
                return 'T'  # Body cell (classifier decides T vs T4 from content)
        
        # Infer from position
        if is_header:
            return 'T2'  # Header row
        
        # Check if text looks like a list item
        if text.startswith(('•', '-', '●', '○', '\t•', '\t-', '	')):
            return 'TBL-MID'
        
        # Default to body cell
        return 'T'
    
    def _truncate(self, text: str) -> str:
        """Truncate text for display purposes."""
        if len(text) <= self.max_text_length:
            return text
        return text[:self.max_text_length] + "..."
    
    def format_for_prompt(self, paragraphs: list[dict]) -> str:
        """
        Format paragraphs for the Gemini prompt.
        
        Args:
            paragraphs: List of paragraph dictionaries
            
        Returns:
            Formatted string for the prompt
        """
        lines = []
        for para in paragraphs:
            lines.append(f"[{para['id']}] {para['text_truncated']}")
        return "\n".join(lines)
    
    def get_document_stats(self, paragraphs: list[dict]) -> dict:
        """
        Get statistics about the extracted document.
        
        Args:
            paragraphs: List of paragraph dictionaries
            
        Returns:
            Dictionary of statistics
        """
        total_chars = sum(len(p['text']) for p in paragraphs)
        table_paras = sum(1 for p in paragraphs if p['metadata'].get('is_table'))
        numbered = sum(1 for p in paragraphs if p['metadata'].get('has_numbering'))
        bulleted = sum(1 for p in paragraphs if p['metadata'].get('has_bullet'))
        
        return {
            'total_paragraphs': len(paragraphs),
            'total_characters': total_chars,
            'estimated_tokens': total_chars // 4,  # Rough estimate
            'table_paragraphs': table_paras,
            'numbered_items': numbered,
            'bulleted_items': bulleted,
        }


def extract_document(docx_path: str | Path) -> tuple[list[dict], dict]:
    """
    Convenience function to extract a document.
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        Tuple of (paragraphs, stats)
    """
    ingestion = DocumentIngestion()
    paragraphs = ingestion.extract_paragraphs(docx_path)
    stats = ingestion.get_document_stats(paragraphs)
    return paragraphs, stats


if __name__ == "__main__":
    # Test with sample file
    import sys
    if len(sys.argv) > 1:
        path = sys.argv[1]
        paragraphs, stats = extract_document(path)
        print(f"Document Statistics: {stats}")
        print(f"\nFirst 10 paragraphs:")
        for p in paragraphs[:10]:
            print(f"  [{p['id']}] {p['text_truncated']}")
