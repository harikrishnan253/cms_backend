"""
Structure Guard: Deterministic enforcement of style-only mutations.

Validates that the processor ONLY changes paragraph styles and does NOT mutate:
- Paragraph count
- Paragraph order
- Paragraph text
- List structure
- Table structure
- Section structure

NO LLM USAGE - 100% deterministic validation.
"""

from __future__ import annotations

import hashlib
import logging
import re
import unicodedata
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document

logger = logging.getLogger(__name__)


def _para_debug_snapshot(para_info: Dict | None) -> Dict | None:
    """Small diagnostic snapshot for first-difference logs."""
    if not para_info:
        return None
    return {
        "index": para_info.get("index"),
        "style": para_info.get("style_name"),
        "text": (para_info.get("normalized_text") or "")[:80],
        "is_list": para_info.get("is_list_item"),
        "list_level": para_info.get("list_level"),
        "list_id": para_info.get("list_id"),
        "in_table": para_info.get("in_table"),
        "section_index": para_info.get("section_index"),
    }


def _style_chain_names(style) -> list[str]:
    """Return paragraph style name chain (current -> base styles)."""
    names: list[str] = []
    seen = set()
    current = style
    while current is not None:
        name = getattr(current, "name", "") or ""
        if name and name not in seen:
            names.append(name)
            seen.add(name)
        current = getattr(current, "base_style", None)
    return names


def _list_level_from_style_name(style_name: str) -> int:
    """Infer zero-based list level from built-in list style names."""
    match = re.search(r"\bList (?:Bullet|Number)(?: (\d+))?$", style_name, re.IGNORECASE)
    if not match:
        return 0
    suffix = match.group(1)
    if not suffix:
        return 0
    try:
        return max(int(suffix) - 1, 0)
    except ValueError:
        return 0


def _normalize_text(text: str) -> str:
    """
    Normalize text for comparison.

    Rules:
    - Strip whitespace
    - Normalize Unicode (NFKC)
    - Collapse multiple spaces
    - DO NOT remove markers

    Parameters
    ----------
    text : str
        Raw text to normalize

    Returns
    -------
    str
        Normalized text
    """
    # Strip whitespace
    text = text.strip()

    # Normalize Unicode (NFKC)
    text = unicodedata.normalize('NFKC', text)

    # Collapse multiple spaces into one
    import re
    text = re.sub(r' +', ' ', text)

    return text


def _get_list_info(paragraph) -> Tuple[bool, int | None, int | None]:
    """
    Extract list metadata from paragraph.

    Checks both:
    1. XML numbering properties (numPr, ilvl, numId)
    2. Paragraph style name (List Bullet, List Number, etc.)

    Parameters
    ----------
    paragraph : docx.text.paragraph.Paragraph
        Paragraph to inspect

    Returns
    -------
    tuple
        (is_list_item, list_level, list_id)
    """
    try:
        # Access paragraph XML
        p_element = paragraph._element

        # Look for numbering properties (numPr)
        numPr = p_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')

        if numPr is not None:
            # Extract ilvl (list level)
            ilvl_element = numPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')
            list_level = int(ilvl_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 0)) if ilvl_element is not None else 0

            # Extract numId (list ID)
            numId_element = numPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')
            list_id = int(numId_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 0)) if numId_element is not None else None

            return True, list_level, list_id

        # Fallback: Check paragraph style and base-style chain for list indicators
        style_chain = _style_chain_names(getattr(paragraph, "style", None))
        for style_name in style_chain:
            lowered = style_name.lower()
            is_list_style = any(indicator in lowered for indicator in ['list', 'bullet', 'number'])
            if is_list_style:
                return True, _list_level_from_style_name(style_name), -1

        return False, None, None
    except Exception as e:
        logger.debug(f"Failed to extract list info: {e}")
        return False, None, None


def _extract_paragraphs(doc: Document) -> List[Dict]:
    """
    Extract ordered list of paragraphs with metadata.

    Parameters
    ----------
    doc : Document
        DOCX document

    Returns
    -------
    list of dict
        Paragraph metadata with structure:
        {
            'index': int,
            'raw_text': str,
            'normalized_text': str,
            'style_name': str,
            'is_list_item': bool,
            'list_level': int | None,
            'list_id': int | None,
            'in_table': bool,
            'section_index': int
        }
    """
    paragraphs = []
    section_index = 0

    for i, para in enumerate(doc.paragraphs):
        raw_text = para.text
        normalized_text = _normalize_text(raw_text)

        # Get style name
        style_name = para.style.name if para.style else ""

        # Get list info
        is_list_item, list_level, list_id = _get_list_info(para)

        # Check if paragraph is in a table
        in_table = False
        try:
            # Walk up XML tree to check for table cell parent
            parent = para._element.getparent()
            while parent is not None:
                if parent.tag.endswith('}tc'):  # table cell
                    in_table = True
                    break
                parent = parent.getparent()
        except:
            pass

        para_info = {
            'index': i,
            'raw_text': raw_text,
            'normalized_text': normalized_text,
            'style_name': style_name,
            'is_list_item': is_list_item,
            'list_level': list_level,
            'list_id': list_id,
            'in_table': in_table,
            'section_index': section_index,
        }

        paragraphs.append(para_info)

    return paragraphs


def _extract_tables(doc: Document) -> List[Dict]:
    """
    Extract table structure metadata.

    Parameters
    ----------
    doc : Document
        DOCX document

    Returns
    -------
    list of dict
        Table metadata with structure:
        {
            'table_index': int,
            'rows': int,
            'cols': int,
            'nested_tables_count': int
        }
    """
    tables = []

    for t, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns) if table.rows else 0

        # Count nested tables
        nested_count = 0
        for row in table.rows:
            for cell in row.cells:
                nested_count += len(cell.tables)

        table_info = {
            'table_index': t,
            'rows': rows,
            'cols': cols,
            'nested_tables_count': nested_count,
        }

        tables.append(table_info)

    return tables


def _extract_sections(doc: Document) -> List[Dict]:
    """
    Extract section structure metadata.

    Parameters
    ----------
    doc : Document
        DOCX document

    Returns
    -------
    list of dict
        Section metadata with structure:
        {
            'index': int,
            'break_type': str
        }
    """
    sections = []

    for i, section in enumerate(doc.sections):
        # Get section break type via python-docx start_type property.
        # Raw XML parsing (w:type element) is unreliable because python-docx
        # omits the element when the value is the document default (NEW_PAGE).
        try:
            break_type = section.start_type.name.lower()
        except Exception:
            break_type = 'new_page'

        section_info = {
            'index': i,
            'break_type': break_type,
        }

        sections.append(section_info)

    return sections


def _create_structural_signature(paragraphs: List[Dict], tables: List[Dict], sections: List[Dict]) -> str:
    """
    Create canonical structural signature for hashing.

    Format:
    - Paragraph: "P|{idx}|L{list_level or -1}|T{int(in_table)}|S{section_index}"
    - Table: "T|{idx}|R{rows}|C{cols}"
    - Section: "SEC|{idx}|{break_type}"

    Parameters
    ----------
    paragraphs : list of dict
        Paragraph metadata
    tables : list of dict
        Table metadata
    sections : list of dict
        Section metadata

    Returns
    -------
    str
        SHA256 hash of structural signature
    """
    signatures = []

    # Paragraph signatures
    for p in paragraphs:
        sig = f"P|{p['index']}|L{p['list_level'] if p['list_level'] is not None else -1}|T{int(p['in_table'])}|S{p['section_index']}"
        signatures.append(sig)

    # Table signatures
    for t in tables:
        sig = f"T|{t['table_index']}|R{t['rows']}|C{t['cols']}"
        signatures.append(sig)

    # Section signatures
    for s in sections:
        sig = f"SEC|{s['index']}|{s['break_type']}"
        signatures.append(sig)

    # Hash the joined signature
    full_signature = '\n'.join(signatures)
    signature_hash = hashlib.sha256(full_signature.encode('utf-8')).hexdigest()

    return signature_hash


def _validate_paragraphs(input_paras: List[Dict], output_paras: List[Dict]) -> Tuple[bool, List[str]]:
    """
    Validate paragraph count, order, and text.

    Parameters
    ----------
    input_paras : list of dict
        Input paragraph metadata
    output_paras : list of dict
        Output paragraph metadata

    Returns
    -------
    tuple
        (all_match, differences)
    """
    differences = []

    # Check paragraph count
    if len(input_paras) != len(output_paras):
        differences.append(
            f"Paragraph count mismatch: input={len(input_paras)}, output={len(output_paras)}"
        )
        return False, differences

    # Check each paragraph at same index
    for i in range(len(input_paras)):
        inp = input_paras[i]
        out = output_paras[i]

        # Check normalized text
        if inp['normalized_text'] != out['normalized_text']:
            differences.append(
                f"Paragraph {i} text differs: "
                f"input='{inp['normalized_text'][:50]}...', "
                f"output='{out['normalized_text'][:50]}...' "
                f"| input_meta={_para_debug_snapshot(inp)} output_meta={_para_debug_snapshot(out)}"
            )

        # Check list structure
        if inp['is_list_item'] != out['is_list_item']:
            differences.append(
                f"Paragraph {i} list status changed: "
                f"input_is_list={inp['is_list_item']}, output_is_list={out['is_list_item']} "
                f"| input_meta={_para_debug_snapshot(inp)} output_meta={_para_debug_snapshot(out)}"
            )

        if inp['list_level'] != out['list_level']:
            differences.append(
                f"Paragraph {i} list level changed: "
                f"input_level={inp['list_level']}, output_level={out['list_level']} "
                f"| input_meta={_para_debug_snapshot(inp)} output_meta={_para_debug_snapshot(out)}"
            )

        # ``-1`` is a sentinel for style-based lists where a concrete list_id is
        # not available. Treat it as "unknown" rather than a hard mismatch.
        if (
            inp['list_id'] != out['list_id']
            and inp['list_id'] != -1
            and out['list_id'] != -1
        ):
            differences.append(
                f"Paragraph {i} list ID changed: "
                f"input_id={inp['list_id']}, output_id={out['list_id']} "
                f"| input_meta={_para_debug_snapshot(inp)} output_meta={_para_debug_snapshot(out)}"
            )

    all_match = len(differences) == 0
    return all_match, differences


def _validate_tables(input_tables: List[Dict], output_tables: List[Dict]) -> Tuple[bool, List[str]]:
    """
    Validate table structure.

    Parameters
    ----------
    input_tables : list of dict
        Input table metadata
    output_tables : list of dict
        Output table metadata

    Returns
    -------
    tuple
        (all_match, differences)
    """
    differences = []

    # Check table count
    if len(input_tables) != len(output_tables):
        differences.append(
            f"Table count mismatch: input={len(input_tables)}, output={len(output_tables)}"
        )
        return False, differences

    # Check each table
    for i in range(len(input_tables)):
        inp = input_tables[i]
        out = output_tables[i]

        if inp['rows'] != out['rows']:
            differences.append(
                f"Table {i} row count differs: input={inp['rows']}, output={out['rows']}"
            )

        if inp['cols'] != out['cols']:
            differences.append(
                f"Table {i} column count differs: input={inp['cols']}, output={out['cols']}"
            )

        if inp['nested_tables_count'] != out['nested_tables_count']:
            differences.append(
                f"Table {i} nested table count differs: "
                f"input={inp['nested_tables_count']}, output={out['nested_tables_count']}"
            )

    all_match = len(differences) == 0
    return all_match, differences


def _validate_sections(input_sections: List[Dict], output_sections: List[Dict]) -> Tuple[bool, List[str]]:
    """
    Validate section structure.

    Parameters
    ----------
    input_sections : list of dict
        Input section metadata
    output_sections : list of dict
        Output section metadata

    Returns
    -------
    tuple
        (all_match, differences)
    """
    differences = []

    # Check section count
    if len(input_sections) != len(output_sections):
        differences.append(
            f"Section count mismatch: input={len(input_sections)}, output={len(output_sections)}"
        )
        return False, differences

    # Check section break types
    for i in range(len(input_sections)):
        inp = input_sections[i]
        out = output_sections[i]

        if inp['break_type'] != out['break_type']:
            differences.append(
                f"Section {i} break type differs: "
                f"input={inp['break_type']}, output={out['break_type']}"
            )

    all_match = len(differences) == 0
    return all_match, differences


def enforce_style_only_mutation(input_path: str | Path, output_path: str | Path) -> dict:
    """
    Enforce that processor only changes paragraph styles and does NOT mutate structure.

    Validates that output preserves:
    - Paragraph count
    - Paragraph order
    - Paragraph text (normalized)
    - List structure
    - Table structure
    - Section structure

    Only style names may change.

    NO LLM USAGE - 100% deterministic validation.

    Parameters
    ----------
    input_path : str or Path
        Path to input DOCX file
    output_path : str or Path
        Path to output/processed DOCX file

    Returns
    -------
    dict
        Validation result with structure:
        {
            "status": "PASS" | "FAIL",
            "paragraph_count_match": bool,
            "list_structure_match": bool,
            "table_structure_match": bool,
            "section_structure_match": bool,
            "structural_hash_match": bool,
            "differences": list[str]
        }

    Raises
    ------
    RuntimeError
        If validation fails (status == "FAIL")

    Examples
    --------
    >>> result = enforce_style_only_mutation("input.docx", "output.docx")
    >>> # Raises RuntimeError if structural mutation detected
    """
    input_path = Path(input_path)
    output_path = Path(output_path)

    logger.info("STRUCTURE_GUARD: Starting validation")
    logger.info(f"  Input: {input_path.name}")
    logger.info(f"  Output: {output_path.name}")

    # Load documents
    logger.debug("Loading input document")
    input_doc = Document(input_path)

    logger.debug("Loading output document")
    output_doc = Document(output_path)

    # Extract structures
    logger.debug("Extracting paragraph structures")
    input_paras = _extract_paragraphs(input_doc)
    output_paras = _extract_paragraphs(output_doc)

    logger.debug("Extracting table structures")
    input_tables = _extract_tables(input_doc)
    output_tables = _extract_tables(output_doc)

    logger.debug("Extracting section structures")
    input_sections = _extract_sections(input_doc)
    output_sections = _extract_sections(output_doc)

    # Validate structures
    logger.debug("Validating paragraph structure")
    para_match, para_diffs = _validate_paragraphs(input_paras, output_paras)

    logger.debug("Validating table structure")
    table_match, table_diffs = _validate_tables(input_tables, output_tables)

    logger.debug("Validating section structure")
    section_match, section_diffs = _validate_sections(input_sections, output_sections)

    # Combine all differences
    all_differences = para_diffs + table_diffs + section_diffs

    # Limit to first 50 differences
    all_differences = all_differences[:50]

    # Create structural signatures
    logger.debug("Computing structural signatures")
    input_hash = _create_structural_signature(input_paras, input_tables, input_sections)
    output_hash = _create_structural_signature(output_paras, output_tables, output_sections)
    hash_match = (input_hash == output_hash)

    # Check list structure match (no list-related diffs)
    list_match = not any('list' in d.lower() for d in all_differences)

    # Determine overall status
    overall_status = "PASS" if (
        para_match and
        table_match and
        section_match and
        hash_match
    ) else "FAIL"

    # Build result
    result = {
        "status": overall_status,
        "paragraph_count_match": len(input_paras) == len(output_paras),
        "list_structure_match": list_match,
        "table_structure_match": table_match,
        "section_structure_match": section_match,
        "structural_hash_match": hash_match,
        "differences": all_differences,
        "first_difference": all_differences[0] if all_differences else None,
    }

    # Log result
    logger.info(
        f"STRUCTURE_GUARD status={overall_status} "
        f"para_match={result['paragraph_count_match']} "
        f"list_match={result['list_structure_match']} "
        f"table_match={result['table_structure_match']} "
        f"section_match={result['section_structure_match']} "
        f"hash_match={result['structural_hash_match']} "
        f"differences_count={len(all_differences)}"
    )

    # If FAIL, raise RuntimeError
    if overall_status == "FAIL":
        if all_differences:
            logger.error(
                "STRUCTURE_GUARD_FIRST_DIFF stage=structure_guard detail=%s",
                all_differences[0],
            )
        error_parts = [f"STRUCTURE_GUARD_FAIL: {len(all_differences)} structural violations detected"]

        # Add first 20 differences to error message
        for diff in all_differences[:20]:
            error_parts.append(f"  - {diff}")

        if len(all_differences) > 20:
            error_parts.append(f"  ... and {len(all_differences) - 20} more differences")

        error_message = "\n".join(error_parts)
        logger.error(error_message)

        raise RuntimeError(error_message)

    logger.info("STRUCTURE_GUARD: PASS - Only style changes detected")
    return result
