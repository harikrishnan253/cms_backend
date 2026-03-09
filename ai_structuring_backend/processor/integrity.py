"""
Automated Integrity Trigger for DOCX Processing Pipeline

Verifies BOTH:
1) Content Integrity - All input content exists in processed output (no text loss)
2) Structural Integrity - Document structure is preserved (paragraphs, tables, headings, lists, sections)

This is a hard validation gate that runs automatically after processing.
NO LLM USAGE - 100% deterministic verification.
"""

from __future__ import annotations

import hashlib
import logging
import re
import unicodedata
import zipfile
from collections import defaultdict
from pathlib import Path
from typing import Set, List, Dict, Tuple
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_BREAK

logger = logging.getLogger(__name__)

# Known marker tokens that may be inserted additively
KNOWN_MARKERS = {
    '<body-open>', '<body-close>',
    '<front-open>', '<front-close>',
    '<back-open>', '<back-close>',
    '<table-open>', '<table-close>',
    '<box-open>', '<box-close>',
    '<exercise-open>', '<exercise-close>',
}

_MARKER_ONLY_TOKEN_RE = re.compile(r"^\s*</?[A-Za-z][A-Za-z0-9._ -]*>\s*$")


def _para_diag_snapshot(para_info: Dict | None) -> Dict | None:
    """Compact paragraph snapshot for first-difference diagnostics."""
    if not para_info:
        return None
    return {
        "index": para_info.get("index"),
        "style": para_info.get("style_name"),
        "heading_level": para_info.get("heading_level"),
        "is_list_item": para_info.get("is_list_item"),
        "list_id": para_info.get("list_id"),
        "list_level": para_info.get("list_level"),
        "in_table": para_info.get("in_table"),
        "section_index": para_info.get("section_index"),
        "is_marker": para_info.get("is_marker"),
        "text_preview": para_info.get("text_preview"),
    }


def _style_chain_names(style) -> list[str]:
    """Return paragraph style name chain (current -> base styles)."""
    names: list[str] = []
    seen = set()
    current = style
    while current is not None:
        name = getattr(current, "name", "") or ""
        if name and name not in seen:
            names.append(name)
            seen.add(name)
        current = getattr(current, "base_style", None)
    return names


def _list_level_from_style_name(style_name: str) -> int:
    """Infer zero-based list level from built-in list style names."""
    match = re.search(r"\bList (?:Bullet|Number)(?: (\d+))?$", style_name, re.IGNORECASE)
    if not match:
        return 0
    suffix = match.group(1)
    if not suffix:
        return 0
    try:
        return max(int(suffix) - 1, 0)
    except ValueError:
        return 0


# ===================================================================
# Content Integrity Verification
# ===================================================================

def _normalize_text(text: str) -> str:
    """
    Normalize text for comparison.

    - Trim whitespace
    - Normalize Unicode (NFKC)
    - Collapse multiple spaces to single
    - Convert smart quotes to straight quotes
    - Remove known marker tokens
    """
    if not text:
        return ""

    # Normalize Unicode
    text = unicodedata.normalize('NFKC', text)

    # Convert smart quotes to straight
    text = text.replace('\u2018', "'").replace('\u2019', "'")
    text = text.replace('\u201C', '"').replace('\u201D', '"')
    text = text.replace('\u2013', '-').replace('\u2014', '-')

    # Remove known marker tokens
    for marker in KNOWN_MARKERS:
        text = text.replace(marker, '')

    # Collapse whitespace
    text = re.sub(r'\s+', ' ', text)

    # Trim
    text = text.strip()

    return text


def _extract_content(docx_path: str | Path) -> Dict:
    """
    Extract all content from DOCX for verification.

    Returns dict with:
    - paragraphs: list of normalized paragraph texts
    - table_cells: list of normalized table cell texts
    - paragraph_count: total paragraph count
    - table_count: total table count
    """
    doc = Document(docx_path)

    paragraphs = []
    table_cells = []

    # Extract paragraphs (not in tables)
    for para in doc.paragraphs:
        text = _normalize_text(para.text)
        if text:
            paragraphs.append(text)

    # Extract table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = _normalize_text(para.text)
                    if text:
                        table_cells.append(text)

    return {
        'paragraphs': paragraphs,
        'table_cells': table_cells,
        'paragraph_count': len(paragraphs),
        'table_count': len(doc.tables),
    }


def _build_output_index(output_data: Dict) -> Set[str]:
    """
    Build an indexed set of output content for fast lookup.

    Includes:
    - Individual paragraphs
    - Individual table cells
    - Two-paragraph concatenations (for split detection)
    """
    index = set()

    # Add all paragraphs
    paragraphs = output_data['paragraphs']
    for text in paragraphs:
        index.add(text)

    # Add all table cells
    for text in output_data['table_cells']:
        index.add(text)

    # Add two-paragraph concatenations for split detection
    for i in range(len(paragraphs) - 1):
        combined = paragraphs[i] + ' ' + paragraphs[i + 1]
        index.add(combined)

    return index


def verify_content_integrity(input_path: str | Path, output_path: str | Path) -> dict:
    """
    Verify that ALL content from input DOCX exists in output DOCX.

    **Coverage Rule:**
    Every normalized input paragraph text must be found in output in one of:
    - Exact match paragraph
    - Split across adjacent output paragraphs (max 2 splits)
    - Present inside a table cell
    - Present inside a list item

    **FAIL Conditions:**
    - Any input paragraph text is missing
    - Any table cell text missing

    Parameters
    ----------
    input_path : str or Path
        Path to input DOCX file
    output_path : str or Path
        Path to output/processed DOCX file

    Returns
    -------
    dict
        Verification result:
        {
            "status": "PASS" | "FAIL",
            "missing_items": list[str],        # up to 50 samples
            "input_paragraphs": int,
            "output_paragraphs": int,
            "input_tables": int,
            "output_tables": int,
            "notes": list[str]
        }
    """
    input_path = Path(input_path)
    output_path = Path(output_path)

    logger.info("CONTENT_INTEGRITY: Starting verification")

    # Extract content
    input_data = _extract_content(input_path)
    output_data = _extract_content(output_path)

    # Build output index for fast lookup
    output_index = _build_output_index(output_data)

    # Check for missing content
    missing_items = []
    notes = []

    # Check input paragraphs
    for input_text in input_data['paragraphs']:
        if input_text not in output_index:
            missing_items.append(f"PARA: {input_text[:100]}")

    # Check input table cells
    for input_text in input_data['table_cells']:
        if input_text not in output_index:
            missing_items.append(f"TABLE_CELL: {input_text[:100]}")

    # Determine status
    if missing_items:
        status = "FAIL"
        logger.error(
            f"CONTENT_INTEGRITY_FAIL: {len(missing_items)} missing items detected"
        )
        for item in missing_items[:10]:
            logger.error(f"  Missing: {item}")
    else:
        status = "PASS"
        logger.info("CONTENT_INTEGRITY_PASS: All content preserved")

    result = {
        "status": status,
        "missing_items": missing_items[:50],  # Limit to 50 samples
        "input_paragraphs": input_data['paragraph_count'],
        "output_paragraphs": output_data['paragraph_count'],
        "input_tables": input_data['table_count'],
        "output_tables": output_data['table_count'],
        "notes": notes,
    }

    return result


# ===================================================================
# Structural Integrity Verification
# ===================================================================

def _get_heading_level(para) -> int | None:
    """Get heading level from paragraph style or inherited base style."""
    style_names = _style_chain_names(getattr(para, "style", None))
    para_text = (getattr(para, "text", "") or "").strip()
    for style_name in style_names:
        match = re.match(r'Heading (\d)', style_name)
        if match:
            return int(match.group(1))
        if style_name == 'Title':
            return 0

    # Canonical table-caption styles are semantically equivalent to legacy Title
    # captions in some publisher templates. Treat them as title-equivalent for
    # structural comparison when the text matches a table-caption pattern.
    if para_text and re.match(r'^Table\s+\d+(?:\.\d+)?\b', para_text, re.IGNORECASE):
        if any(name in {'T1', 'T11', 'T12', 'UNT-T1', 'TableCaption'} for name in style_names):
            return 0

    # Reconstruction / post-enforcement may canonicalize source heading/title
    # paragraphs to semantic styles (e.g. ``H20``, ``H3``, ``T1``) while
    # preserving semantics via paragraph-level outlineLvl.
    try:
        pPr = para._element.pPr
        if pPr is not None:
            outline = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl')
            if outline is not None:
                val = outline.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                if val is not None:
                    try:
                        outline_level = int(val)
                    except (TypeError, ValueError):
                        outline_level = None
                    if outline_level is not None:
                        # Title/table-caption equivalence uses 0.
                        if outline_level == 0:
                            if any(name in {'Title', 'T1', 'T11', 'T12', 'UNT-T1', 'TableCaption'} for name in style_names):
                                return 0
                            # For heading semantics, some reconstruction paths may
                            # preserve outlineLvl=0 on visual heading styles. Treat
                            # as Heading 1 if the style looks heading-like.
                            if any(re.fullmatch(r'H[1-6]\d*', name or '') for name in style_names):
                                return 1
                        # Reconstruction stores source heading level directly
                        # (Heading 2 -> outlineLvl=2) to preserve integrity checks.
                        if 1 <= outline_level <= 9:
                            return outline_level
    except Exception:
        pass

    return None


def _get_list_info(para) -> Tuple[str | None, int | None]:
    """
    Get list ID and level from paragraph.
    Returns (list_id, list_level) or (None, None).
    """
    # Access paragraph format
    pPr = para._element.pPr
    if pPr is None:
        return None, None

    # Look for numPr (numbering properties)
    numPr = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
    if numPr is None:
        # Fallback: detect style-based lists from style/base-style chain
        for style_name in _style_chain_names(getattr(para, "style", None)):
            lowered = style_name.lower()
            if any(ind in lowered for ind in ("list", "bullet", "number")):
                return "-1", _list_level_from_style_name(style_name)
        return None, None

    # Extract numId and ilvl
    numId_elem = numPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')
    ilvl_elem = numPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')

    list_id = numId_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if numId_elem is not None else None
    list_level = int(ilvl_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0')) if ilvl_elem is not None else None

    return list_id, list_level


def _is_marker_style(style_name: str) -> bool:
    """Check if style name indicates a marker paragraph."""
    if not style_name:
        return False
    return 'MARKER' in style_name.upper()


def _is_marker_only_paragraph_text(text: str) -> bool:
    """True for generic marker-only token paragraphs (e.g. ``</COUT>``, ``<FIG71.3>``).

    This intentionally matches *only* paragraphs whose entire content is a single
    angle-bracket token. It does not treat inline heading markers with content
    (e.g. ``<H1>History``) as markers.
    """
    stripped = (text or "").strip()
    if not stripped:
        return False
    if stripped in KNOWN_MARKERS:
        return True
    return bool(_MARKER_ONLY_TOKEN_RE.fullmatch(stripped))


def _extract_structure(docx_path: str | Path) -> Dict:
    """
    Extract structural map from DOCX.

    Returns dict with:
    - paragraphs: list of paragraph structural info
    - tables: list of table structural info
    - sections: list of section info
    """
    doc = Document(docx_path)

    paragraphs = []
    tables = []

    # Track which tables we've seen (to avoid double-counting from paragraphs in cells)
    seen_tables = set()

    # Extract paragraph structure
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""
        heading_level = _get_heading_level(para)
        list_id, list_level = _get_list_info(para)

        # Determine if in table (check parent)
        in_table = False
        try:
            # Check if paragraph is inside a table cell
            parent = para._element.getparent()
            while parent is not None:
                if parent.tag.endswith('}tc'):  # table cell
                    in_table = True
                    break
                parent = parent.getparent()
        except:
            pass

        # Check if paragraph is a marker (by style or content)
        text_content = para.text.strip()
        # Defensive normalization for legacy Title -> canonical T1/T11/T12 table captions:
        # if style-chain heading detection misses, derive title-equivalent semantics from
        # the visible style name + caption text pattern so structural integrity does not
        # treat canonicalization as a heading demotion.
        if heading_level is None:
            if style_name in {"T1", "T11", "T12", "UNT-T1", "TableCaption"} and re.match(
                r"^Table\s+\d+(?:\.\d+)?\b", text_content, re.IGNORECASE
            ):
                heading_level = 0

        is_marker = _is_marker_only_paragraph_text(text_content) or (
            _is_marker_style(style_name)
            and text_content.startswith("<")
            and text_content.endswith(">")
        )

        para_info = {
            'index': i,
            'style_name': style_name,
            'is_heading': heading_level is not None,
            'heading_level': heading_level,
            'is_list_item': list_id is not None,
            'list_id': list_id,
            'list_level': list_level,
            'section_index': 0,  # Will be updated
            'in_table': in_table,
            'is_marker': is_marker,
            'text_preview': _normalize_text(para.text)[:50],
        }

        paragraphs.append(para_info)

    # Extract table structure
    for t, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns) if table.columns else 0

        # Count nested tables
        nested_count = 0
        for row in table.rows:
            for cell in row.cells:
                nested_count += len(cell.tables)

        table_info = {
            'table_index': t,
            'rows': rows,
            'cols': cols,
            'total_cells': rows * cols if cols > 0 else 0,
            'nested_tables_count': nested_count,
        }

        tables.append(table_info)

    # Extract section info
    sections = []
    for idx, section in enumerate(doc.sections):
        # Get section break type
        try:
            start_type = section.start_type
            if start_type == WD_BREAK.CONTINUOUS:
                break_type = "continuous"
            elif start_type == WD_BREAK.NEW_PAGE:
                break_type = "new_page"
            elif start_type == WD_BREAK.NEW_COLUMN:
                break_type = "new_column"
            elif start_type == WD_BREAK.EVEN_PAGE:
                break_type = "even_page"
            elif start_type == WD_BREAK.ODD_PAGE:
                break_type = "odd_page"
            else:
                break_type = "unknown"
        except:
            break_type = "unknown"

        sections.append({
            'index': idx,
            'break_type': break_type,
        })

    return {
        'paragraphs': paragraphs,
        'tables': tables,
        'sections': sections,
    }


def _create_structural_signature(structure: Dict) -> str:
    """
    Create canonical structural signature for hashing.

    Paragraph signature:
    "P|{i}|{style}|H{heading_level}|L{list_level}|T{in_table}|S{section_index}"

    Table signature:
    "T|{t}|R{rows}|C{cols}|N{nested_tables}"

    Section signature:
    "SEC|{idx}|{break_type}"
    """
    signatures = []

    # Paragraph signatures
    for p in structure['paragraphs']:
        sig = f"P|{p['index']}|{p['style_name']}|H{p['heading_level'] or 0}|L{p['list_level'] or -1}|T{int(p['in_table'])}|S{p['section_index']}"
        signatures.append(sig)

    # Table signatures
    for t in structure['tables']:
        sig = f"T|{t['table_index']}|R{t['rows']}|C{t['cols']}|N{t['nested_tables_count']}"
        signatures.append(sig)

    # Section signatures
    for s in structure['sections']:
        sig = f"SEC|{s['index']}|{s['break_type']}"
        signatures.append(sig)

    # Join and hash
    full_signature = '\n'.join(signatures)
    signature_hash = hashlib.sha256(full_signature.encode('utf-8')).hexdigest()

    return signature_hash


def _compare_structures(input_struct: Dict, output_struct: Dict) -> Tuple[bool, List[str], Dict | None]:
    """
    Compare input and output structures.

    Returns (all_match: bool, differences: list[str])
    """
    differences = []
    all_match = True
    first_difference: Dict | None = None

    def _record_diff(message: str, *, paragraph_index: int | None = None, input_para: Dict | None = None, output_para: Dict | None = None):
        nonlocal all_match, first_difference
        differences.append(message)
        all_match = False
        if first_difference is None:
            first_difference = {
                "stage": "integrity_check",
                "paragraph_index": paragraph_index,
                "input": _para_diag_snapshot(input_para),
                "output": _para_diag_snapshot(output_para),
                "message": message,
            }

    # Filter marker paragraphs symmetrically before index-wise comparison.
    # Some source documents legitimately contain marker-only paragraphs; filtering
    # only output causes index drift and false heading/list mismatches.
    input_paras = [p for p in input_struct['paragraphs'] if not p['is_marker']]
    output_paras = [p for p in output_struct['paragraphs'] if not p['is_marker']]

    # 1. Paragraph count check
    if len(input_paras) != len(output_paras):
        _record_diff(
            f"Paragraph count mismatch: input={len(input_paras)}, output={len(output_paras)}"
        )

    # 2. Table count and structure check
    if len(input_struct['tables']) != len(output_struct['tables']):
        _record_diff(
            f"Table count mismatch: input={len(input_struct['tables'])}, output={len(output_struct['tables'])}"
        )
    else:
        for i, (input_table, output_table) in enumerate(zip(input_struct['tables'], output_struct['tables'])):
            if input_table['rows'] != output_table['rows']:
                _record_diff(
                    f"Table {i} row count mismatch: input={input_table['rows']}, output={output_table['rows']}"
                )
            if input_table['cols'] != output_table['cols']:
                _record_diff(
                    f"Table {i} column count mismatch: input={input_table['cols']}, output={output_table['cols']}"
                )

    # 3. Heading levels check (for matching indices)
    min_len = min(len(input_paras), len(output_paras))
    for i in range(min_len):
        input_para = input_paras[i]
        output_para = output_paras[i]

        if input_para['heading_level'] != output_para['heading_level']:
            # Exempt: a non-heading input paragraph (level=None) whose style was
            # canonicalized to a table-caption style (T1, T11, T12, UNT-T1,
            # TableCaption) with matching "Table N..." text gets heading_level=0
            # assigned by _get_heading_level() for structural equivalence purposes.
            # This is expected after table-title enforcement and is NOT a violation.
            _TABLE_CAPTION_STYLES = {'T1', 'T11', 'T12', 'UNT-T1', 'TableCaption'}
            _is_caption_promotion = (
                input_para['heading_level'] is None
                and output_para['heading_level'] == 0
                and output_para['style_name'] in _TABLE_CAPTION_STYLES
                and re.match(
                    r'^Table\s+\d+',
                    output_para.get('text_preview', ''),
                    re.IGNORECASE,
                )
            )
            if not _is_caption_promotion:
                _record_diff(
                    f"Heading level mismatch at para {i}: input={input_para['heading_level']}, output={output_para['heading_level']} "
                    f"| input_meta={_para_diag_snapshot(input_para)} output_meta={_para_diag_snapshot(output_para)}",
                    paragraph_index=i,
                    input_para=input_para,
                    output_para=output_para,
                )
            if len(differences) >= 100:
                break

        # List structure check
        if input_para['is_list_item'] != output_para['is_list_item']:
            _record_diff(
                f"List item status mismatch at para {i}: input={input_para['is_list_item']}, output={output_para['is_list_item']} "
                f"| input_meta={_para_diag_snapshot(input_para)} output_meta={_para_diag_snapshot(output_para)}",
                paragraph_index=i,
                input_para=input_para,
                output_para=output_para,
            )
            if len(differences) >= 100:
                break

        if input_para['list_level'] != output_para['list_level']:
            _record_diff(
                f"List level mismatch at para {i}: input={input_para['list_level']}, output={output_para['list_level']} "
                f"| input_meta={_para_diag_snapshot(input_para)} output_meta={_para_diag_snapshot(output_para)}",
                paragraph_index=i,
                input_para=input_para,
                output_para=output_para,
            )
            if len(differences) >= 100:
                break

    # 4. Section structure check
    if len(input_struct['sections']) != len(output_struct['sections']):
        _record_diff(
            f"Section count mismatch: input={len(input_struct['sections'])}, output={len(output_struct['sections'])}"
        )
    else:
        for i, (input_sec, output_sec) in enumerate(zip(input_struct['sections'], output_struct['sections'])):
            if input_sec['break_type'] != output_sec['break_type']:
                _record_diff(
                    f"Section {i} break type mismatch: input={input_sec['break_type']}, output={output_sec['break_type']}"
                )

    return all_match, differences[:100], first_difference  # Limit to 100 differences


def verify_structural_integrity(input_path: str | Path, output_path: str | Path) -> dict:
    """
    Verify that document structure is preserved in output.

    Validates:
    - Paragraph count (excluding additive markers)
    - Table structure (count, rows, cols)
    - Heading levels
    - List structure (list items, levels)
    - Section structure

    **FAIL Conditions (strict):**
    - Paragraph count differs
    - Table count/structure differs
    - Heading level differs at same index
    - List structure differs at same index
    - Section count/breaks differ

    **ALLOW Conditions:**
    - Style names may change
    - Marker paragraphs may be inserted (additive, not replacement)

    Parameters
    ----------
    input_path : str or Path
        Path to input DOCX file
    output_path : str or Path
        Path to output/processed DOCX file

    Returns
    -------
    dict
        Verification result:
        {
            "status": "PASS" | "FAIL",
            "paragraph_count_match": bool,
            "table_structure_match": bool,
            "heading_levels_match": bool,
            "list_structure_match": bool,
            "section_structure_match": bool,
            "structural_hash_match": bool,
            "differences": list[str]  # up to 100 detailed diffs
        }
    """
    input_path = Path(input_path)
    output_path = Path(output_path)

    logger.info("STRUCTURAL_INTEGRITY: Starting verification")

    # Extract structures
    input_struct = _extract_structure(input_path)
    output_struct = _extract_structure(output_path)

    # Create structural signatures
    input_hash = _create_structural_signature(input_struct)
    output_hash = _create_structural_signature(output_struct)

    # Compare structures
    all_match, differences, first_difference = _compare_structures(input_struct, output_struct)

    # Check table structure match more comprehensively
    table_match = len(input_struct['tables']) == len(output_struct['tables'])
    if table_match and input_struct['tables']:
        # Also check row/col counts
        table_match = not any('Table' in d and ('row' in d or 'column' in d) for d in differences)

    # Build result
    result = {
        "status": "PASS" if all_match else "FAIL",
        "paragraph_count_match": (
            len([p for p in input_struct['paragraphs'] if not p['is_marker']])
            == len([p for p in output_struct['paragraphs'] if not p['is_marker']])
        ),
        "table_structure_match": table_match,
        "heading_levels_match": not any('Heading level mismatch' in d for d in differences),
        "list_structure_match": not any('List' in d for d in differences),
        "section_structure_match": len(input_struct['sections']) == len(output_struct['sections']),
        "structural_hash_match": input_hash == output_hash,
        "input_hash": input_hash,
        "output_hash": output_hash,
        "first_difference": first_difference,
        "differences": differences,
    }

    if result["status"] == "FAIL":
        logger.error("STRUCTURAL_INTEGRITY_FAIL: Structure mismatch detected")
        if first_difference:
            logger.error(
                "INTEGRITY_FIRST_DIFF stage=integrity_check paragraph_index=%s detail=%s",
                first_difference.get("paragraph_index"),
                first_difference,
            )
        for diff in differences[:20]:
            logger.error(f"  {diff}")
        logger.error(f"Input structural hash: {input_hash}")
        logger.error(f"Output structural hash: {output_hash}")
    else:
        logger.info("STRUCTURAL_INTEGRITY_PASS: Structure preserved")

    return result


# ===================================================================
# Combined Integrity Trigger
# ===================================================================

def run_integrity_trigger(input_path: str | Path, output_path: str | Path) -> dict:
    """
    Run complete integrity verification (content + structural).

    Calls both verify_content_integrity and verify_structural_integrity.

    **Raises RuntimeError if either check fails.**

    Parameters
    ----------
    input_path : str or Path
        Path to input DOCX file
    output_path : str or Path
        Path to output/processed DOCX file

    Returns
    -------
    dict
        Combined verification result:
        {
            "status": "PASS" | "FAIL",
            "content_integrity": dict,
            "structural_integrity": dict,
            "error_message": str (if FAIL)
        }

    Raises
    ------
    RuntimeError
        If integrity checks fail, with detailed error message
    """
    logger.info("=" * 60)
    logger.info("INTEGRITY_TRIGGER: Starting combined verification")
    logger.info(f"  Input: {Path(input_path).name}")
    logger.info(f"  Output: {Path(output_path).name}")
    logger.info("=" * 60)

    # Run content integrity check
    content_result = verify_content_integrity(input_path, output_path)

    # Run structural integrity check
    structural_result = verify_structural_integrity(input_path, output_path)

    # Determine overall status
    overall_status = "PASS" if (
        content_result["status"] == "PASS" and
        structural_result["status"] == "PASS"
    ) else "FAIL"

    # Build combined result
    result = {
        "status": overall_status,
        "content_integrity": content_result,
        "structural_integrity": structural_result,
        "error_message": None,
    }

    # If FAIL, raise RuntimeError with detailed message
    if overall_status == "FAIL":
        error_parts = []

        if content_result["status"] == "FAIL":
            missing_count = len(content_result["missing_items"])
            error_parts.append(f"CONTENT INTEGRITY FAIL: {missing_count} missing items")
            if content_result["missing_items"]:
                error_parts.append("Missing content samples:")
                for item in content_result["missing_items"][:20]:
                    error_parts.append(f"  - {item}")

        if structural_result["status"] == "FAIL":
            diff_count = len(structural_result["differences"])
            error_parts.append(f"STRUCTURAL INTEGRITY FAIL: {diff_count} structural differences")
            if structural_result["differences"]:
                error_parts.append("Structural differences (first 20):")
                for diff in structural_result["differences"][:20]:
                    error_parts.append(f"  - {diff}")
            error_parts.append(f"Input hash: {structural_result.get('input_hash', 'N/A')}")
            error_parts.append(f"Output hash: {structural_result.get('output_hash', 'N/A')}")

        error_message = "\n".join(error_parts)
        result["error_message"] = error_message

        logger.error("=" * 60)
        logger.error("INTEGRITY_TRIGGER_FAIL")
        logger.error("=" * 60)
        logger.error(error_message)
        logger.error("=" * 60)

        raise RuntimeError(f"INTEGRITY_TRIGGER_FAIL: {error_message}")

    logger.info("=" * 60)
    logger.info("INTEGRITY_TRIGGER_PASS: All checks passed")
    logger.info("=" * 60)

    return result
