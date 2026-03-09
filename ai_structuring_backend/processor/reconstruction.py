"""
STAGE 4: Document Reconstruction
- Style Applicator (applies Word paragraph styles)
- DOCX Builder
- Review Report Generator

Applies proper Word styles to paragraphs for professional output.
NO CONTENT MODIFICATION - only styles and formatting are applied.

Key implementation notes (critical for golden tests):
- Paragraph IDs must match ingestion order exactly:
  1) non-empty body paragraphs in document order
  2) then non-empty table cell paragraphs in document order (table->row->cell->paragraph)
- When processing tables, increment para_id per non-empty paragraph (NOT per cell).
"""

from __future__ import annotations

from datetime import datetime
import json
import logging
import os
import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

_INLINE_HEADING_MARKER_RE = re.compile(r"^\s*<h[1-6]\b[^>]*>", re.IGNORECASE)
_VISUAL_HEADING_TAG_RE = re.compile(r"^(?:H[1-6]\d*|SP-H1|EOC-H1|REFH1|REF-H1|APX-H[1-3]|CS-H1|REFH2|REFH2A|REF-H2)$", re.IGNORECASE)

# ---------------------------------------------------------------------------
# Table-highlight review threshold
# ---------------------------------------------------------------------------
# Table-related paragraphs with classification confidence < this value receive
# a yellow highlight so human reviewers can spot uncertain table content.
# Configure via TABLE_REVIEW_HIGHLIGHT_THRESHOLD env var (integer percentage).
# Default: 80  (i.e., highlight when confidence < 80%).
_TABLE_HIGHLIGHT_THRESHOLD: int = int(os.getenv("TABLE_REVIEW_HIGHLIGHT_THRESHOLD", "80"))
_REFERENCE_HIGHLIGHT_THRESHOLD: int = int(
    os.getenv("REFERENCE_REVIEW_HIGHLIGHT_THRESHOLD", str(_TABLE_HIGHLIGHT_THRESHOLD))
)

# Style tags that are "table-related" when they appear as BODY paragraphs.
# These use _TABLE_HIGHLIGHT_THRESHOLD rather than the general < 85 threshold.
_TABLE_RELATED_BODY_TAGS: frozenset[str] = frozenset({
    "T1", "T11", "T12", "UNT-T1",  # table captions
    "TFN", "TSN",                    # table footnote / source note
})
_REFERENCE_RELATED_TAG_PREFIXES: tuple[str, ...] = ("REF", "APX-REF")
_REFERENCE_RELATED_TAGS: frozenset[str] = frozenset({"BIB", "SR", "SRH1"})

# Style definitions with formatting properties - extracted from tagged documents
# Note: Text flow variants (TXT1, TXT2, H11, H12, etc.) inherit from base styles
STYLE_DEFINITIONS = {
    # Document Structure
    "CN": {"font_size": 24, "bold": True, "alignment": "center", "space_after": 6},
    "CT": {"font_size": 28, "bold": True, "alignment": "center", "space_after": 12},
    "CAU": {"font_size": 10, "italic": True, "alignment": "center", "space_after": 6},
    "PN": {"font_size": 10, "alignment": "center"},
    "PT": {"font_size": 14, "bold": True, "alignment": "center"},

    # Headings
    "H1": {"font_size": 16, "bold": True, "space_before": 12, "space_after": 6},
    "H10": {"font_size": 16, "bold": True, "space_before": 12, "space_after": 6},
    "H2": {"font_size": 14, "bold": True, "space_before": 10, "space_after": 4},
    "H20": {"font_size": 14, "bold": True, "space_before": 10, "space_after": 4},
    "H3": {"font_size": 12, "bold": True, "space_before": 8, "space_after": 4},
    "H4": {"font_size": 11, "bold": True, "space_before": 6, "space_after": 2},
    "H5": {"font_size": 11, "bold": True, "italic": True, "space_before": 6, "space_after": 2},
    "H6": {"font_size": 10, "bold": True, "space_before": 4, "space_after": 2},
    "SP-H1": {"font_size": 14, "bold": True, "caps": True, "space_before": 12, "space_after": 6},
    "EOC-H1": {"font_size": 14, "bold": True, "space_before": 12, "space_after": 6},
    # Heading text flow variants
    "H11": {"font_size": 16, "bold": True, "space_before": 12, "space_after": 6},
    "H12": {"font_size": 16, "bold": True, "space_before": 12, "space_after": 6},
    "H21": {"font_size": 14, "bold": True, "space_before": 10, "space_after": 4},

    # Reference Headings
    "REFH1": {"font_size": 14, "bold": True, "space_before": 12, "space_after": 6},
    "REFH2": {"font_size": 12, "bold": True, "space_before": 8, "space_after": 4},
    "REFH2a": {"font_size": 12, "bold": True, "space_before": 8, "space_after": 4},
    "Ref-H1": {"font_size": 14, "bold": True, "space_before": 12, "space_after": 6},
    "Ref-H2": {"font_size": 12, "bold": True, "space_before": 8, "space_after": 4},

    # Body Text
    "TXT": {"font_size": 11, "first_line_indent": 0.5, "space_after": 6},
    "TXT-FLUSH": {"font_size": 11, "space_after": 6},
    "TXT-DC": {"font_size": 11, "space_after": 6},  # Drop cap would need special handling
    "TXT-AU": {"font_size": 10, "italic": True, "space_after": 6},
    "T": {"font_size": 10, "space_after": 2},  # Table cell body text
    # Body text flow variants
    "TXT1": {"font_size": 11, "first_line_indent": 0.5, "space_after": 6},
    "TXT2": {"font_size": 11, "first_line_indent": 0.5, "space_after": 6},
    "TXT3": {"font_size": 11, "first_line_indent": 0.5, "space_after": 6},
    "TXT4": {"font_size": 11, "first_line_indent": 0.5, "space_after": 6},
    "TXT-FLUSH1": {"font_size": 11, "space_after": 6},
    "TXT-FLUSH2": {"font_size": 11, "space_after": 6},
    "TXT-FLUSH4": {"font_size": 11, "space_after": 6},

    # Bulleted Lists
    "BL-FIRST": {"font_size": 11, "left_indent": 0.5, "space_after": 2, "bullet": True},
    "BL-MID": {"font_size": 11, "left_indent": 0.5, "space_after": 2, "bullet": True},
    "BL-LAST": {"font_size": 11, "left_indent": 0.5, "space_after": 6, "bullet": True},
    "UL-FIRST": {"font_size": 11, "left_indent": 0.75, "space_after": 2, "bullet": True},
    "UL-MID": {"font_size": 11, "left_indent": 0.75, "space_after": 2, "bullet": True},
    "UL-LAST": {"font_size": 11, "left_indent": 0.75, "space_after": 6, "bullet": True},

    # Numbered Lists
    "NL-FIRST": {"font_size": 11, "left_indent": 0.5, "space_after": 2, "numbered": True},
    "NL-MID": {"font_size": 11, "left_indent": 0.5, "space_after": 2, "numbered": True},
    "NL-LAST": {"font_size": 11, "left_indent": 0.5, "space_after": 6, "numbered": True},

    # End of Chapter Lists
    "EOC-NL-FIRST": {"font_size": 11, "left_indent": 0.5, "space_after": 2, "numbered": True},
    "EOC-NL-MID": {"font_size": 11, "left_indent": 0.5, "space_after": 2, "numbered": True},
    "EOC-NL-LAST": {"font_size": 11, "left_indent": 0.5, "space_after": 6, "numbered": True},
    "EOC-LL2-MID": {"font_size": 11, "left_indent": 0.75, "space_after": 2},

    # Tables - Titles
    "T1": {"font_size": 10, "bold": True, "space_before": 6, "space_after": 2},
    "T11": {"font_size": 10, "bold": True, "space_before": 6, "space_after": 2},
    "T12": {"font_size": 10, "bold": True, "space_before": 6, "space_after": 2},

    # Tables - Headers
    "T2": {"font_size": 10, "bold": True, "alignment": "center"},
    "T2-C": {"font_size": 10, "bold": True, "alignment": "center"},
    "T21": {"font_size": 10, "bold": True},  # Category/row headers
    "T22": {"font_size": 10, "bold": True, "alignment": "center"},  # Column headers
    "T23": {"font_size": 10, "bold": True, "alignment": "center"},  # Specific headers

    # Tables - Body Cells
    "T3": {"font_size": 10, "bold": True},  # Row header/subhead
    "T5": {"font_size": 10},  # Table body cell (data values)
    "T6": {"font_size": 10},  # Table body cell variant

    # Tables - Lists inside cells
    "TBL-FIRST": {"font_size": 10, "left_indent": 0.25, "space_after": 2, "bullet": True},
    "TBL-MID": {"font_size": 10, "left_indent": 0.25, "bullet": True},
    "TBL-MID0": {"font_size": 10, "left_indent": 0.25, "bullet": True},
    "TBL-LAST": {"font_size": 10, "left_indent": 0.25, "space_after": 4, "bullet": True},
    "TBL-LAST1": {"font_size": 10, "left_indent": 0.25, "space_after": 4, "bullet": True},
    "TBL2-MID": {"font_size": 10, "left_indent": 0.5, "bullet": True},
    "TBL3-MID": {"font_size": 10, "left_indent": 0.75, "bullet": True},
    "TBL4-MID": {"font_size": 10, "left_indent": 1.0, "bullet": True},
    "TUL-MID": {"font_size": 10, "left_indent": 0.25, "bullet": True},

    # Tables - Footnotes
    "TFN": {"font_size": 9, "italic": True, "space_before": 2},
    "TFN1": {"font_size": 9, "italic": True, "space_before": 2},
    "TSN": {"font_size": 9, "space_before": 2},

    # Figures
    "FIG-LEG": {"font_size": 10, "space_before": 6, "space_after": 6},
    "PMI": {"font_size": 9, "italic": True},

    # References
    "REF-N": {"font_size": 10, "hanging_indent": 0.5, "space_after": 2},
    "REF-N0": {"font_size": 10, "hanging_indent": 0.5, "space_after": 2},

    # Chapter Outline
    "COUT-1": {"font_size": 11, "left_indent": 0.25, "space_after": 2},
    "COUT-2": {"font_size": 11, "left_indent": 0.5, "space_after": 2},

    # Equations
    "EQ-ONLY": {"font_size": 11, "alignment": "center", "space_before": 6, "space_after": 6},
    "EQ-MID": {"font_size": 11, "alignment": "center", "space_after": 4},  # Multi-line equation continuation

    # Appendix Styles
    "APX-TYPE": {"font_size": 12, "bold": True, "caps": True, "space_before": 12, "space_after": 4},
    "APX-TTL": {"font_size": 14, "bold": True, "space_after": 6},
    "APX-H1": {"font_size": 16, "bold": True, "space_before": 12, "space_after": 6},
    "APX-H2": {"font_size": 14, "bold": True, "space_before": 10, "space_after": 4},
    "APX-H3": {"font_size": 12, "bold": True, "space_before": 8, "space_after": 4},
    "APX-TXT": {"font_size": 11, "first_line_indent": 0.5, "space_after": 6},
    "APX-TXT-FLUSH": {"font_size": 11, "space_after": 6},
    "APX-CAU": {"font_size": 10, "italic": True, "space_after": 6},
    "APX-REF-N": {"font_size": 10, "hanging_indent": 0.5, "space_after": 2},
    "APX-REFH1": {"font_size": 14, "bold": True, "space_before": 12, "space_after": 6},

    # Box Content
    "NBX1-TTL": {"font_size": 12, "bold": True, "space_after": 4},
    "NBX1-TXT": {"font_size": 10, "first_line_indent": 0.25, "space_after": 4},
    "NBX1-TXT-FLUSH": {"font_size": 10, "space_after": 4},
    "NBX1-BL-FIRST": {"font_size": 10, "left_indent": 0.5, "space_after": 2, "bullet": True},
    "NBX1-BL-MID": {"font_size": 10, "left_indent": 0.5, "space_after": 2, "bullet": True},
    "NBX1-BL-LAST": {"font_size": 10, "left_indent": 0.5, "space_after": 4, "bullet": True},
    "NBX1-BL2-MID": {"font_size": 10, "left_indent": 0.75, "space_after": 2, "bullet": True},
    "NBX1-NL-FIRST": {"font_size": 10, "left_indent": 0.5, "space_after": 2, "numbered": True},
    "NBX1-NL-MID": {"font_size": 10, "left_indent": 0.5, "space_after": 2, "numbered": True},
    "NBX1-NL-LAST": {"font_size": 10, "left_indent": 0.5, "space_after": 4, "numbered": True},
    "NBX1-DIA-FIRST": {"font_size": 10, "left_indent": 0.5, "space_after": 2},
    "NBX1-DIA-MID": {"font_size": 10, "left_indent": 0.5, "space_after": 2},
    "NBX1-DIA-LAST": {"font_size": 10, "left_indent": 0.5, "space_after": 4},
    "NBX1-UNT": {"font_size": 10, "space_after": 4},
    "NBX1-UNT-T2": {"font_size": 10, "bold": True, "space_after": 2},
    "NBX1-SRC": {"font_size": 9, "italic": True, "space_after": 4},
    "BX1-TXT-FIRST": {"font_size": 10, "space_after": 4},

    # Case Studies
    "CS-H1": {"font_size": 14, "bold": True, "space_before": 10, "space_after": 4},
    "CS-TTL": {"font_size": 12, "bold": True, "space_after": 4},
    "CS-TXT": {"font_size": 10, "first_line_indent": 0.25, "space_after": 4},
    "CS-TXT-FLUSH": {"font_size": 10, "space_after": 4},
    "CS-QUES-TXT": {"font_size": 10, "bold": True, "space_after": 2},
    "CS-ANS-TXT": {"font_size": 10, "space_after": 4},

    # Learning Objectives
    "OBJ1": {"font_size": 12, "bold": True, "space_after": 4},
    "OBJ-TXT": {"font_size": 10, "space_after": 4},
    "OBJ-BL-FIRST": {"font_size": 10, "left_indent": 0.5, "space_after": 2, "bullet": True},
    "OBJ-BL-MID": {"font_size": 10, "left_indent": 0.5, "space_after": 2, "bullet": True},
    "OBJ-BL-LAST": {"font_size": 10, "left_indent": 0.5, "space_after": 4, "bullet": True},

    # Special/Miscellaneous
    "SUMHD": {"font_size": 12, "bold": True, "space_before": 8, "space_after": 4},
    "EXT-ONLY": {"font_size": 10, "italic": True, "left_indent": 0.5, "space_after": 6},
}

# ---------------------------------------------------------------------------
# List numbering support
# ---------------------------------------------------------------------------

_BULLET_ABSTRACT_NUM_ID = 9900
_NUMBERED_ABSTRACT_NUM_ID = 9901

_BULLET_TAG_RE = re.compile(r"(?:^|-)(?:BL|UL|TBL|TUL|BUL)\d*(?:-|$)", re.IGNORECASE)
_NUMBERED_TAG_RE = re.compile(r"(?:^|-)NL\d*(?:-|$)", re.IGNORECASE)
_LETTERED_TAG_RE = re.compile(r"(?:^|-)LL(\d*)(?:-|$)", re.IGNORECASE)


def _list_props_for_tag(tag: str) -> tuple[str, int] | None:
    """Return ``(list_kind, level)`` if *tag* is a list style, else *None*.

    *list_kind* is ``"bullet"`` or ``"numbered"``.  *level* is the nesting
    depth (0-based).
    """
    style_def = STYLE_DEFINITIONS.get(tag)
    if style_def:
        if style_def.get("bullet"):
            return ("bullet", _bullet_level(tag))
        if style_def.get("numbered"):
            return ("numbered", _numbered_level(tag))
        # Some semantic list tags (e.g. EOC-LL2-MID) are defined for formatting
        # but omit explicit bullet/numbered flags; fall through to pattern matching.

    # Pattern-based fallback for styles not in STYLE_DEFINITIONS
    # (e.g. BX1-BL-MID, BUL1, etc.)
    if _BULLET_TAG_RE.search(tag):
        return ("bullet", _bullet_level(tag))
    if _NUMBERED_TAG_RE.search(tag):
        return ("numbered", _numbered_level(tag))
    if _LETTERED_TAG_RE.search(tag):
        return ("numbered", _numbered_level(tag))
    return None


def _bullet_level(tag: str) -> int:
    """Derive bullet nesting level from a tag name."""
    upper = tag.upper()
    if "TBL4" in upper or "BL4" in upper:
        return 3
    if "TBL3" in upper or "BL3" in upper:
        return 2
    # Match standalone UL (UL-FIRST, BX1-UL-MID) but not BUL1
    if re.search(r"(?:^|-)UL", upper) or "TBL2" in upper or "-BL2" in upper:
        return 1
    return 0


def _numbered_level(tag: str) -> int:
    """Derive numbered/lettered nesting level from a tag name."""
    upper = (tag or "").upper()
    # Handle explicit nested numbered forms (NL2/NL3/NL4, LL2/LL3/LL4)
    for prefix in ("NL", "LL"):
        for n in (4, 3, 2):
            if f"{prefix}{n}" in upper:
                return n - 1
    return 0


def _heading_base_style_for_tag(tag: str) -> str | None:
    """Return optional built-in base style for a style name.

    Structural rule: custom classifier heading tags (``H1``, ``H3``, ``REFH1``...)
    are visual styles only and must not create new Word heading semantics on source
    non-heading paragraphs. Source heading semantics are preserved separately by
    keeping the original source style during reconstruction.
    """
    if (tag or "") == "Title":
        return "Title"
    return None


def _heading_level_for_style_name(style_name: str) -> int | None:
    """Return heading level for built-in heading/title style names."""
    if not style_name:
        return None
    if style_name == "Title":
        return 0
    match = re.fullmatch(r"Heading (\d)", style_name)
    if match:
        return int(match.group(1))
    return None


def _style_has_heading_semantics(style) -> bool:
    """Return True if a style (or one of its bases) is a Word heading/title style."""
    seen = set()
    current = style
    while current is not None:
        name = getattr(current, "name", "") or ""
        if name in seen:
            break
        seen.add(name)
        if _heading_level_for_style_name(name) is not None:
            return True
        current = getattr(current, "base_style", None)
    return False


def _is_visual_heading_tag(tag: str) -> bool:
    """Classifier heading-like tags that should be visual-only in reconstruction."""
    return bool(_VISUAL_HEADING_TAG_RE.fullmatch((tag or "").strip()))


def _heading_level_for_tag(tag: str) -> int | None:
    """Return heading-equivalent level implied by a classifier tag, if any."""
    base_style = _heading_base_style_for_tag(tag or "")
    return _heading_level_for_style_name(base_style or "")


def _is_table_semantic_tag(tag: str) -> bool:
    """True when *tag* is a table-family semantic style (caption/cell/list/note)."""
    t = (tag or "").upper()
    return (
        t in {"T", "T1", "T11", "T12", "T2", "T2-C", "T21", "T22", "T23", "T3", "T4", "T5", "T6", "TD", "TFN", "TSN", "T0", "T10"}
        or t.startswith(("TH", "TBL", "TNL", "TUL"))
    )


def _make_abstract_num(abstract_num_id: int, fmt: str):
    """Build a ``<w:abstractNum>`` element with levels 0-3."""
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_num_id))
    bullets = ["\u2022", "\u25e6", "\u25aa", "\u2013"]
    for i in range(4):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(i))

        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)

        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)

        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), bullets[i] if fmt == "bullet" else f"%{i + 1}.")
        lvl.append(lvl_text)

        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)

        abstract_num.append(lvl)
    return abstract_num


def _get_numbering_part(doc):
    """Return the document's numbering part, creating it when absent."""
    try:
        return doc.part.numbering_part
    except NotImplementedError:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.opc.packuri import PackURI
        from docx.parts.numbering import NumberingPart as _NP

        numbering_elm = OxmlElement("w:numbering")
        part = _NP(
            PackURI("/word/numbering.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
            numbering_elm,
            doc.part.package,
        )
        doc.part.relate_to(part, RT.NUMBERING)
        return part


def _get_or_create_numbering_ids(doc) -> dict[str, int]:
    """Return ``{"bullet": numId, "numbered": numId}`` for *doc*.

    Creates abstract and concrete numbering definitions when they do not
    already exist.  Results are cached on the numbering element.
    """
    numbering_part = _get_numbering_part(doc)
    numbering_elm = numbering_part.element

    cache_attr = "_pipeline_num_ids"
    cached = getattr(numbering_elm, cache_attr, None)
    if cached is not None:
        return cached

    existing_abstract = set()
    for el in numbering_elm.iterchildren(qn("w:abstractNum")):
        existing_abstract.add(int(el.get(qn("w:abstractNumId"))))

    if _BULLET_ABSTRACT_NUM_ID not in existing_abstract:
        numbering_elm.append(_make_abstract_num(_BULLET_ABSTRACT_NUM_ID, "bullet"))
    if _NUMBERED_ABSTRACT_NUM_ID not in existing_abstract:
        numbering_elm.append(_make_abstract_num(_NUMBERED_ABSTRACT_NUM_ID, "decimal"))

    def _find_or_add_num(abstract_id: int) -> int:
        for num_el in numbering_elm.iterchildren(qn("w:num")):
            abs_ref = num_el.find(qn("w:abstractNumId"))
            if abs_ref is not None and int(abs_ref.get(qn("w:val"))) == abstract_id:
                return int(num_el.get(qn("w:numId")))
        num = numbering_elm.add_num(abstract_id)
        return num.numId

    result = {
        "bullet": _find_or_add_num(_BULLET_ABSTRACT_NUM_ID),
        "numbered": _find_or_add_num(_NUMBERED_ABSTRACT_NUM_ID),
    }
    setattr(numbering_elm, cache_attr, result)
    return result


def ensure_numbering(para, list_kind: str, level: int, doc) -> None:
    """Add ``w:numPr`` to *para* if it does not already have one.

    Safety: any existing ``numId``/``ilvl`` is preserved.
    """
    pPr = para._element.get_or_add_pPr()
    if pPr.numPr is not None:
        return  # preserve existing numbering

    num_ids = _get_or_create_numbering_ids(doc)
    num_id = num_ids[list_kind]

    numPr = pPr._add_numPr()

    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    numPr.append(ilvl)

    numId_elm = OxmlElement("w:numId")
    numId_elm.set(qn("w:val"), str(num_id))
    numPr.append(numId_elm)


class DocumentReconstructor:
    """
    Apply style tags to a DOCX and generate outputs.
    IMPORTANT: This module must not change content; only apply styles/markers.
    """

    def __init__(self, output_dir: str | Path):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True, parents=True)

    def _get_or_create_style(self, doc: Document, style_name: str) -> None:
        """
        Create paragraph style if missing and apply formatting from STYLE_DEFINITIONS (when available).
        """
        styles = doc.styles

        # Already exists
        try:
            existing = styles[style_name]
            # Some source templates already define styles named H1/H2/H3 as semantic
            # heading styles. Our classifier tags with these names are visual-only and
            # must not create heading semantics on source non-heading paragraphs.
            if _is_visual_heading_tag(style_name) and _style_has_heading_semantics(existing):
                try:
                    existing.base_style = styles["Normal"]
                    logger.info(
                        "Rebased existing style '%s' to Normal to prevent heading semantic promotion",
                        style_name,
                    )
                except Exception as exc:
                    logger.debug("Could not rebase existing style '%s': %s", style_name, exc)
            return
        except KeyError:
            pass

        style_def = STYLE_DEFINITIONS.get(style_name, {})

        try:
            new_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            base_style_name = (
                _heading_base_style_for_tag(style_name)
                or "Normal"
            )
            try:
                new_style.base_style = styles[base_style_name]
            except KeyError:
                new_style.base_style = styles["Normal"]

            font = new_style.font
            font.size = Pt(style_def.get("font_size", 11))
            font.bold = style_def.get("bold", False)
            font.italic = style_def.get("italic", False)
            font.all_caps = style_def.get("caps", False)

            if "color" in style_def:
                font.color.rgb = RGBColor.from_string(style_def["color"])

            pf = new_style.paragraph_format

            if "alignment" in style_def:
                alignment_map = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
                }
                pf.alignment = alignment_map.get(style_def["alignment"], WD_ALIGN_PARAGRAPH.LEFT)

            if "space_before" in style_def:
                pf.space_before = Pt(style_def["space_before"])
            if "space_after" in style_def:
                pf.space_after = Pt(style_def["space_after"])

            if "first_line_indent" in style_def:
                pf.first_line_indent = Inches(style_def["first_line_indent"])
            if "left_indent" in style_def:
                pf.left_indent = Inches(style_def["left_indent"])
            if "hanging_indent" in style_def:
                pf.first_line_indent = Inches(-style_def["hanging_indent"])
                pf.left_indent = Inches(style_def["hanging_indent"])

            logger.debug("Created style: %s", style_name)
        except Exception as e:
            logger.warning("Could not create style '%s': %s", style_name, e)

    def ensure_paragraph_style(self, doc: Document, style_name: str):
        """
        Return a paragraph style object for `style_name`.
        If it doesn't exist, create it (with formatting if available).
        """
        self._get_or_create_style(doc, style_name)
        return doc.styles[style_name]

    @staticmethod
    def _iter_nonempty_body_paragraphs(doc: Document):
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                yield para

    @staticmethod
    def _iter_nonempty_table_paragraphs(doc: Document):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text and para.text.strip():
                            yield para

    @staticmethod
    def _iter_all_table_paragraphs(doc: Document):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para

    @staticmethod
    def _style_chain_names(style) -> list[str]:
        """Return style name chain (current -> base styles)."""
        names: list[str] = []
        seen = set()
        current = style
        while current is not None:
            name = getattr(current, "name", "") or ""
            if name and name not in seen:
                names.append(name)
                seen.add(name)
            current = getattr(current, "base_style", None)
        return names

    def _source_heading_level(self, para) -> int | None:
        """Return heading level from paragraph style/base-style chain."""
        for style_name in self._style_chain_names(getattr(para, "style", None)):
            level = _heading_level_for_style_name(style_name)
            if level is not None:
                return level
        return None

    @staticmethod
    def _set_paragraph_outline_level(para, level: int) -> None:
        """Set paragraph-level ``w:outlineLvl`` (0-based) to preserve Title semantics."""
        pPr = para._element.get_or_add_pPr()
        outline = pPr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            pPr.append(outline)
        outline.set(qn("w:val"), str(level))

    def _snapshot_paragraph_counts(self, doc: Document) -> dict[str, int]:
        """Capture body/table paragraph counts (total + non-empty) for mutation guard."""
        return {
            "body_total": len(doc.paragraphs),
            "body_nonempty": sum(1 for _ in self._iter_nonempty_body_paragraphs(doc)),
            "table_total": sum(1 for _ in self._iter_all_table_paragraphs(doc)),
            "table_nonempty": sum(1 for _ in self._iter_nonempty_table_paragraphs(doc)),
        }

    @staticmethod
    def _assert_paragraph_counts_unchanged(before: dict[str, int], after: dict[str, int]) -> None:
        """Fail fast if any paragraph count changed during reconstruction."""
        changed = [
            f"{k} {before[k]}->{after[k]}"
            for k in sorted(before.keys())
            if before.get(k) != after.get(k)
        ]
        if changed:
            raise ValueError(
                "Reconstruction paragraph-count guard failed: " + ", ".join(changed)
            )

    @staticmethod
    def _source_list_state(para) -> tuple[bool, bool]:
        """Return ``(is_list, has_numpr)`` for *para* BEFORE style mutation."""
        pPr = para._element.find(qn("w:pPr"))
        numPr = pPr.find(qn("w:numPr")) if pPr is not None else None
        has_numpr = numPr is not None
        if has_numpr:
            return True, True

        # Fallback for documents that encode list semantics only via paragraph style.
        style_name = (para.style.name if getattr(para, "style", None) else "") or ""
        is_list_style = any(ind in style_name.lower() for ind in ("list", "bullet", "number"))
        return is_list_style, False

    @staticmethod
    def _preserve_source_list_style(
        source_is_list: bool,
        source_has_numpr: bool,
        *,
        target_tag: str = "",
        source_text: str = "",
    ) -> bool:
        """Return True when changing paragraph style would destroy style-based list semantics.

        Applies only to paragraphs whose list status is encoded entirely via a
        named paragraph style (no ``numPr`` XML element).  For these paragraphs
        the source style *must* be preserved to avoid a ``STRUCTURE_GUARD_FAIL``:
        the structure guard reads list status from the style chain on both the
        input and output documents, so any style change from a list-named style
        to a non-list style (e.g. Heading 1) is flagged as a list mutation.

        Inline heading markers in ``source_text`` do NOT override list
        preservation for style-based lists: the paragraph's list semantics,
        as declared by its style name, take precedence over any embedded marker.
        (For numPr-based lists the guard is moot because ``source_has_numpr``
        is True and this function returns False on the first gate.)
        """
        if not (source_is_list and not source_has_numpr):
            return False
        # If the target style is itself a semantic list tag, allow re-styling and
        # preserve list structure by upgrading to explicit numPr in _apply_list_numbering.
        if _list_props_for_tag(target_tag or ""):
            return False
        return True

    @staticmethod
    def _preserve_source_heading_style(
        source_heading_level: int | None,
        tag: str,
        *,
        in_table: bool = False,
        source_text: str = "",
    ) -> bool:
        target_heading_level = _heading_level_for_tag(tag)
        # Source headings always keep their source paragraph style so Word heading
        # semantics (outline level) cannot drift during reconstruction.
        #
        # Exception: source ``Title`` (level 0) is often used as a legacy visual
        # style for table content/captions. Allow canonical table tagging while
        # preserving title-equivalent semantics via paragraph-level outlineLvl.
        if source_heading_level == 0:
            if _INLINE_HEADING_MARKER_RE.match(source_text or "") and _is_visual_heading_tag(tag):
                return False
            if in_table or _is_table_semantic_tag(tag):
                return False
            return True
        if source_heading_level is not None:
            if _INLINE_HEADING_MARKER_RE.match(source_text or "") and _is_visual_heading_tag(tag):
                return False
            # Publisher tables sometimes use built-in heading styles (e.g. Heading 5)
            # as visual formatting inside table cells. Allow semantic table tags to
            # override those source heading styles within table contexts while
            # preserving list/table structure via the existing reconstruction guards.
            if in_table and _is_table_semantic_tag(tag):
                return False
            return True
        # Source non-headings may receive visual H* styles, but those styles are
        # created as non-semantic (base Normal) to avoid structural promotion.
        _ = target_heading_level
        return False

    @staticmethod
    def _apply_list_numbering(doc, para, tag, *, source_is_list: bool, source_has_numpr: bool):
        """Preserve existing list numbering XML only for paragraphs that were already lists."""
        props = _list_props_for_tag(tag)
        if not props:
            return

        # Structural safety: never introduce list XML onto a paragraph that was not
        # a list in the source document.
        if not source_is_list:
            logger.debug("Skipping numPr injection for non-list source paragraph (tag=%s)", tag)
            return

        # Style-based lists (no numPr in source) are upgraded to explicit numPr when the
        # target tag is a semantic list style so we can apply semantic style names
        # without losing list structure.
        if not source_has_numpr:
            logger.debug("Upgrading style-based list to explicit numPr (tag=%s)", tag)

        ensure_numbering(para, props[0], props[1], doc)

    @staticmethod
    def _is_reference_review_candidate(clf: dict, tag: str) -> bool:
        if (tag or "").startswith(_REFERENCE_RELATED_TAG_PREFIXES) or (tag or "") in _REFERENCE_RELATED_TAGS:
            return True
        if clf.get("is_reference_zone"):
            return True
        zone = (clf.get("context_zone") or "").upper()
        return zone == "REFERENCE"

    def apply_styles(
        self,
        source_path: str | Path,
        classifications: list[dict],
        output_name: Optional[str] = None,
        table_highlight_threshold: Optional[int] = None,
    ) -> Path:
        """
        Apply classification tags as Word paragraph styles to BOTH body and table paragraphs.
        MUST mirror ingestion paragraph order.

        Parameters
        ----------
        table_highlight_threshold : int or None
            Confidence threshold (0-100) below which table-related paragraphs and
            reference-zone paragraphs receive a yellow highlight for human review.
            Defaults to the module-level constant
            ``_TABLE_HIGHLIGHT_THRESHOLD`` (reads ``TABLE_REVIEW_HIGHLIGHT_THRESHOLD``
            env var; fallback 80).
        """
        source_path = Path(source_path)
        doc = Document(source_path)

        # Resolve table highlight threshold (parameter overrides module default)
        _tbl_thresh = _TABLE_HIGHLIGHT_THRESHOLD if table_highlight_threshold is None else table_highlight_threshold
        _ref_thresh = _REFERENCE_HIGHLIGHT_THRESHOLD if table_highlight_threshold is None else table_highlight_threshold

        clf_lookup = {int(c["id"]): c for c in classifications}

        # Pre-create all styles referenced by classifications
        for tag in {c["tag"] for c in classifications}:
            self._get_or_create_style(doc, tag)

        # Reconstruction mutation guard: paragraph counts must not change.
        counts_before = self._snapshot_paragraph_counts(doc)
        logger.info(
            "Reconstruction counts (before): body_total=%s body_nonempty=%s table_total=%s table_nonempty=%s",
            counts_before["body_total"],
            counts_before["body_nonempty"],
            counts_before["table_total"],
            counts_before["table_nonempty"],
        )

        para_id = 1

        # Body paragraphs
        for para in self._iter_nonempty_body_paragraphs(doc):
            clf = clf_lookup.get(para_id)
            if clf:
                tag = clf["tag"]
                conf = int(clf.get("confidence", 85))
                source_is_list, source_has_numpr = self._source_list_state(para)
                source_heading_level = self._source_heading_level(para)
                preserve_list_style = self._preserve_source_list_style(
                    source_is_list,
                    source_has_numpr,
                    target_tag=tag,
                    source_text=para.text or "",
                )
                preserve_heading_style = self._preserve_source_heading_style(
                    source_heading_level, tag, in_table=False, source_text=para.text or ""
                )
                if not preserve_list_style and not preserve_heading_style:
                    para.style = self.ensure_paragraph_style(doc, tag)
                    if source_heading_level is not None:
                        self._set_paragraph_outline_level(para, source_heading_level)
                self._apply_list_numbering(
                    doc, para, tag,
                    source_is_list=source_is_list,
                    source_has_numpr=source_has_numpr,
                )
                # Table-related body paragraphs (captions, footnotes, source notes)
                # use the configurable table threshold; all others use the general < 85.
                if tag in _TABLE_RELATED_BODY_TAGS:
                    if conf < _tbl_thresh:
                        self._highlight_for_review(para)
                elif self._is_reference_review_candidate(clf, tag):
                    if conf < _ref_thresh:
                        self._highlight_for_review(para)
                else:
                    if conf < 85:
                        self._highlight_for_review(para)
            para_id += 1

        # Table paragraphs (IMPORTANT: para_id increments per non-empty table paragraph)
        for para in self._iter_nonempty_table_paragraphs(doc):
            clf = clf_lookup.get(para_id)
            if clf:
                tag = clf["tag"]
                conf = int(clf.get("confidence", 85))
                source_is_list, source_has_numpr = self._source_list_state(para)
                source_heading_level = self._source_heading_level(para)
                preserve_list_style = self._preserve_source_list_style(
                    source_is_list,
                    source_has_numpr,
                    target_tag=tag,
                    source_text=para.text or "",
                )
                preserve_heading_style = self._preserve_source_heading_style(
                    source_heading_level, tag, in_table=True, source_text=para.text or ""
                )
                if not preserve_list_style and not preserve_heading_style:
                    para.style = self.ensure_paragraph_style(doc, tag)
                    if source_heading_level is not None:
                        self._set_paragraph_outline_level(para, source_heading_level)
                self._apply_list_numbering(
                    doc, para, tag,
                    source_is_list=source_is_list,
                    source_has_numpr=source_has_numpr,
                )
                # All in-table paragraphs use the configurable table threshold.
                if conf < _tbl_thresh:
                    self._highlight_for_review(para)
            para_id += 1

        counts_after = self._snapshot_paragraph_counts(doc)
        logger.info(
            "Reconstruction counts (after): body_total=%s body_nonempty=%s table_total=%s table_nonempty=%s",
            counts_after["body_total"],
            counts_after["body_nonempty"],
            counts_after["table_total"],
            counts_after["table_nonempty"],
        )
        self._assert_paragraph_counts_unchanged(counts_before, counts_after)

        if output_name is None:
            output_name = f"{source_path.stem}_tagged_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        output_path = self.output_dir / output_name
        doc.save(output_path)
        logger.info("Saved styled document to %s", output_path)
        return output_path

    @staticmethod
    def _highlight_for_review(para):
        """
        Yellow highlight for low-confidence items.
        """
        try:
            for run in para.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        except Exception:
            pass

    def apply_tags_with_markers(
        self, source_path: str | Path, classifications: list[dict], output_name: Optional[str] = None
    ) -> Path:
        """
        Apply tags as <TAG> markers at the start of EACH non-empty paragraph (body + table).
        Also applies paragraph style equal to the tag name.
        """
        source_path = Path(source_path)
        doc = Document(source_path)

        clf_lookup = {int(c["id"]): c for c in classifications}

        # Pre-create styles referenced by classifications (so "para.style = tag" always works)
        for tag in {c["tag"] for c in classifications}:
            self._get_or_create_style(doc, tag)

        counts_before = self._snapshot_paragraph_counts(doc)
        logger.info(
            "Reconstruction counts (before markers): body_total=%s body_nonempty=%s table_total=%s table_nonempty=%s",
            counts_before["body_total"],
            counts_before["body_nonempty"],
            counts_before["table_total"],
            counts_before["table_nonempty"],
        )

        def _has_any_marker(text: str) -> bool:
            return text.lstrip().startswith("<")

        para_id = 1

        # Body paragraphs
        for para in self._iter_nonempty_body_paragraphs(doc):
            clf = clf_lookup.get(para_id)
            if clf:
                tag = clf["tag"]
                source_is_list, source_has_numpr = self._source_list_state(para)
                source_heading_level = self._source_heading_level(para)
                preserve_list_style = self._preserve_source_list_style(
                    source_is_list,
                    source_has_numpr,
                    target_tag=tag,
                    source_text=para.text or "",
                )
                preserve_heading_style = self._preserve_source_heading_style(
                    source_heading_level, tag, in_table=False, source_text=para.text or ""
                )
                if not preserve_list_style and not preserve_heading_style:
                    para.style = self.ensure_paragraph_style(doc, tag)
                    if source_heading_level is not None:
                        self._set_paragraph_outline_level(para, source_heading_level)
                self._apply_list_numbering(
                    doc, para, tag,
                    source_is_list=source_is_list,
                    source_has_numpr=source_has_numpr,
                )

                if not _has_any_marker(para.text):
                    # Prefix marker without removing existing text
                    if para.runs:
                        para.runs[0].text = f"<{tag}> " + para.runs[0].text
                    else:
                        para.text = f"<{tag}> {para.text}"
            para_id += 1

        # Table paragraphs (apply to EACH non-empty table paragraph)
        for para in self._iter_nonempty_table_paragraphs(doc):
            clf = clf_lookup.get(para_id)
            if clf:
                tag = clf["tag"]
                source_is_list, source_has_numpr = self._source_list_state(para)
                source_heading_level = self._source_heading_level(para)
                preserve_list_style = self._preserve_source_list_style(
                    source_is_list,
                    source_has_numpr,
                    target_tag=tag,
                    source_text=para.text or "",
                )
                preserve_heading_style = self._preserve_source_heading_style(
                    source_heading_level, tag, in_table=True, source_text=para.text or ""
                )
                if not preserve_list_style and not preserve_heading_style:
                    para.style = self.ensure_paragraph_style(doc, tag)
                    if source_heading_level is not None:
                        self._set_paragraph_outline_level(para, source_heading_level)
                self._apply_list_numbering(
                    doc, para, tag,
                    source_is_list=source_is_list,
                    source_has_numpr=source_has_numpr,
                )

                if not _has_any_marker(para.text):
                    if para.runs:
                        para.runs[0].text = f"<{tag}> " + para.runs[0].text
                    else:
                        para.text = f"<{tag}> {para.text}"
            para_id += 1

        counts_after = self._snapshot_paragraph_counts(doc)
        logger.info(
            "Reconstruction counts (after markers): body_total=%s body_nonempty=%s table_total=%s table_nonempty=%s",
            counts_after["body_total"],
            counts_after["body_nonempty"],
            counts_after["table_total"],
            counts_after["table_nonempty"],
        )
        self._assert_paragraph_counts_unchanged(counts_before, counts_after)

        if output_name is None:
            output_name = f"{source_path.stem}_tagged_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        output_path = self.output_dir / output_name
        doc.save(output_path)
        logger.info("Saved tagged document with markers to %s", output_path)
        return output_path

    def generate_review_report(self, document_name: str, filtered_results: dict, output_name: Optional[str] = None) -> Path:
        doc = Document()
        doc.add_heading("Pre-Editor Review Report", level=0)
        doc.add_paragraph(f"Source Document: {document_name}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()

        doc.add_heading("Summary", level=1)
        summary = filtered_results.get("summary", {})
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = "Table Grid"

        rows_data = [
            ("Total Paragraphs", str(summary.get("total_paragraphs", 0))),
            ("Auto-Applied", str(summary.get("auto_applied", 0))),
            ("Needs Review", str(summary.get("needs_review", 0))),
            ("Auto-Apply Rate", f"{summary.get('auto_apply_percentage', 0):.1f}%"),
        ]
        for i, (label, value) in enumerate(rows_data):
            summary_table.rows[i].cells[0].text = label
            summary_table.rows[i].cells[1].text = value

        doc.add_paragraph()

        needs_review = filtered_results.get("needs_review", [])
        if needs_review:
            doc.add_heading("Items Requiring Review", level=1)
            doc.add_paragraph(
                f"The following {len(needs_review)} items have confidence below 85% and require human review."
            )
            doc.add_paragraph()

            for item in needs_review:
                p = doc.add_paragraph()
                run = p.add_run(f"Paragraph {item['id']}")
                run.bold = True

                txt = item.get("original_text", "")
                doc.add_paragraph(f'Text: "{(txt[:100] + "...") if len(txt) > 100 else txt}"')
                doc.add_paragraph(f"Suggested Tag: {item.get('tag', 'Unknown')} (Confidence: {item.get('confidence', 0)}%)")

                if item.get("reasoning"):
                    doc.add_paragraph(f"Reasoning: {item['reasoning']}")
                if item.get("alternatives"):
                    doc.add_paragraph(f"Alternative Tags: {', '.join(item['alternatives'])}")
                doc.add_paragraph()
        else:
            doc.add_heading("All Items Auto-Applied", level=1)
            doc.add_paragraph("All paragraphs were classified with high confidence (≥85%). No manual review required.")

        if output_name is None:
            output_name = f"review_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        output_path = self.output_dir / output_name
        doc.save(output_path)
        logger.info("Saved review report to %s", output_path)
        return output_path

    def generate_json_output(
        self,
        document_name: str,
        classifications: list[dict],
        filtered_results: dict,
        output_name: Optional[str] = None,
    ) -> Path:
        output_data = {
            "document_name": document_name,
            "processed_at": datetime.now().isoformat(),
            "summary": filtered_results.get("summary", {}),
            "classifications": classifications,
            "flagged_items": filtered_results.get("needs_review", []),
        }
        if output_name is None:
            output_name = f"classification_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"

        output_path = self.output_dir / output_name
        output_path.write_text(json.dumps(output_data, indent=2), encoding="utf-8")
        logger.info("Saved JSON results to %s", output_path)
        return output_path

    def generate_html_report(
        self,
        document_name: str,
        classifications: list[dict],
        filtered_results: dict,
        output_name: Optional[str] = None,
    ) -> Path:
        from .html_report import generate_html_report

        if output_name is None:
            output_name = f"classification_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"

        output_path = self.output_dir / output_name
        generate_html_report(document_name, classifications, filtered_results, output_path)
        logger.info("Saved HTML report to %s", output_path)
        return output_path


def reconstruct_document(
    source_path: str | Path,
    classifications: list[dict],
    filtered_results: dict,
    output_dir: str | Path,
    use_markers: bool = False,
    output_base: str | None = None,
) -> dict:
    """
    Convenience function to generate all outputs.
    """
    source_path = Path(source_path)
    reconstructor = DocumentReconstructor(output_dir)

    if output_base is None:
        output_base = f"{source_path.stem}_processed"

    tagged_name = f"{output_base}.docx"
    report_name = f"{output_base}_review.docx"
    json_name = f"{output_base}_results.json"
    html_name = f"{output_base}_report.html"

    if use_markers:
        tagged_path = reconstructor.apply_tags_with_markers(source_path, classifications, tagged_name)
    else:
        tagged_path = reconstructor.apply_styles(source_path, classifications, tagged_name)

    report_path = reconstructor.generate_review_report(source_path.name, filtered_results, report_name)
    json_path = reconstructor.generate_json_output(source_path.name, classifications, filtered_results, json_name)
    html_path = reconstructor.generate_html_report(source_path.name, classifications, filtered_results, html_name)

    return {
        "tagged_document": tagged_path.name,
        "review_report": report_path.name,
        "json_results": json_path.name,
        "html_report": html_path.name,
    }
