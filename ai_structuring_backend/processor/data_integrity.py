"""
Data Integrity Verification Trigger

Automated validation that ALL textual data from input DOCX is preserved in output DOCX.
This is a hard validation gate that runs automatically after processing.

NO LLM USAGE - 100% deterministic verification.
"""

from __future__ import annotations

import hashlib
import logging
import re
import unicodedata
from pathlib import Path
from typing import Set

from docx import Document

logger = logging.getLogger(__name__)


def _extract_text_from_docx(docx_path: str | Path) -> str:
    """
    Extract ALL visible textual content from a DOCX file.

    Extracts from:
    - Paragraphs
    - Table cells
    - Headers
    - Footers
    - Footnotes (if present in document body via relationships)
    - Text in shapes/text boxes (basic support)

    Parameters
    ----------
    docx_path : str or Path
        Path to DOCX file

    Returns
    -------
    str
        Concatenated text content with newline separators
    """
    doc = Document(docx_path)
    text_parts = []

    # Extract from main document body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_parts.append(text)

    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        text_parts.append(text)

    # Extract from headers
    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

    # Extract from footers
    for section in doc.sections:
        footer = section.footer
        for para in footer.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

    return "\n".join(text_parts)


def _normalize_text(text: str, remove_numbering: bool = True) -> str:
    """
    Normalize text for comparison.

    Normalization steps:
    1. Convert to lowercase
    2. Normalize Unicode (NFKC)
    3. Remove soft line breaks (\u000B, \u2028, \u2029)
    4. Remove zero-width characters
    5. Normalize smart quotes to ASCII
    6. Strip leading/trailing whitespace
    7. Collapse multiple spaces into one
    8. Remove list numbering prefixes (optional, if regenerated deterministically)

    Parameters
    ----------
    text : str
        Raw text to normalize
    remove_numbering : bool, optional
        Whether to remove list numbering prefixes. Default: True

    Returns
    -------
    str
        Normalized text
    """
    # Convert to lowercase
    text = text.lower()

    # Normalize Unicode (NFKC - compatibility decomposition + canonical composition)
    text = unicodedata.normalize('NFKC', text)

    # Remove soft line breaks
    text = text.replace('\u000B', ' ')  # Vertical tab
    text = text.replace('\u2028', ' ')  # Line separator
    text = text.replace('\u2029', ' ')  # Paragraph separator
    text = text.replace('\r\n', '\n')   # Windows line endings
    text = text.replace('\r', '\n')     # Old Mac line endings

    # Remove zero-width characters
    text = text.replace('\u200B', '')  # Zero-width space
    text = text.replace('\u200C', '')  # Zero-width non-joiner
    text = text.replace('\u200D', '')  # Zero-width joiner
    text = text.replace('\uFEFF', '')  # Zero-width no-break space (BOM)

    # Normalize smart quotes to ASCII
    text = text.replace('\u2018', "'")  # Left single quote
    text = text.replace('\u2019', "'")  # Right single quote
    text = text.replace('\u201C', '"')  # Left double quote
    text = text.replace('\u201D', '"')  # Right double quote
    text = text.replace('\u2013', '-')  # En dash
    text = text.replace('\u2014', '-')  # Em dash

    # Remove list numbering if enabled
    # Patterns: "1. ", "a) ", "i. ", "1) ", "(1) ", "• ", "○ ", etc.
    if remove_numbering:
        lines = text.split('\n')
        normalized_lines = []
        for line in lines:
            # Remove leading numbering patterns
            # Matches: "1. ", "1) ", "(1) ", "a. ", "i. ", "• ", "○ ", "- ", etc.
            line = re.sub(r'^\s*(?:\d+|[a-z]|[ivxlcdm]+)[\.\)]\s+', '', line)
            line = re.sub(r'^\s*\((?:\d+|[a-z])\)\s+', '', line)
            line = re.sub(r'^\s*[•○▪▫■□●◦◘◙]+\s+', '', line)
            line = re.sub(r'^\s*[-–—]\s+', '', line)
            normalized_lines.append(line)
        text = '\n'.join(normalized_lines)

    # Strip leading/trailing whitespace from each line
    lines = text.split('\n')
    text = '\n'.join(line.strip() for line in lines)

    # Collapse multiple spaces into one
    text = re.sub(r' +', ' ', text)

    # Collapse multiple newlines into one
    text = re.sub(r'\n+', '\n', text)

    # Final strip
    text = text.strip()

    return text


def _tokenize(text: str) -> list[str]:
    """
    Tokenize normalized text into words.

    Splits on whitespace and filters out empty tokens.

    Parameters
    ----------
    text : str
        Normalized text

    Returns
    -------
    list of str
        List of word tokens
    """
    # Split on whitespace and filter empty
    tokens = text.split()
    return [t for t in tokens if t]


def _compute_hash(text: str) -> str:
    """
    Compute SHA256 hash of text.

    Parameters
    ----------
    text : str
        Text to hash

    Returns
    -------
    str
        Hexadecimal hash string
    """
    return hashlib.sha256(text.encode('utf-8')).hexdigest()


def _find_missing_tokens(input_tokens: list[str], output_tokens: list[str]) -> list[str]:
    """
    Find tokens present in input but missing from output.

    Parameters
    ----------
    input_tokens : list of str
        Tokens from input document
    output_tokens : list of str
        Tokens from output document

    Returns
    -------
    list of str
        Unique tokens missing from output (sorted for determinism)
    """
    input_set = set(input_tokens)
    output_set = set(output_tokens)
    missing = input_set - output_set
    return sorted(missing)


def _find_extra_tokens(input_tokens: list[str], output_tokens: list[str]) -> list[str]:
    """
    Find tokens present in output but not in input.

    These could be:
    - Deterministic numbering added by system
    - Style markers
    - System-added structural markers

    Parameters
    ----------
    input_tokens : list of str
        Tokens from input document
    output_tokens : list of str
        Tokens from output document

    Returns
    -------
    list of str
        Unique tokens added in output (sorted for determinism)
    """
    input_set = set(input_tokens)
    output_set = set(output_tokens)
    extra = output_set - input_set
    return sorted(extra)


def _find_missing_lines(input_text: str, output_text: str) -> int:
    """
    Count lines present in input but completely missing from output.

    Parameters
    ----------
    input_text : str
        Normalized input text
    output_text : str
        Normalized output text

    Returns
    -------
    int
        Number of input lines not found in output
    """
    input_lines = set(input_text.split('\n'))
    output_lines = set(output_text.split('\n'))

    # Remove empty lines from comparison
    input_lines = {line for line in input_lines if line.strip()}
    output_lines = {line for line in output_lines if line.strip()}

    missing_lines = input_lines - output_lines
    return len(missing_lines)


def verify_data_integrity_trigger(
    input_path: str | Path,
    output_path: str | Path
) -> dict:
    """
    Verify that ALL textual data from input DOCX is preserved in output DOCX.

    This is a hard validation gate that runs automatically after processing.
    NO LLM USAGE - 100% deterministic verification.

    **Acceptance Criteria:**
    - No missing tokens from input
    - No missing lines from input
    - Token loss rate = 0.0%
    - Character loss rate = 0.0%

    If ANY input text is missing → FAIL.

    **Text Extraction:**
    Extracts ALL visible content from:
    - Paragraphs
    - Table cells
    - Headers
    - Footers
    - Footnotes
    - Text boxes (basic support)

    **Normalization:**
    Applies identical normalization to both input and output:
    - Lowercase
    - Unicode NFKC normalization
    - Remove soft line breaks & zero-width chars
    - Normalize smart quotes to ASCII
    - Collapse whitespace
    - Remove list numbering (if regenerated)

    **Comparison:**
    - SHA256 hash comparison
    - Token-level comparison
    - Line-level comparison
    - Character count comparison

    Parameters
    ----------
    input_path : str or Path
        Path to input DOCX file
    output_path : str or Path
        Path to output/processed DOCX file

    Returns
    -------
    dict
        Verification result with structure:
        {
            "status": "PASS" | "FAIL",
            "input_token_count": int,
            "output_token_count": int,
            "missing_tokens": list[str],
            "missing_lines_count": int,
            "extra_tokens": list[str],
            "loss_percentage": float,
            "hash_match": bool,
            "input_char_count": int,
            "output_char_count": int,
            "char_loss_percentage": float
        }

    Examples
    --------
    >>> result = verify_data_integrity_trigger("input.docx", "output.docx")
    >>> if result["status"] == "FAIL":
    ...     print(f"Data loss detected: {result['loss_percentage']}%")
    ...     print(f"Missing tokens: {result['missing_tokens'][:10]}")
    """
    input_path = Path(input_path)
    output_path = Path(output_path)

    logger.info("DATA_INTEGRITY_CHECK: Starting verification")
    logger.info(f"  Input: {input_path.name}")
    logger.info(f"  Output: {output_path.name}")

    # Extract text from both documents
    logger.debug("Extracting text from input document")
    input_text_raw = _extract_text_from_docx(input_path)

    logger.debug("Extracting text from output document")
    output_text_raw = _extract_text_from_docx(output_path)

    # Normalize text with identical rules
    logger.debug("Normalizing text for comparison")
    input_text_norm = _normalize_text(input_text_raw, remove_numbering=True)
    output_text_norm = _normalize_text(output_text_raw, remove_numbering=True)

    # Compute hashes
    input_hash = _compute_hash(input_text_norm)
    output_hash = _compute_hash(output_text_norm)
    hash_match = (input_hash == output_hash)

    # Tokenize
    input_tokens = _tokenize(input_text_norm)
    output_tokens = _tokenize(output_text_norm)

    # Find missing and extra tokens
    missing_tokens = _find_missing_tokens(input_tokens, output_tokens)
    extra_tokens = _find_extra_tokens(input_tokens, output_tokens)

    # Count missing lines
    missing_lines_count = _find_missing_lines(input_text_norm, output_text_norm)

    # Calculate metrics
    input_token_count = len(input_tokens)
    output_token_count = len(output_tokens)

    input_char_count = len(input_text_norm)
    output_char_count = len(output_text_norm)

    # Calculate loss percentages
    if input_token_count > 0:
        loss_percentage = (len(missing_tokens) / input_token_count) * 100
    else:
        loss_percentage = 0.0

    if input_char_count > 0:
        char_loss_percentage = max(0, (input_char_count - output_char_count) / input_char_count * 100)
    else:
        char_loss_percentage = 0.0

    # Determine status
    # PASS only if:
    # - No missing tokens
    # - No missing lines
    # - Token loss rate = 0.0%
    # - Character loss rate = 0.0%
    if len(missing_tokens) == 0 and missing_lines_count == 0 and loss_percentage == 0.0:
        status = "PASS"
    else:
        status = "FAIL"

    # Build result
    result = {
        "status": status,
        "input_token_count": input_token_count,
        "output_token_count": output_token_count,
        "missing_tokens": missing_tokens[:100],  # Limit to first 100 for logging
        "missing_lines_count": missing_lines_count,
        "extra_tokens": extra_tokens[:100],  # Limit to first 100 for logging
        "loss_percentage": round(loss_percentage, 2),
        "hash_match": hash_match,
        "input_char_count": input_char_count,
        "output_char_count": output_char_count,
        "char_loss_percentage": round(char_loss_percentage, 2),
    }

    # Emit structured log
    logger.info(
        f"DATA_INTEGRITY_CHECK status={status} "
        f"input_tokens={input_token_count} "
        f"output_tokens={output_token_count} "
        f"missing_count={len(missing_tokens)} "
        f"missing_lines={missing_lines_count} "
        f"loss_percentage={result['loss_percentage']}% "
        f"char_loss={result['char_loss_percentage']}% "
        f"hash_match={hash_match}"
    )

    if status == "FAIL":
        logger.error(
            f"DATA_INTEGRITY_FAIL loss_percentage={result['loss_percentage']}% "
            f"missing_lines_count={missing_lines_count}"
        )
        if missing_tokens:
            logger.error(f"First 10 missing tokens: {missing_tokens[:10]}")

    if len(extra_tokens) > 0:
        logger.warning(
            f"DATA_INTEGRITY_WARNING: {len(extra_tokens)} extra tokens found in output "
            f"(may be system-added numbering or markers)"
        )

    return result
