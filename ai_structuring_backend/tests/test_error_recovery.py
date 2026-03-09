"""
Error recovery tests for structure_guard.py and integrity.py.

Covers TEST-003 from KNOWN_ISSUES.md:
- Invalid / non-existent DOCX files
- Corrupted document structures (bad ZIP, malformed XML)
- Permission / access failures

These tests document the expected exception behaviour for each
failure mode so callers know what to catch.

Note on exception types
-----------------------
python-docx raises PackageNotFoundError when a *string* path is invalid.
When a *Path* object is passed (as structure_guard and integrity do internally),
the lower-level stdlib exceptions propagate directly:
  - Non-existent file      → FileNotFoundError
  - Non-ZIP / empty file   → zipfile.BadZipFile
  - Bad ZIP structure      → KeyError
  - Malformed XML in DOCX  → lxml.etree.XMLSyntaxError
"""

import os
import zipfile
from pathlib import Path

import pytest
from docx import Document

from backend.processor.structure_guard import enforce_style_only_mutation
from backend.processor.integrity import run_integrity_trigger


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_minimal_docx(path: Path, text: str = "Hello world") -> Path:
    """Create a minimal valid DOCX at *path* with a single paragraph."""
    doc = Document()
    doc.add_paragraph(text)
    doc.save(str(path))
    return path


def _make_text_file_as_docx(path: Path) -> Path:
    """Write plain text content but give it a .docx extension."""
    path.write_text("this is NOT a DOCX file", encoding="utf-8")
    return path


def _make_empty_file(path: Path) -> Path:
    """Create a zero-byte file."""
    path.write_bytes(b"")
    return path


def _make_bad_zip_docx(path: Path) -> Path:
    """Create a ZIP file that is missing the required [Content_Types].xml."""
    with zipfile.ZipFile(str(path), "w") as zf:
        zf.writestr("word/document.xml", "<CORRUPTED NOT XML")
    return path


def _make_malformed_xml_docx(path: Path) -> Path:
    """Create a structurally valid ZIP/DOCX, but with malformed XML inside."""
    with zipfile.ZipFile(str(path), "w") as zf:
        zf.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            "</Types>",
        )
        zf.writestr(
            "_rels/.rels",
            '<?xml version="1.0"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
            "</Relationships>",
        )
        zf.writestr("word/document.xml", "<CORRUPTED NOT XML")
    return path


# ===========================================================================
# Structure Guard — error recovery
# ===========================================================================


class TestStructureGuardInvalidInput:
    """enforce_style_only_mutation raises predictable exceptions for bad inputs."""

    def test_nonexistent_input_raises(self, tmp_path):
        """Non-existent input path → FileNotFoundError."""
        missing = tmp_path / "missing_input.docx"
        output = _make_minimal_docx(tmp_path / "output.docx")

        with pytest.raises(FileNotFoundError):
            enforce_style_only_mutation(missing, output)

    def test_nonexistent_output_raises(self, tmp_path):
        """Non-existent output path → FileNotFoundError."""
        input_doc = _make_minimal_docx(tmp_path / "input.docx")
        missing = tmp_path / "missing_output.docx"

        with pytest.raises(FileNotFoundError):
            enforce_style_only_mutation(input_doc, missing)

    def test_plain_text_as_docx_raises(self, tmp_path):
        """Plain text file with .docx extension → BadZipFile."""
        fake_docx = _make_text_file_as_docx(tmp_path / "fake.docx")
        output = _make_minimal_docx(tmp_path / "output.docx")

        with pytest.raises(zipfile.BadZipFile):
            enforce_style_only_mutation(fake_docx, output)

    def test_empty_file_as_docx_raises(self, tmp_path):
        """Zero-byte .docx file → BadZipFile."""
        empty = _make_empty_file(tmp_path / "empty.docx")
        output = _make_minimal_docx(tmp_path / "output.docx")

        with pytest.raises(zipfile.BadZipFile):
            enforce_style_only_mutation(empty, output)

    def test_corrupt_zip_raises(self, tmp_path):
        """ZIP without [Content_Types].xml → KeyError from python-docx."""
        bad_zip = _make_bad_zip_docx(tmp_path / "bad.docx")
        output = _make_minimal_docx(tmp_path / "output.docx")

        with pytest.raises(KeyError):
            enforce_style_only_mutation(bad_zip, output)

    def test_malformed_xml_raises(self, tmp_path):
        """Valid ZIP structure but malformed XML content → XMLSyntaxError."""
        from lxml.etree import XMLSyntaxError

        malformed = _make_malformed_xml_docx(tmp_path / "malformed.docx")
        output = _make_minimal_docx(tmp_path / "output.docx")

        with pytest.raises(XMLSyntaxError):
            enforce_style_only_mutation(malformed, output)


class TestStructureGuardValidDocRaisesOnMutation:
    """Regression: valid mutation still raises RuntimeError (not swallowed)."""

    def test_structural_mutation_raises_runtime_error(self, tmp_path):
        """A document with an extra paragraph in the output raises RuntimeError."""
        input_doc = _make_minimal_docx(tmp_path / "input.docx", "Original paragraph")

        # Output with an extra paragraph — structural mutation
        out_path = tmp_path / "output.docx"
        doc = Document()
        doc.add_paragraph("Original paragraph")
        doc.add_paragraph("Extra paragraph — structural mutation")
        doc.save(str(out_path))

        with pytest.raises(RuntimeError, match="STRUCTURE_GUARD_FAIL"):
            enforce_style_only_mutation(input_doc, out_path)


# ===========================================================================
# Integrity Trigger — error recovery
# ===========================================================================


class TestIntegrityTriggerInvalidInput:
    """run_integrity_trigger raises predictable exceptions for bad inputs."""

    def test_nonexistent_input_raises(self, tmp_path):
        """Non-existent input path → FileNotFoundError."""
        missing = tmp_path / "missing_input.docx"
        output = _make_minimal_docx(tmp_path / "output.docx")

        with pytest.raises(FileNotFoundError):
            run_integrity_trigger(missing, output)

    def test_nonexistent_output_raises(self, tmp_path):
        """Non-existent output path → FileNotFoundError."""
        input_doc = _make_minimal_docx(tmp_path / "input.docx")
        missing = tmp_path / "missing_output.docx"

        with pytest.raises(FileNotFoundError):
            run_integrity_trigger(input_doc, missing)

    def test_plain_text_as_docx_raises(self, tmp_path):
        """Plain text file with .docx extension → BadZipFile."""
        fake_docx = _make_text_file_as_docx(tmp_path / "fake.docx")
        output = _make_minimal_docx(tmp_path / "output.docx")

        with pytest.raises(zipfile.BadZipFile):
            run_integrity_trigger(fake_docx, output)

    def test_corrupt_zip_raises(self, tmp_path):
        """ZIP without [Content_Types].xml → KeyError from python-docx."""
        bad_zip = _make_bad_zip_docx(tmp_path / "bad.docx")
        output = _make_minimal_docx(tmp_path / "output.docx")

        with pytest.raises(KeyError):
            run_integrity_trigger(bad_zip, output)

    def test_malformed_xml_raises(self, tmp_path):
        """Valid ZIP structure but malformed XML content → XMLSyntaxError."""
        from lxml.etree import XMLSyntaxError

        malformed = _make_malformed_xml_docx(tmp_path / "malformed.docx")
        output = _make_minimal_docx(tmp_path / "output.docx")

        with pytest.raises(XMLSyntaxError):
            run_integrity_trigger(malformed, output)


class TestIntegrityTriggerValidDocRaisesOnViolation:
    """Regression: content loss still raises RuntimeError (not swallowed)."""

    def test_missing_content_raises_runtime_error(self, tmp_path):
        """Output with missing paragraph content triggers RuntimeError."""
        input_doc = _make_minimal_docx(tmp_path / "input.docx", "Important text")

        # Output is completely empty — content missing
        out_path = tmp_path / "output.docx"
        Document().save(str(out_path))

        with pytest.raises(RuntimeError, match="INTEGRITY_TRIGGER_FAIL"):
            run_integrity_trigger(input_doc, out_path)
