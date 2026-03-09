import json
import re
import sys
import uuid
import shutil
from pathlib import Path
from types import SimpleNamespace

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT))

import processor.classifier as classifier_mod
import processor.pipeline as pipeline_mod


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Alpha paragraph for style tagging.")
    doc.add_paragraph("Beta paragraph for style tagging.")
    doc.add_paragraph("Gamma paragraph for style tagging.")
    doc.save(str(path))


def _install_stubbed_gemini(monkeypatch, call_counter: dict) -> None:
    class DummyGeminiClient:
        def __init__(self, **kwargs):
            self.provider = "gemini"
            self.model_name = kwargs.get("model_name", "gemini-2.5-pro")
            self._last_usage = {
                "input_tokens": 12,
                "output_tokens": 8,
                "total_tokens": 20,
                "latency_ms": 1,
                "provider": "gemini",
                "model": self.model_name,
            }
            self._usage = {
                "total_input_tokens": 12,
                "total_output_tokens": 8,
                "total_tokens": 20,
                "total_latency_ms": 1,
                "request_count": 0,
            }

        def generate_content(self, prompt: str, timeout=None):
            call_counter["n"] += 1
            ids = [int(x) for x in re.findall(r"\[(\d+)\]", prompt)]
            if not ids:
                ids = [1]
            payload = [{"id": pid, "tag": "H1", "confidence": 95} for pid in sorted(set(ids))]
            self._usage["request_count"] += 1
            return SimpleNamespace(text=json.dumps(payload), usage_metadata=None)

        def get_last_usage(self):
            return dict(self._last_usage)

        def get_token_usage(self):
            return dict(self._usage)

    monkeypatch.setattr(classifier_mod, "GeminiClient", DummyGeminiClient)
    monkeypatch.setattr(classifier_mod, "get_cache", lambda: None)


def _force_single_pass(monkeypatch):
    # Keep integration test deterministic and fast.
    monkeypatch.setattr(
        pipeline_mod,
        "score_document",
        lambda blocks, allowed_styles, strict_style_validation=True: (100, {"forced": True}, "PASS"),
    )


def test_llm_invoked_in_real_pipeline_when_mode_llm(monkeypatch):
    call_counter = {"n": 0}
    _install_stubbed_gemini(monkeypatch, call_counter)
    _force_single_pass(monkeypatch)

    monkeypatch.setenv("STYLE_TAGGING_MODE", "llm")
    monkeypatch.setenv("LLM_PROVIDER", "gemini")
    monkeypatch.setenv("MODEL_NAME", "gemini-2.5-pro")
    monkeypatch.setenv("GEMINI_API_KEY", "test-key")

    monkeypatch.setattr(pipeline_mod, "STYLE_TAGGING_MODE", "llm")
    monkeypatch.setattr(pipeline_mod, "GEMINI_API_KEY", "test-key")
    monkeypatch.setattr(pipeline_mod, "LLM_ENABLED", True)
    monkeypatch.setattr(pipeline_mod, "LLM_REQUIRED", False)

    base_tmp = ROOT / "backend" / "tests" / ".tmp_llm_pipeline"
    run_dir = base_tmp / f"run_{uuid.uuid4().hex}"
    output_root = run_dir / "out_llm"
    output_root.mkdir(parents=True, exist_ok=True)
    (output_root / "processed").mkdir(parents=True, exist_ok=True)
    (output_root / "review").mkdir(parents=True, exist_ok=True)
    (output_root / "json").mkdir(parents=True, exist_ok=True)
    input_docx = run_dir / "input_llm.docx"
    _make_docx(input_docx)

    try:
        result = pipeline_mod.process_document(
            input_path=str(input_docx),
            output_folder=str(output_root),
            use_markers=False,
            apply_repair=False,
        )

        assert call_counter["n"] > 0
        out_doc = Document(result["output_path"])
        non_empty_styles = [p.style.name for p in out_doc.paragraphs if p.text and p.text.strip()]
        assert "H1" in non_empty_styles
    finally:
        if run_dir.exists():
            shutil.rmtree(run_dir, ignore_errors=True)


def test_llm_not_invoked_when_mode_heuristics(monkeypatch):
    call_counter = {"n": 0}
    _install_stubbed_gemini(monkeypatch, call_counter)
    _force_single_pass(monkeypatch)

    monkeypatch.setenv("STYLE_TAGGING_MODE", "heuristics")
    monkeypatch.setenv("LLM_PROVIDER", "gemini")
    monkeypatch.setenv("MODEL_NAME", "gemini-2.5-pro")
    monkeypatch.setenv("GEMINI_API_KEY", "test-key")

    monkeypatch.setattr(pipeline_mod, "STYLE_TAGGING_MODE", "heuristics")
    monkeypatch.setattr(pipeline_mod, "GEMINI_API_KEY", "test-key")
    monkeypatch.setattr(pipeline_mod, "LLM_ENABLED", True)
    monkeypatch.setattr(pipeline_mod, "LLM_REQUIRED", False)

    base_tmp = ROOT / "backend" / "tests" / ".tmp_llm_pipeline"
    run_dir = base_tmp / f"run_{uuid.uuid4().hex}"
    output_root = run_dir / "out_heuristics"
    output_root.mkdir(parents=True, exist_ok=True)
    (output_root / "processed").mkdir(parents=True, exist_ok=True)
    (output_root / "review").mkdir(parents=True, exist_ok=True)
    (output_root / "json").mkdir(parents=True, exist_ok=True)
    input_docx = run_dir / "input_heuristics.docx"
    _make_docx(input_docx)

    try:
        pipeline_mod.process_document(
            input_path=str(input_docx),
            output_folder=str(output_root),
            use_markers=False,
            apply_repair=False,
        )

        assert call_counter["n"] == 0
    finally:
        if run_dir.exists():
            shutil.rmtree(run_dir, ignore_errors=True)


def test_non_canonical_marker_like_text_not_treated_as_canonical_marker():
    assert pipeline_mod._is_canonical_marker_text("<body-open>") is False
    assert pipeline_mod._is_llm_eligible_block({"id": 1, "text": "<body-open>", "metadata": {}}) is True


def test_pipeline_uses_paragraph_fallback_when_blocks_empty(monkeypatch):
    call_counter = {"n": 0}
    _install_stubbed_gemini(monkeypatch, call_counter)
    _force_single_pass(monkeypatch)

    monkeypatch.setenv("STYLE_TAGGING_MODE", "llm")
    monkeypatch.setenv("LLM_PROVIDER", "gemini")
    monkeypatch.setenv("MODEL_NAME", "gemini-2.5-pro")
    monkeypatch.setenv("GEMINI_API_KEY", "test-key")

    monkeypatch.setattr(pipeline_mod, "STYLE_TAGGING_MODE", "llm")
    monkeypatch.setattr(pipeline_mod, "GEMINI_API_KEY", "test-key")
    monkeypatch.setattr(pipeline_mod, "LLM_ENABLED", True)
    monkeypatch.setattr(pipeline_mod, "LLM_REQUIRED", False)

    def _fake_extract_blocks(_input_path):
        blocks = [
            {"id": 1, "text": "<body-matter-open>", "metadata": {"is_marker": True}},
            {"id": 2, "text": "   ", "metadata": {}},
        ]
        paragraphs = [
            {"id": 1, "text": "Figure 1. Example legend", "text_truncated": "Figure 1. Example legend", "metadata": {"context_zone": "FLOAT"}},
            {"id": 2, "text": "1. Introduction", "text_truncated": "1. Introduction", "metadata": {"context_zone": "BODY"}},
            {"id": 3, "text": "Smith J. Clinical Medicine. 2020.", "text_truncated": "Smith J. Clinical Medicine. 2020.", "metadata": {"context_zone": "BACK_MATTER"}},
        ]
        return blocks, paragraphs, {"total_paragraphs": 3}

    monkeypatch.setattr(pipeline_mod, "extract_blocks", _fake_extract_blocks)

    base_tmp = ROOT / "backend" / "tests" / ".tmp_llm_pipeline"
    run_dir = base_tmp / f"run_{uuid.uuid4().hex}"
    output_root = run_dir / "out_fallback"
    output_root.mkdir(parents=True, exist_ok=True)
    (output_root / "processed").mkdir(parents=True, exist_ok=True)
    (output_root / "review").mkdir(parents=True, exist_ok=True)
    (output_root / "json").mkdir(parents=True, exist_ok=True)
    input_docx = run_dir / "input_fallback.docx"
    _make_docx(input_docx)

    try:
        pipeline_mod.process_document(
            input_path=str(input_docx),
            output_folder=str(output_root),
            use_markers=False,
            apply_repair=False,
        )
        assert call_counter["n"] > 0
    finally:
        if run_dir.exists():
            shutil.rmtree(run_dir, ignore_errors=True)


def test_pipeline_passes_strict_style_validation_flag(monkeypatch):
    seen = {}

    def _capture_score(blocks, allowed_styles, strict_style_validation=True):
        seen["strict_style_validation"] = strict_style_validation
        return 100, {"forced": True}, "PASS"

    monkeypatch.setattr(pipeline_mod, "score_document", _capture_score)
    monkeypatch.setattr(pipeline_mod, "STRICT_STYLE_VALIDATION", False)

    base_tmp = ROOT / "backend" / "tests" / ".tmp_llm_pipeline"
    run_dir = base_tmp / f"run_{uuid.uuid4().hex}"
    output_root = run_dir / "out_strict_flag"
    output_root.mkdir(parents=True, exist_ok=True)
    (output_root / "processed").mkdir(parents=True, exist_ok=True)
    (output_root / "review").mkdir(parents=True, exist_ok=True)
    (output_root / "json").mkdir(parents=True, exist_ok=True)
    input_docx = run_dir / "input_strict_flag.docx"
    _make_docx(input_docx)

    def _override_classifier(blocks, _paragraphs):
        return [{"id": int(b["id"]), "tag": "H1", "confidence": 0.9} for b in blocks]

    try:
        pipeline_mod.process_document(
            input_path=str(input_docx),
            output_folder=str(output_root),
            use_markers=False,
            apply_repair=False,
            classifier_override=_override_classifier,
        )
        assert seen.get("strict_style_validation") is False
    finally:
        if run_dir.exists():
            shutil.rmtree(run_dir, ignore_errors=True)
