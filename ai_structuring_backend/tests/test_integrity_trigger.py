"""
Tests for integrity.py - Automated Integrity Trigger

Tests both content integrity (no text loss) and structural integrity (structure preservation).
"""

import os
import sys
from pathlib import Path
from unittest.mock import patch, Mock

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "backend"))

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

from processor.integrity import (
    run_integrity_trigger,
    verify_content_integrity,
    verify_structural_integrity,
    _normalize_text,
    _extract_content,
    _build_output_index,
    _extract_structure,
    _create_structural_signature,
)
from processor.reconstruction import DocumentReconstructor


# ===================================================================
# Test Helpers
# ===================================================================

def _create_test_docx(
    file_path: Path,
    paragraphs: list[str] = None,
    tables: list[list[list[str]]] = None,
    headings: list[tuple[str, int]] = None,  # (text, level)
    list_items: list[tuple[str, int]] = None,  # (text, level)
):
    """Create a test DOCX file with specified content and structure."""
    doc = Document()

    # Add headings
    if headings:
        for text, level in headings:
            doc.add_heading(text, level=level)

    # Add regular paragraphs
    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)

    # Add list items
    if list_items:
        for text, level in list_items:
            para = doc.add_paragraph(text, style='List Number')

    # Add tables
    if tables:
        for table_data in tables:
            rows = len(table_data)
            cols = len(table_data[0]) if table_data else 0
            table = doc.add_table(rows=rows, cols=cols)
            for i, row_data in enumerate(table_data):
                for j, cell_text in enumerate(row_data):
                    table.rows[i].cells[j].text = cell_text

    doc.save(str(file_path))


# ===================================================================
# Test Text Normalization
# ===================================================================

class TestTextNormalization:
    """Test _normalize_text function."""

    def test_normalize_whitespace(self):
        """Collapse multiple spaces to single."""
        text = "Hello    world   test"
        normalized = _normalize_text(text)
        assert normalized == "Hello world test"

    def test_normalize_smart_quotes(self):
        """Convert smart quotes to straight."""
        text = "\u2018single\u2019 \u201Cdouble\u201D"
        normalized = _normalize_text(text)
        assert "'" in normalized
        assert '"' in normalized

    def test_remove_marker_tokens(self):
        """Remove known marker tokens."""
        text = "<body-open> Main content <body-close>"
        normalized = _normalize_text(text)
        assert "<body-open>" not in normalized
        assert "<body-close>" not in normalized
        assert "Main content" in normalized

    def test_unicode_normalization(self):
        """Apply NFKC normalization."""
        text = "café"
        normalized = _normalize_text(text)
        # Should be consistently normalized
        assert len(normalized) > 0


# ===================================================================
# Test Content Integrity - PASS Cases
# ===================================================================

class TestContentIntegrityPass:
    """Test cases where content integrity should PASS."""

    def test_pass_on_identical_docs(self, tmp_path):
        """Identical documents pass content integrity check."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs = ["First paragraph", "Second paragraph", "Third paragraph"]
        _create_test_docx(input_path, paragraphs=paragraphs)
        _create_test_docx(output_path, paragraphs=paragraphs)

        result = verify_content_integrity(input_path, output_path)

        assert result["status"] == "PASS"
        assert result["missing_items"] == []
        assert result["input_paragraphs"] == 3
        assert result["output_paragraphs"] == 3

    def test_pass_with_style_changes(self, tmp_path):
        """Style changes don't affect content integrity."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        doc_in = Document()
        doc_in.add_paragraph("Content text")
        doc_in.save(str(input_path))

        doc_out = Document()
        para = doc_out.add_paragraph("Content text")
        para.runs[0].bold = True
        doc_out.save(str(output_path))

        result = verify_content_integrity(input_path, output_path)

        assert result["status"] == "PASS"

    def test_pass_with_added_markers(self, tmp_path):
        """Added marker paragraphs don't cause failure."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs_in = ["Content paragraph"]
        _create_test_docx(input_path, paragraphs=paragraphs_in)

        # Output has additional marker
        doc_out = Document()
        doc_out.add_paragraph("<body-open>")
        doc_out.add_paragraph("Content paragraph")
        doc_out.add_paragraph("<body-close>")
        doc_out.save(str(output_path))

        result = verify_content_integrity(input_path, output_path)

        assert result["status"] == "PASS"

    def test_pass_with_content_in_tables(self, tmp_path):
        """Content moved to tables is found."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs_in = ["Table content"]
        _create_test_docx(input_path, paragraphs=paragraphs_in)

        # Output has content in table
        tables_out = [[["Table content"]]]
        _create_test_docx(output_path, tables=tables_out)

        result = verify_content_integrity(input_path, output_path)

        assert result["status"] == "PASS"


# ===================================================================
# Test Content Integrity - FAIL Cases
# ===================================================================

class TestContentIntegrityFail:
    """Test cases where content integrity should FAIL."""

    def test_fail_on_missing_paragraph(self, tmp_path):
        """Missing paragraph causes failure."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs_in = ["First", "Second", "Third"]
        paragraphs_out = ["First", "Third"]  # Second missing

        _create_test_docx(input_path, paragraphs=paragraphs_in)
        _create_test_docx(output_path, paragraphs=paragraphs_out)

        result = verify_content_integrity(input_path, output_path)

        assert result["status"] == "FAIL"
        assert len(result["missing_items"]) > 0
        assert any("Second" in item for item in result["missing_items"])

    def test_fail_on_table_content_removed(self, tmp_path):
        """Removed table cell content causes failure."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        tables_in = [[["Cell A", "Cell B"]]]
        tables_out = [[["Cell A", ""]]]  # Cell B removed

        _create_test_docx(input_path, tables=tables_in)
        _create_test_docx(output_path, tables=tables_out)

        result = verify_content_integrity(input_path, output_path)

        assert result["status"] == "FAIL"
        assert len(result["missing_items"]) > 0

    def test_fail_on_partial_text_loss(self, tmp_path):
        """Partial text loss is detected."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs_in = ["Full content here"]
        paragraphs_out = ["Full content"]  # "here" missing

        _create_test_docx(input_path, paragraphs=paragraphs_in)
        _create_test_docx(output_path, paragraphs=paragraphs_out)

        result = verify_content_integrity(input_path, output_path)

        assert result["status"] == "FAIL"


# ===================================================================
# Test Structural Integrity - PASS Cases
# ===================================================================

class TestStructuralIntegrityPass:
    """Test cases where structural integrity should PASS."""

    def test_pass_on_identical_structure(self, tmp_path):
        """Identical structure passes."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        headings = [("Introduction", 1), ("Methods", 1)]
        paragraphs = ["Body text"]
        tables = [[["A", "B"], ["C", "D"]]]

        _create_test_docx(input_path, headings=headings, paragraphs=paragraphs, tables=tables)
        _create_test_docx(output_path, headings=headings, paragraphs=paragraphs, tables=tables)

        result = verify_structural_integrity(input_path, output_path)

        assert result["status"] == "PASS"
        assert result["paragraph_count_match"] is True
        assert result["table_structure_match"] is True
        assert result["heading_levels_match"] is True

    def test_pass_with_style_name_changes(self, tmp_path):
        """Style name changes are allowed."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        doc_in = Document()
        doc_in.add_paragraph("Text", style='Normal')
        doc_in.save(str(input_path))

        doc_out = Document()
        para = doc_out.add_paragraph("Text", style='Body Text')
        doc_out.save(str(output_path))

        result = verify_structural_integrity(input_path, output_path)

        # Style names can change, but paragraph count should match
        assert result["paragraph_count_match"] is True


# ===================================================================
# Test Structural Integrity - FAIL Cases
# ===================================================================

class TestStructuralIntegrityFail:
    """Test cases where structural integrity should FAIL."""

    def test_fail_on_table_row_removed(self, tmp_path):
        """Removed table row causes failure."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        tables_in = [[["A", "B"], ["C", "D"], ["E", "F"]]]  # 3 rows
        tables_out = [[["A", "B"], ["C", "D"]]]  # 2 rows

        _create_test_docx(input_path, tables=tables_in)
        _create_test_docx(output_path, tables=tables_out)

        result = verify_structural_integrity(input_path, output_path)

        assert result["status"] == "FAIL"
        assert result["table_structure_match"] is False
        assert any("row count mismatch" in d.lower() for d in result["differences"])

    def test_fail_on_heading_level_changed(self, tmp_path):
        """Changed heading level causes failure."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        headings_in = [("Title", 1)]
        headings_out = [("Title", 2)]  # Level changed

        _create_test_docx(input_path, headings=headings_in)
        _create_test_docx(output_path, headings=headings_out)

        result = verify_structural_integrity(input_path, output_path)

        assert result["status"] == "FAIL"
        assert result["heading_levels_match"] is False

    def test_fail_on_paragraph_count_mismatch(self, tmp_path):
        """Different paragraph counts cause failure."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs_in = ["Para 1", "Para 2", "Para 3"]
        paragraphs_out = ["Para 1", "Para 2"]

        _create_test_docx(input_path, paragraphs=paragraphs_in)
        _create_test_docx(output_path, paragraphs=paragraphs_out)

        result = verify_structural_integrity(input_path, output_path)

        assert result["status"] == "FAIL"
        assert result["paragraph_count_match"] is False


# ===================================================================
# Test Combined Integrity Trigger
# ===================================================================

class TestIntegrityTrigger:
    """Test run_integrity_trigger (combined checks)."""

    def test_pass_when_marker_inserted_additively(self, tmp_path):
        """Additive marker insertion passes."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs_in = ["Content"]
        _create_test_docx(input_path, paragraphs=paragraphs_in)

        # Output has additive marker with style "MARKER"
        doc_out = Document()
        marker_para = doc_out.add_paragraph("<body-open>")
        # Note: Would need to set style to MARKER, but for simplicity we rely on marker detection
        doc_out.add_paragraph("Content")
        doc_out.save(str(output_path))

        result = run_integrity_trigger(input_path, output_path)

        assert result["status"] == "PASS"
        assert result["content_integrity"]["status"] == "PASS"

    def test_trigger_raises_on_content_fail(self, tmp_path):
        """Trigger raises RuntimeError on content failure."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs_in = ["Important data"]
        paragraphs_out = []  # All content missing

        _create_test_docx(input_path, paragraphs=paragraphs_in)
        _create_test_docx(output_path, paragraphs=paragraphs_out)

        with pytest.raises(RuntimeError) as exc_info:
            run_integrity_trigger(input_path, output_path)

        assert "INTEGRITY_TRIGGER_FAIL" in str(exc_info.value)
        assert "CONTENT INTEGRITY FAIL" in str(exc_info.value)

    def test_trigger_raises_on_structural_fail(self, tmp_path):
        """Trigger raises RuntimeError on structural failure."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        tables_in = [[["A", "B"], ["C", "D"]]]
        tables_out = [[["A", "B"]]]  # Row removed

        _create_test_docx(input_path, tables=tables_in)
        _create_test_docx(output_path, tables=tables_out)

        with pytest.raises(RuntimeError) as exc_info:
            run_integrity_trigger(input_path, output_path)

        assert "INTEGRITY_TRIGGER_FAIL" in str(exc_info.value)

    def test_trigger_passes_when_both_pass(self, tmp_path):
        """Trigger passes when both checks pass."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs = ["Content"]
        _create_test_docx(input_path, paragraphs=paragraphs)
        _create_test_docx(output_path, paragraphs=paragraphs)

        result = run_integrity_trigger(input_path, output_path)

        assert result["status"] == "PASS"
        assert result["content_integrity"]["status"] == "PASS"
        assert result["structural_integrity"]["status"] == "PASS"
        assert result["error_message"] is None


# ===================================================================
# Test Specific Scenarios
# ===================================================================

class TestSpecificScenarios:
    """Test specific real-world scenarios."""

    def test_fail_when_paragraphs_merged(self, tmp_path):
        """Merging paragraphs causes structural failure."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs_in = ["First paragraph.", "Second paragraph."]
        paragraphs_out = ["First paragraph. Second paragraph."]  # Merged

        _create_test_docx(input_path, paragraphs=paragraphs_in)
        _create_test_docx(output_path, paragraphs=paragraphs_out)

        result = verify_structural_integrity(input_path, output_path)

        assert result["status"] == "FAIL"
        assert result["paragraph_count_match"] is False

    def test_fail_when_list_item_flattened(self, tmp_path):
        """Flattening list structure causes failure (paragraph count check)."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Input with list items
        list_items_in = [("Item 1", 0), ("Item 2", 0)]
        _create_test_docx(input_path, list_items=list_items_in)

        # Output with regular paragraphs (no list) but also different count
        paragraphs_out = ["Item 1"]  # Removed one item to force structural failure
        _create_test_docx(output_path, paragraphs=paragraphs_out)

        result = verify_structural_integrity(input_path, output_path)

        # Should fail due to paragraph count mismatch
        assert result["status"] == "FAIL"

    def test_academic_paper_structure(self, tmp_path):
        """Academic paper structure is preserved."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Create academic paper structure
        doc_in = Document()
        doc_in.add_heading("Research Paper Title", 0)
        doc_in.add_heading("Abstract", 1)
        doc_in.add_paragraph("This paper presents...")
        doc_in.add_heading("1. Introduction", 1)
        doc_in.add_paragraph("Machine learning has...")
        table = doc_in.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Method"
        table.rows[0].cells[1].text = "Accuracy"
        table.rows[1].cells[0].text = "SVM"
        table.rows[1].cells[1].text = "92%"
        doc_in.save(str(input_path))

        # Create processed version with same structure
        doc_out = Document()
        doc_out.add_heading("Research Paper Title", 0)
        doc_out.add_heading("Abstract", 1)
        doc_out.add_paragraph("This paper presents...")
        doc_out.add_heading("1. Introduction", 1)
        doc_out.add_paragraph("Machine learning has...")
        table_out = doc_out.add_table(rows=2, cols=2)
        table_out.rows[0].cells[0].text = "Method"
        table_out.rows[0].cells[1].text = "Accuracy"
        table_out.rows[1].cells[0].text = "SVM"
        table_out.rows[1].cells[1].text = "92%"
        doc_out.save(str(output_path))

        result = run_integrity_trigger(input_path, output_path)

        assert result["status"] == "PASS"


class TestReconstructionStructuralRegressions:
    def test_integrity_filters_marker_only_paragraphs_symmetrically(self, tmp_path):
        """Source docs may contain marker-only paragraphs; integrity should not drift indices."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        doc = Document()
        doc.add_paragraph("Before")
        doc.add_paragraph("<body-open>")
        doc.add_paragraph("After")
        doc.save(str(input_path))
        doc.save(str(output_path))

        result = verify_structural_integrity(input_path, output_path)

        assert result["status"] == "PASS"
        assert result["paragraph_count_match"] is True

    def test_integrity_filters_generic_marker_only_tokens(self, tmp_path):
        """Generic marker-only tokens (not just KNOWN_MARKERS) should be excluded symmetrically."""
        input_path = tmp_path / "input_generic_markers.docx"
        output_path = tmp_path / "output_generic_markers.docx"

        doc_in = Document()
        doc_in.add_paragraph("</COUT>")
        doc_in.add_paragraph("<FIG71.3>")
        doc_in.add_paragraph("<H1>History")  # inline heading marker with content: not marker-only
        doc_in.add_paragraph("Body")
        doc_in.save(str(input_path))

        # Simulate output where generic marker-only tokens are absent (or replaced elsewhere),
        # but real content paragraphs remain in order.
        doc_out = Document()
        doc_out.add_paragraph("<H1>History")
        doc_out.add_paragraph("Body")
        doc_out.save(str(output_path))

        result = verify_structural_integrity(input_path, output_path)

        assert result["status"] == "PASS"
        assert result["paragraph_count_match"] is True

    def test_pmi_styled_paragraph_is_not_filtered_from_output_count(self, tmp_path):
        """PMI style alone must not be treated as additive marker in structural integrity."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        _create_test_docx(input_path, paragraphs=["Para 1", "Para 2", "Para 3"])

        doc_out = Document()
        p1 = doc_out.add_paragraph("Para 1")
        p2 = doc_out.add_paragraph("Para 2")
        p3 = doc_out.add_paragraph("Para 3")
        if "PMI" not in [s.name for s in doc_out.styles]:
            doc_out.styles.add_style("PMI", WD_STYLE_TYPE.PARAGRAPH)
        p2.style = doc_out.styles["PMI"]
        doc_out.save(str(output_path))

        result = verify_structural_integrity(input_path, output_path)

        assert result["status"] == "PASS"
        assert result["paragraph_count_match"] is True

    def test_reconstruction_preserves_heading_levels_for_custom_h_tags(self, tmp_path):
        """Reconstruction H1/H2 tags should preserve heading outline semantics."""
        input_path = tmp_path / "input.docx"
        out_dir = tmp_path / "out"
        out_dir.mkdir()

        doc_in = Document()
        doc_in.add_heading("Intro", level=1)
        doc_in.add_paragraph("Body para")
        doc_in.add_heading("Details", level=2)
        doc_in.add_paragraph("More body")
        doc_in.save(str(input_path))

        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(
            input_path,
            classifications=[
                {"id": 1, "tag": "H1", "confidence": 99},
                {"id": 2, "tag": "TXT", "confidence": 99},
                {"id": 3, "tag": "H2", "confidence": 99},
                {"id": 4, "tag": "TXT", "confidence": 99},
            ],
            output_name="out.docx",
        )

        result = verify_structural_integrity(input_path, output_path)
        assert result["status"] == "PASS"
        assert result["heading_levels_match"] is True

    def test_heading1_source_remains_heading_when_predicted_non_heading(self, tmp_path):
        """Regression: source Heading 1 must not be downgraded to non-heading by style rewrite."""
        input_path = tmp_path / "input.docx"
        out_dir = tmp_path / "out"
        out_dir.mkdir()

        doc_in = Document()
        doc_in.add_heading("Section title", level=1)
        doc_in.add_paragraph("Body text")
        doc_in.save(str(input_path))

        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(
            input_path,
            classifications=[
                {"id": 1, "tag": "TXT", "confidence": 99},  # would normally lose heading semantics
                {"id": 2, "tag": "TXT", "confidence": 99},
            ],
            output_name="out.docx",
        )

        result = verify_structural_integrity(input_path, output_path)
        assert result["status"] == "PASS"
        assert result["heading_levels_match"] is True

    def test_non_heading_source_not_promoted_to_heading_by_predicted_h_tags(self, tmp_path):
        """Regression: source Normal paragraphs may get H* styles but must not gain heading semantics."""
        input_path = tmp_path / "input.docx"
        out_dir = tmp_path / "out"
        out_dir.mkdir()

        doc_in = Document()
        doc_in.add_paragraph("INTRODUCTION")  # style Normal in source, not a heading
        doc_in.add_paragraph("Time Has Structure and Perception")
        doc_in.add_paragraph("Time On a Continuum")
        doc_in.add_paragraph("Time Involves \"Humanization\"")
        doc_in.save(str(input_path))

        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(
            input_path,
            classifications=[
                {"id": 1, "tag": "H1", "confidence": 99},
                {"id": 2, "tag": "H3", "confidence": 99},
                {"id": 3, "tag": "H3", "confidence": 99},
                {"id": 4, "tag": "H3", "confidence": 99},
            ],
            output_name="out.docx",
        )

        result = verify_structural_integrity(input_path, output_path)
        assert result["status"] == "PASS"
        assert result["heading_levels_match"] is True

        # Visual styles should still be applied (non-semantic custom H* styles).
        doc_out = Document(str(output_path))
        assert [doc_out.paragraphs[i].style.name for i in range(4)] == ["H1", "H3", "H3", "H3"]

    def test_existing_semantic_h1_style_is_rebased_for_visual_heading_tags(self, tmp_path):
        """If source defines H1 as heading-semantic, applying H1 tag must not create heading semantics."""
        input_path = tmp_path / "input.docx"
        out_dir = tmp_path / "out"
        out_dir.mkdir()

        doc_in = Document()
        if "H1" not in [s.name for s in doc_in.styles]:
            h1 = doc_in.styles.add_style("H1", WD_STYLE_TYPE.PARAGRAPH)
            h1.base_style = doc_in.styles["Heading 1"]  # semantic style conflict
        doc_in.add_paragraph("<H1>History")  # source non-heading
        doc_in.save(str(input_path))

        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(
            input_path,
            [{"id": 1, "tag": "H1", "confidence": 99}],
            output_name="out.docx",
        )

        result = verify_structural_integrity(input_path, output_path)
        assert result["status"] == "PASS"
        assert result["heading_levels_match"] is True

        doc_out = Document(str(output_path))
        assert doc_out.paragraphs[0].style.name == "H1"

    def test_deep_and_custom_heading_levels_preserved(self, tmp_path):
        """Regression: deeper headings and custom heading-like styles remain heading-equivalent."""
        input_path = tmp_path / "input.docx"
        out_dir = tmp_path / "out"
        out_dir.mkdir()

        doc_in = Document()
        doc_in.add_heading("Deep heading", level=7)

        # Create custom style inheriting from Heading 2
        custom_style_name = "Custom Heading 2 Variant"
        if custom_style_name not in [s.name for s in doc_in.styles]:
            custom = doc_in.styles.add_style(custom_style_name, WD_STYLE_TYPE.PARAGRAPH)
            custom.base_style = doc_in.styles["Heading 2"]
        p_custom = doc_in.add_paragraph("Template heading", style=custom_style_name)
        doc_in.add_paragraph("Body text")
        doc_in.save(str(input_path))

        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(
            input_path,
            classifications=[
                {"id": 1, "tag": "TXT", "confidence": 99},  # preserve Heading 7
                {"id": 2, "tag": "TXT", "confidence": 99},  # preserve custom Heading 2-derived style
                {"id": 3, "tag": "TXT", "confidence": 99},
            ],
            output_name="out.docx",
        )

        result = verify_structural_integrity(input_path, output_path)
        assert result["status"] == "PASS"
        assert result["heading_levels_match"] is True

    def test_reconstruction_preserves_style_based_list_semantics(self, tmp_path):
        """Style-based lists should remain detectable as lists after custom BL-* styles are applied."""
        input_path = tmp_path / "input.docx"
        out_dir = tmp_path / "out"
        out_dir.mkdir()

        doc_in = Document()
        doc_in.add_paragraph("List item 1", style="List Bullet")
        doc_in.add_paragraph("List item 2", style="List Bullet")
        doc_in.add_paragraph("List item 3", style="List Bullet")
        doc_in.save(str(input_path))

        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(
            input_path,
            classifications=[
                {"id": 1, "tag": "BL-FIRST", "confidence": 99},
                {"id": 2, "tag": "BL-MID", "confidence": 99},
                {"id": 3, "tag": "BL-LAST", "confidence": 99},
            ],
            output_name="out.docx",
        )

        result = verify_structural_integrity(input_path, output_path)
        assert result["status"] == "PASS"
        assert result["list_structure_match"] is True

    def test_title_styled_table_paragraph_can_be_table_tagged_with_semantics_preserved(self, tmp_path):
        """Regression: source Title inside tables should accept table tags without integrity heading drift."""
        input_path = tmp_path / "input.docx"
        out_dir = tmp_path / "out"
        out_dir.mkdir()

        doc_in = Document()
        tbl = doc_in.add_table(rows=1, cols=1)
        p = tbl.rows[0].cells[0].paragraphs[0]
        p.text = "Latency"
        p.style = doc_in.styles["Title"]  # legacy template quirk seen in real files
        doc_in.save(str(input_path))

        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(
            input_path,
            classifications=[{"id": 1, "tag": "T", "confidence": 99}],
            output_name="out.docx",
        )

        result = verify_structural_integrity(input_path, output_path)
        assert result["status"] == "PASS"
        assert result["heading_levels_match"] is True

        doc_out = Document(str(output_path))
        out_p = doc_out.tables[0].rows[0].cells[0].paragraphs[0]
        assert out_p.style.name == "T"

    def test_title_table_caption_and_t1_caption_are_heading_equivalent(self, tmp_path):
        """Legacy Title table captions and canonical T1 captions should compare as equivalent."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        doc_in = Document()
        p_in = doc_in.add_paragraph("Table 71.1. EBV-encoded Proteins.")
        p_in.style = doc_in.styles["Title"]
        doc_in.save(str(input_path))

        doc_out = Document()
        recon = DocumentReconstructor(output_dir=str(tmp_path))
        # Ensure T1 exists in output and apply it to the matching caption text.
        recon._get_or_create_style(doc_out, "T1")
        p_out = doc_out.add_paragraph("Table 71.1. EBV-encoded Proteins.")
        p_out.style = doc_out.styles["T1"]
        doc_out.save(str(output_path))

        result = verify_structural_integrity(input_path, output_path)
        assert result["status"] == "PASS"
        assert result["heading_levels_match"] is True

    def test_source_heading2_can_restyle_to_h20_with_outline_level_preserved(self, tmp_path):
        """Regression: visual H20 output should preserve source Heading 2 semantics via outlineLvl."""
        input_path = tmp_path / "input.docx"
        out_dir = tmp_path / "out"
        out_dir.mkdir()

        doc_in = Document()
        p = doc_in.add_paragraph("<H2>Overview")
        p.style = doc_in.styles["Heading 2"]
        doc_in.save(str(input_path))

        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(
            input_path,
            classifications=[{"id": 1, "tag": "H20", "confidence": 99}],
            output_name="out.docx",
        )

        result = verify_structural_integrity(input_path, output_path)
        assert result["status"] == "PASS"
        assert result["heading_levels_match"] is True

    def test_mixed_list_and_non_list_roundtrip_preserves_list_structure(self, tmp_path):
        """Regression: mixed documents must not create lists or remove existing list semantics."""
        input_path = tmp_path / "input.docx"
        out_dir = tmp_path / "out"
        out_dir.mkdir()

        doc_in = Document()
        # Non-list block (classifier may predict list-like tags)
        doc_in.add_paragraph("Intro line")
        doc_in.add_paragraph("Another non-list line")
        doc_in.add_paragraph("Third non-list line")
        # Style-based lists (classifier may predict non-list tags)
        doc_in.add_paragraph("List A", style="List Bullet")
        doc_in.add_paragraph("List B", style="List Bullet")
        doc_in.add_paragraph("List C", style="List Bullet")
        doc_in.save(str(input_path))

        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(
            input_path,
            classifications=[
                {"id": 1, "tag": "BL-FIRST", "confidence": 99},  # must stay non-list
                {"id": 2, "tag": "BL-MID", "confidence": 99},    # must stay non-list
                {"id": 3, "tag": "BL-LAST", "confidence": 99},   # must stay non-list
                {"id": 4, "tag": "TXT", "confidence": 99},       # must stay list
                {"id": 5, "tag": "TXT", "confidence": 99},       # must stay list
                {"id": 6, "tag": "TXT", "confidence": 99},       # must stay list
            ],
            output_name="out.docx",
        )

        result = verify_structural_integrity(input_path, output_path)
        assert result["status"] == "PASS"
        assert result["list_structure_match"] is True


# ===================================================================
# Test Performance
# ===================================================================

class TestPerformance:
    """Test performance requirements."""

    def test_handles_large_document(self, tmp_path):
        """Handles 1000+ paragraph document efficiently."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Create document with 1000 paragraphs
        paragraphs = [f"Paragraph {i} with content" for i in range(1000)]
        _create_test_docx(input_path, paragraphs=paragraphs)
        _create_test_docx(output_path, paragraphs=paragraphs)

        import time
        start = time.time()
        result = run_integrity_trigger(input_path, output_path)
        elapsed = time.time() - start

        assert result["status"] == "PASS"
        # Should complete in under 10 seconds for 1000 paragraphs
        assert elapsed < 10.0, f"Took {elapsed:.2f}s, should be <10s"


# ===================================================================
# Test Edge Cases
# ===================================================================

class TestEdgeCases:
    """Test edge cases and boundary conditions."""

    def test_empty_documents(self, tmp_path):
        """Empty documents pass."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        _create_test_docx(input_path)
        _create_test_docx(output_path)

        result = run_integrity_trigger(input_path, output_path)

        assert result["status"] == "PASS"


class TestDiagnostics:
    """Diff-first diagnostics for faster failure localization."""

    def test_verify_structural_integrity_includes_first_difference_snapshot(self, tmp_path):
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        doc_in = Document()
        doc_in.add_heading("Heading A", level=1)
        doc_in.save(str(input_path))

        doc_out = Document()
        doc_out.add_paragraph("Heading A")  # heading demoted -> structural mismatch
        doc_out.save(str(output_path))

        result = verify_structural_integrity(input_path, output_path)

        assert result["status"] == "FAIL"
        first = result.get("first_difference")
        assert isinstance(first, dict)
        assert first.get("stage") == "integrity_check"
        assert first.get("paragraph_index") == 0
        assert "Heading level mismatch" in (first.get("message") or "")
        assert isinstance(first.get("input"), dict)
        assert isinstance(first.get("output"), dict)

    def test_only_tables(self, tmp_path):
        """Document with only tables."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        tables = [[["A", "B"], ["C", "D"]]]
        _create_test_docx(input_path, tables=tables)
        _create_test_docx(output_path, tables=tables)

        result = run_integrity_trigger(input_path, output_path)

        assert result["status"] == "PASS"

    def test_special_characters_preserved(self, tmp_path):
        """Special characters are preserved."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs = ["Text with @#$%^&*() special chars"]
        _create_test_docx(input_path, paragraphs=paragraphs)
        _create_test_docx(output_path, paragraphs=paragraphs)

        result = run_integrity_trigger(input_path, output_path)

        assert result["status"] == "PASS"
