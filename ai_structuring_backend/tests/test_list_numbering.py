"""Tests for bullet/list numbering preservation in DOCX output."""

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "backend"))

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement

from processor.reconstruction import (
    DocumentReconstructor,
    _list_props_for_tag,
    _bullet_level,
    ensure_numbering,
)


# ===================================================================
# _list_props_for_tag
# ===================================================================

class TestListPropsForTag:
    # --- Bullet styles (level 0) ---
    def test_bl_first(self):
        assert _list_props_for_tag("BL-FIRST") == ("bullet", 0)

    def test_bl_mid(self):
        assert _list_props_for_tag("BL-MID") == ("bullet", 0)

    def test_bl_last(self):
        assert _list_props_for_tag("BL-LAST") == ("bullet", 0)

    def test_tbl_first(self):
        assert _list_props_for_tag("TBL-FIRST") == ("bullet", 0)

    def test_tbl_mid0(self):
        assert _list_props_for_tag("TBL-MID0") == ("bullet", 0)

    def test_tbl_last(self):
        assert _list_props_for_tag("TBL-LAST") == ("bullet", 0)

    def test_obj_bl_first(self):
        assert _list_props_for_tag("OBJ-BL-FIRST") == ("bullet", 0)

    def test_nbx1_bl_first(self):
        assert _list_props_for_tag("NBX1-BL-FIRST") == ("bullet", 0)

    # --- Bullet styles (level 1) ---
    def test_ul_first(self):
        assert _list_props_for_tag("UL-FIRST") == ("bullet", 1)

    def test_ul_mid(self):
        assert _list_props_for_tag("UL-MID") == ("bullet", 1)

    def test_ul_last(self):
        assert _list_props_for_tag("UL-LAST") == ("bullet", 1)

    def test_tbl2_mid(self):
        assert _list_props_for_tag("TBL2-MID") == ("bullet", 1)

    def test_nbx1_bl2_mid(self):
        assert _list_props_for_tag("NBX1-BL2-MID") == ("bullet", 1)

    # --- Bullet styles (level 2-3) ---
    def test_tbl3_mid(self):
        assert _list_props_for_tag("TBL3-MID") == ("bullet", 2)

    def test_tbl4_mid(self):
        assert _list_props_for_tag("TBL4-MID") == ("bullet", 3)

    # --- Numbered styles ---
    def test_nl_first(self):
        assert _list_props_for_tag("NL-FIRST") == ("numbered", 0)

    def test_nl_mid(self):
        assert _list_props_for_tag("NL-MID") == ("numbered", 0)

    def test_nl_last(self):
        assert _list_props_for_tag("NL-LAST") == ("numbered", 0)

    def test_eoc_nl_first(self):
        assert _list_props_for_tag("EOC-NL-FIRST") == ("numbered", 0)

    def test_eoc_nl_mid(self):
        assert _list_props_for_tag("EOC-NL-MID") == ("numbered", 0)

    def test_nbx1_nl_mid(self):
        assert _list_props_for_tag("NBX1-NL-MID") == ("numbered", 0)

    def test_rq_ll2_mid(self):
        assert _list_props_for_tag("RQ-LL2-MID") == ("numbered", 1)

    # --- Pattern-based fallback (styles NOT in STYLE_DEFINITIONS) ---
    def test_bul1_pattern(self):
        assert _list_props_for_tag("BUL1") == ("bullet", 0)

    def test_bx1_bl_mid_pattern(self):
        assert _list_props_for_tag("BX1-BL-MID") == ("bullet", 0)

    def test_bx1_nl_first_pattern(self):
        assert _list_props_for_tag("BX1-NL-FIRST") == ("numbered", 0)

    def test_bx2_ul_last_pattern(self):
        assert _list_props_for_tag("BX2-UL-LAST") == ("bullet", 1)

    # --- Non-list styles ---
    def test_txt_not_list(self):
        assert _list_props_for_tag("TXT") is None

    def test_h1_not_list(self):
        assert _list_props_for_tag("H1") is None

    def test_pmi_not_list(self):
        assert _list_props_for_tag("PMI") is None

    def test_t_not_list(self):
        assert _list_props_for_tag("T") is None

    def test_unknown_not_list(self):
        assert _list_props_for_tag("XXXXXX") is None


# ===================================================================
# _bullet_level
# ===================================================================

class TestBulletLevel:
    def test_bl_first_level0(self):
        assert _bullet_level("BL-FIRST") == 0

    def test_ul_first_level1(self):
        assert _bullet_level("UL-FIRST") == 1

    def test_tbl2_mid_level1(self):
        assert _bullet_level("TBL2-MID") == 1

    def test_tbl3_mid_level2(self):
        assert _bullet_level("TBL3-MID") == 2

    def test_tbl4_mid_level3(self):
        assert _bullet_level("TBL4-MID") == 3

    def test_nbx1_bl2_mid_level1(self):
        assert _bullet_level("NBX1-BL2-MID") == 1


# ===================================================================
# ensure_numbering  (unit-level, operates on live Document)
# ===================================================================

class TestEnsureNumbering:
    def _make_doc_with_para(self, text="Test"):
        doc = Document()
        para = doc.add_paragraph(text)
        return doc, para

    def test_bullet_adds_numpr(self):
        doc, para = self._make_doc_with_para("First item")
        ensure_numbering(para, "bullet", 0, doc)

        pPr = para._element.find(qn("w:pPr"))
        assert pPr is not None
        numPr = pPr.find(qn("w:numPr"))
        assert numPr is not None
        ilvl = numPr.find(qn("w:ilvl"))
        assert ilvl is not None
        assert ilvl.get(qn("w:val")) == "0"
        numId = numPr.find(qn("w:numId"))
        assert numId is not None
        assert int(numId.get(qn("w:val"))) > 0

    def test_numbered_adds_numpr(self):
        doc, para = self._make_doc_with_para("1. First")
        ensure_numbering(para, "numbered", 0, doc)

        numPr = para._element.find(qn("w:pPr")).find(qn("w:numPr"))
        assert numPr is not None

    def test_level1_sets_ilvl_1(self):
        doc, para = self._make_doc_with_para("Sub item")
        ensure_numbering(para, "bullet", 1, doc)

        numPr = para._element.find(qn("w:pPr")).find(qn("w:numPr"))
        ilvl = numPr.find(qn("w:ilvl"))
        assert ilvl.get(qn("w:val")) == "1"

    def test_preserves_existing_numpr(self):
        """If paragraph already has numbering, ensure_numbering is a no-op."""
        doc, para = self._make_doc_with_para("Item")

        # Manually add a numPr with unusual values
        pPr = para._element.get_or_add_pPr()
        numPr = pPr._add_numPr()
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "5")
        numId_elm = OxmlElement("w:numId")
        numId_elm.set(qn("w:val"), "42")
        numPr.append(ilvl)
        numPr.append(numId_elm)

        # Call ensure_numbering — should be a no-op
        ensure_numbering(para, "bullet", 0, doc)

        numPr = pPr.numPr
        assert numPr.find(qn("w:ilvl")).get(qn("w:val")) == "5"
        assert numPr.find(qn("w:numId")).get(qn("w:val")) == "42"

    def test_bullet_and_numbered_different_numids(self):
        doc = Document()
        para_b = doc.add_paragraph("bullet")
        para_n = doc.add_paragraph("numbered")

        ensure_numbering(para_b, "bullet", 0, doc)
        ensure_numbering(para_n, "numbered", 0, doc)

        numId_b = (
            para_b._element.find(qn("w:pPr"))
            .find(qn("w:numPr"))
            .find(qn("w:numId"))
            .get(qn("w:val"))
        )
        numId_n = (
            para_n._element.find(qn("w:pPr"))
            .find(qn("w:numPr"))
            .find(qn("w:numId"))
            .get(qn("w:val"))
        )
        assert numId_b != numId_n

    def test_multiple_bullets_share_numid(self):
        doc = Document()
        p1 = doc.add_paragraph("first")
        p2 = doc.add_paragraph("second")

        ensure_numbering(p1, "bullet", 0, doc)
        ensure_numbering(p2, "bullet", 0, doc)

        nid1 = p1._element.find(qn("w:pPr")).find(qn("w:numPr")).find(qn("w:numId")).get(qn("w:val"))
        nid2 = p2._element.find(qn("w:pPr")).find(qn("w:numPr")).find(qn("w:numId")).get(qn("w:val"))
        assert nid1 == nid2


# ===================================================================
# Integration: apply_styles / apply_tags_with_markers
# ===================================================================

class TestApplyStylesWithNumbering:
    def _create_test_docx(self, tmp_path, paragraphs, numpr_specs=None):
        doc = Document()
        para_objs = []
        for text in paragraphs:
            para_objs.append(doc.add_paragraph(text))
        for idx, (list_kind, level) in (numpr_specs or {}).items():
            ensure_numbering(para_objs[idx], list_kind, level, doc)
        docx_path = tmp_path / "input.docx"
        doc.save(str(docx_path))
        return docx_path

    def test_bullet_paragraphs_get_numpr(self, tmp_path):
        docx_path = self._create_test_docx(tmp_path, [
            "Chapter 1",
            "First item",
            "Second item",
            "Third item",
        ], numpr_specs={1: ("bullet", 0), 2: ("bullet", 0), 3: ("bullet", 0)})
        classifications = [
            {"id": 1, "tag": "H1", "confidence": 99},
            {"id": 2, "tag": "BL-FIRST", "confidence": 99},
            {"id": 3, "tag": "BL-MID", "confidence": 99},
            {"id": 4, "tag": "BL-LAST", "confidence": 99},
        ]

        out_dir = tmp_path / "output"
        out_dir.mkdir()
        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(docx_path, classifications, "out.docx")

        result_doc = Document(str(output_path))
        paras = [p for p in result_doc.paragraphs if p.text.strip()]

        # H1 should NOT have numPr
        pPr_h1 = paras[0]._element.find(qn("w:pPr"))
        if pPr_h1 is not None:
            assert pPr_h1.find(qn("w:numPr")) is None

        # BL-FIRST, BL-MID, BL-LAST should have numPr with ilvl=0
        for i in range(1, 4):
            pPr = paras[i]._element.find(qn("w:pPr"))
            assert pPr is not None, f"para {i+1} missing pPr"
            numPr = pPr.find(qn("w:numPr"))
            assert numPr is not None, f"para {i+1} missing numPr"
            assert numPr.find(qn("w:ilvl")).get(qn("w:val")) == "0"

    def test_numbered_paragraphs_get_numpr(self, tmp_path):
        docx_path = self._create_test_docx(tmp_path, [
            "Step one",
            "Step two",
        ], numpr_specs={0: ("numbered", 0), 1: ("numbered", 0)})
        classifications = [
            {"id": 1, "tag": "NL-FIRST", "confidence": 99},
            {"id": 2, "tag": "NL-LAST", "confidence": 99},
        ]

        out_dir = tmp_path / "output"
        out_dir.mkdir()
        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(docx_path, classifications, "out.docx")

        result_doc = Document(str(output_path))
        paras = [p for p in result_doc.paragraphs if p.text.strip()]

        for i, para in enumerate(paras):
            pPr = para._element.find(qn("w:pPr"))
            numPr = pPr.find(qn("w:numPr"))
            assert numPr is not None, f"para {i+1} missing numPr"

    def test_non_list_paragraph_no_numpr(self, tmp_path):
        docx_path = self._create_test_docx(tmp_path, ["Normal text"])
        classifications = [{"id": 1, "tag": "TXT", "confidence": 99}]

        out_dir = tmp_path / "output"
        out_dir.mkdir()
        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(docx_path, classifications, "out.docx")

        result_doc = Document(str(output_path))
        paras = [p for p in result_doc.paragraphs if p.text.strip()]
        pPr = paras[0]._element.find(qn("w:pPr"))
        if pPr is not None:
            assert pPr.find(qn("w:numPr")) is None

    def test_apply_tags_with_markers_adds_numpr(self, tmp_path):
        docx_path = self._create_test_docx(tmp_path, [
            "First item",
            "Second item",
        ], numpr_specs={0: ("bullet", 0), 1: ("bullet", 0)})
        classifications = [
            {"id": 1, "tag": "BL-FIRST", "confidence": 99},
            {"id": 2, "tag": "BL-LAST", "confidence": 99},
        ]

        out_dir = tmp_path / "output"
        out_dir.mkdir()
        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_tags_with_markers(docx_path, classifications, "out.docx")

        result_doc = Document(str(output_path))
        paras = [p for p in result_doc.paragraphs if p.text.strip()]

        for i, para in enumerate(paras):
            pPr = para._element.find(qn("w:pPr"))
            numPr = pPr.find(qn("w:numPr"))
            assert numPr is not None, f"para {i+1} missing numPr"

    def test_sub_bullet_gets_level1(self, tmp_path):
        docx_path = self._create_test_docx(tmp_path, [
            "Main item",
            "Sub item",
        ], numpr_specs={0: ("bullet", 0), 1: ("bullet", 1)})
        classifications = [
            {"id": 1, "tag": "BL-FIRST", "confidence": 99},
            {"id": 2, "tag": "UL-LAST", "confidence": 99},
        ]

        out_dir = tmp_path / "output"
        out_dir.mkdir()
        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(docx_path, classifications, "out.docx")

        result_doc = Document(str(output_path))
        paras = [p for p in result_doc.paragraphs if p.text.strip()]

        # BL-FIRST → level 0
        numPr_0 = paras[0]._element.find(qn("w:pPr")).find(qn("w:numPr"))
        assert numPr_0.find(qn("w:ilvl")).get(qn("w:val")) == "0"

        # UL-LAST → level 1
        numPr_1 = paras[1]._element.find(qn("w:pPr")).find(qn("w:numPr"))
        assert numPr_1.find(qn("w:ilvl")).get(qn("w:val")) == "1"

    def test_saved_docx_roundtrip_preserves_numpr(self, tmp_path):
        """numPr survives save → reopen cycle."""
        docx_path = self._create_test_docx(tmp_path, ["First item"], numpr_specs={0: ("bullet", 0)})
        classifications = [{"id": 1, "tag": "BL-FIRST", "confidence": 99}]

        out_dir = tmp_path / "output"
        out_dir.mkdir()
        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(docx_path, classifications, "out.docx")

        # Re-open saved file
        reopened = Document(str(output_path))
        paras = [p for p in reopened.paragraphs if p.text.strip()]
        pPr = paras[0]._element.find(qn("w:pPr"))
        numPr = pPr.find(qn("w:numPr"))
        assert numPr is not None, "numPr lost after save/reopen"
        assert numPr.find(qn("w:ilvl")).get(qn("w:val")) == "0"

    def test_non_list_source_list_tag_does_not_add_numpr(self, tmp_path):
        """Regression: list-like tags must not inject numPr on non-list source paragraphs."""
        docx_path = self._create_test_docx(tmp_path, ["First item", "Second item"])
        classifications = [
            {"id": 1, "tag": "BL-FIRST", "confidence": 99},
            {"id": 2, "tag": "BL-LAST", "confidence": 99},
        ]

        out_dir = tmp_path / "output"
        out_dir.mkdir()
        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(docx_path, classifications, "out.docx")

        result_doc = Document(str(output_path))
        paras = [p for p in result_doc.paragraphs if p.text.strip()]

        for i, para in enumerate(paras):
            pPr = para._element.find(qn("w:pPr"))
            if pPr is not None:
                assert pPr.find(qn("w:numPr")) is None, f"para {i+1} unexpectedly gained numPr"

    def test_existing_list_numpr_level_and_id_preserved(self, tmp_path):
        """Regression: existing list XML semantics (numId/ilvl) must be preserved exactly."""
        docx_path = self._create_test_docx(
            tmp_path,
            ["Item 1", "Item 2"],
            numpr_specs={0: ("bullet", 0), 1: ("bullet", 1)},
        )
        classifications = [
            {"id": 1, "tag": "BL-FIRST", "confidence": 99},
            {"id": 2, "tag": "UL-LAST", "confidence": 99},
        ]

        def _numpr_tuple(para):
            pPr = para._element.find(qn("w:pPr"))
            numPr = pPr.find(qn("w:numPr")) if pPr is not None else None
            assert numPr is not None
            return (
                numPr.find(qn("w:numId")).get(qn("w:val")),
                numPr.find(qn("w:ilvl")).get(qn("w:val")),
            )

        input_doc = Document(str(docx_path))
        input_paras = [p for p in input_doc.paragraphs if p.text.strip()]
        before = [_numpr_tuple(input_paras[0]), _numpr_tuple(input_paras[1])]

        out_dir = tmp_path / "output"
        out_dir.mkdir()
        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(docx_path, classifications, "out.docx")

        result_doc = Document(str(output_path))
        result_paras = [p for p in result_doc.paragraphs if p.text.strip()]
        after = [_numpr_tuple(result_paras[0]), _numpr_tuple(result_paras[1])]

        assert after == before

    def test_inline_heading_marker_overrides_style_based_list_preservation(self, tmp_path):
        """Legacy list-styled <H1> source paragraphs should be allowed to become headings."""
        doc = Document()
        p = doc.add_paragraph("<H1>FURTHER EDUCATIONAL OPPORTUNITIES", style="List Bullet")
        # Confirm style-based list (no explicit numPr) to mirror Wheeler pattern.
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None:
            assert pPr.find(qn("w:numPr")) is None
        docx_path = tmp_path / "input.docx"
        doc.save(str(docx_path))

        out_dir = tmp_path / "output"
        out_dir.mkdir()
        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(
            docx_path,
            [{"id": 1, "tag": "H1", "confidence": 99}],
            "out.docx",
        )

        result_doc = Document(str(output_path))
        out_para = result_doc.paragraphs[0]
        assert out_para.style.name == "H1"
        out_pPr = out_para._element.find(qn("w:pPr"))
        if out_pPr is not None:
            assert out_pPr.find(qn("w:numPr")) is None

    def test_style_based_list_source_can_restyle_to_semantic_list_and_gain_numpr(self, tmp_path):
        """Style-only source lists should accept semantic list styles without losing list structure."""
        doc = Document()
        p1 = doc.add_paragraph("• First", style="List Bullet")
        p2 = doc.add_paragraph("• Second", style="List Bullet")
        for p in (p1, p2):
            pPr = p._element.find(qn("w:pPr"))
            if pPr is not None:
                assert pPr.find(qn("w:numPr")) is None
        docx_path = tmp_path / "input.docx"
        doc.save(str(docx_path))

        out_dir = tmp_path / "output"
        out_dir.mkdir()
        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(
            docx_path,
            [
                {"id": 1, "tag": "BL-FIRST", "confidence": 99},
                {"id": 2, "tag": "BL-LAST", "confidence": 99},
            ],
            "out.docx",
        )

        result_doc = Document(str(output_path))
        out_paras = result_doc.paragraphs[:2]
        assert [p.style.name for p in out_paras] == ["BL-FIRST", "BL-LAST"]
        for para in out_paras:
            pPr = para._element.find(qn("w:pPr"))
            assert pPr is not None
            assert pPr.find(qn("w:numPr")) is not None

    def test_table_semantic_tag_overrides_source_heading_style_in_table(self, tmp_path):
        """Table rows styled as Heading 5 in source should still accept TBL-* styles."""
        doc = Document()
        doc.add_paragraph("Body intro")
        tbl = doc.add_table(rows=1, cols=1)
        cell_para = tbl.cell(0, 0).paragraphs[0]
        cell_para.text = "Vitamin D: Example table bullet row"
        cell_para.style = doc.styles["Heading 5"]
        ensure_numbering(cell_para, "bullet", 0, doc)

        docx_path = tmp_path / "input.docx"
        doc.save(str(docx_path))

        out_dir = tmp_path / "output"
        out_dir.mkdir()
        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(
            docx_path,
            [
                {"id": 1, "tag": "TXT", "confidence": 99},
                {"id": 2, "tag": "TBL-FIRST", "confidence": 99},
            ],
            "out.docx",
        )

        result_doc = Document(str(output_path))
        out_cell_para = result_doc.tables[0].cell(0, 0).paragraphs[0]
        assert out_cell_para.style.name == "TBL-FIRST"
        out_pPr = out_cell_para._element.find(qn("w:pPr"))
        assert out_pPr is not None
        assert out_pPr.find(qn("w:numPr")) is not None

    def test_inline_heading_marker_allows_custom_heading_variant_over_source_heading_style(self, tmp_path):
        """Source built-in Heading 2 should restyle to H20 when inline marker/tag agree, preserving outlineLvl."""
        doc = Document()
        p = doc.add_paragraph("<H2>Overview")
        p.style = doc.styles["Heading 2"]

        docx_path = tmp_path / "input.docx"
        doc.save(str(docx_path))

        out_dir = tmp_path / "output"
        out_dir.mkdir()
        recon = DocumentReconstructor(output_dir=str(out_dir))
        output_path = recon.apply_styles(
            docx_path,
            [{"id": 1, "tag": "H20", "confidence": 99}],
            "out.docx",
        )

        result_doc = Document(str(output_path))
        out_para = result_doc.paragraphs[0]
        assert out_para.style.name == "H20"
        pPr = out_para._element.find(qn("w:pPr"))
        assert pPr is not None
        outline = pPr.find(qn("w:outlineLvl"))
        assert outline is not None
        # Heading 2 -> outlineLvl 1 (0-based)
        assert outline.get(qn("w:val")) == "1"
