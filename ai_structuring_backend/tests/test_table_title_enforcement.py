"""
Tests for table_title_enforcement.py — post-reconstruction enforcement
of table-title house rules on actual DOCX documents.
"""

import os
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "backend"))

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

from processor.table_title_enforcement import (
    enforce_table_title_house_rules,
    _para_text,
    _TITLE_STRICT_RE,
    _TITLE_BROAD_RE,
    _PLACEHOLDER_RE,
)


# ===================================================================
# Helpers
# ===================================================================

def _make_docx(tmp_path: Path, build_fn) -> Path:
    """Create a DOCX via *build_fn(doc)* and return the saved path."""
    path = tmp_path / "test.docx"
    doc = Document()
    build_fn(doc)
    doc.save(str(path))
    return path


def _add_table(doc: Document, rows: int = 2, cols: int = 2):
    """Append a small table to the document body."""
    tbl = doc.add_table(rows=rows, cols=cols)
    for r in range(rows):
        for c in range(cols):
            tbl.rows[r].cells[c].text = f"R{r}C{c}"
    return tbl


def _body_paragraph_styles(docx_path: Path) -> list[tuple[str, str]]:
    """Return list of (text, style_name) for body <w:p> elements.

    Reads style ID directly from XML to avoid Paragraph-wrapper issues.
    """
    doc = Document(str(docx_path))
    # Build a style-id → style-name lookup from the document styles.
    id_to_name: dict[str, str] = {}
    for s in doc.styles:
        if s.style_id:
            id_to_name[s.style_id] = s.name

    body = doc.element.body
    result = []
    for child in body:
        if child.tag == qn("w:p"):
            text = _para_text(child).strip()
            style_id = _para_style_id(child)
            style_name = id_to_name.get(style_id, style_id) if style_id else "Normal"
            result.append((text, style_name))
    return result


def _para_text(p_element) -> str:
    """Extract text from a <w:p> XML element."""
    return "".join(t.text for t in p_element.iter(qn("w:t")) if t.text)


def _para_style_id(p_element) -> str | None:
    """Read the style ID from <w:pPr><w:pStyle w:val='…'/>."""
    pPr = p_element.find(qn("w:pPr"))
    if pPr is None:
        return None
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return None
    return pStyle.get(qn("w:val"))


def _para_outline_level(p_element) -> str | None:
    """Read paragraph-level outlineLvl (0-based) if present."""
    pPr = p_element.find(qn("w:pPr"))
    if pPr is None:
        return None
    outline = pPr.find(qn("w:outlineLvl"))
    if outline is None:
        return None
    return outline.get(qn("w:val"))


def _style_of_paragraph(docx_path: Path, text_prefix: str) -> str | None:
    """Find the first body paragraph starting with *text_prefix* and return its style."""
    for text, style in _body_paragraph_styles(docx_path):
        if text.startswith(text_prefix):
            return style
    return None


# ===================================================================
# Test Pattern Matching (sanity checks)
# ===================================================================

class TestPatterns:

    def test_strict_colon(self):
        assert _TITLE_STRICT_RE.match("Table 1: Demographics")

    def test_strict_period(self):
        assert _TITLE_STRICT_RE.match("Table 2.1.Title text here")

    def test_strict_with_subsection(self):
        assert _TITLE_STRICT_RE.match("Table 10.3: Some title")

    def test_broad_without_punctuation(self):
        assert _TITLE_BROAD_RE.match("Table 5 with description")

    def test_broad_number_only(self):
        assert _TITLE_BROAD_RE.match("Table 7")

    def test_placeholder_tab(self):
        assert _PLACEHOLDER_RE.match("<TAB5.1>")

    def test_placeholder_insert(self):
        assert _PLACEHOLDER_RE.match("<INSERT TAB 3>")

    def test_no_match_body_text(self):
        assert not _TITLE_BROAD_RE.match("The table below shows results")

    def test_no_match_placeholder_wrong(self):
        assert not _PLACEHOLDER_RE.match("<FIGURE 1>")


# ===================================================================
# Test: Title directly preceding table → T1
# ===================================================================

class TestTitleDirectlyPrecedingTable:

    def test_title_colon_styled_t1(self, tmp_path):
        """Title 'Table 1: Demographics' directly above a table → T1."""
        def build(doc):
            doc.add_paragraph("Some intro text.")
            doc.add_paragraph("Table 1: Demographics")
            _add_table(doc)

        path = _make_docx(tmp_path, build)
        metrics = enforce_table_title_house_rules(str(path))

        assert metrics["tables"] == 1
        assert metrics["titles_fixed"] == 1
        assert _style_of_paragraph(path, "Table 1:") == "T1"

    def test_title_period_styled_t1(self, tmp_path):
        """Title 'Table 2.Results summary' → T1."""
        def build(doc):
            doc.add_paragraph("Table 2.Results summary")
            _add_table(doc)

        path = _make_docx(tmp_path, build)
        metrics = enforce_table_title_house_rules(str(path))

        assert metrics["titles_fixed"] == 1
        assert _style_of_paragraph(path, "Table 2.") == "T1"

    def test_broad_title_styled_t1(self, tmp_path):
        """Title 'Table 3' (no colon/period) → T1."""
        def build(doc):
            doc.add_paragraph("Table 3")
            _add_table(doc)

        path = _make_docx(tmp_path, build)
        metrics = enforce_table_title_house_rules(str(path))

        assert metrics["titles_fixed"] == 1
        assert _style_of_paragraph(path, "Table 3") == "T1"

    def test_table_without_title_no_change(self, tmp_path):
        """Table with non-title paragraph above → no T1 applied."""
        def build(doc):
            doc.add_paragraph("This is just body text.")
            _add_table(doc)

        path = _make_docx(tmp_path, build)
        metrics = enforce_table_title_house_rules(str(path))

        assert metrics["titles_fixed"] == 0
        assert _style_of_paragraph(path, "This is just") != "T1"


# ===================================================================
# Test: Blank paragraphs between title and table → collapsed
# ===================================================================

class TestBlankLineRemoval:

    def test_single_blank_collapsed(self, tmp_path):
        """Blank paragraph between title and table → PMI style."""
        def build(doc):
            doc.add_paragraph("Table 1: Data")
            doc.add_paragraph("")          # blank
            _add_table(doc)

        path = _make_docx(tmp_path, build)
        metrics = enforce_table_title_house_rules(str(path))

        assert metrics["titles_fixed"] == 1
        assert metrics["blank_lines_removed"] == 1

        styles = _body_paragraph_styles(path)
        # Find the blank paragraph
        blank_styles = [s for t, s in styles if t == ""]
        assert any(s == "PMI" for s in blank_styles)

    def test_multiple_blanks_collapsed(self, tmp_path):
        """Two blank paragraphs between title and table → both PMI."""
        def build(doc):
            doc.add_paragraph("Table 1: Values")
            doc.add_paragraph("")
            doc.add_paragraph("")
            _add_table(doc)

        path = _make_docx(tmp_path, build)
        metrics = enforce_table_title_house_rules(str(path))

        assert metrics["blank_lines_removed"] == 2

    def test_title_with_no_blanks_unchanged(self, tmp_path):
        """Title directly adjacent to table → zero blanks removed."""
        def build(doc):
            doc.add_paragraph("Table 5: Clean")
            _add_table(doc)

        path = _make_docx(tmp_path, build)
        metrics = enforce_table_title_house_rules(str(path))

        assert metrics["blank_lines_removed"] == 0
        assert metrics["titles_fixed"] == 1


# ===================================================================
# Test: Multiple candidates → closest wins, others demoted
# ===================================================================

class TestMultipleCandidates:

    def test_closest_candidate_wins(self, tmp_path):
        """Two title candidates: closest to table becomes T1, farther → TXT."""
        def build(doc):
            doc.add_paragraph("Table 1: First caption")
            doc.add_paragraph("Table 1: Second caption")
            _add_table(doc)

        path = _make_docx(tmp_path, build)
        metrics = enforce_table_title_house_rules(str(path))

        assert metrics["titles_fixed"] == 1
        assert metrics["demoted_to_txt"] == 1

        # Closest (Second caption) → T1, farther (First caption) → TXT
        assert _style_of_paragraph(path, "Table 1: Second") == "T1"
        assert _style_of_paragraph(path, "Table 1: First") == "TXT"

    def test_three_candidates_only_closest_wins(self, tmp_path):
        """Three candidates: only the closest is T1."""
        def build(doc):
            doc.add_paragraph("Table 2: A")
            doc.add_paragraph("Table 2: B")
            doc.add_paragraph("Table 2: C")
            _add_table(doc)

        path = _make_docx(tmp_path, build)
        metrics = enforce_table_title_house_rules(str(path))

        assert metrics["titles_fixed"] == 1
        assert metrics["demoted_to_txt"] == 2
        assert _style_of_paragraph(path, "Table 2: C") == "T1"


# ===================================================================
# Test: Placeholder patterns (<TAB5.1>)
# ===================================================================

class TestPlaceholders:

    def test_placeholder_before_table_becomes_t1(self, tmp_path):
        """<TAB5.1> directly before a table → T1."""
        def build(doc):
            doc.add_paragraph("<TAB5.1>")
            _add_table(doc)

        path = _make_docx(tmp_path, build)
        metrics = enforce_table_title_house_rules(str(path))

        assert metrics["titles_fixed"] == 1
        assert metrics["placeholders_converted"] == 1
        assert _style_of_paragraph(path, "<TAB5.1>") == "T1"

    def test_insert_tab_placeholder(self, tmp_path):
        """<INSERT TAB 3> before a table → T1."""
        def build(doc):
            doc.add_paragraph("<INSERT TAB 3>")
            _add_table(doc)

        path = _make_docx(tmp_path, build)
        metrics = enforce_table_title_house_rules(str(path))

        assert metrics["placeholders_converted"] == 1
        assert _style_of_paragraph(path, "<INSERT TAB 3>") == "T1"

    def test_placeholder_no_table_after_untouched(self, tmp_path):
        """<TAB5.1> with no table following → NOT converted (no candidates found)."""
        def build(doc):
            doc.add_paragraph("<TAB5.1>")
            doc.add_paragraph("Just body text after.")

        path = _make_docx(tmp_path, build)
        original_style = _style_of_paragraph(path, "<TAB5.1>")
        metrics = enforce_table_title_house_rules(str(path))

        assert metrics["placeholders_converted"] == 0
        # Style should be unchanged (no table to enforce against)
        assert _style_of_paragraph(path, "<TAB5.1>") == original_style


# ===================================================================
# Test: Idempotency
# ===================================================================

class TestIdempotency:

    def test_double_run_same_result(self, tmp_path):
        """Running enforcement twice produces the same output."""
        def build(doc):
            doc.add_paragraph("Intro text.")
            doc.add_paragraph("Table 1: Values")
            doc.add_paragraph("")
            _add_table(doc)
            doc.add_paragraph("Table 2: More data")
            _add_table(doc)

        path = _make_docx(tmp_path, build)

        m1 = enforce_table_title_house_rules(str(path))
        styles_after_first = _body_paragraph_styles(path)

        m2 = enforce_table_title_house_rules(str(path))
        styles_after_second = _body_paragraph_styles(path)

        assert styles_after_first == styles_after_second

    def test_already_t1_stays_t1(self, tmp_path):
        """A paragraph already styled T1 before a table stays T1."""
        def build(doc):
            # Pre-create the T1 style
            t1 = doc.styles.add_style("T1", WD_STYLE_TYPE.PARAGRAPH)
            t1.font.size = Pt(10)
            t1.font.bold = True

            p = doc.add_paragraph("Table 1: Already styled")
            p.style = t1
            _add_table(doc)

        path = _make_docx(tmp_path, build)
        metrics = enforce_table_title_house_rules(str(path))

        # Still fixed (style set to T1) but the result is the same
        assert metrics["titles_fixed"] == 1
        assert _style_of_paragraph(path, "Table 1:") == "T1"

    def test_source_title_style_semantics_preserved(self, tmp_path):
        """Source Title-style table captions may canonicalize to T1 with title semantics preserved."""
        def build(doc):
            p = doc.add_paragraph("Table 1: Title semantics caption")
            p.style = doc.styles["Title"]
            _add_table(doc)

        path = _make_docx(tmp_path, build)
        metrics = enforce_table_title_house_rules(str(path))

        assert metrics["titles_fixed"] == 1
        assert _style_of_paragraph(path, "Table 1: Title semantics caption") == "T1"
        doc = Document(str(path))
        body = doc.element.body
        para_el = next(
            p for p in body if p.tag == qn("w:p")
            if _para_text(p).startswith("Table 1: Title semantics caption")
        )
        assert _para_outline_level(para_el) == "0"


# ===================================================================
# Test: Multiple tables
# ===================================================================

class TestMultipleTables:

    def test_two_tables_both_titled(self, tmp_path):
        """Two tables, each with a title → both enforced."""
        def build(doc):
            doc.add_paragraph("Table 1: First")
            _add_table(doc)
            doc.add_paragraph("Table 2: Second")
            _add_table(doc)

        path = _make_docx(tmp_path, build)
        metrics = enforce_table_title_house_rules(str(path))

        assert metrics["tables"] == 2
        assert metrics["titles_fixed"] == 2
        assert _style_of_paragraph(path, "Table 1:") == "T1"
        assert _style_of_paragraph(path, "Table 2:") == "T1"

    def test_first_table_no_title_second_has_title(self, tmp_path):
        """First table has no title, second does → only second fixed."""
        def build(doc):
            doc.add_paragraph("Some regular text.")
            _add_table(doc)
            doc.add_paragraph("Table 2: Data")
            _add_table(doc)

        path = _make_docx(tmp_path, build)
        metrics = enforce_table_title_house_rules(str(path))

        assert metrics["tables"] == 2
        assert metrics["titles_fixed"] == 1
        assert _style_of_paragraph(path, "Table 2:") == "T1"

    def test_adjacent_tables_separate_titles(self, tmp_path):
        """Back-to-back tables: title before first table not confused with second."""
        def build(doc):
            doc.add_paragraph("Table 1: A")
            _add_table(doc)
            # Second table has no title paragraph before it (directly after first table)
            _add_table(doc)

        path = _make_docx(tmp_path, build)
        metrics = enforce_table_title_house_rules(str(path))

        # Only first table has a title candidate
        assert metrics["titles_fixed"] == 1


# ===================================================================
# Test: No tables in document
# ===================================================================

class TestNoTables:

    def test_no_tables_no_changes(self, tmp_path):
        """Document without tables → zero metrics."""
        def build(doc):
            doc.add_paragraph("Just text.")
            doc.add_paragraph("Table 1: Looks like a title but no actual table.")

        path = _make_docx(tmp_path, build)
        metrics = enforce_table_title_house_rules(str(path))

        assert metrics["tables"] == 0
        assert metrics["titles_fixed"] == 0


# ===================================================================
# Test: Table content never modified
# ===================================================================

class TestTableContentPreserved:

    def test_table_cells_unchanged(self, tmp_path):
        """Enforcement must never alter table cell content."""
        def build(doc):
            doc.add_paragraph("Table 1: Values")
            tbl = doc.add_table(rows=2, cols=2)
            tbl.rows[0].cells[0].text = "Header1"
            tbl.rows[0].cells[1].text = "Header2"
            tbl.rows[1].cells[0].text = "Data1"
            tbl.rows[1].cells[1].text = "Data2"

        path = _make_docx(tmp_path, build)
        enforce_table_title_house_rules(str(path))

        # Re-read and check table content
        doc = Document(str(path))
        table = doc.tables[0]
        assert table.rows[0].cells[0].text == "Header1"
        assert table.rows[0].cells[1].text == "Header2"
        assert table.rows[1].cells[0].text == "Data1"
        assert table.rows[1].cells[1].text == "Data2"

    def test_paragraph_text_unchanged(self, tmp_path):
        """Enforcement must never alter paragraph text, only style."""
        def build(doc):
            doc.add_paragraph("Table 1: Original title text")
            _add_table(doc)

        path = _make_docx(tmp_path, build)
        enforce_table_title_house_rules(str(path))

        # Title text should be preserved
        doc = Document(str(path))
        texts = [p.text for p in doc.paragraphs if "Table 1:" in p.text]
        assert len(texts) == 1
        assert texts[0] == "Table 1: Original title text"


# ===================================================================
# Test: Lookback stops at non-title paragraphs
# ===================================================================

class TestLookbackBehavior:

    def test_non_title_paragraph_stops_search(self, tmp_path):
        """Body text between a title candidate and the table blocks the candidate."""
        def build(doc):
            doc.add_paragraph("Table 1: Far away title")
            doc.add_paragraph("Regular body text blocks search.")
            _add_table(doc)

        path = _make_docx(tmp_path, build)
        metrics = enforce_table_title_house_rules(str(path))

        # The non-title paragraph stops the search; no title found
        assert metrics["titles_fixed"] == 0

    def test_another_table_stops_search(self, tmp_path):
        """A table between the candidate and the target table stops search."""
        def build(doc):
            doc.add_paragraph("Table 1: Title for first table")
            _add_table(doc)
            # Second table has no title directly above
            _add_table(doc)

        path = _make_docx(tmp_path, build)
        metrics = enforce_table_title_house_rules(str(path))

        # Only first table gets a title
        assert metrics["titles_fixed"] == 1
        # Second table: lookback hits first table → stops → no candidate
