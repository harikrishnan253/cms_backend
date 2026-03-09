"""
Tests for list-semantics hardening across ingestion → blocks → reconstruction.

Covers:
  - TestIngestionStyleInference   : has_bullet/has_numbering inferred from style names
  - TestBlocksStyleFallback       : _is_list_item / _list_kind / candidate detection
  - TestPreserveSourceListStyle   : reconstruction no longer loses list on <H1> text
  - TestStructureGuardListStyled  : end-to-end: List Bullet + <H1> text → guard PASS
"""

from __future__ import annotations

import sys
from pathlib import Path
from unittest.mock import MagicMock, patch, PropertyMock

import pytest

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "backend"))

from docx import Document
from docx.oxml.ns import qn

from processor.reconstruction import DocumentReconstructor
from processor.structure_guard import enforce_style_only_mutation


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_mock_para(style_name: str, text: str = "Item", has_numpr: bool = False):
    """Return a lightweight mock of a python-docx Paragraph for ingestion tests."""
    para = MagicMock()
    # Style
    style = MagicMock()
    style.name = style_name
    para.style = style
    # Text
    para.text = text
    # Runs
    run = MagicMock()
    run.bold = False
    run.italic = False
    run.font.size = None
    run.font.all_caps = False
    para.runs = [run]
    # Alignment
    para.alignment = None
    # Paragraph format indentation
    pf = MagicMock()
    pf.left_indent = None
    para.paragraph_format = pf
    # XML (_p.pPr.numPr)
    p_elem = MagicMock()
    pPr = MagicMock()
    if has_numpr:
        numPr = MagicMock()
        numPr.ilvl = MagicMock()
        numPr.ilvl.val = 0
        numPr.numId = MagicMock()
        numPr.numId.val = 1
        pPr.numPr = numPr
    else:
        pPr.numPr = None
    p_elem.pPr = pPr
    para._p = p_elem
    return para


# ---------------------------------------------------------------------------
# 1. Ingestion: style-name inference
# ---------------------------------------------------------------------------

class TestIngestionStyleInference:
    """_extract_formatting() must infer has_bullet/has_numbering from style name
    when numPr is absent.  This class exercises the no-numPr inference path."""

    def _extract(self, style_name: str, text: str = "Item text") -> dict:
        from processor.ingestion import DocumentIngestion
        ingestion = DocumentIngestion.__new__(DocumentIngestion)
        para = _make_mock_para(style_name, text=text, has_numpr=False)
        return ingestion._extract_formatting(para)

    # --- "List Bullet" family → has_bullet ---

    def test_list_bullet_sets_has_bullet(self):
        meta = self._extract("List Bullet")
        assert meta["has_bullet"] is True
        assert meta["has_numbering"] is False

    def test_list_bullet_2_sets_has_bullet(self):
        meta = self._extract("List Bullet 2")
        assert meta["has_bullet"] is True

    def test_list_bullet_3_sets_has_bullet(self):
        meta = self._extract("List Bullet 3")
        assert meta["has_bullet"] is True

    def test_bulletlist_publisher_style_sets_has_bullet(self):
        meta = self._extract("BulletList")
        assert meta["has_bullet"] is True

    def test_bulletlist_numbered_variant_sets_has_bullet(self):
        meta = self._extract("BulletList 2")
        assert meta["has_bullet"] is True

    # --- "List Number" family → has_numbering ---

    def test_list_number_sets_has_numbering(self):
        meta = self._extract("List Number")
        assert meta["has_numbering"] is True
        assert meta["has_bullet"] is False

    def test_list_number_2_sets_has_numbering(self):
        meta = self._extract("List Number 2")
        assert meta["has_numbering"] is True

    def test_numberlist_publisher_style_sets_has_numbering(self):
        meta = self._extract("NumberList")
        assert meta["has_numbering"] is True

    def test_numberlist_variant_sets_has_numbering(self):
        meta = self._extract("NumberList 3")
        assert meta["has_numbering"] is True

    # --- Non-list styles must NOT set flags ---

    def test_normal_style_no_flags(self):
        meta = self._extract("Normal")
        assert meta["has_bullet"] is False
        assert meta["has_numbering"] is False

    def test_heading_style_no_flags(self):
        meta = self._extract("Heading 1")
        assert meta["has_bullet"] is False
        assert meta["has_numbering"] is False

    def test_list_paragraph_no_flags(self):
        # "List Paragraph" is a generic indented style — no bullet/number keyword
        meta = self._extract("List Paragraph")
        assert meta["has_bullet"] is False
        assert meta["has_numbering"] is False

    def test_body_text_no_flags(self):
        meta = self._extract("Body Text")
        assert meta["has_bullet"] is False
        assert meta["has_numbering"] is False

    # --- Text-regex still works (unchanged) ---

    def test_text_bullet_still_detected(self):
        meta = self._extract("Normal", text="• First item")
        assert meta["has_bullet"] is True

    def test_text_number_still_detected(self):
        meta = self._extract("Normal", text="1. First item")
        assert meta["has_numbering"] is True

    # --- numPr path unchanged: when numPr present, text determines has_bullet/numbering ---

    def test_numpr_path_unchanged_for_bullet_style(self):
        from processor.ingestion import DocumentIngestion
        ingestion = DocumentIngestion.__new__(DocumentIngestion)
        para = _make_mock_para("List Bullet", text="Item", has_numpr=True)
        meta = ingestion._extract_formatting(para)
        # numPr-present path sets has_bullet via style name inside numPr block
        assert meta["has_bullet"] is True

    def test_numpr_path_unchanged_for_number_style(self):
        from processor.ingestion import DocumentIngestion
        ingestion = DocumentIngestion.__new__(DocumentIngestion)
        para = _make_mock_para("List Number", text="Item", has_numpr=True)
        meta = ingestion._extract_formatting(para)
        assert meta["has_numbering"] is True

    # --- Guard: no double-setting when text already provides the flag ---

    def test_style_inference_skipped_when_already_set_by_text(self):
        """has_bullet already set by text regex → style inference guard fires."""
        meta = self._extract("List Bullet", text="• Explicit bullet text")
        # Still True (both text-regex and style agree), no contradiction
        assert meta["has_bullet"] is True

    # --- style_name key is present in returned metadata ---

    def test_style_name_key_present_in_metadata(self):
        meta = self._extract("List Bullet")
        assert "style_name" in meta
        assert meta["style_name"] == "List Bullet"


# ---------------------------------------------------------------------------
# 2. Blocks: style_name fallback
# ---------------------------------------------------------------------------

class TestBlocksStyleFallback:
    """_is_list_item, _list_kind, and _enrich_list_metadata candidate gate
    must recognise style-only lists even when has_bullet/has_numbering are absent."""

    # ---------- _is_list_item ----------

    def test_is_list_item_via_has_bullet(self):
        from processor.blocks import _is_list_item
        assert _is_list_item({"has_bullet": True}, "Item") is True

    def test_is_list_item_via_has_numbering(self):
        from processor.blocks import _is_list_item
        assert _is_list_item({"has_numbering": True}, "Item") is True

    def test_is_list_item_via_style_name_bullet(self):
        from processor.blocks import _is_list_item
        # No has_bullet flag — only style_name
        assert _is_list_item({"style_name": "List Bullet"}, "Item") is True

    def test_is_list_item_via_style_name_number(self):
        from processor.blocks import _is_list_item
        assert _is_list_item({"style_name": "List Number"}, "Item") is True

    def test_is_list_item_via_bulletlist_publisher_style(self):
        from processor.blocks import _is_list_item
        assert _is_list_item({"style_name": "BulletList"}, "Item") is True

    def test_is_list_item_via_numberlist_publisher_style(self):
        from processor.blocks import _is_list_item
        assert _is_list_item({"style_name": "NumberList"}, "Item") is True

    def test_is_list_item_normal_style_returns_false(self):
        from processor.blocks import _is_list_item
        assert _is_list_item({"style_name": "Normal"}, "Item") is False

    def test_is_list_item_list_paragraph_no_bullet_no_number_returns_false(self):
        from processor.blocks import _is_list_item
        # "List Paragraph" has no bullet/number keyword
        assert _is_list_item({"style_name": "List Paragraph"}, "Item") is False

    def test_is_list_item_empty_meta_returns_false(self):
        from processor.blocks import _is_list_item
        assert _is_list_item({}, "Item") is False

    # ---------- _list_kind ----------

    def test_list_kind_bullet_via_has_bullet(self):
        from processor.blocks import _list_kind
        assert _list_kind({"has_bullet": True}, "Item") == "bullet"

    def test_list_kind_numbered_via_has_numbering(self):
        from processor.blocks import _list_kind
        assert _list_kind({"has_numbering": True}, "Item") == "numbered"

    def test_list_kind_bullet_via_style_name(self):
        from processor.blocks import _list_kind
        assert _list_kind({"style_name": "List Bullet"}, "Item") == "bullet"

    def test_list_kind_numbered_via_style_name(self):
        from processor.blocks import _list_kind
        assert _list_kind({"style_name": "List Number"}, "Item") == "numbered"

    def test_list_kind_bullet_via_bulletlist_publisher_style(self):
        from processor.blocks import _list_kind
        assert _list_kind({"style_name": "BulletList 2"}, "Item") == "bullet"

    def test_list_kind_numbered_via_numberlist_publisher_style(self):
        from processor.blocks import _list_kind
        assert _list_kind({"style_name": "NumberList"}, "Item") == "numbered"

    def test_list_kind_none_for_normal_style(self):
        from processor.blocks import _list_kind
        assert _list_kind({"style_name": "Normal"}, "Item") is None

    # ---------- _enrich_list_metadata candidate gate ----------

    def test_enrich_candidate_gate_uses_style_name(self, tmp_path):
        """A paragraph with only style_name="List Bullet" must be treated as
        a list candidate so the hierarchy detector is invoked."""
        from processor.blocks import _enrich_list_metadata

        paragraphs = [
            {
                "id": 1,
                "text": "Style-only list item",
                "metadata": {"style_name": "List Bullet", "context_zone": "BODY"},
            }
        ]

        # ListHierarchyDetector is imported locally inside _enrich_list_metadata;
        # patch at the source module so the local-import picks up the mock.
        with patch("processor.list_hierarchy_detector.ListHierarchyDetector") as MockDet:
            instance = MockDet.return_value
            det_result = MagicMock()
            det_result.is_list = True
            det_result.style_prefix = "BL"
            det_result.semantic_level = 1
            det_result.indent_twips = 360
            det_result.indent_source = "style"
            instance.detect.return_value = det_result

            result = _enrich_list_metadata(paragraphs, tmp_path / "fake.docx")

        # The detector must have been called (candidate gate was open)
        assert instance.detect.called
        meta = result[0]["metadata"]
        assert meta.get("list_style_prefix") == "BL"

    def test_enrich_candidate_gate_skips_non_list_style(self, tmp_path):
        """A paragraph with style_name="Normal" must NOT be sent to detector."""
        from processor.blocks import _enrich_list_metadata

        paragraphs = [
            {
                "id": 1,
                "text": "Regular paragraph",
                "metadata": {"style_name": "Normal", "context_zone": "BODY"},
            }
        ]

        with patch("processor.list_hierarchy_detector.ListHierarchyDetector") as MockDet:
            instance = MockDet.return_value
            result = _enrich_list_metadata(paragraphs, tmp_path / "fake.docx")

        assert not instance.detect.called


# ---------------------------------------------------------------------------
# 3. Reconstruction: _preserve_source_list_style
# ---------------------------------------------------------------------------

class TestPreserveSourceListStyle:
    """_preserve_source_list_style must return True for style-based lists
    regardless of what source_text contains."""

    def _call(self, source_is_list, source_has_numpr, *, target_tag="TXT", source_text=""):
        return DocumentReconstructor._preserve_source_list_style(
            source_is_list,
            source_has_numpr,
            target_tag=target_tag,
            source_text=source_text,
        )

    # --- Core: style-based list without <H1> → True ---

    def test_style_list_no_h1_returns_true(self):
        assert self._call(True, False, target_tag="TXT") is True

    def test_style_list_empty_source_text_returns_true(self):
        assert self._call(True, False, source_text="") is True

    # --- THE FIX: <H1> marker must NOT override list preservation ---

    def test_style_list_h1_source_text_returns_true(self):
        """Style-based list + <H1> in source text → MUST return True (list preserved)."""
        assert self._call(True, False, source_text="<H1>Chapter One") is True

    def test_style_list_h2_source_text_returns_true(self):
        assert self._call(True, False, source_text="<H2>Section title") is True

    def test_style_list_h6_source_text_returns_true(self):
        assert self._call(True, False, source_text="<H6>Deep heading") is True

    def test_style_list_h1_with_attributes_returns_true(self):
        assert self._call(True, False, source_text='<h1 class="foo">Title') is True

    def test_style_list_h1_lowercase_returns_true(self):
        assert self._call(True, False, source_text="<h1>Chapter") is True

    # --- Non-list source → always False ---

    def test_non_list_no_numpr_returns_false(self):
        assert self._call(False, False) is False

    def test_non_list_with_h1_returns_false(self):
        assert self._call(False, False, source_text="<H1>Chapter") is False

    # --- numPr list → always False (numPr path, not style-based path) ---

    def test_numpr_list_returns_false(self):
        assert self._call(True, True) is False

    def test_numpr_list_with_h1_returns_false(self):
        assert self._call(True, True, source_text="<H1>Chapter") is False

    # --- Target is a list tag → False (let _apply_list_numbering handle it) ---

    def test_style_list_target_is_list_tag_returns_false(self):
        # BL-FIRST is a bullet list tag
        assert self._call(True, False, target_tag="BL-FIRST") is False

    def test_style_list_target_is_nl_tag_returns_false(self):
        assert self._call(True, False, target_tag="NL-FIRST") is False

    def test_style_list_target_txt_returns_true(self):
        assert self._call(True, False, target_tag="TXT") is True

    def test_style_list_target_h1_returns_true(self):
        """H1 tag is NOT a list tag → preserve the style-based list."""
        assert self._call(True, False, target_tag="H1") is True

    def test_style_list_target_h1_with_h1_source_text_returns_true(self):
        """Both target=H1 and source text has <H1> → still preserve list."""
        assert self._call(
            True, False, target_tag="H1", source_text="<H1>Chapter"
        ) is True


# ---------------------------------------------------------------------------
# 4. End-to-end: structure guard passes for List Bullet + <H1> text
# ---------------------------------------------------------------------------

class TestStructureGuardListStyled:
    """Reproduction of the exact failure pattern that motivated this PR.

    Before the fix:
      input_is_list=True (List Bullet style), output_is_list=False (Heading 1 style)
      → STRUCTURE_GUARD_FAIL

    After the fix:
      style-based list preserved → output_is_list=True → PASS
    """

    def test_list_bullet_with_h1_text_guard_passes(self, tmp_path):
        """THE KEY REGRESSION TEST.

        Source: "List Bullet" style paragraph, text starts with '<H1>'.
        Classification: H1 tag.
        Expected: structure guard PASS (list style preserved, not converted to Heading 1).
        """
        input_path = tmp_path / "input.docx"
        doc_in = Document()
        # Paragraph styled as list — text starts with an inline heading marker
        doc_in.add_paragraph("<H1>Chapter One list item", style="List Bullet")
        doc_in.save(str(input_path))

        recon = DocumentReconstructor(output_dir=tmp_path / "out")
        output_path = recon.apply_styles(
            input_path,
            classifications=[{"id": 1, "tag": "H1", "confidence": 99}],
            output_name="styled.docx",
        )

        result = enforce_style_only_mutation(input_path, output_path)
        assert result["status"] == "PASS", (
            f"Expected PASS but got {result['status']}. "
            f"Differences: {result.get('differences', [])}"
        )

    def test_list_bullet_without_h1_marker_guard_passes(self, tmp_path):
        """Baseline: normal list item with non-heading tag → guard PASS."""
        input_path = tmp_path / "input.docx"
        doc_in = Document()
        doc_in.add_paragraph("First item", style="List Bullet")
        doc_in.add_paragraph("Second item", style="List Bullet")
        doc_in.save(str(input_path))

        recon = DocumentReconstructor(output_dir=tmp_path / "out")
        output_path = recon.apply_styles(
            input_path,
            classifications=[
                {"id": 1, "tag": "BL-FIRST", "confidence": 99},
                {"id": 2, "tag": "BL-LAST", "confidence": 99},
            ],
            output_name="styled.docx",
        )

        result = enforce_style_only_mutation(input_path, output_path)
        assert result["status"] == "PASS"

    def test_list_number_with_h2_text_guard_passes(self, tmp_path):
        """List Number + <H2> source text classified as H2 → guard PASS."""
        input_path = tmp_path / "input.docx"
        doc_in = Document()
        doc_in.add_paragraph("<H2>Section heading item", style="List Number")
        doc_in.save(str(input_path))

        recon = DocumentReconstructor(output_dir=tmp_path / "out")
        output_path = recon.apply_styles(
            input_path,
            classifications=[{"id": 1, "tag": "H2", "confidence": 99}],
            output_name="styled.docx",
        )

        result = enforce_style_only_mutation(input_path, output_path)
        assert result["status"] == "PASS", result.get("differences", [])

    def test_multiple_list_items_mixed_h1_text_guard_passes(self, tmp_path):
        """Multiple list items, some with <H1> text, classified as H1/TXT → PASS."""
        input_path = tmp_path / "input.docx"
        doc_in = Document()
        doc_in.add_paragraph("<H1>Heading-like list item", style="List Bullet")
        doc_in.add_paragraph("Regular list item", style="List Bullet")
        doc_in.add_paragraph("<H1>Another heading list item", style="List Bullet")
        doc_in.save(str(input_path))

        recon = DocumentReconstructor(output_dir=tmp_path / "out")
        output_path = recon.apply_styles(
            input_path,
            classifications=[
                {"id": 1, "tag": "H1", "confidence": 99},
                {"id": 2, "tag": "TXT", "confidence": 99},
                {"id": 3, "tag": "H1", "confidence": 99},
            ],
            output_name="styled.docx",
        )

        result = enforce_style_only_mutation(input_path, output_path)
        assert result["status"] == "PASS", result.get("differences", [])

    def test_plain_para_with_h1_text_guard_passes(self, tmp_path):
        """Sanity: plain Normal para + H1 tag → guard PASS (no list injection)."""
        input_path = tmp_path / "input.docx"
        doc_in = Document()
        doc_in.add_paragraph("<H1>Plain paragraph with heading marker")
        doc_in.save(str(input_path))

        recon = DocumentReconstructor(output_dir=tmp_path / "out")
        output_path = recon.apply_styles(
            input_path,
            classifications=[{"id": 1, "tag": "H1", "confidence": 99}],
            output_name="styled.docx",
        )

        result = enforce_style_only_mutation(input_path, output_path)
        assert result["status"] == "PASS"

    def test_list_bullet_classified_as_list_tag_guard_passes(self, tmp_path):
        """List Bullet source re-classified as BL-* semantic tag → guard PASS."""
        input_path = tmp_path / "input.docx"
        doc_in = Document()
        doc_in.add_paragraph("Item A", style="List Bullet")
        doc_in.add_paragraph("Item B", style="List Bullet")
        doc_in.add_paragraph("Item C", style="List Bullet")
        doc_in.save(str(input_path))

        recon = DocumentReconstructor(output_dir=tmp_path / "out")
        output_path = recon.apply_styles(
            input_path,
            classifications=[
                {"id": 1, "tag": "BL-FIRST", "confidence": 99},
                {"id": 2, "tag": "BL-MID",   "confidence": 99},
                {"id": 3, "tag": "BL-LAST",  "confidence": 99},
            ],
            output_name="styled.docx",
        )

        result = enforce_style_only_mutation(input_path, output_path)
        assert result["status"] == "PASS"

    def test_output_does_not_inject_numpr_for_style_list_with_h1(self, tmp_path):
        """When style-based list is preserved for H1 tag, no numPr XML is injected."""
        input_path = tmp_path / "input.docx"
        doc_in = Document()
        doc_in.add_paragraph("<H1>List styled item", style="List Bullet")
        doc_in.save(str(input_path))

        recon = DocumentReconstructor(output_dir=tmp_path / "out")
        output_path = recon.apply_styles(
            input_path,
            classifications=[{"id": 1, "tag": "H1", "confidence": 99}],
            output_name="styled.docx",
        )

        out_doc = Document(str(output_path))
        for para in [p for p in out_doc.paragraphs if p.text.strip()]:
            pPr = para._element.find(qn("w:pPr"))
            if pPr is not None:
                # No new numPr should have been added (source had none)
                assert pPr.find(qn("w:numPr")) is None, (
                    f"Unexpected numPr on '{para.text}'"
                )
