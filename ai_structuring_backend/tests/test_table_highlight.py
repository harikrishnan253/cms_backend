"""
Regression tests for table-specific low-confidence yellow highlighting.

Verifies that:
- Table-zone paragraphs with confidence < TABLE_REVIEW_HIGHLIGHT_THRESHOLD
  receive WD_COLOR_INDEX.YELLOW on all runs.
- Table-zone paragraphs with confidence >= threshold do not get highlighted.
- Non-table body paragraphs (tag=TXT etc.) are NOT highlighted by the table
  feature when their confidence is above the general < 85 threshold.
- Body paragraphs with table-caption tags (T1, TFN, TSN) DO get highlighted
  at the table threshold.
- The threshold is controllable via the optional parameter on apply_styles().
"""

import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "backend"))

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from processor.reconstruction import DocumentReconstructor


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_docx_with_table(tmp_path, body_text="Body paragraph", cell_text="Cell paragraph"):
    """Create a minimal DOCX with one body paragraph and one 1×1 table cell."""
    doc = Document()
    doc.add_paragraph(body_text)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.cell(0, 0).paragraphs[0].text = cell_text
    path = tmp_path / "input.docx"
    doc.save(path)
    return path


def _run_apply(tmp_path, input_path, classifications, threshold=None):
    """Apply styles and return the output Document."""
    rec = DocumentReconstructor(output_dir=str(tmp_path))
    kwargs = {"table_highlight_threshold": threshold} if threshold is not None else {}
    out_path = rec.apply_styles(input_path, classifications, output_name="out.docx", **kwargs)
    return Document(str(out_path))


def _body_para(doc):
    """Return the first (body) paragraph of the document."""
    return doc.paragraphs[0]


def _cell_para(doc):
    """Return the first paragraph of the first cell of the first table."""
    return doc.tables[0].cell(0, 0).paragraphs[0]


def _has_yellow(para):
    """True if every non-empty run in para has YELLOW highlight."""
    runs = para.runs
    if not runs:
        return False
    return all(r.font.highlight_color == WD_COLOR_INDEX.YELLOW for r in runs)


def _has_any_highlight(para):
    """True if any run in para has any non-None highlight."""
    return any(r.font.highlight_color is not None for r in para.runs)


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

class TestTableHighlight:
    """Low-confidence table paragraphs receive yellow highlight."""

    def test_table_para_low_conf_highlighted(self, tmp_path):
        """In-table paragraph with confidence 70 (< 80) → YELLOW highlight."""
        path = _make_docx_with_table(tmp_path)
        # body id=1 (high conf so no interference), cell id=2 (low conf)
        clfs = [
            {"id": 1, "tag": "TXT", "confidence": 90},
            {"id": 2, "tag": "T",   "confidence": 70},
        ]
        doc = _run_apply(tmp_path, path, clfs)
        assert _has_yellow(_cell_para(doc)), "Cell para should be YELLOW at conf=70 (< 80)"

    def test_table_para_high_conf_not_highlighted(self, tmp_path):
        """In-table paragraph with confidence 90 (≥ 80) → no highlight."""
        path = _make_docx_with_table(tmp_path)
        clfs = [
            {"id": 1, "tag": "TXT", "confidence": 90},
            {"id": 2, "tag": "T",   "confidence": 90},
        ]
        doc = _run_apply(tmp_path, path, clfs)
        assert not _has_any_highlight(_cell_para(doc)), (
            "Cell para should NOT be highlighted at conf=90 (≥ 80)"
        )

    def test_nontable_para_not_highlighted_by_table_feature(self, tmp_path):
        """Non-table body paragraph above the general 85 threshold → no highlight.

        Uses conf=87 so it is above both the general < 85 body threshold AND the
        new table < 80 threshold, confirming the table feature does not bleed into
        non-table paragraphs.
        """
        path = _make_docx_with_table(tmp_path)
        clfs = [
            {"id": 1, "tag": "TXT", "confidence": 87},  # above 85 → no general highlight
            {"id": 2, "tag": "T",   "confidence": 70},  # table: should be highlighted
        ]
        doc = _run_apply(tmp_path, path, clfs)
        assert not _has_any_highlight(_body_para(doc)), (
            "Non-table body para (TXT, conf=87) should NOT be highlighted"
        )
        assert _has_yellow(_cell_para(doc)), (
            "Cell para (T, conf=70) SHOULD be highlighted"
        )

    def test_table_caption_body_para_low_conf_highlighted(self, tmp_path):
        """Body paragraph with tag T1 (table caption) at conf=70 → YELLOW highlight."""
        path = _make_docx_with_table(tmp_path)
        clfs = [
            {"id": 1, "tag": "T1", "confidence": 70},   # table caption in body zone
            {"id": 2, "tag": "T",  "confidence": 90},
        ]
        doc = _run_apply(tmp_path, path, clfs)
        assert _has_yellow(_body_para(doc)), (
            "Body para with tag T1 (table caption) at conf=70 should be YELLOW"
        )

    def test_table_caption_body_para_high_conf_not_highlighted(self, tmp_path):
        """Body paragraph with tag T1 at conf=90 (≥ 80) → no highlight."""
        path = _make_docx_with_table(tmp_path)
        clfs = [
            {"id": 1, "tag": "T1", "confidence": 90},
            {"id": 2, "tag": "T",  "confidence": 90},
        ]
        doc = _run_apply(tmp_path, path, clfs)
        assert not _has_any_highlight(_body_para(doc)), (
            "Body para with tag T1 at conf=90 should NOT be highlighted"
        )

    def test_reference_zone_low_conf_highlighted_at_reference_threshold(self, tmp_path):
        """Reference-zone body paragraph with low confidence uses yellow highlight."""
        path = _make_docx_with_table(tmp_path, body_text="Reference entry text", cell_text="Cell")
        clfs = [
            {"id": 1, "tag": "TXT", "confidence": 70, "is_reference_zone": True, "context_zone": "REFERENCE"},
            {"id": 2, "tag": "T", "confidence": 90},
        ]
        doc = _run_apply(tmp_path, path, clfs, threshold=80)
        assert _has_yellow(_body_para(doc)), "Low-confidence reference-zone paragraph should be highlighted"

    def test_reference_zone_uses_reference_threshold_not_general_85(self, tmp_path):
        """Reference-zone para at conf=82 should not highlight when threshold=80."""
        path = _make_docx_with_table(tmp_path, body_text="Reference entry text", cell_text="Cell")
        clfs = [
            {"id": 1, "tag": "TXT", "confidence": 82, "is_reference_zone": True, "context_zone": "REFERENCE"},
            {"id": 2, "tag": "T", "confidence": 90},
        ]
        doc = _run_apply(tmp_path, path, clfs, threshold=80)
        assert not _has_any_highlight(_body_para(doc)), (
            "Reference-zone paragraph should use threshold=80 (not general <85) for review highlight"
        )


class TestThresholdConfigurable:
    """The table_highlight_threshold parameter controls highlight boundary."""

    def test_threshold_boundary_below(self, tmp_path):
        """conf=79, threshold=80 → highlighted (79 < 80)."""
        path = _make_docx_with_table(tmp_path)
        clfs = [
            {"id": 1, "tag": "TXT", "confidence": 90},
            {"id": 2, "tag": "T",   "confidence": 79},
        ]
        doc = _run_apply(tmp_path, path, clfs, threshold=80)
        assert _has_yellow(_cell_para(doc)), "conf=79 < threshold=80 → YELLOW"

    def test_threshold_boundary_at(self, tmp_path):
        """conf=80, threshold=80 → NOT highlighted (80 is not < 80)."""
        path = _make_docx_with_table(tmp_path)
        clfs = [
            {"id": 1, "tag": "TXT", "confidence": 90},
            {"id": 2, "tag": "T",   "confidence": 80},
        ]
        doc = _run_apply(tmp_path, path, clfs, threshold=80)
        assert not _has_any_highlight(_cell_para(doc)), "conf=80 == threshold=80 → no highlight"

    def test_custom_lower_threshold(self, tmp_path):
        """Threshold=60: conf=65 → no highlight (65 ≥ 60); conf=55 → highlighted."""
        path = _make_docx_with_table(tmp_path)

        # conf=65, threshold=60 → no highlight
        clfs_no = [
            {"id": 1, "tag": "TXT", "confidence": 90},
            {"id": 2, "tag": "T",   "confidence": 65},
        ]
        doc_no = _run_apply(tmp_path, path, clfs_no, threshold=60)
        assert not _has_any_highlight(_cell_para(doc_no)), "conf=65 ≥ threshold=60 → no highlight"

        # conf=55, threshold=60 → highlighted
        clfs_yes = [
            {"id": 1, "tag": "TXT", "confidence": 90},
            {"id": 2, "tag": "T",   "confidence": 55},
        ]
        doc_yes = _run_apply(tmp_path, path, clfs_yes, threshold=60)
        assert _has_yellow(_cell_para(doc_yes)), "conf=55 < threshold=60 → YELLOW"
