"""
Performance regression tests for structure_guard.py and integrity.py.

Covers TEST-002 from KNOWN_ISSUES.md (Missing large-document coverage).

Thresholds from PERF-001 and project requirements:
  - 1 000 paragraphs:  < 10 s for structure guard
  - 2 000 paragraphs:  < 60 s for structure guard
  - 2 000 paragraphs:  < 60 s for integrity trigger

Run with:
    pytest backend/tests/test_performance_regression.py --slow -v

These tests are skipped by default (marked @pytest.mark.slow) to keep
the regular test suite fast. Pass ``--slow`` to enable them.
"""

import time
from pathlib import Path

import pytest
from docx import Document

from backend.processor.structure_guard import enforce_style_only_mutation
from backend.processor.integrity import run_integrity_trigger


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _create_large_docx(path: Path, n_paragraphs: int) -> Path:
    """Create a DOCX with *n_paragraphs* simple body paragraphs."""
    doc = Document()
    for i in range(n_paragraphs):
        doc.add_paragraph(f"Paragraph {i}: The quick brown fox jumps over the lazy dog.")
    doc.save(str(path))
    return path


def _create_large_docx_with_tables(path: Path, n_paragraphs: int) -> Path:
    """Create a DOCX with paragraphs and occasional tables."""
    doc = Document()
    for i in range(n_paragraphs):
        if i > 0 and i % 50 == 0:
            # Add a small table every 50 paragraphs
            tbl = doc.add_table(rows=2, cols=3)
            for row in tbl.rows:
                for cell in row.cells:
                    cell.text = f"Cell {i}"
        else:
            doc.add_paragraph(f"Paragraph {i}: Body text content for performance testing.")
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Structure Guard performance tests
# ---------------------------------------------------------------------------


@pytest.mark.slow
class TestStructureGuardPerformance:
    """Structure guard must complete within required time bounds."""

    def test_1000_paragraphs_under_10s(self, tmp_path):
        """1 000-paragraph document: structure guard completes in < 10 s."""
        n = 1000
        input_path = _create_large_docx(tmp_path / "input.docx", n)
        # Output is an identical copy — should PASS (no structural differences)
        output_path = _create_large_docx(tmp_path / "output.docx", n)

        start = time.perf_counter()
        result = enforce_style_only_mutation(input_path, output_path)
        elapsed = time.perf_counter() - start

        assert result["status"] == "PASS", (
            f"Structure guard FAILED on identical {n}-para documents: {result.get('differences')}"
        )
        assert elapsed < 10.0, (
            f"Structure guard took {elapsed:.1f}s for {n} paragraphs — exceeds 10 s threshold"
        )

    def test_2000_paragraphs_under_60s(self, tmp_path):
        """2 000-paragraph document: structure guard completes in < 60 s (project requirement)."""
        n = 2000
        input_path = _create_large_docx(tmp_path / "input.docx", n)
        output_path = _create_large_docx(tmp_path / "output.docx", n)

        start = time.perf_counter()
        result = enforce_style_only_mutation(input_path, output_path)
        elapsed = time.perf_counter() - start

        assert result["status"] == "PASS", (
            f"Structure guard FAILED on identical {n}-para documents: {result.get('differences')}"
        )
        assert elapsed < 60.0, (
            f"Structure guard took {elapsed:.1f}s for {n} paragraphs — exceeds 60 s threshold"
        )

    def test_1000_paragraphs_with_tables_under_15s(self, tmp_path):
        """1 000-paragraph document with tables: structure guard completes in < 15 s."""
        n = 1000
        input_path = _create_large_docx_with_tables(tmp_path / "input.docx", n)
        output_path = _create_large_docx_with_tables(tmp_path / "output.docx", n)

        start = time.perf_counter()
        result = enforce_style_only_mutation(input_path, output_path)
        elapsed = time.perf_counter() - start

        assert result["status"] == "PASS", (
            f"Structure guard FAILED on identical {n}-para + tables documents"
        )
        assert elapsed < 15.0, (
            f"Structure guard took {elapsed:.1f}s for {n} paragraphs with tables — exceeds 15 s"
        )


# ---------------------------------------------------------------------------
# Integrity Trigger performance tests
# ---------------------------------------------------------------------------


@pytest.mark.slow
class TestIntegrityTriggerPerformance:
    """Integrity trigger must complete within required time bounds."""

    def test_1000_paragraphs_under_10s(self, tmp_path):
        """1 000-paragraph document: integrity trigger completes in < 10 s."""
        n = 1000
        input_path = _create_large_docx(tmp_path / "input.docx", n)
        output_path = _create_large_docx(tmp_path / "output.docx", n)

        start = time.perf_counter()
        result = run_integrity_trigger(input_path, output_path)
        elapsed = time.perf_counter() - start

        assert result["status"] == "PASS", (
            f"Integrity trigger FAILED on identical {n}-para documents"
        )
        assert elapsed < 10.0, (
            f"Integrity trigger took {elapsed:.1f}s for {n} paragraphs — exceeds 10 s threshold"
        )

    def test_2000_paragraphs_under_60s(self, tmp_path):
        """2 000-paragraph document: integrity trigger completes in < 60 s (project requirement)."""
        n = 2000
        input_path = _create_large_docx(tmp_path / "input.docx", n)
        output_path = _create_large_docx(tmp_path / "output.docx", n)

        start = time.perf_counter()
        result = run_integrity_trigger(input_path, output_path)
        elapsed = time.perf_counter() - start

        assert result["status"] == "PASS", (
            f"Integrity trigger FAILED on identical {n}-para documents"
        )
        assert elapsed < 60.0, (
            f"Integrity trigger took {elapsed:.1f}s for {n} paragraphs — exceeds 60 s threshold"
        )
