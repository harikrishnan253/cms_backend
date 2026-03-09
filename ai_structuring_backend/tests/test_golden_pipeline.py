import sys
import zipfile
from pathlib import Path

import docx
from docx import Document


ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT))

from backend.app.services.style_normalizer import normalize_style

from processor.pipeline import process_document
from processor.ingestion import extract_document

RAW_ZIP = ROOT / "Org files.zip"
TAG_ZIP = ROOT / "manual tagged files.zip"


def _extract_style_sequence(docx_path: Path) -> list[str]:
    """
    Extract a linear sequence of paragraph styles from:
    - document paragraphs (non-empty only)
    - table cell paragraphs (non-empty only)
    """
    doc = Document(str(docx_path))
    styles: list[str] = []

    # Body paragraphs
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            styles.append(para.style.name if para.style else "Normal")

    # Table paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text and para.text.strip():
                        styles.append(para.style.name if para.style else "Normal")

    return styles


def _extract_text_sequence(docx_path: Path) -> list[str]:
    """
    Extract a linear sequence of paragraph texts aligned to _extract_style_sequence ordering.
    Used only for debugging mismatch context.
    """
    doc = Document(str(docx_path))
    texts: list[str] = []

    for para in doc.paragraphs:
        if para.text and para.text.strip():
            texts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text and para.text.strip():
                        texts.append(para.text.strip())

    return texts


def _build_classifications_from_tagged(tagged_path: Path, raw_path: Path) -> list[dict]:
    """
    Override classifier output using the manual tagged file.
    This ensures any mismatch is due to pipeline validation/repair/writing, NOT the model.
    """
    paragraphs, _ = extract_document(raw_path)

    expected_styles = _extract_style_sequence(tagged_path)
    assert len(expected_styles) == len(paragraphs), (
        f"Tagged docx paragraph count mismatch: "
        f"expected_styles={len(expected_styles)} vs raw_paragraphs={len(paragraphs)}"
    )

    return [{"id": p["id"], "tag": expected_styles[i], "confidence": 0.99} for i, p in enumerate(paragraphs)]


def _collect_pairs() -> list[tuple[str, str]]:
    """
    Match *_org.docx in Org files.zip with *_tag*.docx in manual tagged files.zip.
    We search by base name and do not assume folder names inside the zip.
    """
    with zipfile.ZipFile(RAW_ZIP) as raw_zip, zipfile.ZipFile(TAG_ZIP) as tag_zip:
        raw_names = [n for n in raw_zip.namelist() if n.lower().endswith("_org.docx")]

        tag_candidates = [n for n in tag_zip.namelist() if n.lower().endswith(".docx") and ("_tag" in n.lower())]
        tag_by_base = {}
        for n in tag_candidates:
            base = Path(n).name
            # Accept *_tag.docx or *_tagged.docx
            base_key = base.replace("_tagged.docx", "").replace("_tag.docx", "")
            tag_by_base[base_key] = n

        pairs = []
        for raw_name in raw_names:
            base_key = Path(raw_name).name.replace("_org.docx", "")
            if base_key in tag_by_base:
                pairs.append((raw_name, tag_by_base[base_key]))

        return sorted(pairs)


def test_pipeline_matches_manual_tagged(tmp_path):
    pairs = _collect_pairs()
    assert pairs, "No matching raw/tagged pairs found in zip files"

    # Keep small sample for speed (you can increase later)

    processed = 0
    for idx, (raw_name, tag_name) in enumerate(pairs, start=1):
        if processed >= 5:
            break
        with zipfile.ZipFile(RAW_ZIP) as raw_zip, zipfile.ZipFile(TAG_ZIP) as tag_zip:
            raw_path = Path(raw_zip.extract(raw_name, tmp_path))
            tagged_path = Path(tag_zip.extract(tag_name, tmp_path))

        output_root = tmp_path / f"output_{idx}"
        (output_root / "processed").mkdir(parents=True, exist_ok=True)
        (output_root / "review").mkdir(parents=True, exist_ok=True)
        (output_root / "json").mkdir(parents=True, exist_ok=True)

        expected_styles_raw = _extract_style_sequence(tagged_path)
        paragraphs, _ = extract_document(raw_path)
        if len(expected_styles_raw) != len(paragraphs):
            print(f"Skipping {raw_path.name}: tagged count {len(expected_styles_raw)} != raw {len(paragraphs)}")
            continue

        def classifier_override(blocks, _paragraphs):
            return [{"id": p["id"], "tag": expected_styles_raw[i], "confidence": 0.99} for i, p in enumerate(paragraphs)]

        result = process_document(
            input_path=str(raw_path),
            output_folder=str(output_root),
            use_markers=True,
            classifier_override=classifier_override,
            apply_repair=True,
        )

        processed += 1

        output_path = Path(result["output_path"])
        print(f"\nOUTPUT DOCX: {output_path}")

        output_doc = Document(str(output_path))
        output_style_names = {s.name for s in output_doc.styles}
        if "CN" in expected_styles_raw:
            assert "CN" in output_style_names, "Missing CN style in output docx"
        if "CT" in expected_styles_raw:
            assert "CT" in output_style_names, "Missing CT style in output docx"
        if "H1" in expected_styles_raw:
            assert "H1" in output_style_names, "Missing H1 style in output docx"

        expected_styles_raw = _extract_style_sequence(tagged_path)
        actual_styles_raw = _extract_style_sequence(output_path)

        expected_styles = [normalize_style(s) for s in expected_styles_raw]
        actual_styles = [normalize_style(s) for s in actual_styles_raw]

        if expected_styles != actual_styles:
            expected_texts = _extract_text_sequence(tagged_path)
            actual_texts = _extract_text_sequence(output_path)

            # Find first mismatch
            mismatch_i = None
            for i in range(min(len(expected_styles), len(actual_styles))):
                if expected_styles[i] != actual_styles[i]:
                    mismatch_i = i
                    break

            print("\nFIRST MISMATCH")
            print("Doc:", raw_path.name)
            print("Expected length:", len(expected_styles), "Actual length:", len(actual_styles))

            if mismatch_i is not None:
                i = mismatch_i
                print("Index:", i)
                print("Expected style:", expected_styles_raw[i])
                print("Actual style:", actual_styles_raw[i])
                print("Expected text:", (expected_texts[i] if i < len(expected_texts) else "")[:160])
                print("Actual text:", (actual_texts[i] if i < len(actual_texts) else "")[:160])

                lo = max(0, i - 2)
                hi = min(len(expected_styles), i + 3)

                print("\nExpected context:")
                for j in range(lo, hi):
                    t = expected_texts[j] if j < len(expected_texts) else ""
                    print(j, expected_styles_raw[j], t[:120])

                print("\nActual context:")
                for j in range(lo, hi):
                    t = actual_texts[j] if j < len(actual_texts) else ""
                    print(j, actual_styles_raw[j], t[:120])

            raise AssertionError(f"Style sequence mismatch for {raw_path.name}")

    assert processed == 5, f"Expected 5 processed pairs, got {processed}"
