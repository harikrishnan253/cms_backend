"""Generated DOCX structural regression corpus for hard-gate validation.

Each corpus case is generated dynamically (no static fixtures) and run through:
- reconstruction (style application)
- structure_guard
- integrity trigger

This keeps CI coverage on representative structures that previously regressed.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "backend"))

from processor.integrity import run_integrity_trigger
from processor.reconstruction import DocumentReconstructor
from processor.structure_guard import enforce_style_only_mutation


def _iter_reconstruction_nonempty_paragraphs(doc: Document):
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text and para.text.strip():
                        yield para


def _classifications_for_doc(input_path: Path) -> list[dict]:
    """Build a classification set that intentionally stresses heading/list preservation."""
    doc = Document(str(input_path))
    out: list[dict] = []
    para_id = 1
    for para in _iter_reconstruction_nonempty_paragraphs(doc):
        text = para.text.strip()
        style_name = (para.style.name if para.style else "") or ""

        if text.startswith("<INSERT FIGURE") or text.startswith("<TAB"):
            tag = "PMI"
        elif text.startswith("Table "):
            tag = "T1"
        elif "list-like non-list" in text.lower():
            # Should remain non-list despite list-like tag prediction.
            tag = "BL-MID"
        elif style_name.lower().startswith("list "):
            # Should remain list despite non-list tag prediction.
            tag = "TXT"
        elif style_name.startswith("Heading"):
            # Should remain heading (same level) despite non-heading prediction.
            tag = "TXT"
        elif para._element.getparent() is not None and para._element.getparent().tag.endswith("}tc"):
            tag = "T"
        else:
            tag = "TXT"

        out.append({"id": para_id, "tag": tag, "confidence": 99})
        para_id += 1
    return out


def _build_mixed_structural_case(path: Path) -> None:
    doc = Document()

    # Headings (including deeper heading level)
    doc.add_heading("Corpus Heading 1", level=1)
    p_h7 = doc.add_paragraph("Corpus Heading 7")
    p_h7.style = "Heading 7"

    # Regular paragraphs that classifier may mark list-like
    doc.add_paragraph("List-like non-list paragraph A")
    doc.add_paragraph("List-like non-list paragraph B")

    # Real lists (style-based)
    doc.add_paragraph("Bullet item one", style="List Bullet")
    doc.add_paragraph("Bullet item two", style="List Bullet")
    doc.add_paragraph("Numbered item one", style="List Number")

    # Figure markers and table title/marker text
    doc.add_paragraph("<INSERT FIGURE 7.1 HERE>")
    doc.add_paragraph("<INSERT FIGURE 7.2>")
    doc.add_paragraph("Table 7.1: Regression Corpus Table Title")
    doc.add_paragraph("<TAB7.1>")

    # Table content
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Header A"
    table.rows[0].cells[1].text = "Header B"
    table.rows[1].cells[0].text = "Value 1"
    table.rows[1].cells[1].text = "Value 2"

    # Mixed sections
    doc.add_paragraph("Section one text")
    doc.add_section(WD_SECTION.CONTINUOUS)
    doc.add_paragraph("Section two text")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Section three text")

    doc.save(str(path))


def test_structural_regression_corpus_mixed_doc_roundtrip(tmp_path):
    input_path = tmp_path / "corpus_mixed.docx"
    output_dir = tmp_path / "out"
    output_dir.mkdir()

    _build_mixed_structural_case(input_path)
    classifications = _classifications_for_doc(input_path)

    recon = DocumentReconstructor(output_dir=str(output_dir))
    output_path = recon.apply_styles(
        source_path=input_path,
        classifications=classifications,
        output_name="corpus_mixed_processed.docx",
    )

    structure_result = enforce_style_only_mutation(input_path, output_path)
    assert structure_result["status"] == "PASS"

    integrity_result = run_integrity_trigger(input_path, output_path)
    assert integrity_result["status"] == "PASS"
    assert integrity_result["structural_integrity"]["heading_levels_match"] is True
    assert integrity_result["structural_integrity"]["list_structure_match"] is True
    assert integrity_result["structural_integrity"]["paragraph_count_match"] is True

