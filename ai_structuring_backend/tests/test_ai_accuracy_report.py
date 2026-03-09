import os
import sys
import zipfile
from collections import Counter
from pathlib import Path

import pytest
from docx import Document

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT))

from backend.app.services.allowed_styles import load_allowed_styles
from backend.app.services.style_normalizer import normalize_style
from processor.pipeline import process_document
from processor.ingestion import extract_document


RAW_ZIP = ROOT / "Org files.zip"
TAG_ZIP = ROOT / "manual tagged files.zip"


def _extract_style_sequence(docx_path: Path) -> list[str]:
    doc = Document(str(docx_path))
    styles: list[str] = []
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            styles.append(para.style.name if para.style else "Normal")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text and para.text.strip():
                        styles.append(para.style.name if para.style else "Normal")
    return styles


def _collect_pairs() -> list[tuple[str, str]]:
    with zipfile.ZipFile(RAW_ZIP) as raw_zip, zipfile.ZipFile(TAG_ZIP) as tag_zip:
        raw_names = [n for n in raw_zip.namelist() if n.lower().endswith("_org.docx")]
        tag_candidates = [n for n in tag_zip.namelist() if n.lower().endswith(".docx") and ("_tag" in n.lower())]
        tag_by_base = {}
        for n in tag_candidates:
            base = Path(n).name
            base_key = base.replace("_tagged.docx", "").replace("_tag.docx", "")
            tag_by_base[base_key] = n
        pairs = []
        for raw_name in raw_names:
            base_key = Path(raw_name).name.replace("_org.docx", "")
            if base_key in tag_by_base:
                pairs.append((raw_name, tag_by_base[base_key]))
        return sorted(pairs)


@pytest.mark.slow
def test_ai_accuracy_report(tmp_path):
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        pytest.skip("GOOGLE_API_KEY not set")

    allowed = load_allowed_styles()
    pairs = _collect_pairs()
    assert pairs, "No matching raw/tagged pairs found in zip files"

    processed = 0
    overall_matches = 0
    overall_total = 0
    overall_mismatches = Counter()
    overall_txt_gap = Counter()

    for idx, (raw_name, tag_name) in enumerate(pairs, start=1):
        if processed >= 5:
            break

        with zipfile.ZipFile(RAW_ZIP) as raw_zip, zipfile.ZipFile(TAG_ZIP) as tag_zip:
            raw_path = Path(raw_zip.extract(raw_name, tmp_path))
            tagged_path = Path(tag_zip.extract(tag_name, tmp_path))

        expected_styles_raw = _extract_style_sequence(tagged_path)
        paragraphs, _ = extract_document(raw_path)
        if len(expected_styles_raw) != len(paragraphs):
            continue

        output_root = tmp_path / f"ai_output_{idx}"
        (output_root / "processed").mkdir(parents=True, exist_ok=True)
        (output_root / "review").mkdir(parents=True, exist_ok=True)
        (output_root / "json").mkdir(parents=True, exist_ok=True)

        result = process_document(
            input_path=str(raw_path),
            output_folder=str(output_root),
            use_markers=True,
            apply_repair=True,
        )

        output_path = Path(result["output_path"])
        actual_styles_raw = _extract_style_sequence(output_path)

        expected_styles = [normalize_style(s) for s in expected_styles_raw]
        actual_styles = [normalize_style(s) for s in actual_styles_raw]

        assert len(actual_styles) > 50

        unknown_style_count = sum(1 for s in actual_styles if s not in allowed)
        assert unknown_style_count == 0

        matches = 0
        mismatches = Counter()
        txt_gap = Counter()
        for exp, got in zip(expected_styles, actual_styles):
            if exp == got:
                matches += 1
            else:
                mismatches[(exp, got)] += 1
                if got == "TXT" and exp != "TXT":
                    txt_gap[exp] += 1

        total = len(expected_styles)
        accuracy = (matches / total) * 100 if total else 0

        overall_matches += matches
        overall_total += total
        overall_mismatches.update(mismatches)
        overall_txt_gap.update(txt_gap)

        print("\n=== AI Accuracy Report ===")
        print("Doc:", raw_path.name)
        print(f"Total: {total}  Matches: {matches}  Accuracy: {accuracy:.2f}%")
        print("Top 20 mismatches:")
        for (exp, got), count in mismatches.most_common(20):
            print(f"  {exp} -> {got}: {count}")
        print("Top 20 TXT coverage gaps (expected != TXT):")
        for exp, count in txt_gap.most_common(20):
            print(f"  {exp} -> TXT: {count}")

        processed += 1

    assert processed == 5, f"Expected 5 processed pairs, got {processed}"

    overall_accuracy = (overall_matches / overall_total) * 100 if overall_total else 0
    print("\n=== Overall Summary ===")
    print(f"Docs: {processed}")
    print(f"Total: {overall_total}  Matches: {overall_matches}  Accuracy: {overall_accuracy:.2f}%")
    print("Top 20 overall mismatches:")
    for (exp, got), count in overall_mismatches.most_common(20):
        print(f"  {exp} -> {got}: {count}")
    print("Top 20 overall TXT coverage gaps:")
    for exp, count in overall_txt_gap.most_common(20):
        print(f"  {exp} -> TXT: {count}")
