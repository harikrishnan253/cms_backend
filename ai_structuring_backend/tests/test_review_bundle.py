import sys
from pathlib import Path
import zipfile

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT))

from backend.app.services.review_bundle import create_review_bundle


def test_review_bundle_contents(tmp_path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Input")
    doc.save(input_path)

    doc2 = Document()
    doc2.add_paragraph("Output")
    doc2.save(output_path)

    decisions = [
        {"id": 1, "text": "Sample text", "tag": "TXT", "confidence": 0.5, "repaired": False},
    ]
    quality = {"score": 60, "action": "REVIEW", "metrics": {}, "retry_count": 2}

    bundle = create_review_bundle("job123", str(input_path), str(output_path), decisions, quality)
    bundle_path = Path(bundle)
    assert bundle_path.exists()

    with zipfile.ZipFile(bundle_path) as zf:
        names = set(zf.namelist())
        assert "input_original.docx" in names
        assert "output_tagged.docx" in names
        assert "decisions.json" in names
        assert "quality_report.json" in names
        assert "diff_hint.txt" in names
