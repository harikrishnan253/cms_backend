import sys
import uuid
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT))

from processor.ingestion import DocumentIngestion
from processor.zones import parse_and_normalize_marker


def test_split_run_marker_is_preserved_and_detectable():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("<body-")
    p.add_run("open>")
    doc.add_paragraph("Regular paragraph.")

    tmp_dir = ROOT / "backend" / "tests" / ".tmp"
    tmp_dir.mkdir(parents=True, exist_ok=True)
    docx_path = tmp_dir / f"split_marker_{uuid.uuid4().hex}.docx"
    doc.save(str(docx_path))

    try:
        ingestion = DocumentIngestion()
        paragraphs = ingestion.extract_paragraphs(docx_path)
        assert paragraphs[0]["text"] == "<body-open>"
        assert parse_and_normalize_marker(paragraphs[0]["text"]) == "<body-matter-open>"
    finally:
        if docx_path.exists():
            docx_path.unlink()
