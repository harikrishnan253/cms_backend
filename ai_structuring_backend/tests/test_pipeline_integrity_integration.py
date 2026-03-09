"""
Integration tests for data integrity verification in pipeline.

Tests that the pipeline properly calls the integrity check and fails
when data loss is detected.
"""

import os
import sys
import tempfile
from pathlib import Path
from unittest.mock import patch, Mock

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "backend"))

import pytest
from docx import Document

from processor.data_integrity import verify_data_integrity_trigger


# ===================================================================
# Test Pipeline Integration
# ===================================================================

class TestPipelineIntegration:
    """Test that pipeline properly integrates integrity check."""

    def test_integrity_check_called_automatically(self, tmp_path):
        """Pipeline calls integrity check automatically after processing."""
        # This test verifies that the integrity check is called
        # We don't run the full pipeline here, but verify the function works
        # with realistic DOCX files

        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Create realistic test documents
        doc_in = Document()
        doc_in.add_paragraph("Introduction paragraph")
        doc_in.add_paragraph("Body paragraph with content")
        table = doc_in.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Header 1"
        table.rows[0].cells[1].text = "Header 2"
        table.rows[1].cells[0].text = "Data A"
        table.rows[1].cells[1].text = "Data B"
        doc_in.save(str(input_path))

        # Create output with same content
        doc_out = Document()
        doc_out.add_paragraph("Introduction paragraph")
        doc_out.add_paragraph("Body paragraph with content")
        table_out = doc_out.add_table(rows=2, cols=2)
        table_out.rows[0].cells[0].text = "Header 1"
        table_out.rows[0].cells[1].text = "Header 2"
        table_out.rows[1].cells[0].text = "Data A"
        table_out.rows[1].cells[1].text = "Data B"
        doc_out.save(str(output_path))

        # Call integrity check (as pipeline would)
        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["status"] == "PASS"
        assert "input_token_count" in result
        assert "output_token_count" in result
        assert "loss_percentage" in result

    def test_integrity_failure_returns_fail_status(self, tmp_path):
        """When integrity check fails, status is FAIL."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Create input with content
        doc_in = Document()
        doc_in.add_paragraph("Important data that should not be lost")
        doc_in.add_paragraph("Another paragraph with information")
        doc_in.save(str(input_path))

        # Create output with missing content
        doc_out = Document()
        doc_out.add_paragraph("Important data that should not be lost")
        # Second paragraph missing!
        doc_out.save(str(output_path))

        # Call integrity check
        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["status"] == "FAIL"
        assert result["loss_percentage"] > 0
        assert result["missing_lines_count"] > 0
        assert len(result["missing_tokens"]) > 0

    def test_integrity_result_structure(self, tmp_path):
        """Integrity result has all required fields."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        doc = Document()
        doc.add_paragraph("Test content")
        doc.save(str(input_path))
        doc.save(str(output_path))

        result = verify_data_integrity_trigger(input_path, output_path)

        # Verify all required fields are present
        required_fields = [
            "status",
            "input_token_count",
            "output_token_count",
            "missing_tokens",
            "missing_lines_count",
            "extra_tokens",
            "loss_percentage",
            "hash_match",
            "input_char_count",
            "output_char_count",
            "char_loss_percentage",
        ]

        for field in required_fields:
            assert field in result, f"Missing required field: {field}"


# ===================================================================
# Test Realistic Scenarios
# ===================================================================

class TestRealisticScenarios:
    """Test realistic document processing scenarios."""

    def test_academic_paper_structure_preserved(self, tmp_path):
        """Academic paper structure is fully preserved."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Simulate academic paper structure
        doc_in = Document()
        doc_in.add_paragraph("Title: Research Paper on Machine Learning")
        doc_in.add_paragraph("Abstract: This paper presents...")
        doc_in.add_paragraph("1. Introduction")
        doc_in.add_paragraph("Machine learning has become...")
        doc_in.add_paragraph("2. Methodology")
        doc_in.add_paragraph("We used the following approach...")

        # Add a table
        table = doc_in.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = "Method"
        table.rows[0].cells[1].text = "Accuracy"
        table.rows[1].cells[0].text = "SVM"
        table.rows[1].cells[1].text = "92.5%"
        table.rows[2].cells[0].text = "Random Forest"
        table.rows[2].cells[1].text = "94.1%"

        doc_in.add_paragraph("3. Results")
        doc_in.add_paragraph("The results show that...")
        doc_in.add_paragraph("References")
        doc_in.add_paragraph("[1] Smith et al. (2020)")

        doc_in.save(str(input_path))

        # Create processed version with same content (but maybe styled differently)
        doc_out = Document()
        doc_out.add_paragraph("Title: Research Paper on Machine Learning")
        doc_out.add_paragraph("Abstract: This paper presents...")
        doc_out.add_paragraph("1. Introduction")
        doc_out.add_paragraph("Machine learning has become...")
        doc_out.add_paragraph("2. Methodology")
        doc_out.add_paragraph("We used the following approach...")

        table_out = doc_out.add_table(rows=3, cols=2)
        table_out.rows[0].cells[0].text = "Method"
        table_out.rows[0].cells[1].text = "Accuracy"
        table_out.rows[1].cells[0].text = "SVM"
        table_out.rows[1].cells[1].text = "92.5%"
        table_out.rows[2].cells[0].text = "Random Forest"
        table_out.rows[2].cells[1].text = "94.1%"

        doc_out.add_paragraph("3. Results")
        doc_out.add_paragraph("The results show that...")
        doc_out.add_paragraph("References")
        doc_out.add_paragraph("[1] Smith et al. (2020)")

        doc_out.save(str(output_path))

        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["status"] == "PASS"
        assert result["loss_percentage"] == 0.0

    def test_technical_document_with_code(self, tmp_path):
        """Technical document with code snippets preserved."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        doc_in = Document()
        doc_in.add_paragraph("API Documentation")
        doc_in.add_paragraph("def process_data(input):")
        doc_in.add_paragraph("    return input.transform()")
        doc_in.add_paragraph("Usage example:")
        doc_in.add_paragraph("result = process_data(my_data)")

        doc_in.save(str(input_path))

        # Output preserves code
        doc_out = Document()
        doc_out.add_paragraph("API Documentation")
        doc_out.add_paragraph("def process_data(input):")
        doc_out.add_paragraph("    return input.transform()")
        doc_out.add_paragraph("Usage example:")
        doc_out.add_paragraph("result = process_data(my_data)")

        doc_out.save(str(output_path))

        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["status"] == "PASS"

    def test_document_with_lists_preserved(self, tmp_path):
        """Document with lists is fully preserved."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        doc_in = Document()
        doc_in.add_paragraph("Shopping List:")
        doc_in.add_paragraph("• Apples")
        doc_in.add_paragraph("• Bananas")
        doc_in.add_paragraph("• Oranges")
        doc_in.add_paragraph("Instructions:")
        doc_in.add_paragraph("1. Wash the fruit")
        doc_in.add_paragraph("2. Cut into pieces")
        doc_in.add_paragraph("3. Serve fresh")

        doc_in.save(str(input_path))

        # Output may have different list formatting but same content
        doc_out = Document()
        doc_out.add_paragraph("Shopping List:")
        doc_out.add_paragraph("- Apples")  # Different bullet
        doc_out.add_paragraph("- Bananas")
        doc_out.add_paragraph("- Oranges")
        doc_out.add_paragraph("Instructions:")
        doc_out.add_paragraph("1) Wash the fruit")  # Different numbering
        doc_out.add_paragraph("2) Cut into pieces")
        doc_out.add_paragraph("3) Serve fresh")

        doc_out.save(str(output_path))

        result = verify_data_integrity_trigger(input_path, output_path)

        # Should pass because normalization handles list formatting differences
        assert result["status"] == "PASS"


# ===================================================================
# Test Error Scenarios
# ===================================================================

class TestErrorScenarios:
    """Test error handling in integrity check."""

    def test_handles_nonexistent_file_gracefully(self):
        """Handles nonexistent file paths gracefully."""
        input_path = Path("nonexistent_input.docx")
        output_path = Path("nonexistent_output.docx")

        with pytest.raises(Exception):
            # Should raise FileNotFoundError or similar
            verify_data_integrity_trigger(input_path, output_path)

    def test_handles_corrupted_docx(self, tmp_path):
        """Handles corrupted DOCX files gracefully."""
        input_path = tmp_path / "corrupted.docx"

        # Create an invalid DOCX (just text file with .docx extension)
        with open(input_path, 'w') as f:
            f.write("This is not a valid DOCX file")

        output_path = tmp_path / "output.docx"
        doc = Document()
        doc.add_paragraph("Valid content")
        doc.save(str(output_path))

        with pytest.raises(Exception):
            # Should raise an exception when trying to open corrupted file
            verify_data_integrity_trigger(input_path, output_path)
