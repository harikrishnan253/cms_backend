"""
Tests for data_integrity.py - Data Integrity Verification Trigger

Tests the automated verification that ALL textual data from input DOCX
is preserved in processed output DOCX.
"""

import os
import sys
import tempfile
from pathlib import Path
from unittest.mock import patch, Mock

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "backend"))

import pytest
from docx import Document
from docx.shared import Pt

from processor.data_integrity import (
    verify_data_integrity_trigger,
    _extract_text_from_docx,
    _normalize_text,
    _tokenize,
    _compute_hash,
    _find_missing_tokens,
    _find_extra_tokens,
    _find_missing_lines,
)


# ===================================================================
# Test Helpers
# ===================================================================

def _create_test_docx(file_path: Path, paragraphs: list[str] = None, tables: list[list[list[str]]] = None):
    """Create a test DOCX file with specified content.

    Parameters
    ----------
    file_path : Path
        Path where DOCX will be created
    paragraphs : list of str, optional
        List of paragraph texts
    tables : list of list of list of str, optional
        Tables structure: [[[cell1, cell2], [cell3, cell4]], ...]
    """
    doc = Document()

    # Add paragraphs
    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)

    # Add tables
    if tables:
        for table_data in tables:
            rows = len(table_data)
            cols = len(table_data[0]) if table_data else 0
            table = doc.add_table(rows=rows, cols=cols)
            for i, row_data in enumerate(table_data):
                for j, cell_text in enumerate(row_data):
                    table.rows[i].cells[j].text = cell_text

    doc.save(str(file_path))


# ===================================================================
# Test Text Extraction
# ===================================================================

class TestTextExtraction:
    """Test _extract_text_from_docx function."""

    def test_extract_simple_paragraphs(self, tmp_path):
        """Extract text from document with only paragraphs."""
        doc_path = tmp_path / "test.docx"
        paragraphs = ["First paragraph", "Second paragraph", "Third paragraph"]
        _create_test_docx(doc_path, paragraphs=paragraphs)

        text = _extract_text_from_docx(doc_path)

        assert "First paragraph" in text
        assert "Second paragraph" in text
        assert "Third paragraph" in text

    def test_extract_from_tables(self, tmp_path):
        """Extract text from table cells."""
        doc_path = tmp_path / "test.docx"
        tables = [[
            ["Header 1", "Header 2"],
            ["Cell A", "Cell B"],
            ["Cell C", "Cell D"]
        ]]
        _create_test_docx(doc_path, tables=tables)

        text = _extract_text_from_docx(doc_path)

        assert "Header 1" in text
        assert "Header 2" in text
        assert "Cell A" in text
        assert "Cell B" in text
        assert "Cell C" in text
        assert "Cell D" in text

    def test_extract_mixed_content(self, tmp_path):
        """Extract text from document with paragraphs and tables."""
        doc_path = tmp_path / "test.docx"
        paragraphs = ["Introduction text"]
        tables = [[["Data 1", "Data 2"]]]
        _create_test_docx(doc_path, paragraphs=paragraphs, tables=tables)

        text = _extract_text_from_docx(doc_path)

        assert "Introduction text" in text
        assert "Data 1" in text
        assert "Data 2" in text

    def test_extract_empty_document(self, tmp_path):
        """Extract from empty document returns empty string."""
        doc_path = tmp_path / "empty.docx"
        _create_test_docx(doc_path)

        text = _extract_text_from_docx(doc_path)

        assert text == ""


# ===================================================================
# Test Text Normalization
# ===================================================================

class TestTextNormalization:
    """Test _normalize_text function."""

    def test_normalize_lowercase(self):
        """Convert text to lowercase."""
        text = "Hello WORLD Test"
        normalized = _normalize_text(text, remove_numbering=False)
        assert normalized == "hello world test"

    def test_normalize_whitespace(self):
        """Collapse multiple spaces into one."""
        text = "Hello    world   test"
        normalized = _normalize_text(text, remove_numbering=False)
        assert normalized == "hello world test"

    def test_normalize_smart_quotes(self):
        """Normalize smart quotes to ASCII."""
        text = "\u2018single\u2019 \u201Cdouble\u201D"
        normalized = _normalize_text(text, remove_numbering=False)
        assert "'" in normalized
        assert '"' in normalized
        assert "\u2018" not in normalized
        assert "\u201C" not in normalized

    def test_normalize_dashes(self):
        """Normalize en-dash and em-dash to hyphen."""
        text = "en\u2013dash em\u2014dash"
        normalized = _normalize_text(text, remove_numbering=False)
        assert "en-dash" in normalized
        assert "em-dash" in normalized

    def test_normalize_zero_width_chars(self):
        """Remove zero-width characters."""
        text = "hello\u200Bworld\u200Ctest\u200D"
        normalized = _normalize_text(text, remove_numbering=False)
        assert "\u200B" not in normalized
        assert "\u200C" not in normalized
        assert "\u200D" not in normalized

    def test_normalize_unicode_nfkc(self):
        """Apply NFKC normalization."""
        text = "café"  # Can be composed differently
        normalized = _normalize_text(text, remove_numbering=False)
        # Should be consistent regardless of input composition
        assert "cafe" in normalized.lower() or "café" in normalized.lower()

    def test_remove_numbered_list_prefix(self):
        """Remove list numbering prefixes."""
        text = "1. First item\n2. Second item\n3. Third item"
        normalized = _normalize_text(text, remove_numbering=True)
        assert "first item" in normalized
        assert "second item" in normalized
        assert "third item" in normalized
        # Number prefixes should be removed
        assert not normalized.startswith("1.")
        assert not normalized.startswith("2.")

    def test_remove_lettered_list_prefix(self):
        """Remove lettered list prefixes."""
        text = "a. First item\nb) Second item\n(c) Third item"
        normalized = _normalize_text(text, remove_numbering=True)
        assert "first item" in normalized
        assert "second item" in normalized
        assert "third item" in normalized

    def test_remove_bullet_prefix(self):
        """Remove bullet point prefixes."""
        text = "• Bullet one\n○ Bullet two\n■ Bullet three"
        normalized = _normalize_text(text, remove_numbering=True)
        assert "bullet one" in normalized
        assert "bullet two" in normalized
        assert "bullet three" in normalized

    def test_preserve_numbering_when_disabled(self):
        """Preserve numbering when remove_numbering=False."""
        text = "1. First item"
        normalized = _normalize_text(text, remove_numbering=False)
        assert "1." in normalized


# ===================================================================
# Test Tokenization
# ===================================================================

class TestTokenization:
    """Test _tokenize function."""

    def test_tokenize_simple_text(self):
        """Tokenize simple text into words."""
        text = "hello world test"
        tokens = _tokenize(text)
        assert tokens == ["hello", "world", "test"]

    def test_tokenize_with_punctuation(self):
        """Tokenize text with punctuation."""
        text = "hello, world! test?"
        tokens = _tokenize(text)
        # Note: simple split() keeps punctuation attached
        assert "hello," in tokens
        assert "world!" in tokens
        assert "test?" in tokens

    def test_tokenize_empty_string(self):
        """Tokenize empty string returns empty list."""
        text = ""
        tokens = _tokenize(text)
        assert tokens == []

    def test_tokenize_whitespace_only(self):
        """Tokenize whitespace-only string returns empty list."""
        text = "   \n\t  "
        tokens = _tokenize(text)
        assert tokens == []


# ===================================================================
# Test Hash Computation
# ===================================================================

class TestHashComputation:
    """Test _compute_hash function."""

    def test_hash_deterministic(self):
        """Same text produces same hash."""
        text = "hello world"
        hash1 = _compute_hash(text)
        hash2 = _compute_hash(text)
        assert hash1 == hash2

    def test_hash_different_for_different_text(self):
        """Different text produces different hash."""
        hash1 = _compute_hash("hello world")
        hash2 = _compute_hash("hello world!")
        assert hash1 != hash2

    def test_hash_returns_hex_string(self):
        """Hash returns hexadecimal string."""
        text = "test"
        hash_val = _compute_hash(text)
        # SHA256 produces 64-character hex string
        assert len(hash_val) == 64
        assert all(c in "0123456789abcdef" for c in hash_val)


# ===================================================================
# Test Token Comparison
# ===================================================================

class TestTokenComparison:
    """Test token comparison functions."""

    def test_find_missing_tokens_none(self):
        """No missing tokens when sets are identical."""
        input_tokens = ["hello", "world", "test"]
        output_tokens = ["hello", "world", "test"]
        missing = _find_missing_tokens(input_tokens, output_tokens)
        assert missing == []

    def test_find_missing_tokens_some(self):
        """Detect missing tokens."""
        input_tokens = ["hello", "world", "test", "data"]
        output_tokens = ["hello", "test"]
        missing = _find_missing_tokens(input_tokens, output_tokens)
        assert set(missing) == {"world", "data"}

    def test_find_extra_tokens(self):
        """Detect extra tokens in output."""
        input_tokens = ["hello", "world"]
        output_tokens = ["hello", "world", "extra", "new"]
        extra = _find_extra_tokens(input_tokens, output_tokens)
        assert set(extra) == {"extra", "new"}

    def test_find_missing_lines(self):
        """Detect missing lines."""
        input_text = "line one\nline two\nline three"
        output_text = "line one\nline three"
        missing_count = _find_missing_lines(input_text, output_text)
        assert missing_count == 1


# ===================================================================
# Test verify_data_integrity_trigger - PASS Cases
# ===================================================================

class TestIntegrityVerificationPass:
    """Test cases where integrity check should PASS."""

    def test_identical_documents(self, tmp_path):
        """Identical documents pass integrity check."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs = ["First paragraph", "Second paragraph"]
        _create_test_docx(input_path, paragraphs=paragraphs)
        _create_test_docx(output_path, paragraphs=paragraphs)

        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["status"] == "PASS"
        assert result["loss_percentage"] == 0.0
        assert result["missing_tokens"] == []
        assert result["missing_lines_count"] == 0
        assert result["hash_match"] is True

    def test_only_style_changed(self, tmp_path):
        """Same content with different styles passes."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Create input with plain text
        doc_in = Document()
        para = doc_in.add_paragraph("Test paragraph with some content")
        doc_in.save(str(input_path))

        # Create output with same text but styled
        doc_out = Document()
        para = doc_out.add_paragraph("Test paragraph with some content")
        para.runs[0].bold = True
        para.runs[0].font.size = Pt(14)
        doc_out.save(str(output_path))

        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["status"] == "PASS"
        assert result["loss_percentage"] == 0.0

    def test_numbering_changed_but_text_same(self, tmp_path):
        """Numbering changes but text preserved passes."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Input with numbered list
        paragraphs_in = ["1. First item", "2. Second item"]
        _create_test_docx(input_path, paragraphs=paragraphs_in)

        # Output with different numbering but same content
        paragraphs_out = ["a) First item", "b) Second item"]
        _create_test_docx(output_path, paragraphs=paragraphs_out)

        result = verify_data_integrity_trigger(input_path, output_path)

        # Should pass because normalization removes numbering
        assert result["status"] == "PASS"

    def test_whitespace_differences_ignored(self, tmp_path):
        """Whitespace differences are normalized."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs_in = ["Hello    world", "Test   paragraph"]
        paragraphs_out = ["Hello world", "Test paragraph"]

        _create_test_docx(input_path, paragraphs=paragraphs_in)
        _create_test_docx(output_path, paragraphs=paragraphs_out)

        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["status"] == "PASS"

    def test_case_differences_ignored(self, tmp_path):
        """Case differences are normalized."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs_in = ["HELLO WORLD"]
        paragraphs_out = ["hello world"]

        _create_test_docx(input_path, paragraphs=paragraphs_in)
        _create_test_docx(output_path, paragraphs=paragraphs_out)

        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["status"] == "PASS"


# ===================================================================
# Test verify_data_integrity_trigger - FAIL Cases
# ===================================================================

class TestIntegrityVerificationFail:
    """Test cases where integrity check should FAIL."""

    def test_one_word_removed(self, tmp_path):
        """Removing one word causes failure."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs_in = ["Hello world test data"]
        paragraphs_out = ["Hello world test"]  # "data" removed

        _create_test_docx(input_path, paragraphs=paragraphs_in)
        _create_test_docx(output_path, paragraphs=paragraphs_out)

        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["status"] == "FAIL"
        assert result["loss_percentage"] > 0
        assert "data" in result["missing_tokens"]

    def test_one_paragraph_removed(self, tmp_path):
        """Removing entire paragraph causes failure."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs_in = ["First paragraph", "Second paragraph", "Third paragraph"]
        paragraphs_out = ["First paragraph", "Third paragraph"]  # Second removed

        _create_test_docx(input_path, paragraphs=paragraphs_in)
        _create_test_docx(output_path, paragraphs=paragraphs_out)

        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["status"] == "FAIL"
        assert result["missing_lines_count"] > 0
        assert "second" in result["missing_tokens"]

    def test_table_cell_removed(self, tmp_path):
        """Removing table cell content causes failure."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        tables_in = [[["Cell A", "Cell B"], ["Cell C", "Cell D"]]]
        tables_out = [[["Cell A", "Cell B"], ["Cell C", ""]]]  # Cell D removed

        _create_test_docx(input_path, tables=tables_in)
        _create_test_docx(output_path, tables=tables_out)

        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["status"] == "FAIL"
        assert "d" in result["missing_tokens"]  # "cell" and "d"

    def test_partial_text_loss(self, tmp_path):
        """Partial text loss is detected."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs_in = [
            "Paragraph one with content",
            "Paragraph two with data",
            "Paragraph three with information"
        ]
        paragraphs_out = [
            "Paragraph one with",  # "content" removed
            "Paragraph two with data",
            "Paragraph three"  # "with information" removed
        ]

        _create_test_docx(input_path, paragraphs=paragraphs_in)
        _create_test_docx(output_path, paragraphs=paragraphs_out)

        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["status"] == "FAIL"
        assert "content" in result["missing_tokens"]
        assert "information" in result["missing_tokens"]

    def test_empty_output_fails(self, tmp_path):
        """Empty output when input has content fails."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs_in = ["Content exists"]
        _create_test_docx(input_path, paragraphs=paragraphs_in)
        _create_test_docx(output_path)  # Empty

        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["status"] == "FAIL"
        assert result["loss_percentage"] == 100.0


# ===================================================================
# Test Extra Content Handling
# ===================================================================

class TestExtraContentHandling:
    """Test handling of extra content in output."""

    def test_extra_content_with_warning(self, tmp_path):
        """Extra content generates warning but can still pass if no loss."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs_in = ["Original content"]
        paragraphs_out = ["Original content", "Extra added content"]

        _create_test_docx(input_path, paragraphs=paragraphs_in)
        _create_test_docx(output_path, paragraphs=paragraphs_out)

        result = verify_data_integrity_trigger(input_path, output_path)

        # Should still PASS because no input content is lost
        assert result["status"] == "PASS"
        assert len(result["extra_tokens"]) > 0
        assert "added" in result["extra_tokens"]


# ===================================================================
# Test Edge Cases
# ===================================================================

class TestEdgeCases:
    """Test edge cases and boundary conditions."""

    def test_empty_input_and_output(self, tmp_path):
        """Both empty documents pass."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        _create_test_docx(input_path)
        _create_test_docx(output_path)

        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["status"] == "PASS"
        assert result["input_token_count"] == 0
        assert result["output_token_count"] == 0

    def test_special_characters_preserved(self, tmp_path):
        """Special characters are preserved correctly."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs = ["Test with @#$%^&*() special chars"]
        _create_test_docx(input_path, paragraphs=paragraphs)
        _create_test_docx(output_path, paragraphs=paragraphs)

        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["status"] == "PASS"

    def test_unicode_characters_preserved(self, tmp_path):
        """Unicode characters are preserved correctly."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs = ["Test with émojis 😀 and café"]
        _create_test_docx(input_path, paragraphs=paragraphs)
        _create_test_docx(output_path, paragraphs=paragraphs)

        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["status"] == "PASS"

    def test_large_document_simulation(self, tmp_path):
        """Simulate larger document processing."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Create document with many paragraphs
        paragraphs = [f"Paragraph number {i} with content" for i in range(100)]
        _create_test_docx(input_path, paragraphs=paragraphs)
        _create_test_docx(output_path, paragraphs=paragraphs)

        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["status"] == "PASS"
        assert result["input_token_count"] > 0


# ===================================================================
# Test Metrics Calculation
# ===================================================================

class TestMetricsCalculation:
    """Test metrics calculation accuracy."""

    def test_loss_percentage_calculation(self, tmp_path):
        """Loss percentage is calculated correctly."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # 10 words in input, remove 2 = 20% loss
        paragraphs_in = ["one two three four five six seven eight nine ten"]
        paragraphs_out = ["one two three four five six seven eight"]  # 2 removed

        _create_test_docx(input_path, paragraphs=paragraphs_in)
        _create_test_docx(output_path, paragraphs=paragraphs_out)

        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["status"] == "FAIL"
        assert result["loss_percentage"] == 20.0

    def test_token_counts_accurate(self, tmp_path):
        """Token counts are accurate."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs = ["hello world test data"]  # 4 tokens
        _create_test_docx(input_path, paragraphs=paragraphs)
        _create_test_docx(output_path, paragraphs=paragraphs)

        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["input_token_count"] == 4
        assert result["output_token_count"] == 4

    def test_char_count_metrics(self, tmp_path):
        """Character count metrics are calculated."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs_in = ["hello world"]
        paragraphs_out = ["hello"]

        _create_test_docx(input_path, paragraphs=paragraphs_in)
        _create_test_docx(output_path, paragraphs=paragraphs_out)

        result = verify_data_integrity_trigger(input_path, output_path)

        assert result["input_char_count"] > 0
        assert result["output_char_count"] < result["input_char_count"]
        assert result["char_loss_percentage"] > 0
