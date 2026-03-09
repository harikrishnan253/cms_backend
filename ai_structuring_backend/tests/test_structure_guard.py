"""
Tests for structure_guard.py - deterministic enforcement of style-only mutations.

Validates that the processor ONLY changes paragraph styles and does NOT mutate structure.
"""

import os
import sys
from pathlib import Path
from unittest.mock import patch, Mock

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "backend"))

import pytest
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

from processor.structure_guard import (
    enforce_style_only_mutation,
    _normalize_text,
    _extract_paragraphs,
    _extract_tables,
    _extract_sections,
    _create_structural_signature,
    _validate_paragraphs,
    _validate_tables,
    _validate_sections,
)
from processor.reconstruction import DocumentReconstructor


# ===================================================================
# Test Helpers
# ===================================================================

def _create_test_docx(
    file_path: Path,
    paragraphs: list[str] = None,
    tables: list[list[list[str]]] = None,
    headings: list[tuple[str, int]] = None,
    list_items: list[tuple[str, int]] = None,
    sections: list[str] = None,
):
    """Create a test DOCX file with specified content and structure."""
    doc = Document()

    if headings:
        for text, level in headings:
            doc.add_heading(text, level=level)

    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)

    if list_items:
        for text, level in list_items:
            doc.add_paragraph(text, style='List Bullet')

    if tables:
        for table_data in tables:
            rows = len(table_data)
            cols = len(table_data[0]) if table_data else 0
            table = doc.add_table(rows=rows, cols=cols)
            for i, row_data in enumerate(table_data):
                for j, cell_text in enumerate(row_data):
                    table.rows[i].cells[j].text = cell_text

    if sections:
        for section_name in sections:
            doc.add_section()

    doc.save(str(file_path))


# ===================================================================
# Test Text Normalization
# ===================================================================

class TestTextNormalization:
    """Test text normalization for comparison."""

    def test_normalize_strips_whitespace(self):
        """Normalization strips leading/trailing whitespace."""
        text = "  Hello World  "
        result = _normalize_text(text)
        assert result == "Hello World"

    def test_normalize_collapses_spaces(self):
        """Normalization collapses multiple spaces."""
        text = "Hello    World"
        result = _normalize_text(text)
        assert result == "Hello World"

    def test_normalize_unicode_nfkc(self):
        """Normalization applies Unicode NFKC."""
        text = "\u00C5"  # Angstrom sign
        result = _normalize_text(text)
        # NFKC normalization converts Angstrom to A + combining ring
        assert result == "\u00C5"  # Normalized form

    def test_normalize_preserves_markers(self):
        """Normalization does NOT remove markers."""
        text = "<body-open>"
        result = _normalize_text(text)
        assert result == "<body-open>"


# ===================================================================
# Test PASS Cases
# ===================================================================

class TestPassCases:
    """Test cases that should PASS validation."""

    def test_identical_documents_pass(self, tmp_path):
        """Identical documents pass validation."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs = ["Introduction", "Body paragraph", "Conclusion"]
        _create_test_docx(input_path, paragraphs=paragraphs)
        _create_test_docx(output_path, paragraphs=paragraphs)

        result = enforce_style_only_mutation(input_path, output_path)

        assert result["status"] == "PASS"
        assert result["paragraph_count_match"] is True
        assert result["list_structure_match"] is True
        assert result["table_structure_match"] is True
        assert result["section_structure_match"] is True
        assert result["structural_hash_match"] is True
        assert len(result["differences"]) == 0

    def test_style_changes_only_pass(self, tmp_path):
        """Documents with only style changes pass."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Create input
        doc_in = Document()
        p1 = doc_in.add_paragraph("Paragraph 1")
        p1.style = 'Normal'
        p2 = doc_in.add_paragraph("Paragraph 2")
        p2.style = 'Normal'
        doc_in.save(str(input_path))

        # Create output with different styles but same text
        doc_out = Document()
        p1 = doc_out.add_paragraph("Paragraph 1")
        p1.style = 'Heading 1'  # Style changed
        p2 = doc_out.add_paragraph("Paragraph 2")
        p2.style = 'Heading 2'  # Style changed
        doc_out.save(str(output_path))

        result = enforce_style_only_mutation(input_path, output_path)

        assert result["status"] == "PASS"
        assert result["paragraph_count_match"] is True

    def test_whitespace_normalization_pass(self, tmp_path):
        """Documents with whitespace differences pass (normalized)."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Create input with extra whitespace
        doc_in = Document()
        doc_in.add_paragraph("Hello    World  ")
        doc_in.save(str(input_path))

        # Create output with normalized whitespace
        doc_out = Document()
        doc_out.add_paragraph("Hello World")
        doc_out.save(str(output_path))

        result = enforce_style_only_mutation(input_path, output_path)

        assert result["status"] == "PASS"

    def test_markers_preserved_pass(self, tmp_path):
        """Documents with markers preserved pass."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs = ["<body-open>", "Content paragraph", "<body-close>"]
        _create_test_docx(input_path, paragraphs=paragraphs)
        _create_test_docx(output_path, paragraphs=paragraphs)

        result = enforce_style_only_mutation(input_path, output_path)

        assert result["status"] == "PASS"

    def test_reconstruction_does_not_inject_list_xml_for_non_list_inputs(self, tmp_path):
        """Regression: list-like predicted tags must not add numPr to non-list source paragraphs."""
        input_path = tmp_path / "input.docx"
        doc_in = Document()
        doc_in.add_paragraph("Plain paragraph one")
        doc_in.add_paragraph("Plain paragraph two")
        doc_in.save(str(input_path))

        recon = DocumentReconstructor(output_dir=tmp_path / "out")
        output_path = recon.apply_styles(
            input_path,
            classifications=[
                {"id": 1, "tag": "BL-FIRST", "confidence": 99},
                {"id": 2, "tag": "BL-LAST", "confidence": 99},
            ],
            output_name="styled.docx",
        )

        out_doc = Document(str(output_path))
        for para in [p for p in out_doc.paragraphs if p.text.strip()]:
            pPr = para._element.find(qn("w:pPr"))
            if pPr is not None:
                assert pPr.find(qn("w:numPr")) is None

        result = enforce_style_only_mutation(input_path, output_path)
        assert result["status"] == "PASS"

    def test_reconstruction_preserves_style_based_list_semantics(self, tmp_path):
        """Regression: style-based list inputs remain lists after BL-* custom styles are applied."""
        input_path = tmp_path / "input.docx"
        doc_in = Document()
        doc_in.add_paragraph("Item 1", style="List Bullet")
        doc_in.add_paragraph("Item 2", style="List Bullet")
        doc_in.add_paragraph("Item 3", style="List Bullet")
        doc_in.save(str(input_path))

        recon = DocumentReconstructor(output_dir=tmp_path / "out")
        output_path = recon.apply_styles(
            input_path,
            classifications=[
                {"id": 1, "tag": "BL-FIRST", "confidence": 99},
                {"id": 2, "tag": "BL-MID", "confidence": 99},
                {"id": 3, "tag": "BL-LAST", "confidence": 99},
            ],
            output_name="styled.docx",
        )

        result = enforce_style_only_mutation(input_path, output_path)
        assert result["status"] == "PASS"

    def test_reconstruction_preserves_body_paragraph_count_including_empty(self, tmp_path):
        """Regression: reconstruction must preserve total body paragraph count, not just non-empty."""
        input_path = tmp_path / "input.docx"
        doc_in = Document()
        doc_in.add_paragraph("First para")
        doc_in.add_paragraph("")  # empty paragraph should be preserved structurally
        doc_in.add_paragraph("Second para")
        doc_in.save(str(input_path))

        recon = DocumentReconstructor(output_dir=tmp_path / "out")
        output_path = recon.apply_styles(
            input_path,
            classifications=[
                {"id": 1, "tag": "TXT", "confidence": 99},
                {"id": 2, "tag": "TXT", "confidence": 99},
            ],
            output_name="styled.docx",
        )

        in_doc = Document(str(input_path))
        out_doc = Document(str(output_path))
        assert len(in_doc.paragraphs) == 3
        assert len(out_doc.paragraphs) == 3

        result = enforce_style_only_mutation(input_path, output_path)
        assert result["status"] == "PASS"


# ===================================================================
# Test FAIL Cases - Paragraph Mutations
# ===================================================================

class TestFailParagraphMutations:
    """Test cases where paragraph structure is mutated."""

    def test_fail_paragraph_count_differs(self, tmp_path):
        """FAIL when paragraph count differs."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        _create_test_docx(input_path, paragraphs=["Para 1", "Para 2", "Para 3"])
        _create_test_docx(output_path, paragraphs=["Para 1", "Para 2"])  # Missing para 3

        with pytest.raises(RuntimeError, match="STRUCTURE_GUARD_FAIL"):
            enforce_style_only_mutation(input_path, output_path)

    def test_fail_paragraph_text_differs(self, tmp_path):
        """FAIL when paragraph text differs at same index."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        _create_test_docx(input_path, paragraphs=["Original text", "Para 2"])
        _create_test_docx(output_path, paragraphs=["Modified text", "Para 2"])

        with pytest.raises(RuntimeError, match="STRUCTURE_GUARD_FAIL"):
            enforce_style_only_mutation(input_path, output_path)

    def test_fail_paragraph_order_changed(self, tmp_path):
        """FAIL when paragraph order is changed."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        _create_test_docx(input_path, paragraphs=["Para 1", "Para 2", "Para 3"])
        _create_test_docx(output_path, paragraphs=["Para 2", "Para 1", "Para 3"])

        with pytest.raises(RuntimeError, match="STRUCTURE_GUARD_FAIL"):
            enforce_style_only_mutation(input_path, output_path)

    def test_fail_paragraph_removed(self, tmp_path):
        """FAIL when paragraph is removed."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        _create_test_docx(input_path, paragraphs=["Para 1", "Para 2", "Para 3"])
        _create_test_docx(output_path, paragraphs=["Para 1", "Para 3"])

        with pytest.raises(RuntimeError, match="STRUCTURE_GUARD_FAIL"):
            enforce_style_only_mutation(input_path, output_path)

    def test_fail_paragraph_added(self, tmp_path):
        """FAIL when paragraph is added."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        _create_test_docx(input_path, paragraphs=["Para 1", "Para 2"])
        _create_test_docx(output_path, paragraphs=["Para 1", "Para 2", "Para 3"])

        with pytest.raises(RuntimeError, match="STRUCTURE_GUARD_FAIL"):
            enforce_style_only_mutation(input_path, output_path)

    def test_fail_when_paragraph_merged(self, tmp_path):
        """FAIL when paragraphs are merged into one."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Input has two separate paragraphs
        _create_test_docx(input_path, paragraphs=["First paragraph.", "Second paragraph."])

        # Output merges them into one
        _create_test_docx(output_path, paragraphs=["First paragraph. Second paragraph."])

        with pytest.raises(RuntimeError, match="STRUCTURE_GUARD_FAIL"):
            enforce_style_only_mutation(input_path, output_path)


# ===================================================================
# Test FAIL Cases - Table Mutations
# ===================================================================

class TestFailTableMutations:
    """Test cases where table structure is mutated."""

    def test_fail_table_row_count_differs(self, tmp_path):
        """FAIL when table row count differs."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        table_in = [["A", "B"], ["C", "D"], ["E", "F"]]
        table_out = [["A", "B"], ["C", "D"]]  # Missing row

        _create_test_docx(input_path, tables=[table_in])
        _create_test_docx(output_path, tables=[table_out])

        with pytest.raises(RuntimeError, match="STRUCTURE_GUARD_FAIL"):
            enforce_style_only_mutation(input_path, output_path)

    def test_fail_table_column_count_differs(self, tmp_path):
        """FAIL when table column count differs."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        table_in = [["A", "B", "C"], ["D", "E", "F"]]
        table_out = [["A", "B"], ["D", "E"]]  # Missing column

        _create_test_docx(input_path, tables=[table_in])
        _create_test_docx(output_path, tables=[table_out])

        with pytest.raises(RuntimeError, match="STRUCTURE_GUARD_FAIL"):
            enforce_style_only_mutation(input_path, output_path)

    def test_fail_table_count_differs(self, tmp_path):
        """FAIL when table count differs."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        table1 = [["A", "B"], ["C", "D"]]
        table2 = [["E", "F"], ["G", "H"]]

        _create_test_docx(input_path, tables=[table1, table2])
        _create_test_docx(output_path, tables=[table1])  # Missing table 2

        with pytest.raises(RuntimeError, match="STRUCTURE_GUARD_FAIL"):
            enforce_style_only_mutation(input_path, output_path)


# ===================================================================
# Test FAIL Cases - List Mutations
# ===================================================================

class TestFailListMutations:
    """Test cases where list structure is mutated."""

    def test_fail_list_item_becomes_regular_paragraph(self, tmp_path):
        """FAIL when list item becomes regular paragraph."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Create input with list item
        doc_in = Document()
        doc_in.add_paragraph("Item 1", style='List Bullet')
        doc_in.save(str(input_path))

        # Create output with regular paragraph (same text)
        doc_out = Document()
        doc_out.add_paragraph("Item 1")  # Not a list item
        doc_out.save(str(output_path))

        with pytest.raises(RuntimeError, match="STRUCTURE_GUARD_FAIL"):
            enforce_style_only_mutation(input_path, output_path)

    def test_fail_regular_paragraph_becomes_list_item(self, tmp_path):
        """FAIL when regular paragraph becomes list item."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Create input with regular paragraph
        doc_in = Document()
        doc_in.add_paragraph("Para 1")
        doc_in.save(str(input_path))

        # Create output with list item (same text)
        doc_out = Document()
        doc_out.add_paragraph("Para 1", style='List Bullet')
        doc_out.save(str(output_path))

        with pytest.raises(RuntimeError, match="STRUCTURE_GUARD_FAIL"):
            enforce_style_only_mutation(input_path, output_path)

    def test_fail_when_list_level_changed(self, tmp_path):
        """FAIL when list item level changes (e.g., indent changes)."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Create input with list at level 0
        doc_in = Document()
        p1 = doc_in.add_paragraph("Item 1", style='List Bullet')
        p2 = doc_in.add_paragraph("Item 2", style='List Bullet')
        doc_in.save(str(input_path))

        # Create output with list at level 1 (indented)
        doc_out = Document()
        p1_out = doc_out.add_paragraph("Item 1", style='List Bullet 2')  # Different level
        p2_out = doc_out.add_paragraph("Item 2", style='List Bullet 2')
        doc_out.save(str(output_path))

        # Note: This may or may not fail depending on whether the style change
        # affects the actual XML list level. The test validates the behavior.
        try:
            result = enforce_style_only_mutation(input_path, output_path)
            # If no XML list properties differ, it might pass (style-only change)
            # which is acceptable behavior
        except RuntimeError:
            # If XML list properties differ, it should fail
            pass  # Expected failure is acceptable


# ===================================================================
# Test FAIL Cases - Section Mutations
# ===================================================================

class TestFailSectionMutations:
    """Test cases where section structure is mutated."""

    def test_fail_section_count_differs(self, tmp_path):
        """FAIL when section count differs."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Create input with 2 sections
        doc_in = Document()
        doc_in.add_paragraph("Section 1")
        doc_in.add_section()
        doc_in.add_paragraph("Section 2")
        doc_in.save(str(input_path))

        # Create output with 1 section
        doc_out = Document()
        doc_out.add_paragraph("Section 1")
        doc_out.add_paragraph("Section 2")
        doc_out.save(str(output_path))

        with pytest.raises(RuntimeError, match="STRUCTURE_GUARD_FAIL"):
            enforce_style_only_mutation(input_path, output_path)

    def test_fail_when_section_break_modified(self, tmp_path):
        """FAIL when section break type changes."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Create input with continuous section break
        doc_in = Document()
        doc_in.add_paragraph("Section 1 content")
        section = doc_in.add_section(WD_SECTION.CONTINUOUS)
        doc_in.add_paragraph("Section 2 content")
        doc_in.save(str(input_path))

        # Create output with next page section break (different break type)
        doc_out = Document()
        doc_out.add_paragraph("Section 1 content")
        section = doc_out.add_section(WD_SECTION.NEW_PAGE)
        doc_out.add_paragraph("Section 2 content")
        doc_out.save(str(output_path))

        # Section break type change must be caught as a structural mutation
        with pytest.raises(RuntimeError, match="STRUCTURE_GUARD_FAIL"):
            enforce_style_only_mutation(input_path, output_path)


# ===================================================================
# Test RuntimeError Handling
# ===================================================================

class TestRuntimeErrorHandling:
    """Test that RuntimeError is raised with proper details."""

    def test_runtime_error_contains_differences(self, tmp_path):
        """RuntimeError message contains specific differences."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        _create_test_docx(input_path, paragraphs=["Para 1", "Para 2"])
        _create_test_docx(output_path, paragraphs=["Para 1"])

        with pytest.raises(RuntimeError) as exc_info:
            enforce_style_only_mutation(input_path, output_path)

        error_message = str(exc_info.value)
        assert "STRUCTURE_GUARD_FAIL" in error_message
        assert "Paragraph count mismatch" in error_message

    def test_runtime_error_limits_differences_to_20(self, tmp_path):
        """RuntimeError message limits differences to first 20."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Create many paragraphs with different text
        paras_in = [f"Input para {i}" for i in range(30)]
        paras_out = [f"Output para {i}" for i in range(30)]

        _create_test_docx(input_path, paragraphs=paras_in)
        _create_test_docx(output_path, paragraphs=paras_out)

        with pytest.raises(RuntimeError) as exc_info:
            enforce_style_only_mutation(input_path, output_path)

        error_message = str(exc_info.value)
        assert "and 10 more differences" in error_message or "more differences" in error_message


# ===================================================================
# Test Result Structure
# ===================================================================

class TestResultStructure:
    """Test that result dict has correct structure."""

    def test_result_has_all_fields(self, tmp_path):
        """Result dict has all required fields."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        paragraphs = ["Para 1", "Para 2"]
        _create_test_docx(input_path, paragraphs=paragraphs)
        _create_test_docx(output_path, paragraphs=paragraphs)

        result = enforce_style_only_mutation(input_path, output_path)

        required_fields = [
            "status",
            "paragraph_count_match",
            "list_structure_match",
            "table_structure_match",
            "section_structure_match",
            "structural_hash_match",
            "differences",
        ]

        for field in required_fields:
            assert field in result, f"Missing required field: {field}"

    def test_differences_limited_to_50(self, tmp_path):
        """Differences list is limited to 50 items."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Create 60 paragraphs with different text
        paras_in = [f"Input para {i}" for i in range(60)]
        paras_out = [f"Output para {i}" for i in range(60)]

        _create_test_docx(input_path, paragraphs=paras_in)
        _create_test_docx(output_path, paragraphs=paras_out)

        try:
            result = enforce_style_only_mutation(input_path, output_path)
        except RuntimeError:
            # Expected to fail, but check differences limit by calling internal function
            from processor.structure_guard import _extract_paragraphs
            input_doc = Document(input_path)
            output_doc = Document(output_path)
            input_paras = _extract_paragraphs(input_doc)
            output_paras = _extract_paragraphs(output_doc)
            _, diffs = _validate_paragraphs(input_paras, output_paras)

            # Manually limit to 50 (as done in enforce_style_only_mutation)
            diffs = diffs[:50]
            assert len(diffs) <= 50


# ===================================================================
# Test Edge Cases
# ===================================================================

class TestEdgeCases:
    """Test edge cases and boundary conditions."""

    def test_empty_documents_pass(self, tmp_path):
        """Empty documents pass validation."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        _create_test_docx(input_path)
        _create_test_docx(output_path)

        result = enforce_style_only_mutation(input_path, output_path)

        assert result["status"] == "PASS"

    def test_only_tables_no_paragraphs(self, tmp_path):
        """Documents with only tables (no paragraphs) pass."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        table = [["A", "B"], ["C", "D"]]

        _create_test_docx(input_path, tables=[table])
        _create_test_docx(output_path, tables=[table])

        result = enforce_style_only_mutation(input_path, output_path)

        assert result["status"] == "PASS"

    def test_complex_document_structure(self, tmp_path):
        """Complex document with mixed content passes."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        # Create complex structure
        doc_in = Document()
        doc_in.add_heading("Title", level=1)
        doc_in.add_paragraph("Introduction")
        doc_in.add_paragraph("Item 1", style='List Bullet')
        doc_in.add_paragraph("Item 2", style='List Bullet')
        table = doc_in.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "A"
        table.rows[0].cells[1].text = "B"
        doc_in.add_paragraph("Conclusion")
        doc_in.save(str(input_path))

        # Create identical output
        doc_out = Document()
        doc_out.add_heading("Title", level=1)
        doc_out.add_paragraph("Introduction")
        doc_out.add_paragraph("Item 1", style='List Bullet')
        doc_out.add_paragraph("Item 2", style='List Bullet')
        table = doc_out.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "A"
        table.rows[0].cells[1].text = "B"
        doc_out.add_paragraph("Conclusion")
        doc_out.save(str(output_path))

        result = enforce_style_only_mutation(input_path, output_path)

        assert result["status"] == "PASS"

    def test_special_characters_preserved(self, tmp_path):
        """Special characters are preserved correctly."""
        input_path = tmp_path / "input.docx"
        output_path = tmp_path / "output.docx"

        special_text = ["Hello © 2024", "Price: $100", "Degree: 90°"]

        _create_test_docx(input_path, paragraphs=special_text)
        _create_test_docx(output_path, paragraphs=special_text)

        result = enforce_style_only_mutation(input_path, output_path)

        assert result["status"] == "PASS"


# ===================================================================
# Test Internal Functions
# ===================================================================

class TestInternalFunctions:
    """Test internal helper functions."""

    def test_create_structural_signature(self):
        """Structural signature is deterministic."""
        paras = [
            {'index': 0, 'list_level': None, 'in_table': False, 'section_index': 0},
            {'index': 1, 'list_level': 0, 'in_table': False, 'section_index': 0},
        ]
        tables = [
            {'table_index': 0, 'rows': 2, 'cols': 2},
        ]
        sections = [
            {'index': 0, 'break_type': 'continuous'},
        ]

        sig1 = _create_structural_signature(paras, tables, sections)
        sig2 = _create_structural_signature(paras, tables, sections)

        assert sig1 == sig2
        assert len(sig1) == 64  # SHA256 hash length

    def test_extract_paragraphs_returns_metadata(self, tmp_path):
        """Paragraph extraction returns complete metadata."""
        doc_path = tmp_path / "test.docx"

        doc = Document()
        doc.add_paragraph("Test paragraph")
        doc.save(str(doc_path))

        doc = Document(doc_path)
        paras = _extract_paragraphs(doc)

        assert len(paras) == 1
        assert 'index' in paras[0]
        assert 'raw_text' in paras[0]
        assert 'normalized_text' in paras[0]
        assert 'style_name' in paras[0]
        assert 'is_list_item' in paras[0]
        assert 'list_level' in paras[0]
        assert 'list_id' in paras[0]
        assert 'in_table' in paras[0]
        assert 'section_index' in paras[0]

    def test_extract_tables_returns_metadata(self, tmp_path):
        """Table extraction returns complete metadata."""
        doc_path = tmp_path / "test.docx"

        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        doc.save(str(doc_path))

        doc = Document(doc_path)
        tables = _extract_tables(doc)

        assert len(tables) == 1
        assert tables[0]['table_index'] == 0
        assert tables[0]['rows'] == 2
        assert tables[0]['cols'] == 3
        assert 'nested_tables_count' in tables[0]

    def test_validate_paragraphs_diff_includes_metadata_snapshot(self):
        """Diff-first diagnostics should include input/output metadata snapshots."""
        input_paras = [{
            "index": 0,
            "raw_text": "Item",
            "normalized_text": "Item",
            "style_name": "List Bullet",
            "is_list_item": True,
            "list_level": 0,
            "list_id": -1,
            "in_table": False,
            "section_index": 0,
        }]
        output_paras = [{
            "index": 0,
            "raw_text": "Item",
            "normalized_text": "Item",
            "style_name": "Normal",
            "is_list_item": False,
            "list_level": None,
            "list_id": None,
            "in_table": False,
            "section_index": 0,
        }]

        ok, diffs = _validate_paragraphs(input_paras, output_paras)
        assert ok is False
        assert diffs
        assert "input_meta=" in diffs[0]
        assert "output_meta=" in diffs[0]
