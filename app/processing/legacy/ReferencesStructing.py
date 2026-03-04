#!/usr/bin/env python3
"""
Reference fixer (updated)

Major changes:
- Improved APA title parsing
- Prevent "Published YEAR" overwrites (detect and reject bad metadata)
- Keep original text when metadata is invalid; add a Word comment if possible,
  otherwise append an inline bracketed comment. Always log the event.
- Mix-source selection policy:
    * If DOI exists -> prefer API (PubMed/CrossRef) with lower threshold (0.60)
    * If no DOI -> prefer original unless API confidence >= 0.75
- Safer URL handling and improved fallbacks
- Wrapped low-level Word comment insertion in try/except; fallback to inline note
- Some minor refactors for readability
"""

import re
import requests
from pathlib import Path
from datetime import datetime
import logging
from requests.adapters import HTTPAdapter
from requests.exceptions import RequestException
from urllib3.util.retry import Retry
from xml.etree import ElementTree as ET
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from difflib import SequenceMatcher
from typing import Optional, Tuple, Dict, Any, List
import urllib.parse
import time
import uuid
import concurrent.futures
import threading
import html
from docx.shared import Pt
import difflib

import os
import sys

# Add the 'app' directory to sys.path so 'utils' can be imported
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..')))

# Import utils
try:
    from utils import track_changes
    TRACK_CHANGES_ENABLED = True
except ImportError:
    TRACK_CHANGES_ENABLED = False
    logging.warning("track_changes module not found. Track changes disabled.")

# -------------------------
# CONFIG
# -------------------------
# INPUT_DOCX = Path("...") # Removed hardcoded path

# Timeouts and parameters (tweakable)
CROSSREF_TIMEOUT = 12
PUBMED_TIMEOUT = 30
CROSSREF_ROWS = 6

# thresholds
# thresholds
SIMILARITY_MIN = 0.60      # lowered for better matching
PREF_DOI_THRESHOLD = 0.5   # prefer DOI source if similarity >= this

# -------------------------
# CACHE LOGIC
# -------------------------
import json
import atexit

from app.core.paths import REF_CACHE_PATH
REF_CACHE_FILE = REF_CACHE_PATH
REF_CACHE = {
    "crossref_doi": {},
    "crossref_search": {},
    "pubmed_search": {},
    "pubmed_fetch": {},
    "journal_abbrev": {}
}
CACHE_LOCK = threading.Lock()

# --- GOOGLE GEMINI CONFIG ---
# Load API Key from environment or hardcode if necessary
# We will check this dynamically inside the function to support late env setting

# convert_reference_with_gemini has been moved to gemini_ref_converter.py

def load_ref_cache():
    global REF_CACHE
    if REF_CACHE_FILE.exists():
        try:
            with open(REF_CACHE_FILE, 'r', encoding='utf-8') as f:
                loaded = json.load(f)
                # merge to ensure all keys exist
                for k in REF_CACHE:
                    if k in loaded:
                        REF_CACHE[k] = loaded[k]
            print(f"[Info] Loaded ref_cache.json with {sum(len(v) for v in REF_CACHE.values())} entries.")
        except Exception as e:
            print(f"[Warning] Failed to load ref_cache.json: {e}")

def save_ref_cache():
    try:
        with open(REF_CACHE_FILE, 'w', encoding='utf-8') as f:
            json.dump(REF_CACHE, f, indent=2)
        print("[Info] Saved ref_cache.json.")
    except Exception as e:
        print(f"[Warning] Failed to save ref_cache.json: {e}")

# Load cache on startup
load_ref_cache()
# Save on exit
atexit.register(save_ref_cache)

# NCBI E-utilities base
NCBI_BASE = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"

# Logging
logger = logging.getLogger(__name__)
if not logger.handlers:
    logging.basicConfig(level=logging.INFO, format='%(asctime)s %(levelname)s: %(message)s')

# -------------------------
# HTTP session with retries
# -------------------------
def get_requests_session() -> requests.Session:
    session = requests.Session()
    retries = Retry(total=5, backoff_factor=1.0,
                    status_forcelist=(429, 500, 502, 503, 504),
                    allowed_methods=frozenset(['GET', 'POST', 'HEAD']))
    adapter = HTTPAdapter(max_retries=retries, pool_connections=50, pool_maxsize=50)
    session.mount('https://', adapter)
    session.mount('http://', adapter)
    session.headers.update({'User-Agent': 'refboth/1.0 (+https://example.org)'})
    return session

SESSION = get_requests_session()

# -------------------------
# Utility helpers
# -------------------------
doi_regex = re.compile(r"10\.\d{4,9}/[-._;()/:A-Za-z0-9]+", re.IGNORECASE)
doi_full_regex = re.compile(r"^10\.\d{4,9}/[-._;()/:A-Za-z0-9]+$", re.IGNORECASE)

def normalize_whitespace(s: Optional[str]) -> str:
    if not s:
        return ""
    return re.sub(r'\s+', ' ', s).strip()

def extract_doi_from_text(s: str) -> Optional[str]:
    m = doi_regex.search(s)
    return m.group(0) if m else None

def is_valid_doi(s: str) -> bool:
    return bool(doi_full_regex.match(s.strip()))

def is_url(s: str) -> bool:
    try:
        p = urllib.parse.urlparse(s)
        return bool(p.scheme and p.netloc)
    except:
        return False

def validate_url(url: str, timeout: int = 5) -> bool:
    if not url:
        return False
    try:
        r = SESSION.head(url, timeout=timeout, allow_redirects=True)
        if r.status_code == 405:
            r = SESSION.get(url, timeout=timeout, stream=True)
            r.close()
        return 200 <= r.status_code < 400
    except (RequestException, Exception):
        # On error, assume valid to avoid blocking
        return True

def similarity(a: str, b: str) -> float:
    if not a or not b:
        return 0.0
    return SequenceMatcher(None, normalize_whitespace(a.lower()), normalize_whitespace(b.lower())).ratio()

# Ported from Referencenumvalidation.py
def find_duplicates(ref_objects):
        """
        Finds duplicate references using fuzzy matching (difflib).
        Returns a list of dicts: {'id': int, 'text': str, 'duplicate_of': int, 'score': float}
        """
        import difflib
        
        duplicates = []
        processed_refs = [] # list of (id, clean_text)
        
        # 1. Pre-process all candidates
        for obj in ref_objects:
            # ReferencesStructing handles 'ref_objects' differently?
            # If this is called within a class that has 'ref_objects', great.
            # If standalone, we need to know what 'obj' looks like.
            # Assuming obj has 'para' attribute like in Referencenumvalidation or simple text
            if hasattr(obj, 'para'):
                 full_text = obj.para.text.strip()
            elif isinstance(obj, dict) and 'para' in obj:
                 full_text = obj['para'].text.strip()
            elif isinstance(obj, str):
                 full_text = obj # Just string list
            else:
                 continue

            # Remove leading numbering like "1. ", "[1] "
            clean_text = re.sub(r'^\[?\d+\]?[\.\s]*', '', full_text)
            
            # Use index or explicit ID
            rid = obj.get('id') if isinstance(obj, dict) else len(processed_refs) + 1
            processed_refs.append({'id': rid, 'text': clean_text})
            
        # 2. Compare O(N^2)
        n = len(processed_refs)
        for i in range(n):
            ref_a = processed_refs[i]
            for j in range(i + 1, n):
                ref_b = processed_refs[j]
                
                len_a = len(ref_a['text'])
                len_b = len(ref_b['text'])
                if len_a == 0 or len_b == 0: continue
                    
                if min(len_a, len_b) / max(len_a, len_b) < 0.6:
                    continue
                    
                ratio = difflib.SequenceMatcher(None, ref_a['text'], ref_b['text']).ratio()
                
                if ratio > 0.85:
                    duplicates.append({
                        'id': ref_b['id'], 
                        'text': ref_b['text'][:100] + "...",
                        'duplicate_of': ref_a['id'],
                        'score': round(ratio * 100, 1)
                    })
                    
        return duplicates

# -------------------------
# CrossRef helpers
# -------------------------
def crossref_get_by_doi(doi: str) -> Optional[Dict[str, Any]]:
    # CACHE CHECK
    if doi in REF_CACHE['crossref_doi']:
        return REF_CACHE['crossref_doi'][doi]

    try:
        url = f"https://api.crossref.org/works/{requests.utils.requote_uri(doi)}"
        r = SESSION.get(url, timeout=CROSSREF_TIMEOUT)
        if r.status_code == 200:
            data = r.json().get('message')
            if data:
                with CACHE_LOCK:
                    REF_CACHE['crossref_doi'][doi] = data
            return data
    except RequestException:
        logger.debug("CrossRef DOI fetch failed for %s", doi)
        return None
    return None

def crossref_search(title: str, journal: Optional[str] = None, year: Optional[str] = None, rows: int = CROSSREF_ROWS) -> List[Dict[str, Any]]:
    # Cache Key Construction
    k_year = f"({year})" if year else ""
    k_jnl = str(journal).strip().lower() if journal else "none"
    k_rows = str(rows)
    # Key: (Year)|Title|Journal|Rows
    cache_key = f"{k_year}|{normalize_whitespace(title).lower()}|{k_jnl}|{k_rows}"
    
    if cache_key in REF_CACHE['crossref_search']:
        return REF_CACHE['crossref_search'][cache_key]

    params = {'query.title': title, 'rows': rows}
    if journal:
        params['query.container-title'] = journal
    if year:
        params['filter'] = f'from-pub-date:{year},until-pub-date:{year}'
    try:
        r = SESSION.get("https://api.crossref.org/works", params=params, timeout=CROSSREF_TIMEOUT)
        r.raise_for_status()
        msg = r.json().get('message', {})
        results = msg.get('items', []) or []
        
        # Update Cache
        with CACHE_LOCK:
            REF_CACHE['crossref_search'][cache_key] = results
        return results
    except RequestException:
        logger.debug("CrossRef search failed for title: %s", title)
        return []

def crossref_pick_best(title: str, candidates: List[Dict[str, Any]]) -> Tuple[Optional[Dict[str, Any]], float]:
    tnorm = normalize_whitespace(title).lower()
    best = None
    best_score = 0.0
    for item in candidates:
        ititle = (item.get('title') or [''])[0]
        sc = similarity(tnorm, ititle or '')
        if sc > best_score:
            best_score = sc
            best = item
        if normalize_whitespace(ititle).lower() == tnorm:
            return item, 1.0
    return best, best_score

# -------------------------
# PubMed helpers
# -------------------------
def pubmed_search_ids(title: str, journal: Optional[str] = None, year: Optional[str] = None, max_results: int = 5) -> List[str]:
    # Cache Key
    k_year = f"({year})" if year else ""
    k_jnl = str(journal).strip().lower() if journal else "none"
    k_max = str(max_results)
    cache_key = f"{k_year}|{normalize_whitespace(title).lower()}|{k_jnl}|{k_max}"

    if cache_key in REF_CACHE['pubmed_search']:
        return REF_CACHE['pubmed_search'][cache_key]

    results = set()
    def truncate_title(t: str, max_words: int = 10) -> str:
        words = t.split()
        if len(words) > max_words:
            return ' '.join(words[:max_words])
        return t

    # Strategy 1: Full query with title, journal, year
    if journal and year:
        q = f'{title}[ti] AND {journal}[ta] AND {year}[dp]'
        params = {'db': 'pubmed', 'term': q, 'retmax': max_results, 'retmode': 'json'}
        try:
            r = SESSION.get(f"{NCBI_BASE}/esearch.fcgi", params=params, timeout=PUBMED_TIMEOUT)
            r.raise_for_status()
            ids = r.json().get('esearchresult', {}).get('idlist', []) or []
            results.update(ids)
        except RequestException:
            logger.debug("PubMed esearch failed (journal+year) for: %s", title)

    # Strategy 2: Title + year
    if year and len(results) < max_results:
        q = f'{title}[ti] AND {year}[dp]'
        params = {'db': 'pubmed', 'term': q, 'retmax': max_results, 'retmode': 'json'}
        try:
            r = SESSION.get(f"{NCBI_BASE}/esearch.fcgi", params=params, timeout=PUBMED_TIMEOUT)
            r.raise_for_status()
            ids = r.json().get('esearchresult', {}).get('idlist', []) or []
            results.update(ids)
        except RequestException:
            logger.debug("PubMed esearch failed (title+year) for: %s", title)

    # Strategy 3: Just title
    if len(results) < max_results:
        q = f'{title}[ti]'
        params = {'db': 'pubmed', 'term': q, 'retmax': max_results * 2, 'retmode': 'json'}
        try:
            r = SESSION.get(f"{NCBI_BASE}/esearch.fcgi", params=params, timeout=PUBMED_TIMEOUT)
            r.raise_for_status()
            ids = r.json().get('esearchresult', {}).get('idlist', []) or []
            results.update(ids)
        except RequestException:
            logger.debug("PubMed esearch failed (title only) for: %s", title)

    # Strategy 4: Truncated title if long (previously dead code — now active)
    title_words = title.split()
    if len(results) < max_results and len(title_words) > 15:
        short_title = truncate_title(title, 10)
        q = f'{short_title}[ti]'
        if year:
            q += f' AND {year}[dp]'
        params = {'db': 'pubmed', 'term': q, 'retmax': max_results * 2, 'retmode': 'json'}
        try:
            r = SESSION.get(f"{NCBI_BASE}/esearch.fcgi", params=params, timeout=PUBMED_TIMEOUT)
            r.raise_for_status()
            ids = r.json().get('esearchresult', {}).get('idlist', []) or []
            results.update(ids)
        except RequestException:
            logger.debug("PubMed esearch failed (truncated title) for: %s", short_title)

    # Strategy 5: keyword fallback (previously dead code — now active)
    if len(results) < max_results:
        significant_words = [
            w for w in title.split()
            if len(w) > 4 and w.lower() not in (
                'that', 'with', 'from', 'have', 'this', 'their', 'which', 'viral', 'virus'
            )
        ]
        if significant_words:
            key_phrase = ' '.join(significant_words[:6])
            q = f'{key_phrase}[ti]'
            if year:
                q += f' AND {year}[dp]'
            params = {'db': 'pubmed', 'term': q, 'retmax': max_results * 2, 'retmode': 'json'}
            try:
                r = SESSION.get(f"{NCBI_BASE}/esearch.fcgi", params=params, timeout=PUBMED_TIMEOUT)
                r.raise_for_status()
                ids = r.json().get('esearchresult', {}).get('idlist', []) or []
                results.update(ids)
            except RequestException:
                logger.debug("PubMed keyword fallback search failed for: %s", key_phrase)

    # Cache and return — limit to max_results (not max_results*2 which was a bug)
    final_ids = list(results)[:max_results]
    with CACHE_LOCK:
        REF_CACHE['pubmed_search'][cache_key] = final_ids
    return final_ids

def pubmed_fetch_xml(pubmed_id: str) -> Optional[ET.Element]:
    if pubmed_id in REF_CACHE['pubmed_fetch']:
        try:
            return ET.fromstring(REF_CACHE['pubmed_fetch'][pubmed_id])
        except Exception:
            pass

    params = {'db': 'pubmed', 'id': pubmed_id, 'retmode': 'xml'}
    try:
        r = SESSION.get(f"{NCBI_BASE}/efetch.fcgi", params=params, timeout=PUBMED_TIMEOUT)
        r.raise_for_status()
        # Verify XML before caching
        root = ET.fromstring(r.text)
        with CACHE_LOCK:
            REF_CACHE['pubmed_fetch'][pubmed_id] = r.text
        return root
    except RequestException:
        logger.debug("PubMed efetch failed for id: %s", pubmed_id)
        return None
    except ET.ParseError:
        logger.debug("Failed to parse PubMed XML for id: %s", pubmed_id)
        return None

def search_google_books(query: str, author: Optional[str] = None, api_key: Optional[str] = None) -> Optional[Dict[str, Any]]:
    base_url = "https://www.googleapis.com/books/v1/volumes"
    params = {'q': query, 'maxResults': 1, 'printType': 'books'}
    if author:
        params['q'] += f"+inauthor:{author}"
    if api_key:
        params['key'] = api_key
    try:
        resp = SESSION.get(base_url, params=params, timeout=10)
        if resp.status_code == 200:
            data = resp.json()
            if 'items' in data and len(data['items']) > 0:
                book = data['items'][0]['volumeInfo']
                parts = book.get('authors', [])
                authors_list = []
                for a in parts:
                   a_parts = a.split()
                   if len(a_parts)>1: authors_list.append({'given':" ".join(a_parts[:-1]), 'family':a_parts[-1]})
                   else: authors_list.append({'literal':a})
                   
                return {
                    'title': [book.get('title')],
                    'subtitle': [book.get('subtitle','')],
                    'author': authors_list,
                    'published-print': {'date-parts': [[book.get('publishedDate')[:4]]]} if book.get('publishedDate') else {},
                    'publisher': book.get('publisher'),
                    'type': 'book',
                    'URL': book.get('infoLink',''),
                    'DOI': ''
                }
    except Exception:
        pass
    return None

def expand_abbreviated_page_range(pages: str) -> str:
    """
    Expand abbreviated PubMed page ranges to full format.
    PubMed uses MedlinePgn which abbreviates ranges by showing only differing digits.
    E.g., "863-8" should be "863-868", "2259-66" should be "2259-2266"
    """
    if not pages or '-' not in pages:
        return pages
    
    try:
        parts = pages.split('-')
        if len(parts) != 2:
            return pages
        
        first_page = parts[0].strip()
        last_page_abbrev = parts[1].strip()
        
        # If last page is longer than or equal to first page, it's likely already full
        if len(last_page_abbrev) >= len(first_page):
            return pages
        
        # Expand abbreviated last page
        # Take the necessary prefix from first page and append the abbreviated last page
        # E.g., first="863", abbrev="8" -> take "86" from first + "8" = "868"
        prefix_len = len(first_page) - len(last_page_abbrev)
        if prefix_len > 0:
            expanded_last_page = first_page[:prefix_len] + last_page_abbrev
            return f"{first_page}-{expanded_last_page}"
    except Exception:
        pass
    
    return pages

def pubmed_parse_article_from_xml(root: ET.Element) -> Optional[Dict[str, Any]]:
    pa = root.find('.//PubmedArticle')
    if pa is None:
        return None

    article_title_el = pa.find('.//ArticleTitle')
    if article_title_el is None:
        return None
    title = ''.join(article_title_el.itertext()).strip()

    journal_el = pa.find('.//Journal/Title')
    journal = journal_el.text.strip() if journal_el is not None and journal_el.text else ''
    
    # Extract ISO Abbreviation
    iso_abbrev_el = pa.find('.//Journal/ISOAbbreviation')
    iso_abbrev = iso_abbrev_el.text.strip() if iso_abbrev_el is not None and iso_abbrev_el.text else ''

    volume_el = pa.find('.//JournalIssue/Volume')
    issue_el = pa.find('.//JournalIssue/Issue')
    pages_el = pa.find('.//Pagination/MedlinePgn')

    volume = volume_el.text.strip() if volume_el is not None and volume_el.text else ''
    issue = issue_el.text.strip() if issue_el is not None and issue_el.text else ''
    pages = pages_el.text.strip() if pages_el is not None and pages_el.text else ''
    
    # EXPAND abbreviated PubMed page ranges
    pages = expand_abbreviated_page_range(pages)
    
    # DEBUG: Log PubMed page extraction
    if pages:
        logger.debug(f"PubMed - Extracted pages: '{pages}' (Length: {len(pages)})")

    year = None
    pubdate_year = pa.find('.//Journal/JournalIssue/PubDate/Year')
    if pubdate_year is not None and pubdate_year.text:
        year = pubdate_year.text.strip()
    else:
        artdate_year = pa.find('.//ArticleDate/Year')
        if artdate_year is not None and artdate_year.text:
            year = artdate_year.text.strip()
        else:
            medline_date = pa.find('.//Journal/JournalIssue/PubDate/MedlineDate')
            if medline_date is not None and medline_date.text:
                m = re.search(r'\b(19|20)\d{2}\b', medline_date.text)
                year = m.group(0) if m else None

    authors_list = []
    for author_el in pa.findall('.//AuthorList/Author'):
        last = author_el.find('LastName')
        fore = author_el.find('ForeName')
        initial = author_el.find('Initials')
        if last is not None:
            family = last.text.strip() if last.text else ''
            given = ''
            if fore is not None and fore.text:
                given = fore.text.strip()
            elif initial is not None and initial.text:
                given = initial.text.strip()
            authors_list.append({'given': given, 'family': family})

    doi = None
    for aid in pa.findall('.//ArticleIdList/ArticleId'):
        if aid.attrib.get('IdType', '').lower() == 'doi' and aid.text:
            doi = aid.text.strip()
            break

    unified = {
        'author': authors_list,
        'title': [title],
        'container-title': [journal],
        'iso_abbrev': iso_abbrev,
        'volume': volume or '',
        'issue': issue or '',
        'page': pages or '',
        'DOI': doi or '',
        'created': {'date-parts': [[int(year) if year and year.isdigit() else None]]}
    }
    return unified

# -------------------------
# Parsing heuristics
# -------------------------
def parse_authors_to_list(author_str: str) -> List[Dict[str, str]]:
    """Parse an author string into a list of {'family': ..., 'given': ...} dicts.
    Handles APA format: Surname, I. I., & Surname, I. I.
    Also handles AMA format: Surname II, Surname II
    And organizational authors: American Nurses Association
    """
    if not author_str: return []
    author_str = author_str.strip().rstrip('.')
    
    # Organizational authors (no commas separating surname/initials)
    if any(k in author_str.lower() for k in ('association', 'institute', 'organization', 'society', 'academy', 'university', 'nurses', 'center', 'committee')):
        return [{'family': author_str.strip()}]
    
    res = []
    
    # APA format: "Surname, I. I., & Surname, I. I."
    # Step 1: Split by " & " or ", &" to get individual author groups
    author_groups = re.split(r'\s*&\s*', author_str)
    
    for group in author_groups:
        group = group.strip().rstrip(',')
        if not group:
            continue
        
        # Try to match "Surname, I. I." pattern (APA)
        # E.g. "Ayers, S." or "Barrett, L. F." or "Arseneault, L., Cannon, M., Fisher, H. L."
        # This group may contain multiple comma-separated authors
        
        # Check if this looks like multiple authors separated by commas
        # Pattern: each author is "Surname, Initials" 
        # We look for: word(s), single-letter-dots pattern
        apa_authors = re.findall(
            r'([A-Z][a-z]+(?:[\-\s]+[A-Z][a-z]+)*)\s*,\s*([A-Z]\.?(?:\s*[A-Z]\.?)*)',
            group
        )
        
        if apa_authors:
            for surname, initials in apa_authors:
                res.append({'family': surname.strip(), 'given': initials.strip()})
        else:
            # AMA format or single name: "Smith JB" or just "Smith"
            # Try to match "Surname Initials" where initials are uppercase letters
            m = re.match(r'^(.+?)\s+([A-Z]{1,4})$', group.strip())
            if m:
                res.append({'family': m.group(1).strip(), 'given': m.group(2).strip()})
            elif group.strip():
                # Just treat the whole thing as a family name
                res.append({'family': group.strip()})
    
    return res

def parse_ama_reference_raw(raw: str) -> Dict[str, Any]:
    s = normalize_whitespace(raw)
    s = re.sub(r'^\d+\.\s*', '', s)
    parts = [p.strip() for p in re.split(r'\.\s+', s) if p.strip()]
    authors = parts[0] if len(parts) > 0 else ''
    title = parts[1] if len(parts) > 1 else ''
    journal = parts[2] if len(parts) > 2 else ''
    year_match = re.search(r'\b(19|20)\d{2}\b', raw)
    year = year_match.group(0) if year_match else None
    journal = journal.rstrip('. ')
    return {
        'authors': authors, 
        'author': parse_authors_to_list(authors),
        'title': [title], 
        'container-title': [journal], 
        'year': year,
        'type': 'journal-article'
    }

def parse_apa_reference_raw(raw: str) -> Dict[str, Any]:
    s = normalize_whitespace(raw)
    m = re.match(
        r'^(?P<authors>.+?)\s*\((?P<year>\d{4})\)\.\s*'
        r'(?P<title>.+?)\.\s*'
        r'(?P<journal>[^,\.]+)',
        s
    )
    if m:
        auths = m.group('authors').strip()
        return {
            'authors': auths,
            'author': parse_authors_to_list(auths),
            'year': m.group('year').strip(),
            'title': [m.group('title').strip()],
            'container-title': [m.group('journal').strip()],
            'type': 'journal-article'
        }
    parts = [p.strip() for p in s.split('.') if p.strip()]
    year_match = re.search(r'\b(19|20)\d{2}\b', s)
    auths = parts[0] if len(parts) > 0 else ''
    return {
        'authors': auths,
        'author': parse_authors_to_list(auths),
        'title': [parts[1] if len(parts) > 1 else ''],
        'container-title': [parts[2] if len(parts) > 2 else ''],
        'year': year_match.group(0) if year_match else None,
        'type': 'journal-article'
    }

# -------------------------
# Unified selection logic
# -------------------------
def pick_best_between_pubmed_crossref(title: str,
                                     pubmed_items: List[Dict[str, Any]],
                                     crossref_items: List[Dict[str, Any]]) -> Tuple[Optional[Dict[str, Any]], str, float]:
    candidates = []
    for cr in crossref_items:
        ititle = (cr.get('title') or [''])[0]
        doi = cr.get('DOI') or ''
        candidates.append(('crossref', cr, ititle, doi))

    for pm in pubmed_items:
        ititle = (pm.get('title') or [''])[0]
        doi = pm.get('DOI') or ''
        candidates.append(('pubmed', pm, ititle, doi))

    if not candidates:
        return None, '', 0.0

    tnorm = normalize_whitespace(title).lower()
    best_score = 0.0
    best = None
    best_source = ''
    for src, obj, ititle, doi in candidates:
        sc = similarity(tnorm, ititle or '')
        if doi:
            sc += 0.06
        if sc > best_score:
            best_score = sc
            best = (src, obj)
            best_source = src

    if best_score < SIMILARITY_MIN:
        return None, '', best_score

    # If pubmed chosen but crossref has DOI and comparable score, prefer crossref
    if best_source == 'pubmed' and crossref_items:
        cr_best, cr_best_score = None, 0.0
        for cr in crossref_items:
            sc = similarity(tnorm, (cr.get('title') or [''])[0] or '')
            if sc > cr_best_score:
                cr_best_score = sc
                cr_best = cr
        if cr_best and cr_best_score + 0.02 >= best_score and cr_best.get('DOI'):
            return cr_best, 'crossref', cr_best_score

    if best_source == 'crossref':
        return best[1], 'crossref', best_score
    else:
        return best[1], 'pubmed', best_score

def pubmed_to_crossref_like(pubmed_item: Dict[str, Any]) -> Dict[str, Any]:
    page_val = pubmed_item.get('page', '')
    # DEBUG: Log page data transformation
    if page_val:
        logger.debug(f"PubMed->CrossRef - Pages: '{page_val}' (Length: {len(page_val)})")
    cr_like = {
        'author': pubmed_item.get('author', []),
        'title': pubmed_item.get('title', []),
        'container-title': pubmed_item.get('container-title', []),
        'short-container-title': [pubmed_item.get('iso_abbrev')] if pubmed_item.get('iso_abbrev') else [],
        'volume': pubmed_item.get('volume', ''),
        'issue': pubmed_item.get('issue', ''),
        'page': page_val,
        'DOI': pubmed_item.get('DOI', '') or None,
        'created': pubmed_item.get('created', {})
    }
    return cr_like

# -------------------------
# Citation generators (APA/AMA) - kept as in original but with slight safety checks
# -------------------------
def format_authors_apa(authors: List[Dict[str, str]]) -> str:
    if not authors:
        return "Unknown authors"
    def initial(given):
        return (given[0] + '.') if given else ''
    if len(authors) == 1:
        a = authors[0]; return f"{a.get('family','')}, {initial(a.get('given',''))}".strip()
    if len(authors) == 2:
        a1,a2 = authors[0], authors[1]
        return f"{a1.get('family','')}, {initial(a1.get('given',''))}, & {a2.get('family','')}, {initial(a2.get('given',''))}"
    parts = []
    for i,a in enumerate(authors):
        node = f"{a.get('family','')}, {initial(a.get('given',''))}".strip()
        if i == len(authors)-1:
            parts.append(f"& {node}")
        else:
            parts.append(node)
    return ", ".join(parts)

def extract_initials(given_name: str) -> str:
    """Extract all initials from a given name string (e.g. 'John B.' -> 'JB')."""
    if not given_name:
        return ""
    # Filter for uppercase letters that are likely initials
    # or just take the first letter of each part separated by space/hyphen
    parts = re.split(r'[ \-]', given_name)
    initials = []
    for p in parts:
        if p and p[0].isalpha():
            initials.append(p[0].upper())
    return "".join(initials)

def format_authors_ama(authors: List[Dict[str, str]]) -> str:
    if not authors:
        return "Unknown authors"
    
    # AMA 11th Edition Rule:
    # If <= 6 authors, list all.
    # If > 6 authors, list first 3 followed by "et al."
    
    if len(authors) <= 6:
        subset = authors
        suffix = ""
    else:
        subset = authors[:3]
        suffix = ", et al."
    
    formatted_list = []
    for a in subset:
        family = a.get('family', '').strip()
        given = a.get('given', '').strip()
        initials = extract_initials(given)
        # AMA style: "Family I" (no dots usually, but user asked for standard? Standard AMA is no dots, tight)
        # However user example: "Assayag, E., ..." -> This looks like APA style actually?
        # User REQ: "Assayag, E., ... & Hallevi, H." 
        # WAIT: The User provided output shows "Assayag, E., ... & Hallevi, H."
        # That is APA style structure (comma, & before last).
        # AMA style is "Assayag EB, Tene O, ..." (No 'and', no dots).
        # But the function name is `format_authors_ama`.
        # Let's check usage. ReferenceStructing uses `generate_ama_citation` -> `format_authors_ama`?
        # NO, `generate_ama_citation` manually builds segments. `format_authors_ama` seems unused or helper?
        # A grep showed `generate_ama_citation` calls manual building.
        # Let's fix `generate_ama_citation` logic block instead, AND `generate_apa_citation`.
        
        # We will keep `extract_initials` helper and use it in the main generation functions.
        pass
    return "" # This function seems to be unused by main logic, we will check `generate_ama_citation`

def generate_apa_citation(item: Dict[str, Any]) -> List[Tuple[str, Optional[str]]]:
    ctype = item.get('type', 'journal-article')
    # Year extraction priority: published-print > published-online > issued > year > created
    # 'created' is the CrossRef record creation date, NOT publication date (caused 1996→2003 errors)
    date_parts = (
        item.get('published-print', {}).get('date-parts') or
        item.get('published-online', {}).get('date-parts') or
        item.get('issued', {}).get('date-parts') or
        [[None]]
    )
    year = str(date_parts[0][0]) if date_parts[0][0] else (item.get('year') or 'n.d.')
    # Last resort: use 'created' only if nothing else worked and year is still 'n.d.'
    if year == 'n.d.' and item.get('created', {}).get('date-parts'):
        cd = item['created']['date-parts']
        if cd and cd[0] and cd[0][0]:
            year = str(cd[0][0])
    
    # Clean title
    def clean_text(t):
        if not t: return t
        return html.unescape(t).replace('&amp;', '&')

    title = clean_text((item.get('title') or ['No title available'])[0])
    container = (item.get('container-title') or [''])[0]
    container = clean_text(container)
    
    volume = item.get('volume', '')
    issue = item.get('issue', '')
    pages = item.get('page', '')
    # DEBUG: Log page data for troubleshooting
    if pages:
        logger.debug(f"APA Citation - Pages from API: '{pages}' (Length: {len(pages)})")
    doi = item.get('DOI', '')
    url = item.get('URL', '')

    segments = []
    
    # --- Author Formatting (APA 7th) ---
    # Rule: <= 20 authors: list all. > 20: list first 19 ... last.
    # Format: Surname, I. I., & Surname, I. I. (Space between initials)
    authors_list = item.get('author', [])
    
    def format_apa_author(auth_obj):
        family = auth_obj.get('family', '').strip()
        given = auth_obj.get('given', '').strip()
        if not family: return ""
        
        s = family
        if given:
            # Initials with dots and spaces: "John B." -> "J. B."
            parts = [p[0].upper() + '.' for p in re.split(r'[ \-]', given) if p and p[0].isalpha()]
            if parts:
                s += ", " + " ".join(parts)
        return s

    if not authors_list:
        segments.append(("Unknown authors", 'bib_unpubl'))
    else:
        # Pre-format all authors
        formatted_authors = [format_apa_author(a) for a in authors_list if a.get('family')]
        count = len(formatted_authors)
        
        if count == 0:
             segments.append(("Unknown authors", 'bib_unpubl'))
        elif count <= 20:
            for i, auth_str in enumerate(formatted_authors):
                if i > 0:
                    segments.append((", ", None))
                    if i == count - 1:
                        segments.append(("& ", None))
                segments.append((auth_str, 'bib_surname')) # Keeping single style for simplicity or split? using bib_surname for whole author string for APA usually ok or split for First/Last? 
                # Original used bib_fname.
                # Let's split for semantic styling if possible, but the string is pre-built.
                # Actually, let's use the object to be granular if required.
                # But for strict punctuation control, string building is safer.
                # Let's revert to granular for styles:
        else:
             # > 20: First 19, ellipsis, last
             for i in range(19):
                 if i > 0: segments.append((", ", None))
                 segments.append((formatted_authors[i], 'bib_surname'))
             segments.append((", ... ", None))
             segments.append((formatted_authors[-1], 'bib_surname'))
             
    # Re-do Granular Loop for <= 20 to ensure styles apply to Surname vs Initials? 
    # User requirement is "Strict Punctuation". Granular styles might interfere if we aren't careful.
    # Let's stick to the granular loop logic but with strict punctuation.
    segments = []
    
    def add_granular_apa_author(a):
        fam = a.get('family', '').strip()
        giv = a.get('given', '').strip()
        if fam:
            segments.append((fam, 'bib_surname'))
            if giv:
                segments.append((", ", None))
                # Initials with dots and spaces
                parts = [p[0].upper() + '.' for p in re.split(r'[ \-]', giv) if p and p[0].isalpha()]
                if parts:
                    segments.append((" ".join(parts), 'bib_fname'))

    if not authors_list:
        segments.append(("Unknown authors", 'bib_unpubl'))
    else:
        count = len(authors_list)
        if count <= 20:
             for i, a in enumerate(authors_list):
                 if i > 0:
                     segments.append((", ", None))
                     if i == count - 1:
                         segments.append(("& ", None))
                 add_granular_apa_author(a)
        else:
             for i in range(19):
                 if i > 0: segments.append((", ", None))
                 add_granular_apa_author(authors_list[i])
             segments.append((", ... ", None))
             add_granular_apa_author(authors_list[-1])

    segments.append((" (", None))
    segments.append((year, 'bib_year')) 
    segments.append(("). ", None))
    
    # Title styling
    # APA: Article title not italicized.    # Title styling based on type
    if ctype == 'book':
         segments.append((title, 'bib_book'))
    elif ctype == 'web':
         segments.append((title, 'bib_title'))
    elif ctype == 'book-chapter':
         segments.append((title, 'bib_chaptertitle'))
    elif ctype == 'thesis':
         segments.append((title, 'bib_book')) # Thesis title italic in APA
    elif ctype in ('conference-paper', 'proceedings-article'):
         segments.append((title, 'bib_confpaper'))
    else:
         segments.append((title, 'bib_article')) # Not italic (Journal article)
    
    segments.append((". ", None))

    # Container / Publisher
    if ctype == 'book':
        publisher = item.get('publisher', '')
        if publisher:
            segments.append((publisher, 'bib_publisher'))
            segments.append((".", None))
            
    elif ctype in ('proceedings-article', 'conference-paper', 'book-chapter'):
        if container:
            segments.append(("In ", None))
            # Editor info usually goes here for APA "In A. Editor (Ed.), Book..."
            # For now, just container.
            segments.append((container, 'bib_book')) # Italic
            if pages:
                 segments.append((" (pp. ", None))
                 segments.append((pages.replace('-', '–'), 'bib_fpage')) # En-dash
                 segments.append((")", None))
            segments.append((".", None))
        if item.get('publisher'):
            segments.append((" ", None))
            segments.append((item.get('publisher'), 'bib_publisher'))
            segments.append((".", None))
            
    elif ctype == 'web':
        if container:
            segments.append((container, 'bib_journal')) # Site name
            segments.append((". ", None))
        if url:
            segments.append((url, 'bib_url'))
            
    elif ctype == 'thesis':
        # APA Thesis: Author (Year). Title (Publication No.) [Doctoral dissertation, Institution]. Database.
        # Minimal: Title [Doctoral dissertation, Institution].
        
        # Determine genre (Doctoral dissertation / Master's thesis)
        genre = item.get('genre', 'Doctoral dissertation')
        publisher = item.get('publisher', '') # Institution
        
        segments.append((" [", None))
        segments.append((genre, 'bib_genre'))
        if publisher:
             segments.append((", ", None))
             segments.append((publisher, 'bib_publisher'))
        segments.append(("].", None))
        
        # Archive / Database? usually container-title
        if container:
             segments.append((" ", None))
             segments.append((container, 'bib_journal'))
             segments.append((".", None))
        if url:
             segments.append((" ", None))
             segments.append((url, 'bib_url'))
             
    else:
        # Journal Article
        if container:
            segments.append((container, 'bib_journal')) # Italic
        if volume:
            segments.append((", ", None))
            segments.append((volume, 'bib_volume')) # Italic
        if issue:
            segments.append(("(", None))
            segments.append((issue, 'bib_issue')) # No Italic
            segments.append((")", None))
        # NOTE: Full page range is preserved (not truncated) - APA style uses complete page numbers
        if pages:
            segments.append((", ", None))
            # Preserve full page range (e.g., 123-456) with en-dash conversion
            segments.append((pages.replace('-', '–'), 'bib_fpage')) # En-dash
        segments.append((".", None))

    if doi:
        segments.append((" ", None))
        # strictly https://doi.org/
        clean_doi = doi.replace('http://dx.doi.org/', '').replace('https://doi.org/', '')
        segments.append(("https://doi.org/", 'bib_doi'))
        segments.append((clean_doi, 'bib_doi'))
    elif url and ctype != 'web':
        segments.append((" ", None))
        segments.append((url, 'bib_url'))

    return segments

def abbreviate_journal_name_basic(name: str) -> str:
    if not name:
        return "No journal available"
        
    name_norm = normalize_whitespace(name).lower()
    if name_norm in REF_CACHE['journal_abbrev']:
        return REF_CACHE['journal_abbrev'][name_norm]

    mapping = {
    "Journal of Virology": "J Virol",
    "The Journal of Virology": "J Virol",
    "journal of virology": "J Virol",
    "The journal of virology": "J Virol",
    "Journal of General Virology": "J Gen Virol",
    "The Journal of General Virology": "J Gen Virol",
    "journal of general virology": "J Gen Virol",
    "The journal of general virology": "J Gen Virol",
    "Nature Immunology": "Nat Immunol",
    "nature immunology": "Nat Immunol",
    "Thorax": "Thorax",
    "thorax": "Thorax",
    "Ultrasound in Obstetrics & Gynecology": "Ultrasound Obstet Gynecol",
    "Ultrasound in Obstetrics and Gynecology": "Ultrasound Obstet Gynecol",
    "Ultrasound Obstet Gynecol": "Ultrasound Obstet Gynecol",
    "Prenatal Diagnosis": "Prenat Diagn",
    "Prenatal diagnosis": "Prenat Diagn",
    "American Journal of Perinatology": "Am J Perinat",
    "American Journal of Obstetrics and Gynecology": "Am J Obstet Gynecol",
    "American Journal of Obstetrics & Gynecology": "Am J Obstet Gynecol",
    "Obstetrics and Gynecology": "Obstet Gynecol",
    "Obstetrics &amp; Gynecology": "Obstet Gynecol",
    "Journal of Ultrasound in Medicine": "J Ultrasound Med",
    "Journal of Clinical Ultrasound": "J Clin Ultrasound",
    "Current Opinion in Obstetrics &amp; Gynecology": "Curr Opin Obstet Gynecol",
    "Current Opinion in Obstetrics & Gynecology": "Curr Opin Obstet Gynecol",
    "Cochrane Database of Systematic Reviews": "Cochrane Database Syst Rev",
    "New England Journal of Medicine": "N Engl J Med",
    "The New England Journal of Medicine": "N Engl J Med",
    "Journal of the American Medical Association": "JAMA",
    "JAMA": "JAMA",
    "The Lancet": "Lancet",
    "British Medical Journal": "BMJ",
    "Annals of Internal Medicine": "Ann Intern Med",
    "Journal of Clinical Investigation": "J Clin Invest",
    "Nature": "Nature",
    "Science": "Science",
    "Cell": "Cell",
    "Proceedings of the National Academy of Sciences": "Proc Natl Acad Sci U S A",
    "PNAS": "Proc Natl Acad Sci U S A",
    "Journal of Immunology": "J Immunol",
    "Blood": "Blood",
    "Circulation": "Circulation",
    "Journal of the American College of Cardiology": "J Am Coll Cardiol",
    "Journal of the American Society of Nephrology": "J Am Soc Nephrol",
    "JASN": "J Am Soc Nephrol",
    "Frontiers in Neurology": "Front Neurol",
    "Frontiers in Immunology": "Front Immunol",
    "Neurology": "Neurology",
    "Annals of Neurology": "Ann Neurol",
    "JAMA Neurology": "JAMA Neurol",
    "JAMA Psychiatry": "JAMA Psychiatry",
    "JAMA Internal Medicine": "JAMA Intern Med",
    "JAMA Surgery": "JAMA Surg",
    "JAMA Pediatrics": "JAMA Pediatr",
    "JAMA Ophthalmology": "JAMA Ophthalmol",
    "JAMA Otolaryngology–Head & Neck Surgery": "JAMA Otolaryngol Head Neck Surg",
    "JAMA Dermatology": "JAMA Dermatol",
    "JAMA Cardiology": "JAMA Cardiol",
    "JAMA Oncology": "JAMA Oncol",
    "Journal of Neuroscience": "J Neurosci",
    "Journal of Neuroinflammation": "J Neuroinflammation",
    "Brain": "Brain",
    "Brain Research": "Brain Res",
    "Neuroscience": "Neuroscience",
    "Nature Neuroscience": "Nat Neurosci",
    "Nature Medicine": "Nat Med",
    "Nature Reviews Neuroscience": "Nat Rev Neurosci",
    "Science Translational Medicine": "Sci Transl Med",
    "Journal of Clinical Oncology": "J Clin Oncol",
    "Journal of the National Cancer Institute": "J Natl Cancer Inst",
    "Cancer Research": "Cancer Res",
    "Clinical Cancer Research": "Clin Cancer Res",
    "American Journal of Respiratory and Critical Care Medicine": "Am J Respir Crit Care Med",
    "Chest": "Chest",
    "European Respiratory Journal": "Eur Respir J",
    "Diabetes": "Diabetes",
    "Diabetes Care": "Diabetes Care",
    "Diabetologia": "Diabetologia",
    "Journal of Clinical Endocrinology & Metabolism": "J Clin Endocrinol Metab",
    "Endocrinology": "Endocrinology",
    "Hepatology": "Hepatology",
    "Journal of Hepatology": "J Hepatol",
    "Gastroenterology": "Gastroenterology",
    "Gut": "Gut",
    "American Journal of Gastroenterology": "Am J Gastroenterol",
    "Rheumatology": "Rheumatology",
    "Annals of the Rheumatic Diseases": "Ann Rheum Dis",
    "Arthritis & Rheumatology": "Arthritis Rheumatol",
    "Kidney International": "Kidney Int",
    "American Journal of Kidney Diseases": "Am J Kidney Dis",
    "Clinical Journal of the American Society of Nephrology": "Clin J Am Soc Nephrol",
    "Nephrology Dialysis Transplantation": "Nephrol Dial Transplant",
    "Journal of the American Geriatrics Society": "J Am Geriatr Soc",
    "Journal of Gerontology": "J Gerontol",
    "Aging Cell": "Aging Cell",
    "PLOS ONE": "PLoS One",
    "PLOS Medicine": "PLoS Med",
    "PLOS Biology": "PLoS Biol",
    "BMC Medicine": "BMC Med",
    "BMC Biology": "BMC Biol",
    "Molecular Psychiatry": "Mol Psychiatry",
    "Biological Psychiatry": "Biol Psychiatry",
    "American Journal of Psychiatry": "Am J Psychiatry",
    "Journal of Affective Disorders": "J Affect Disord",
    "Schizophrenia Research": "Schizophr Res",
    "Journal of Infectious Diseases": "J Infect Dis",
    "Clinical Infectious Diseases": "Clin Infect Dis",
    "The Journal of Infectious Diseases": "J Infect Dis",
    "Emerging Infectious Diseases": "Emerg Infect Dis",
    "Infection and Immunity": "Infect Immun",
    "Vaccine": "Vaccine",
    "Antimicrobial Agents and Chemotherapy": "Antimicrob Agents Chemother",
    "Journal of Antimicrobial Chemotherapy": "J Antimicrob Chemother",
    "Clinical Microbiology Reviews": "Clin Microbiol Rev",
    "Journal of Clinical Microbiology": "J Clin Microbiol",
    "journal of virology": "J Virol",
    "the journal of virology": "J Virol",
    "journal of general virology": "J Gen Virol",
    "nature immunology": "Nat Immunol",
    "thorax": "Thorax",
    "ultrasound in obstetrics & gynecology": "Ultrasound Obstet Gynecol",
    "ultrasound in obstetrics and gynecology": "Ultrasound Obstet Gynecol",
    "ultrasound obstet gynecol": "Ultrasound Obstet Gynecol",
    "prenatal diagnosis": "Prenat Diagn",
    "american journal of perinatology": "Am J Perinat",
    "american journal of obstetrics and gynecology": "Am J Obstet Gynecol",
    "american journal of obstetrics & gynecology": "Am J Obstet Gynecol",
    "obstetrics and gynecology": "Obstet Gynecol",
    "obstetrics & gynecology": "Obstet Gynecol",
    "journal of ultrasound in medicine": "J Ultrasound Med",
    "journal of clinical ultrasound": "J Clin Ultrasound",
    "current opinion in obstetrics & gynecology": "Curr Opin Obstet Gynecol",
    "current opinion in obstetrics and gynecology": "Curr Opin Obstet Gynecol",
    "cochrane database of systematic reviews": "Cochrane Database Syst Rev",
    "new england journal of medicine": "N Engl J Med",
    "journal of the american medical association": "JAMA",
    "jama": "JAMA",
    "the lancet": "Lancet",
    "british medical journal": "BMJ",
    "annals of internal medicine": "Ann Intern Med",
    "journal of clinical investigation": "J Clin Invest",
    "nature": "Nature",
    "science": "Science",
    "cell": "Cell",
    "proceedings of the national academy of sciences": "Proc Natl Acad Sci U S A",
    "pnas": "Proc Natl Acad Sci U S A",
    "journal of immunology": "J Immunol",
    "blood": "Blood",
    "circulation": "Circulation",
    "journal of the american college of cardiology": "J Am Coll Cardiol",
    "journal of the american society of nephrology": "J Am Soc Nephrol",
    "jasn": "J Am Soc Nephrol",
    "frontiers in neurology": "Front Neurol",
    "frontiers in immunology": "Front Immunol",
    "neurology": "Neurology",
    "annals of neurology": "Ann Neurol",
    "jama neurology": "JAMA Neurol",
    "jama psychiatry": "JAMA Psychiatry",
    "jama internal medicine": "JAMA Intern Med",
    "jama surgery": "JAMA Surg",
    "jama pediatrics": "JAMA Pediatr",
    "jama ophthalmology": "JAMA Ophthalmol",
    "jama otolaryngology–head & neck surgery": "JAMA Otolaryngol Head Neck Surg",
    "jama dermatology": "JAMA Dermatol",
    "jama cardiology": "JAMA Cardiol",
    "jama oncology": "JAMA Oncol",
    "journal of neuroscience": "J Neurosci",
    "journal of neuroinflammation": "J Neuroinflammation",
    "brain": "Brain",
    "brain research": "Brain Res",
    "neuroscience": "Neuroscience",
    "nature neuroscience": "Nat Neurosci",
    "nature medicine": "Nat Med",
    "nature reviews neuroscience": "Nat Rev Neurosci",
    "science translational medicine": "Sci Transl Med",
    "journal of clinical oncology": "J Clin Oncol",
    "journal of the national cancer institute": "J Natl Cancer Inst",
    "cancer research": "Cancer Res",
    "clinical cancer research": "Clin Cancer Res",
    "american journal of respiratory and critical care medicine": "Am J Respir Crit Care Med",
    "chest": "Chest",
    "european respiratory journal": "Eur Respir J",
    "diabetes": "Diabetes",
    "diabetes care": "Diabetes Care",
    "diabetologia": "Diabetologia",
    "journal of clinical endocrinology & metabolism": "J Clin Endocrinol Metab",
    "endocrinology": "Endocrinology",
    "hepatology": "Hepatology",
    "journal of hepatology": "J Hepatol",
    "gastroenterology": "Gastroenterology",
    "gut": "Gut",
    "american journal of gastroenterology": "Am J Gastroenterol",
    "rheumatology": "Rheumatology",
    "annals of the rheumatic diseases": "Ann Rheum Dis",
    "arthritis & rheumatology": "Arthritis Rheumatol",
    "kidney international": "Kidney Int",
    "american journal of kidney diseases": "Am J Kidney Dis",
    "clinical journal of the american society of nephrology": "Clin J Am Soc Nephrol",
    "nephrology dialysis transplantation": "Nephrol Dial Transplant",
    "journal of the american geriatrics society": "J Am Geriatr Soc",
    "journal of gerontology": "J Gerontol",
    "aging cell": "Aging Cell",
    "plos one": "PLoS One",
    "plos medicine": "PLoS Med",
    "plos biology": "PLoS Biol",
    "bmc medicine": "BMC Med",
    "bmc biology": "BMC Biol",
    "molecular psychiatry": "Mol Psychiatry",
    "biological psychiatry": "Biol Psychiatry",
    "american journal of psychiatry": "Am J Psychiatry",
    "journal of affective disorders": "J Affect Disord",
    "schizophrenia research": "Schizophr Res",
    "journal of infectious diseases": "J Infect Dis",
    "clinical infectious diseases": "Clin Infect Dis",
    "emerging infectious diseases": "Emerg Infect Dis",
    "infection and immunity": "Infect Immun",
    "vaccine": "Vaccine",
    "antimicrobial agents and chemotherapy": "Antimicrob Agents Chemother",
    "journal of antimicrobial chemotherapy": "J Antimicrob Chemother",
    "clinical microbiology reviews": "Clin Microbiol Rev",
    "journal of clinical microbiology": "J Clin Microbiol",
    "cold spring harbor perspectives in medicine": "Cold Spring Harb Perspect Med",
    "cold spring harbor perspectives in biology": "Cold Spring Harb Perspect Biol"
    }
    name_lower = name.lower()
    for key, abbr in mapping.items():
        if key.lower() == name_lower:
            with CACHE_LOCK:
                REF_CACHE['journal_abbrev'][name_norm] = abbr
            return abbr
    sorted_items = sorted(mapping.items(), key=lambda x: len(x[0]), reverse=True)
    for key, abbr in sorted_items:
        key_lower = key.lower()
        # Clean entities thoroughly
        name_clean = html.unescape(name_lower).replace('–', '-')
        key_clean = html.unescape(key_lower).replace('–', '-')

        name_clean = name_clean.replace('&amp;', '&')
        key_clean = key_clean.replace('&amp;', '&')

        # Only match if:
        # 1. Exact match (already handled above)
        # 2. The key is contained as complete words in the name
        # This prevents "cancer" from matching "journal of the national cancer institute"
        
        # For multi-word keys, allow substring matching
        if ' ' in key_clean:
            if key_clean in name_clean or name_clean in key_clean:
                with CACHE_LOCK:
                    REF_CACHE['journal_abbrev'][name_norm] = abbr
                return abbr
        else:
            # For single-word keys, require word boundary match
            # Use regex to ensure it's a complete word
            import re as re_module
            pattern = r'\b' + re_module.escape(key_clean) + r'\b'
            if re_module.search(pattern, name_clean):
                # But only if it's not a partial match of a longer journal name
                # E.g., "cancer" should not match "journal of cancer research"
                # unless the name IS "cancer"
                if name_clean == key_clean or len(name_clean.split()) == 1:
                    with CACHE_LOCK:
                        REF_CACHE['journal_abbrev'][name_norm] = abbr
                    return abbr
    
    # Store original as fallback if no match found (or should we?)
    # Usually cache success only. But if we want to avoid re-scanning mapping...
    return name

def generate_ama_citation(item: Dict[str, Any]) -> List[Tuple[str, Optional[str]]]:
    ctype = item.get('type', 'journal-article')
    # Year extraction priority: published-print > published-online > issued > year > created
    # 'created' is the CrossRef record creation date, NOT publication date
    date_parts = (
        item.get('published-print', {}).get('date-parts') or
        item.get('published-online', {}).get('date-parts') or
        item.get('issued', {}).get('date-parts') or
        [[None]]
    )
    year = str(date_parts[0][0]) if date_parts[0][0] else (item.get('year') or 'n.d')
    if year == 'n.d' and item.get('created', {}).get('date-parts'):
        cd = item['created']['date-parts']
        if cd and cd[0] and cd[0][0]:
            year = str(cd[0][0])

    # helper to clean text
    def clean_text(t):
        if not t: return t
        return html.unescape(t).replace('&amp;', '&')

    title = clean_text((item.get('title') or ['No title available'])[0])
    container = (item.get('container-title') or [''])[0]
    container_display = clean_text(container)

    # 1. Try Manual Mapping / Abbreviation
    manual_abbr = abbreviate_journal_name_basic(container_display)
    if manual_abbr != container_display:
        journal_abbr = manual_abbr
    else:
        short_titles = item.get('short-container-title')
        if short_titles and isinstance(short_titles, list) and short_titles[0]:
            journal_abbr = clean_text(short_titles[0])
        else:
            journal_abbr = manual_abbr 

    volume = item.get('volume', '')
    issue = item.get('issue', '')
    pages = item.get('page', '')
    # DEBUG: Log page data for troubleshooting
    if pages:
        logger.debug(f"AMA Citation - Pages from API: '{pages}' (Length: {len(pages)})")
    doi = item.get('DOI', '')
    url = item.get('URL', '')

    segments = []
    authors_list = item.get('author', [])
    
    if not authors_list:
        segments.append(("Unknown authors", 'bib_unpubl'))
    else:
        # AMA 11th Edition Rule:
        # If <= 6 authors, list all.
        # If > 6 authors, list first 3 followed by "et al."
        if len(authors_list) <= 6:
            subset = authors_list
            has_etal = False
        else:
            subset = authors_list[:3]
            has_etal = True
            
        for i, author in enumerate(subset):
            if i > 0:
                segments.append((", ", None))
            
            family = author.get('family', '').strip()
            given = author.get('given', '').strip()
            
            if family:
                segments.append((family, 'bib_surname'))
            
            if given:
                segments.append((" ", None))
                # AMA Initials: "AB" (No dots, no spaces between)
                # extract_initials returns "AB" which is perfect.
                initials = extract_initials(given).replace('.', '').replace(' ', '')
                segments.append((initials, 'bib_fname'))
        
        if has_etal:
            segments.append((", et al", 'bib_etal'))

    segments.append((". ", None))
    
    # Title styling (Sentence case usually, but we take API text)
    t_style = 'bib_article'
    if ctype == 'book': t_style = 'bib_book'
    elif ctype == 'web': t_style = 'bib_title'
    elif ctype == 'book-chapter': t_style = 'bib_chaptertitle'
    elif ctype == 'thesis': t_style = 'bib_book' # AMA thesis title italic? No, usually standard. But let's follow book pattern if treated as monograph. AMA manual matches book style.
    elif ctype in ('conference-paper', 'proceedings-article'): t_style = 'bib_confpaper'
    
    # Strip dots from title to avoid double punctuation
    title_clean = title.rstrip('.')
    segments.append((title_clean, t_style))
    segments.append((". ", None))

    if ctype == 'book':
        publisher = item.get('publisher', '')
        if publisher:
            segments.append((publisher, 'bib_publisher'))
            segments.append(("; ", None))
        segments.append((year, 'bib_year'))
        segments.append((".", None))
        
    elif ctype in ('proceedings-article', 'conference-paper', 'book-chapter'):
        if container:
            segments.append(("In: ", None))
            # Editors often listed here for AMA. "In: Editor, ed. Book."
            # We skip editor logic for now as it needs more API data.
            segments.append((container, 'bib_confproceedings' if ctype != 'book-chapter' else 'bib_book'))
            segments.append((". ", None))
        if item.get('publisher'):
            segments.append((item.get('publisher'), 'bib_publisher'))
            segments.append(("; ", None))
        segments.append((year, 'bib_year'))
        if pages:
            segments.append((":", None))
            segments.append((pages, 'bib_fpage')) # Keep hyphen for AMA
        segments.append((".", None))
        
    elif ctype == 'thesis':
        # AMA: Author. Title [dissertation]. Location: Institution; Year.
        segments.append((" [dissertation]. ", None))
        # Location? Publisher often has "City, State: Univ" or just "Univ".
        # We append publisher directly.
        if item.get('publisher'):
             segments.append((item.get('publisher'), 'bib_publisher'))
             segments.append(("; ", None))
        segments.append((year, 'bib_year'))
        segments.append((".", None))
        
    elif ctype == 'web':
        if container:
            segments.append((container, 'bib_journal'))
            segments.append((". ", None))
        segments.append(("Published ", None))
        segments.append((year, 'bib_year'))
        segments.append((".", None))
        if url:
            segments.append((" ", None))
            segments.append((url, 'bib_url'))
            
    else:
        # Journal Article
        # AMA: Journal Years;Vol(Issue):Pages. (Tight punctuation)
        segments.append((journal_abbr, 'bib_journal')) # Italic
        segments.append((". ", None))
        segments.append((year, 'bib_year'))
        segments.append((";", None))
        
        if volume:
            segments.append((volume, 'bib_volume'))
        if issue:
            segments.append(("(", None))
            segments.append((issue, 'bib_issue'))
            segments.append((")", None))
        # NOTE: Full page range is preserved (not truncated) - AMA style uses complete page numbers with hyphens
        if pages:
            segments.append((":", None))
            # Preserve full page range (e.g., 123-456) with hyphen (AMA style)
            segments.append((pages, 'bib_fpage')) # Keep hyphen for AMA
            
        # DOI preference over URL for journals
        if doi:
            segments.append((". ", None))
            clean_doi = doi.replace('http://dx.doi.org/', '').replace('https://doi.org/', '')
            segments.append(("doi:", 'bib_doi'))
            segments.append((clean_doi, 'bib_doi'))
        elif url:
            segments.append((" ", None))
            segments.append((url, 'bib_url'))

    return segments

# -------------------------
# Chicago Style Generator (18th Ed.)
# -------------------------
def generate_chicago_citation(item: Dict[str, Any]) -> List[Tuple[str, Optional[str]]]:
    """
    Generates Chicago Style (Bibliography) segments.
    Format: Author. Title. Container. Publisher, Year.
    """
    ctype = item.get('type', 'journal-article')
    # Year extraction priority: published-print > published-online > issued > year > created
    date_parts = (
        item.get('published-print', {}).get('date-parts') or
        item.get('published-online', {}).get('date-parts') or
        item.get('issued', {}).get('date-parts') or
        [[None]]
    )
    year = str(date_parts[0][0]) if date_parts[0][0] else (item.get('year') or 'n.d.')
    if year == 'n.d.' and item.get('created', {}).get('date-parts'):
        cd = item['created']['date-parts']
        if cd and cd[0] and cd[0][0]:
            year = str(cd[0][0])
    
    # helper to clean text
    def clean_text(t):
        if not t: return t
        return html.unescape(t).replace('&amp;', '&')

    title = clean_text((item.get('title') or ['No title available'])[0])
    container = (item.get('container-title') or [''])[0]
    container = clean_text(container)
    
    publisher = item.get('publisher', '')
    volume = item.get('volume', '')
    issue = item.get('issue', '')
    pages = item.get('page', '')
    doi = item.get('DOI', '')
    url = item.get('URL', '')
    
    segments = []
    
    # --- Authors ---
    authors_list = item.get('author', [])
    if not authors_list:
        segments.append(("Unknown authors", 'bib_unpubl'))
    else:
        # Chicago: First author inverted (Surname, Given), subsequent given surname.
        limit = 10
        count = len(authors_list)
        subset = authors_list[:limit]
        
        for i, a in enumerate(subset):
            fam = a.get('family', '').strip()
            giv = a.get('given', '').strip()
            
            if i == 0:
                segments.append((fam, 'bib_surname'))
                if giv:
                    segments.append((", ", None))
                    segments.append((giv, 'bib_fname'))
            else:
                segments.append((", ", None))
                if i == count - 1 and count <= limit:
                    segments.append(("and ", None))
                
                if giv:
                    segments.append((giv, 'bib_fname'))
                    segments.append((" ", None))
                segments.append((fam, 'bib_surname'))
        
        if count > limit:
            segments.append((", et al", 'bib_etal'))
            
    segments.append((". ", None))
    
    # --- Title & Container ---
    if ctype == 'book':
        # Author. *Title*. Place: Publisher, Year.
        segments.append((title, 'bib_book'))
        segments.append((". ", None))
        if publisher:
            segments.append((publisher, 'bib_publisher'))
            segments.append((", ", None))
        segments.append((year, 'bib_year'))
        segments.append((".", None))
        
    elif ctype == 'chapter' or ctype == 'book-chapter':
        # Author. "Title." In *Book*, edited by..., pages. Pub, Year.
        segments.append(('"', None))
        segments.append((title, 'bib_chaptertitle'))
        segments.append(('."', None))
        segments.append((" In ", None))
        segments.append((container, 'bib_book'))
        segments.append((", ", None))
        if pages:
            segments.append((pages, 'bib_fpage')) 
            segments.append((". ", None))
        if publisher:
            segments.append((publisher, 'bib_publisher'))
            segments.append((", ", None))
        segments.append((year, 'bib_year'))
        segments.append((".", None))

    elif ctype == 'web':
        # Author. "Title." Site. Year. URL.
        segments.append(('"', None))
        segments.append((title, 'bib_title'))
        segments.append(('."', None))
        if container:
             segments.append((" ", None))
             segments.append((container, 'bib_journal'))
        segments.append((". ", None))
        segments.append((year, 'bib_year'))
        segments.append((".", None))
    
    else:
        # Journal
        segments.append(('"', None))
        segments.append((title, 'bib_article'))
        segments.append(('."', None))
        segments.append((" ", None))
        segments.append((container, 'bib_journal'))
        
        if volume:
            segments.append((" ", None))
            segments.append((volume, 'bib_volume'))
        if issue:
            segments.append((", no. ", None))
            segments.append((issue, 'bib_issue'))
            
        segments.append((" (", None))
        segments.append((year, 'bib_year'))
        segments.append(("): ", None))
        
        if pages:
             segments.append((pages, 'bib_fpage'))
        segments.append((".", None))

    # --- Links ---
    if doi:
        segments.append((" https://doi.org/", 'bib_doi'))
        segments.append((doi, 'bib_doi'))
        segments.append((".", None))
    elif url:
         segments.append((" ", None))
         segments.append((url, 'bib_url'))
         segments.append((".", None))
         
    return segments

# -------------------------
# Style-based parsing (character styles)
# -------------------------
def parse_reference_from_styles(para) -> Optional[Dict[str, Any]]:
    data = {
        'author_list': [],
        'year': '',
        'title': [],
        'container-title': [],
        'volume': '',
        'issue': '',
        'page': '',
        'DOI': '',
        'URL': '',
        'publisher': '',
        'type': 'journal-article',
        'has_etal': False
    }
    curr_surname = ''
    curr_fname = ''
    def flush_author():
        nonlocal curr_surname, curr_fname
        if curr_surname or curr_fname:
            data['author_list'].append({'family': curr_surname.strip(), 'given': curr_fname.strip()})
            curr_surname = ''
            curr_fname = ''

    has_data = False
    prev_style = None
    for run in para.runs:
        text = run.text
        if not text:
            continue
        style = run.style.name.lower() if run.style else ''
        if 'bib_surname' in style or 'bib_ed-surname' in style:
            if curr_fname:
                flush_author()
            curr_surname += text
            has_data = True
        elif 'bib_fname' in style or 'bib_ed-fname' in style:
            curr_fname += text
            has_data = True
        elif 'bib_etal' in style or 'bib_ed-etal' in style:
            flush_author()
            data['has_etal'] = True
            has_data = True
        elif 'bib_year' in style or 'bib_confdate' in style:
            flush_author()
            data['year'] += text
            has_data = True
        elif 'bib_title' in style or 'bib_chaptertitle' in style or 'bib_article' in style or 'bib_confpaper' in style:
            flush_author()
            if not data['title']: data['title'] = ['']
            data['title'][0] += text
            has_data = True
        elif 'bib_journal' in style or 'bib_confproceedings' in style or 'bib_conference' in style:
            flush_author()
            if not data['container-title']: data['container-title'] = ['']
            data['container-title'][0] += text
            has_data = True
        elif 'bib_volume' in style or 'bib_volcount' in style:
            flush_author()
            data['volume'] += text
            has_data = True
        elif 'bib_issue' in style or 'bib_number' in style:
            flush_author()
            data['issue'] += text
            has_data = True
        elif 'bib_pages' in style or 'bib_fpage' in style or 'bib_lpage' in style or 'bib_pagecount' in style:
            flush_author()
            data['page'] += text
            has_data = True
        elif 'bib_doi' in style:
            flush_author()
            data['DOI'] += text
            has_data = True
        elif 'bib_url' in style or 'bib_extlink' in style:
            flush_author()
            data['URL'] += text
            has_data = True
        elif 'bib_publisher' in style or 'bib_institution' in style or 'bib_organization' in style or 'bib_school' in style:
            flush_author()
            data['publisher'] += text
            has_data = True
        elif 'bib_book' in style:
            flush_author()
            data['type'] = 'book'
            if not data['title']: data['title'] = ['']
            data['title'][0] += text
            has_data = True
        else:
            # ignore unstyled punctuation mostly
            pass

        if style:
            prev_style = style

    flush_author()
    if not has_data:
        return None
    if data['has_etal'] and len(data['author_list']) < 4:
        while len(data['author_list']) < 4:
            data['author_list'].append({'family': '', 'given': ''})
    data['author'] = data['author_list']
    data['year'] = re.sub(r'[^\d]', '', data['year'])
    data['volume'] = data['volume'].strip()
    data['issue'] = data['issue'].strip().replace('(', '').replace(')', '')
    data['page'] = data['page'].strip()
    data['DOI'] = data['DOI'].strip()
    data['URL'] = data['URL'].strip()
    data['publisher'] = data['publisher'].strip()
    if data['title']: data['title'][0] = data['title'][0].strip()
    if data['container-title']: data['container-title'][0] = data['container-title'][0].strip()
    if data['publisher'] and data['type'] == 'journal-article':
        data['type'] = 'book'
    return data

# -------------------------
# Fallback generator for unvalidated references
# -------------------------
def generate_fallback_citation(parsed: Dict[str, Any], original_raw: str, style_mode='REF-U') -> List[Tuple[str, Optional[str]]]:
    """
    Generates granular segments for non-API references (Book, Web, Thesis, Edited Book)
    or failed validations, adhering to AMA (REF-N) or APA (REF-U) styles.
    """
    segments = []
    manual_type = parsed.get('manual_type') or parsed.get('type')
    
    # helper regexes
    re_year = re.compile(r'\b(19|20)\d{2}\b')
    re_url = re.compile(r'https?://\S+')
    
    def parse_authors_title_rest(text):
        """Splitter that respects ?, !, and . as delimiters."""
        # Capture delimiters
        parts = re.split(r'([.?!])\s+', text)
        chunks = []
        if len(parts) > 1:
            for i in range(0, len(parts)-1, 2):
                c = parts[i].strip()
                d = parts[i+1]
                if c: chunks.append(c + d)
            if len(parts) % 2 == 1 and parts[-1].strip():
                chunks.append(parts[-1].strip())
        else:
            chunks = [text]

        if len(chunks) >= 3:
            return chunks[0], chunks[1], " ".join(chunks[2:])
        elif len(chunks) == 2:
            return chunks[0], chunks[1], ""
        else:
            return text, "", ""

    def granular_authors(auth_str, mode):
        # "Surname I" or "Surname, I."
        s = []
        if not auth_str: return s
        
        # Helper to add parsed author
        def add_single_author(fam, initials):
            s.append((fam, 'bib_surname'))
            if initials:
                 if mode == 'REF-U': # APA wants dots/commas
                      s.append((", ", None))
                      # if initials like "A.B." or "A B", normalize?
                      # Assume initials string is raw.
                      s.append((initials, 'bib_fname'))
                 else: # AMA usually space only
                      s.append((" ", None))
                      s.append((initials.replace('.', '').replace(' ', ''), 'bib_fname'))

        if mode == 'REF-N': # AMA
            alist = [a.strip() for a in auth_str.split(',') if a.strip()]
            for i, a in enumerate(alist):
                if i > 0: s.append((", ", None))
                # Try match "Surname Initials" e.g. "Smith JB"
                m = re.match(r'^(.+)\s+([A-Z]+(?:\.?[A-Z]+)*)$', a)
                if m:
                    add_single_author(m.group(1), m.group(2))
                else:
                    s.append((a, 'bib_surname'))
        else: # APA
             # Pattern: Surname, I., & Surname, I.
             # Clean "&" first to treat list uniformly? Or split by "&" and ","?
             # Heuristic: split by ", &" or "," then parse "Surname, I."
             
             # Step 1: Replace "&" with "," for splitting
             clean = re.sub(r'\s+&\s+', ',', auth_str)
             parts = [p.strip() for p in clean.split(',') if p.strip()]
             
             # Re-assemble "Surname, I." pairs
             authors = []
             curr = ""
             for p in parts:
                 # Check if p is likely initials (1-2 letters, optionally dots)
                 if re.match(r'^[A-Z]\.?\s*[A-Z]?\.?$', p):
                     if curr:
                         authors.append((curr, p)) # Tuple (Surname, Initials)
                         curr = ""
                 else:
                     if curr: # pushing duplicate surname or weird structure
                         authors.append((curr, "")) 
                     curr = p
             if curr: authors.append((curr, "")) # trailing
             
             count = len(authors)
             for i, (fam, init) in enumerate(authors):
                 if i > 0:
                     s.append((", ", None))
                     if i == count - 1:
                         s.append(("& ", None))
                 
                 add_single_author(fam, init)
        
        return s

    # 1. SPECIAL MANUAL HANDLING
    if manual_type in ('book', 'edited_book', 'thesis', 'web', 'other'):
        
        # --- PRE-PROCESS: Extract URL/DOI if Web ---
        url_val = ""
        m_url = re_url.search(original_raw)
        txt_clean = original_raw
        if m_url:
            url_val = m_url.group(0).rstrip('.,;)')
            txt_clean = original_raw.replace(url_val, "").strip().rstrip('.,;')
            
        m_y = re_year.search(txt_clean)
        year_val = m_y.group(0) if m_y else ""
        
        if style_mode == 'REF-N': # AMA
             # AMA Format: Authors. Title. | Journal/City: Pub | ; Year.
             # Thesis: Authors. Title [Dissertation]. City: University; Year.
             # Web: Authors. Title. URL. Published Year. Accessed...
             
             a, t, rest = parse_authors_title_rest(txt_clean)
             
             # Authors
             segments.extend(granular_authors(a, 'REF-N'))
             if segments and not segments[-1][0].endswith('.'):
                  segments.append((". ", None))
             elif segments:
                  segments.append((" ", None))
             
             # Title + Edition Check
             # Check for Edition in title: "Title. 4th ed." or "Title 4th ed."
             edition_match = re.search(r'\b(\d+(?:st|nd|rd|th)?\s+ed\.?)$', t, re.IGNORECASE)
             edition_str = None
             if edition_match:
                 edition_str = edition_match.group(1)
                 t = t[:edition_match.start()].strip().rstrip('.')
             
             t_style = 'bib_book' if manual_type == 'book' else 'bib_chaptertitle' if manual_type == 'edited_book' else 'bib_title'
             
             segments.append((t, t_style))
             segments.append((". ", None))
             
             if edition_str:
                 segments.append((edition_str, 'bib_editionno'))
                 segments.append((". ", None))
             
             # Rest (Editors / Pub / City / Journal)
             if rest:
                 curr_rest = rest
                 
                 # Check if edition is at start of rest? e.g. "4th ed. City: Pub"
                 if not edition_str:
                     m_rest_ed = re.match(r'^(\d+(?:st|nd|rd|th)?\s+ed\.?)\s+', curr_rest, re.IGNORECASE)
                     if m_rest_ed:
                         edition_str = m_rest_ed.group(1)
                         # Consume edition from curr_rest
                         curr_rest = curr_rest[m_rest_ed.end():].strip()
                         
                         # Add edition segment now
                         segments.append((edition_str, 'bib_editionno'))
                         segments.append((". ", None))

                 # 1. Check for Editors ("In: ... eds.")
                 # Pattern: "In: Smith J, ed." or "In: Smith J, Jones B, eds."
                 # We need to extract the editor string.
                 m_eds = re.match(r'^In:\s+(.+?)(?:,\s*eds?\.?|\s+\(eds?\.?\))', curr_rest, re.IGNORECASE)
                 if m_eds:
                     ed_chunk = m_eds.group(1) # "Smith J, Jones B"
                     # consume this part from curr_rest
                     # find where it ended
                     match_len = m_eds.end()
                     curr_rest = curr_rest[match_len:].strip()
                     if curr_rest.startswith('.'): curr_rest = curr_rest[1:].strip()
                     
                     segments.append(("In: ", None))
                     
                     # Parse Editors granularly (similar to authors but BIB_ED_*)
                     # Reuse granular_authors logic but map styles?
                     # Or just custom split here since AMA editors usually "Surname Initial"
                     ed_list = [e.strip() for e in ed_chunk.split(',') if e.strip()]
                     for i, ed in enumerate(ed_list):
                         if i > 0: segments.append((", ", None))
                         m_nm = re.match(r'^(.+)\s+([A-Z]+(?:\.?[A-Z]+)*)$', ed)
                         if m_nm:
                             segments.append((m_nm.group(1), 'bib_ed-surname'))
                             segments.append((" ", None))
                             segments.append((m_nm.group(2), 'bib_ed-fname'))
                         else:
                             segments.append((ed, 'bib_ed-surname'))
                     
                     segments.append((", eds. ", None))
                 
                 # 2. Check for Book Title (after editors) if edited book
                 if manual_type == 'edited_book':
                      # Now we expect Book Title. City: Pub; Year.
                      # "Nursing Research... Science. New York: Springer; 1994:211-215"
                      # Heuristic: split by dot or look for City: Pub
                      # Let's try to detect City: Pub pattern at the end of the non-year part
                      pass
                 
                 # 3. Handle Publisher / Location / Pages
                 # Pattern: "City, State: Publisher; Year:Pages" or "City: Pub; Year."
                 
                 # remove year/pages part first?
                 # curr_rest might be: "Book Title. City: Pub; 1994:211-215" or just "City: Pub..."
                 
                 # Let's extract year+pages from the END first
                 # Match "; 1995" or "; 1994:211-215" or ". 1995"
                 m_end = re.search(r'[;.]\s*(\d{4})(?::(\d+[-–]\d+))?\.?$', curr_rest)
                 pages_val = None
                 
                 rest_main = curr_rest
                 if m_end:
                     year_val = m_end.group(1)
                     pages_val = m_end.group(2)
                     rest_main = curr_rest[:m_end.start()].strip()
                 else:
                     # Year might be extracted earlier by global regex, remove it if present
                     if year_val:
                         rest_main = rest_main.replace(year_val, "").strip().rstrip(';').rstrip('.')

                 # Check for Book Title in rest_main (if Edited Book)
                 if manual_type == 'edited_book':
                     # rest_main is "Book Title. City: Pub"
                     # Split by lookahead for City: Pub?
                     # Or split by first dot?
                     parts = rest_main.split('. ')
                     if len(parts) > 1:
                         bk_title = parts[0]
                         pub_part = ". ".join(parts[1:])
                         segments.append((bk_title, 'bib_book'))
                         segments.append((". ", None))
                         rest_main = pub_part
                     else:
                          # Assume entire thing is book title if no punctuation? Unlikely.
                          pass

                 # Parse Location: Publisher from rest_main
                 # "Baltimore, MD: Williams & Wilkins" -> Loc: Baltimore, MD | Pub: Williams & Wilkins
                 m_pub = re.match(r'^(.+?):\s*(.+)$', rest_main)
                 if m_pub:
                     loc = m_pub.group(1).strip()
                     pub = m_pub.group(2).strip()
                     segments.append((loc, 'bib_location'))
                     segments.append((": ", None))
                     segments.append((pub, 'bib_publisher'))
                 else:
                     if manual_type == 'thesis':
                         segments.append((rest_main, 'bib_publisher'))
                     elif manual_type == 'web':
                         pass
                     else:
                         segments.append((rest_main, 'bib_publisher'))

                 segments.append(("; ", None))
                 
                 if year_val:
                     segments.append((year_val, 'bib_year'))
                 
                 if pages_val:
                     segments.append((":", None))
                     # Split fpage-lpage
                     if '-' in pages_val or '–' in pages_val:
                         p_parts = re.split(r'[-–]', pages_val)
                         segments.append((p_parts[0], 'bib_fpage'))
                         segments.append(("-", None))
                         segments.append((p_parts[1], 'bib_lpage'))
                     else:
                         segments.append((pages_val, 'bib_fpage'))
                         
                 segments.append((".", None))

             if manual_type == 'web':
                 if url_val:
                     segments.append((" ", None))
                     segments.append((url_val, 'bib_url'))
                 if 'accessed' in original_raw.lower():
                     m_acc = re.search(r'(?i)accessed.*$', original_raw)
                     if m_acc:
                         acc_text = m_acc.group(0)
                         if not acc_text.startswith('.'): segments.append((". ", None))
                         segments.append((acc_text, 'bib_comment'))

        else: # APA
             # APA Format: Authors (Year). Title. In... (Eds), Book Title (pp. x-y). Pub.
             # or: Authors (Year). Title. Pub.
             
             # Find the year-in-parens pattern, not just first '(' (to avoid abbreviations like (NACNS))
             m_year_paren = re.search(r'\(\s*(?:\d{4}|n\.d\.)\s*(?:,\s*[^)]+)?\)', txt_clean)
             if m_year_paren and m_year_paren.start() > 5:
                 a_end = m_year_paren.start()
                 a = txt_clean[:a_end].strip()
                 rem = txt_clean[a_end:].strip()
             elif txt_clean.find('(') > 5:
                 a_end = txt_clean.find('(')
                 a = txt_clean[:a_end].strip()
                 rem = txt_clean[a_end:].strip()
             else:
                 a, t, rest = parse_authors_title_rest(txt_clean)
                 rem = ""
            
             segments.extend(granular_authors(a, 'REF-U'))
             if segments and not segments[-1][0].endswith('.'):
                  segments.append((". ", None))
             elif segments:
                  # Force a space before (Year) if not present
                  if not segments[-1][0].endswith(' '):
                      segments.append((" ", None))
             
             # Granular Date Parsing for APA
             # Try to find full date (YYYY, Month Day) or n.d.
             m_date = re.search(r'\(((?:\d{4}|n\.d\.|n\. d\.))(?:,\s*([^)]+))?\)', txt_clean, re.IGNORECASE)
             if m_date:
                 full_y = m_date.group(1)
                 rest_date = m_date.group(2) # "October 16" or "July" or "Spring"
                 
                 segments.append(("(", None))
                 segments.append((full_y, 'bib_year'))
                 
                 if rest_date:
                     segments.append((", ", None))
                     # Try parse Month Day: "October 16"
                     m_md = re.match(r'([A-Za-z]+)\s*(\d+)?', rest_date)
                     if m_md:
                         mon = m_md.group(1)
                         day = m_md.group(2)
                         segments.append((mon, 'bib_month'))
                         if day:
                             segments.append((" ", None))
                             segments.append((day, 'bib_day'))
                     else:
                          segments.append((rest_date, 'bib_month'))
                 
                 segments.append(("). ", None))
             elif year_val:
                 segments.append(("(", None))
                 segments.append((year_val, 'bib_year'))
                 segments.append(("). ", None))
             
             # Extract title from remainder?
             if rem:
                 # remove (Year). or (Year, Month Day). or (n.d.).
                 rem_no_y = re.sub(r'^\(\s*(?:\d{4}|n\.d\.|n\. d\.)\s*(?:,.*?)?\)\s*\.?\s*', '', rem, flags=re.IGNORECASE)
             elif m_date and t:
                 # If rem was empty, it means parse_authors_title_rest put everything in `t`.
                 # We must remove the date string from `t` to avoid duplication.
                 # The date string is `m_date.group(0)`
                 date_str = m_date.group(0)
                 # Expect it to be at the end of authors, start of title? 
                 # Actually, normally `t` starts with the title.
                 # But if `parse_authors_title_rest` failed, `t` might be "Title. (2020)." or "(2020). Title."
                 # Let's try to remove it from `t`.
                 rem_no_y = t.replace(date_str, "").strip()
                 # If it starts with a dot now, strip it.
                 if rem_no_y.startswith('.'): rem_no_y = rem_no_y[1:].strip()
                 # Now use this cleaned string as the source for title parsing logic below
                 # Trick: set 'rem' to something so we enter the block below?
                 # No, better simply set `rem_no_y` and let following logic use it if we adapt it.
                 # But wait, the logic below uses `rem_no_y` only if `if rem:` block is active?
                 # Actually, `rem_no_y` is local to the if block. 
                 # We need to restructure.
                 pass 
             
             # Re-thinking structure to handle both cases
             if rem:
                 rem_no_y = re.sub(r'^\(\s*(?:\d{4}|n\.d\.|n\. d\.)\s*(?:,.*?)?\)\s*\.?\s*', '', rem, flags=re.IGNORECASE)
             elif m_date and t:
                  # Attempt to remove found date from t
                  rem_no_y = t.replace(m_date.group(0), "").strip()
                  if rem_no_y.startswith('.'): rem_no_y = rem_no_y[1:].strip()
             else:
                  rem_no_y = t # Fallback if no date found or no rem match

             # Now unconditional check on rem_no_y is tricky because variables inside `if rem` scope?
             # Let's verify variable scope. `rem_no_y` needs to be defined.
             if 'rem_no_y' not in locals():
                 rem_no_y = rem or t # fallback
                 
             if True: #Indent block wrapper for diff simplicity
                 # Check for Edited Book Logic: "Title. In Editors (Eds.), Book Title (pp. 1-2). Publisher."
                 # Pattern: "In ... (Eds.?),"
                 
                 # Check for Edited Book Logic: "Title. In Editors (Eds.), Book Title (pp. 1-2). Publisher."
                 # Pattern: "In ... (Eds.?),"
                 # Require "In" at sentence boundary: after ". " or at start of string
                 # This prevents splitting titles like "Consultation in Clinical Nurse Specialist Practice"
                 # Check for Edited Book Logic (Eds.) OR Chapter in Book (no editors)
                 m_eds_papa = re.search(r'(?:(?<=\.\s)|^)In\s+(.+?)\s+\(Eds?\.?\),', rem_no_y, re.IGNORECASE)
                 m_just_in = None
                 if not m_eds_papa:
                     # Check for "In " that is NOT "In press"
                     m_tmp = re.search(r'(?:(?<=\.\s)|^)In\s+', rem_no_y, re.IGNORECASE)
                     if m_tmp:
                         # Filter out "In press" etc.
                         look_ahead = rem_no_y[m_tmp.end():].strip().lower()
                         if not any(look_ahead.startswith(bad) for bad in ['press', 'prep', '(press', '(prep']):
                             m_just_in = m_tmp

                 if m_eds_papa or m_just_in:
                      match_obj = m_eds_papa or m_just_in
                      # Chapter Title is before "In"
                      start_in = match_obj.start()
                      chap_tit = rem_no_y[:start_in].strip()
                      # Ensure it ends with dot
                      if not chap_tit.endswith('.'): chap_tit += '.'
                      
                      segments.append((chap_tit, 'bib_chaptertitle'))
                      segments.append((" ", None))
                      segments.append(("In ", None))
                      
                      if m_eds_papa:
                           ed_blob = m_eds_papa.group(1).strip()
                           # Split editors logic
                           has_amp = '&' in ed_blob
                           if has_amp:
                               amp_parts = re.split(r',?\s*&\s*', ed_blob)
                           else:
                               amp_parts = [ed_blob]
                           
                           ed_parts = []
                           for part in amp_parts:
                               part = part.strip()
                               if not part: continue
                               sub_parts = re.split(r',\s+(?=[A-Z][\.\-])', part)
                               ed_parts.extend(sub_parts)
                           
                           # FIX: Merge separated Surname/Initials (e.g. "Cross", "W." -> "Cross, W.")
                           # Helps when editors are listed as "Surname, I." instead of "I. Surname"
                           merged_parts = []
                           buf = None
                           for p in ed_parts:
                               p = p.strip()
                               if not p: continue
                               if buf:
                                   # If p is Initial-like (W. or A.) and buf has no dots (Cross)
                                   # Regex for initial: Start with cap, ends with dot, short length
                                   if len(p) <= 4 and re.match(r'^[A-Z](?:[\.\-][A-Z]?)?\.?$', p) and not re.search(r'[\.\-]', buf):
                                       buf = f"{buf}, {p}"
                                   else:
                                       merged_parts.append(buf)
                                       buf = p
                               else:
                                   buf = p
                           if buf: merged_parts.append(buf)
                           ed_parts = merged_parts
                           
                           for i, ep in enumerate(ed_parts):
                               ep = ep.strip()
                               if not ep: continue
                               if i > 0: segments.append((", ", None))
                               if i == len(ed_parts)-1 and has_amp and len(ed_parts) > 1:
                                   segments.append(("& ", None))
                               
                               # Fix "Surname Initials" (Lyon BL) -> B. L. Lyon
                               m_rev_no_comma = re.match(r'^([A-Z][a-z]+)\s+([A-Z]{1,3})$', ep)
                               if m_rev_no_comma:
                                   sur = m_rev_no_comma.group(1)
                                   inits_raw = m_rev_no_comma.group(2)
                                   # Convert BL -> B. L.
                                   new_inits = ". ".join(list(inits_raw)) + "."
                                   segments.append((new_inits, 'bib_ed-fname'))
                                   segments.append((" ", None))
                                   segments.append((sur, 'bib_ed-surname'))
                                   continue

                               # Try parse "S. Coates" or "A. B. Surname"
                               m_name = re.match(r'^([A-Z][\.\-](?:\s*[A-Z][\.\-])*\.?)\s+(.+)$', ep)
                               if m_name:
                                   segments.append((m_name.group(1), 'bib_ed-fname'))
                                   segments.append((" ", None))
                                   segments.append((m_name.group(2), 'bib_ed-surname'))
                               else:
                                   segments.append((ep, 'bib_ed-surname'))
                           
                           segments.append((" (Eds.), ", None))
                           rem_book = rem_no_y[m_eds_papa.end():].strip()
                      else:
                           # No editors
                           rem_book = rem_no_y[m_just_in.end():].strip()
                      
                      # Extract Pages: (pp. 75-97) or combined edition+pages (3 ed., pp. 75-97)
                      # First try combined edition+pages pattern
                      m_ed_pp = re.search(r'\(([^)]*ed\.),\s*pp\.?\s*(\d+[-–]\d+)\)', rem_book)
                      m_pp = re.search(r'\(pp\.?\s*(\d+[-–]\d+)\)', rem_book)
                      bk_title = rem_book
                      pub_blob = ""
                      
                      if m_ed_pp:
                          # Combined: (3 ed., pp. 129-133)
                          bk_title = rem_book[:m_ed_pp.start()].strip()
                          edition_inner = m_ed_pp.group(1).strip()
                          pages_inner = m_ed_pp.group(2)
                          rem_after_pp = rem_book[m_ed_pp.end():].strip()
                          if rem_after_pp.startswith('.'): rem_after_pp = rem_after_pp[1:].strip()
                          pub_blob = rem_after_pp
                          
                          segments.append((bk_title, 'bib_book'))
                          segments.append((" (", None))
                          segments.append((edition_inner, 'bib_editionno'))
                          segments.append((", pp. ", None))
                          # split fpage/lpage
                          if '-' in pages_inner or '–' in pages_inner:
                              pgs = re.split(r'[-–]', pages_inner)
                              segments.append((pgs[0], 'bib_fpage'))
                              segments.append(("–", None))
                              segments.append((pgs[1], 'bib_lpage'))
                          else:
                              segments.append((pages_inner, 'bib_fpage'))
                          segments.append(("). ", None))
                      elif m_pp:
                          bk_title = rem_book[:m_pp.start()].strip()
                          pages_inner = m_pp.group(1)
                          rem_after_pp = rem_book[m_pp.end():].strip()
                          if rem_after_pp.startswith('.'): rem_after_pp = rem_after_pp[1:].strip()
                          pub_blob = rem_after_pp
                          
                          segments.append((bk_title, 'bib_book'))
                          segments.append((" (pp. ", None))
                          # split fpage/lpage
                          if '-' in pages_inner or '–' in pages_inner:
                              pgs = re.split(r'[-–]', pages_inner)
                              segments.append((pgs[0], 'bib_fpage'))
                              segments.append(("-", None))
                              segments.append((pgs[1], 'bib_lpage'))
                          else:
                              segments.append((pages_inner, 'bib_fpage'))
                          segments.append(("). ", None))
                      else:
                          # No pages found?
                          # Try split by dot for publisher
                          bk_parts = rem_book.split('.')
                          if len(bk_parts) > 1:
                              bk_title = bk_parts[0]
                              pub_blob = ". ".join(bk_parts[1:])
                          segments.append((bk_title, 'bib_book'))
                          segments.append((". ", None))
                      
                      if pub_blob:
                          segments.append((pub_blob.strip().rstrip('.'), 'bib_publisher'))
                          segments.append((".", None))

                 else:
                     # Standard Book or Article
                     # rem_no_y matches "Title. Publisher." or "Title (5th ed.). Publisher."
                     
                     # Check for Edition: (5th ed.) OR (Rev. ed.)
                     m_ed_apa = re.search(r'\(([^)]+ed\.)\)', rem_no_y)
                     edition_str = None
                     
                     if m_ed_apa:
                         edition_str = m_ed_apa.group(1)
                         # remove edition from string to isolate title and publisher
                         # Usually Title (Ed). Publisher.
                         # We need to act carefully around where it is.
                         # It splits title and rest.
                         pre_ed = rem_no_y[:m_ed_apa.start()].strip()
                         post_ed = rem_no_y[m_ed_apa.end():].strip() # includes following dot usually
                         
                         title_part = pre_ed
                         pub_part = post_ed
                         # Clean up leading dot in pub_part if it was "Title (Ed). Publisher"
                         if pub_part.startswith('.'): pub_part = pub_part[1:].strip()
                         
                     else:
                         t_parts = rem_no_y.split('.', 1)
                         title_part = t_parts[0]
                         pub_part = t_parts[1] if len(t_parts)>1 else ""
                     
                     # TAGGING: If manual_type is book, FORCE bib_book. 
                     # If unknown but looks like book (no volume/issue), prefer bib_book? 
                     # User said "some books title are wronly tagged as bib_title".
                     if manual_type == 'book':
                         t_style = 'bib_book'
                     elif manual_type == 'web':
                         t_style = 'bib_title'
                     # Fix: If it has volume/issue pattern in pub_part, it IS a journal
                     elif re.search(r'\d+\(\d+\)', pub_part):
                         t_style = 'bib_article'
                     else:
                         t_style = 'bib_book' # Default fallback, but maybe check for italic heuristics?
                         
                     segments.append((title_part.strip(), t_style))
                     
                     if edition_str:
                         segments.append((" (", None))
                         segments.append((edition_str, 'bib_editionno'))
                         segments.append((")", None))
                         
                     segments.append((". ", None))
                     
                     if pub_part:
                         # Remove trailing ... or dots
                         pub_clean = pub_part.strip().rstrip('.')
                         
                         # Check for malformed "In (Ed)." (Missing book title)
                         if re.match(r'^\s*In\s*\(\d', pub_clean, re.IGNORECASE):
                              segments.append((" [MISSING BOOK TITLE?]", None))
                         
                         # Heuristic: If we decided it's an article above, the rest is Journal info
                         if t_style == 'bib_article': 
                              segments.append((pub_clean, 'bib_journal'))
                         else:
                              segments.append((pub_clean, 'bib_publisher'))
                         
                         segments.append((".", None))
             
             elif t: 
                 t_style = 'bib_book' if manual_type == 'book' else 'bib_title'
                 segments.append((t.strip(), t_style))
                 segments.append((". ", None))

             if url_val:
                 segments.append((" ", None))
                 segments.append((url_val, 'bib_url'))
        
        # FINAL CLEANUP: Remove duplicate punctuation
        # Iterate and merge
        cleaned_segments = []
        for txt, sty in segments:
            if not txt: continue
            if cleaned_segments:
                last_txt = cleaned_segments[-1][0]
                # Avoid ".. " or ".;"
                if txt.strip() in ('.', ';', ',') and last_txt.rstrip()[-1:] in ('.', ';', ','):
                    # If we are skipping, but txt had a space (e.g. ". "), we might need to preserve the space
                    # unless last_txt already ends with space.
                    if txt.endswith(' ') and not last_txt.endswith(' '):
                        cleaned_segments.append((" ", None))
                    continue
                # Avoid ". ."
                if txt.strip() == '.' and last_txt.strip().endswith('.'):
                     if txt.endswith(' ') and not last_txt.endswith(' '):
                        cleaned_segments.append((" ", None))
                     continue
            
            cleaned_segments.append((txt, sty))
            
        return cleaned_segments

    # 2. EMERGENCY FALLBACK (Generic granular parsing for skipped/failed API)
    # Instead of returning original raw text, we attempt to parse it granularly
    # assuming standard formats (Author. Title. Source. Year).
    
    # --- PRE-PROCESS: Extract URL/DOI if Web ---
    url_val = ""
    m_url = re_url.search(original_raw)
    txt_clean = original_raw
    if m_url:
        url_val = m_url.group(0).rstrip('.,;)')
        txt_clean = original_raw.replace(url_val, "").strip().rstrip('.,;')
        
    m_y = re_year.search(txt_clean)
    year_val = m_y.group(0) if m_y else ""

    if style_mode == 'REF-N': # AMA
         a, t, rest = parse_authors_title_rest(txt_clean)
         
         # Authors
         segments.extend(granular_authors(a, 'REF-N'))
         if segments and not segments[-1][0].endswith('.'):
              segments.append((". ", None))
         elif segments:
              segments.append((" ", None))
         
         # Title
         segments.append((t.rstrip('.'), 'bib_article')) # Default to article style for generic
         segments.append((". ", None))
         
         # Rest
         # Rest
         if rest:
             # Try to parse standard AMA tail: Journal. Year;Vol(Issue):Pages
             # e.g. "Ultrasound Obstet Gynecol. 2015;27(2):143-150."
             # Regex allowing for optional dots/spaces after Journal
             m_ama = re.match(r'^(?P<j>.+?)[. ]+(?P<y>\d{4});\s*(?P<v>[\w]+)(?:\((?P<i>[\w]+)\))?:(?P<p>[\w\-–]+)\.?$', rest)
             
             if m_ama:
                 # Granular parse successful
                 j = m_ama.group('j')
                 y = m_ama.group('y')
                 v = m_ama.group('v')
                 i = m_ama.group('i')
                 p = m_ama.group('p')
                 
                 segments.append((j.strip(), 'bib_journal'))
                 segments.append((". ", None))
                 segments.append((y, 'bib_year'))
                 segments.append((";", None))
                 segments.append((v, 'bib_volume'))
                 if i:
                     segments.append(("(", None))
                     segments.append((i, 'bib_issue'))
                     segments.append((")", None))
                 segments.append((":", None))
                 if '-' in p or '–' in p:
                     p_parts = re.split(r'[-–]', p)
                     segments.append((p_parts[0], 'bib_fpage'))
                     segments.append(("-", None))
                     segments.append((p_parts[1], 'bib_lpage'))
                 else:
                     segments.append((p, 'bib_fpage'))
                 segments.append((".", None))
                 
             else:
                 # Naive fallback
                 # Only strip year if it's NOT in a weird place?
                 # Actually, if year_val is present, naive approach moves it to end.
                 # If user says "year move at end how?", they imply it looks wrong.
                 # If we failed regex, maybe year is just part of the string?
                 # e.g. "Some Journal, 2015." -> "Some Journal. 2015."
                 
                 pub_clean = rest.replace(year_val, "").strip().rstrip(';').rstrip('.')
                 segments.append((pub_clean, 'bib_journal')) # Default to journal style
                 
                 if year_val:
                      segments.append(("; ", None))
                      segments.append((year_val, 'bib_year'))
                      segments.append((".", None))
                 else:
                      segments.append((".", None))
             
         if url_val:
             segments.append((" ", None))
             segments.append((url_val, 'bib_url'))

    else: # APA
         # APA Format: Authors (Year). Title. Publisher/URL.
         # Find the year-in-parens pattern, not just first '(' (to avoid (NACNS) etc.)
         m_year_paren = re.search(r'\(\s*(?:\d{4}|n\.d\.)\s*(?:,\s*[^)]+)?\)', txt_clean)
         if m_year_paren and m_year_paren.start() > 5:
             a_end = m_year_paren.start()
             a = txt_clean[:a_end].strip()
             rem = txt_clean[a_end:].strip()
         elif txt_clean.find('(') > 5:
             a_end = txt_clean.find('(')
             a = txt_clean[:a_end].strip()
             rem = txt_clean[a_end:].strip()
         else:
             a, t, rest = parse_authors_title_rest(txt_clean)
             rem = ""
        
         segments.extend(granular_authors(a, 'REF-U'))
         if segments and not segments[-1][0].endswith('.'):
              segments.append((". ", None))
         elif segments:
              segments.append((" ", None))  # Ensure space if parsed authors didn't end with dot
         
         if year_val:
             segments.append(("(", None))
             segments.append((year_val, 'bib_year'))
             segments.append(("). ", None))
         
         # Extract title from remainder?
         if rem:
             # Strip year with optional internal spaces: (2020), (2020 ), etc.
             rem_no_y = re.sub(r'^\(\s*\d{4}\s*\)\s*\.?\s*', '', rem)
             # Also strip any remaining year-in-parens that wasn't at the start
             rem_no_y = re.sub(r'\(\s*\d{4}\s*\)\s*\.?\s*', '', rem_no_y).strip()
             
             if rem_no_y:
                 t_parts = rem_no_y.split('.', 1)
                 title_part = t_parts[0]
                 pub_part = t_parts[1] if len(t_parts)>1 else ""
                 
                 t_style = 'bib_book' # Default to book style for generic fallback (it's usually a book/report if it failed API)
                 segments.append((title_part.strip(), t_style)) # Default to title/italic
                 segments.append((". ", None))
                 
                 if pub_part:
                     pub_clean = pub_part.strip().rstrip('.')
                     segments.append((pub_clean, 'bib_publisher')) # Default to publisher/source
                     segments.append((".", None))
         elif t: 
             segments.append((t.strip(), 'bib_book'))
             segments.append((". ", None))

         if url_val:
             segments.append((" ", None))
             segments.append((url_val, 'bib_url'))
    
    # If segments empty (very unlikely), fallback to raw
    if not segments:
        segments.append((original_raw, None))
        
    return segments

def parse_authors_string(author_str: str) -> List[Dict[str, str]]:
    if not author_str:
        return []
    # Try to return list of dicts for closer compatibility with generators if needed
    return [{'family': author_str, 'given': ''}]

# -------------------------
# Word comment insertion (best-effort)
# -------------------------
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from lxml import etree as _lxml_etree

# Namespace map for lxml — keeps prefix as `w:` in serialised XML
_W_NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
_W14_NS= "http://schemas.microsoft.com/office/word/2010/wordml"
_NSMAP = {
    "w"  : _W_NS,
    "mc" : _MC_NS,
    "w14": _W14_NS,
}
_COMMENTS_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
_COMMENTS_CT  = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"

def _fresh_comments_tree():
    """Return a minimal lxml comments root with correct w: namespace prefix."""
    root = _lxml_etree.Element(f"{{{_W_NS}}}comments", nsmap=_NSMAP)
    root.set(f"{{{_MC_NS}}}Ignorable", "w14 wp14")
    return root

def _comments_part_for(doc):
    """Resolve (or create) the comments part, returning the Part object."""
    doc_part = doc.part
    for rel in doc_part.rels.values():
        if rel.reltype == _COMMENTS_REL:
            return rel.target_part

    # Not yet linked — look in package
    for p in doc_part.package.parts:
        if p.partname == "/word/comments.xml":
            doc_part.relate_to(p, _COMMENTS_REL)
            return p

    # Create from scratch
    blob = _lxml_etree.tostring(
        _fresh_comments_tree(),
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )
    part = Part(PackURI("/word/comments.xml"), _COMMENTS_CT, blob, doc_part.package)
    doc_part.package.add_part(part)
    doc_part.relate_to(part, _COMMENTS_REL)
    return part


def add_comment_to_runs(doc, runs, text, author="RefFix", initials="RF"):
    """
    Append a Word comment that wraps *runs* with correct w: namespace prefixes.
    Uses lxml throughout so the blob is never re-serialised through stdlib ET,
    which would replace the w: prefix with ns0: and corrupt the OOXML schema.
    """
    if not runs:
        return False

    try:
        # ── 1. Get / create comments part ──────────────────────────────
        cp = _comments_part_for(doc)

        # ── 2. Parse existing blob with lxml ───────────────────────────
        try:
            comments_tree = _lxml_etree.fromstring(cp._blob)
        except Exception:
            comments_tree = _fresh_comments_tree()

        # ── 3. Compute next comment id ─────────────────────────────────
        existing_ids = [
            int(c.get(f"{{{_W_NS}}}id"))
            for c in comments_tree.findall(f"{{{_W_NS}}}comment")
            if c.get(f"{{{_W_NS}}}id")
        ]
        next_id = max(existing_ids) + 1 if existing_ids else 0

        # ── 4. Build comment element (proper w: tags via lxml) ──────────
        comment_el = _lxml_etree.SubElement(comments_tree, f"{{{_W_NS}}}comment")
        comment_el.set(f"{{{_W_NS}}}id",       str(next_id))
        comment_el.set(f"{{{_W_NS}}}author",   author)
        comment_el.set(f"{{{_W_NS}}}initials", initials)

        p_el = _lxml_etree.SubElement(comment_el, f"{{{_W_NS}}}p")
        r_el = _lxml_etree.SubElement(p_el,       f"{{{_W_NS}}}r")
        t_el = _lxml_etree.SubElement(r_el,        f"{{{_W_NS}}}t")
        t_el.text = text

        # ── 5. Serialise back with lxml — keeps w: prefix ──────────────
        new_blob = _lxml_etree.tostring(
            comments_tree,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )
        cp._blob = new_blob

        # ── 6. Insert range markers into the paragraph ─────────────────
        p_element = runs[0]._element.getparent()  # w:p

        start = OxmlElement("w:commentRangeStart")
        start.set(qn("w:id"), str(next_id))

        end = OxmlElement("w:commentRangeEnd")
        end.set(qn("w:id"), str(next_id))

        ref   = OxmlElement("w:commentReference")
        ref.set(qn("w:id"), str(next_id))
        ref_r = OxmlElement("w:r")
        ref_r.append(ref)

        pPr = p_element.find(qn("w:pPr"))
        if pPr is not None:
            pPr.addnext(start)
        else:
            p_element.insert(0, start)

        p_element.append(end)
        p_element.append(ref_r)

        return True

    except Exception as e:
        logger.debug("Comment insertion failed: %s", repr(e))
        return False


def try_add_word_comment(doc, para, comment_text, author="RefFix", initials="RF"):
    """
    Wrapper to add comment to a full paragraph after styling.
    """
    if not para.runs:
        return False
    return add_comment_to_runs(doc, para.runs, comment_text, author, initials)

# -------------------------
# Utility: write citation into paragraph with styles
# -------------------------
def write_citation_with_styles(para, segments: List[Tuple[str, Optional[str]]], preserve_original_styles: bool = False, styles=None, doc=None):
    if preserve_original_styles:
        return

    # NEW: Global Post-Processing Cleanup (Segments)
    # Fix double punctuation ("..") and boundary issues (e.g. "Title?." -> "Title?")
    cleaned_segments = []
    for txt, sty in segments:
        if not txt: continue
        
        # 1. Internal cleanup: ".. -> ."
        txt = txt.replace('..', '.')
        
        # 2. Cross-segment boundary cleanup
        if cleaned_segments:
            prev_txt = cleaned_segments[-1][0]
            # Check if previous segment ended with punctuation
            if prev_txt and prev_txt.strip():
                last_char = prev_txt.rstrip()[-1]
                # If duplicative punctuation found across boundary
                if last_char in ('.', '?', '!') and txt.strip().startswith('.'):
                     # Remove the leading dot from current segment
                     # e.g. ". " -> " "
                     txt = re.sub(r'^\s*\.', '', txt, count=1)
        
        if txt:
            cleaned_segments.append((txt, sty))
    segments = cleaned_segments

    # 1. Prepare New Content
    new_full_text = ""
    style_map = [] # List of style names matching new_full_text indices
    
    # Pre-calculate mapping
    for text, style_suffix in segments:
        if not text: continue
        
        # Determine Style Name
        sname = None
        if style_suffix:
            style_map_dict = {
                'surname': 'bib_surname',
                'fname': 'bib_fname',
                'year': 'bib_year',
                'month': 'bib_month',
                'day': 'bib_day',
                'season': 'bib_season',
                'article': 'bib_article',
                'title': 'bib_title',
                'book': 'bib_book',
                'chapter': 'bib_chaptertitle',
                'journal': 'bib_journal',
                'volume': 'bib_volume',
                'issue': 'bib_issue',
                'fpage': 'bib_fpage',
                'lpage': 'bib_lpage',
                'etal': 'bib_etal',
                'doi': 'bib_doi',
                'url': 'bib_url',
                'publisher': 'bib_publisher',
                'institution': 'bib_institution',
                'organization': 'bib_organization',
                'bib_number': 'bib_number',
                'bib_alt-year': 'bib_alt-year',
                'bib_article': 'bib_article',
                'bib_base': 'bib_base',
                'bib_book': 'bib_book',
                'bib_chapterno': 'bib_chapterno',
                'bib_chaptertitle': 'bib_chaptertitle',
                'bib_comment': 'bib_comment',
                'bib_confacronym': 'bib_confacronym',
                'bib_confdate': 'bib_confdate',
                'bib_conference': 'bib_conference',
                'bib_conflocation': 'bib_conflocation',
                'bib_confpaper': 'bib_confpaper',
                'bib_confproceedings': 'bib_confproceedings',
                'bib_day': 'bib_day',
                'bib_deg': 'bib_deg',
                'bib_doi': 'bib_doi',
                'bib_ed-etal': 'bib_ed-etal',
                'bib_ed-fname': 'bib_ed-fname',
                'bib_editionno': 'bib_editionno',
                'bib_ed-organization': 'bib_ed-organization',
                'bib_ed-suffix': 'bib_ed-suffix',
                'bib_ed-surname': 'bib_ed-surname',
                'bib_etal': 'bib_etal',
                'bib_extlink': 'bib_extlink',
                'bib_fname': 'bib_fname',
                'bib_fpage': 'bib_fpage',
                'bib_institution': 'bib_institution',
                'bib_isbn': 'bib_isbn',
                'bib_issue': 'bib_issue',
                'bib_journal': 'bib_journal',
                'bib_location': 'bib_location',
                'bib_lpage': 'bib_lpage',
                'bib_medline': 'bib_medline',
                'bib_month': 'bib_month',
                'bib_number': 'bib_number',
                'bib_organization': 'bib_organization',
                'bib_pagecount': 'bib_pagecount',
                'bib_papernumber': 'bib_papernumber',
                'bib_patent': 'bib_patent',
                'bib_publisher': 'bib_publisher',
                'bib_reportnum': 'bib_reportnum',
                'bib_school': 'bib_school',
                'bib_season': 'bib_season',
                'bib_series': 'bib_series',
                'bib_seriesno': 'bib_seriesno',
                'bib_suffix': 'bib_suffix',
                'bib_suppl': 'bib_suppl',
                'bib_surname': 'bib_surname',
                'bib_title': 'bib_title',
                'bib_trans': 'bib_trans',
                'bib_unpubl': 'bib_unpubl',
                'bib_url': 'bib_url',
                'bib_volcount': 'bib_volcount',
                'bib_volume': 'bib_volume'
            }
            sname = style_suffix if style_suffix.startswith('bib_') else style_map_dict.get(style_suffix)
        
        # DEBUG: Log style mapping
        # print(f"DEBUG: Segment: '{text}' -> Style: {sname}")
        
        new_full_text += text
        style_map.extend([sname] * len(text))
        
    # print(f"DEBUG: Full Style Map Length: {len(style_map)}")

    # 2. Get Original Content
    original_text = para.text
    
    # 3. Diff and Reconstruct (if Track Changes enabled)
    if TRACK_CHANGES_ENABLED:
        try:
            # Clear Paragraph
            para._element.clear_content() # We're rebuilding it
            
            # Diff
            matcher = difflib.SequenceMatcher(None, original_text, new_full_text)
            
            for opcode, i1, i2, j1, j2 in matcher.get_opcodes():
                
                if opcode == 'equal':
                    # EQUAL: Text matches.
                    # We rebuild it as NORMAL runs (not tracked).
                    # But we apply the NEW styles from style_map.
                    
                    segment_text = new_full_text[j1:j2]
                    segment_styles = style_map[j1:j2]
                    
                    # Ensure style existence if needed
                    # We skip the complex checking here for brevity, assuming styles exist or pass str
                    
                    # Chunk by style
                    chunk_start = 0
                    for k in range(len(segment_text) + 1):
                         # Break if end or style changes
                        is_end = (k == len(segment_text))
                        style_changed = (k > 0 and k < len(segment_text) and segment_styles[k] != segment_styles[k-1])
                        
                        if is_end or style_changed:
                            chunk = segment_text[chunk_start:k]
                            if chunk:
                                style = segment_styles[chunk_start]
                                # Add Run
                                run = para.add_run(chunk)
                                if style:
                                    try:
                                        if styles and style not in styles:
                                             try:
                                                 styles.add_style(style, WD_STYLE_TYPE.CHARACTER)
                                             except: pass
                                        
                                        # Force style application
                                        if styles and style in styles:
                                            run.style = styles[style]
                                        else:
                                            run.style = style
                                    except: 
                                        pass
                            
                            chunk_start = k
                            
                elif opcode == 'delete':
                    # DELETE: In original, not in new.
                    # Add Tracked Deletion
                    deleted_chunk = original_text[i1:i2]
                    track_changes.add_tracked_deletion(para, deleted_chunk, doc=doc)
                    
                elif opcode == 'insert':
                    # INSERT: In new, not in original.
                    # Add Tracked Insertion (with styles)
                    
                    segment_text = new_full_text[j1:j2]
                    segment_styles = style_map[j1:j2]
                    
                    chunk_start = 0
                    for k in range(len(segment_text) + 1):
                         # Break if end or style changes
                        is_end = (k == len(segment_text))
                        style_changed = (k > 0 and k < len(segment_text) and segment_styles[k] != segment_styles[k-1])
                        
                        if is_end or style_changed:
                            chunk = segment_text[chunk_start:k]
                            if chunk:
                                style = segment_styles[chunk_start]
                                
                                try:
                                    track_changes.add_tracked_text(para, chunk, style=style, doc=doc)
                                except Exception as e:
                                    # Fallback
                                    # logging.warning(f"Tracked insert failed: {e}")
                                    para.add_run(chunk)
                            
                            chunk_start = k
                            
                elif opcode == 'replace':
                    # REPLACE: Delete + Insert
                    
                    # Delete
                    deleted_chunk = original_text[i1:i2]
                    track_changes.add_tracked_deletion(para, deleted_chunk, doc=doc)
                    
                    # Insert
                    segment_text = new_full_text[j1:j2]
                    segment_styles = style_map[j1:j2]
                    
                    chunk_start = 0
                    for k in range(len(segment_text) + 1):
                         # Break if end or style changes
                        is_end = (k == len(segment_text))
                        style_changed = (k > 0 and k < len(segment_text) and segment_styles[k] != segment_styles[k-1])
                        
                        if is_end or style_changed:
                            chunk = segment_text[chunk_start:k]
                            if chunk:
                                style = segment_styles[chunk_start]
                                
                                try:
                                    track_changes.add_tracked_text(para, chunk, style=style, doc=doc)
                                except Exception as e:
                                    # Fallback
                                    # logging.warning(f"Tracked replace/insert failed: {e}")
                                    para.add_run(chunk)
                            
                            chunk_start = k
                            
                            chunk_start = k

        except Exception as e:
            logging.error(f"Granular Track Changes Failed: {e}. Falling back to overwrite.")
            # Fallback: Wipe and Insert All (Old method but without track changes structure valid? No we cleared content)
            # If we cleared content and failed mid-way, we might have partial text.
            # Ideally try to just dump new_full_text? 
            # Or assume non-granular fallback
            pass
            
    else:
        # Standard Overwrite (No Track Changes)
        try:
            para._element.clear_content()
        except:
             para.text = ''
        
        # Build runs from segments (existing logic)
        for text, style_suffix in segments:
            run = para.add_run(text)
            
            if style_suffix:
                style_map_dict = {
                    'surname': 'bib_surname',
                    'fname': 'bib_fname',
                    'year': 'bib_year',
                    'month': 'bib_month',
                    'day': 'bib_day',
                    'season': 'bib_season',
                    'article': 'bib_article',
                    'title': 'bib_title',
                    'book': 'bib_book',
                    'chapter': 'bib_chaptertitle',
                    'journal': 'bib_journal',
                    'volume': 'bib_volume',
                    'issue': 'bib_issue',
                    'fpage': 'bib_fpage',
                    'lpage': 'bib_lpage',
                    'etal': 'bib_etal',
                    'doi': 'bib_doi',
                    'url': 'bib_url',
                    'publisher': 'bib_publisher',
                    'institution': 'bib_institution',
                    'organization': 'bib_organization',
                    'bib_number': 'bib_number',
                    'bib_alt-year': 'bib_alt-year',
                    'bib_article': 'bib_article',
                    'bib_base': 'bib_base',
                    'bib_book': 'bib_book',
                    'bib_chapterno': 'bib_chapterno',
                    'bib_chaptertitle': 'bib_chaptertitle',
                    'bib_comment': 'bib_comment',
                    'bib_confacronym': 'bib_confacronym',
                    'bib_confdate': 'bib_confdate',
                    'bib_conference': 'bib_conference',
                    'bib_conflocation': 'bib_conflocation',
                    'bib_confpaper': 'bib_confpaper',
                    'bib_confproceedings': 'bib_confproceedings',
                    'bib_day': 'bib_day',
                    'bib_deg': 'bib_deg',
                    'bib_doi': 'bib_doi',
                    'bib_ed-etal': 'bib_ed-etal',
                    'bib_ed-fname': 'bib_ed-fname',
                    'bib_editionno': 'bib_editionno',
                    'bib_ed-organization': 'bib_ed-organization',
                    'bib_ed-suffix': 'bib_ed-suffix',
                    'bib_ed-surname': 'bib_ed-surname',
                    'bib_etal': 'bib_etal',
                    'bib_extlink': 'bib_extlink',
                    'bib_fname': 'bib_fname',
                    'bib_fpage': 'bib_fpage',
                    'bib_institution': 'bib_institution',
                    'bib_isbn': 'bib_isbn',
                    'bib_issue': 'bib_issue',
                    'bib_journal': 'bib_journal',
                    'bib_location': 'bib_location',
                    'bib_lpage': 'bib_lpage',
                    'bib_medline': 'bib_medline',
                    'bib_month': 'bib_month',
                    'bib_number': 'bib_number',
                    'bib_organization': 'bib_organization',
                    'bib_pagecount': 'bib_pagecount',
                    'bib_papernumber': 'bib_papernumber',
                    'bib_patent': 'bib_patent',
                    'bib_publisher': 'bib_publisher',
                    'bib_reportnum': 'bib_reportnum',
                    'bib_school': 'bib_school',
                    'bib_season': 'bib_season',
                    'bib_series': 'bib_series',
                    'bib_seriesno': 'bib_seriesno',
                    'bib_suffix': 'bib_suffix',
                    'bib_suppl': 'bib_suppl',
                    'bib_surname': 'bib_surname',
                    'bib_title': 'bib_title',
                    'bib_trans': 'bib_trans',
                    'bib_unpubl': 'bib_unpubl',
                    'bib_url': 'bib_url',
                    'bib_volcount': 'bib_volcount',
                    'bib_volume': 'bib_volume'
                }
                
                sname = style_suffix if style_suffix.startswith('bib_') else style_map_dict.get(style_suffix)
                
                if sname:
                    try:
                        if styles is not None:
                            try:
                                if sname not in styles:
                                    styles.add_style(sname, WD_STYLE_TYPE.CHARACTER)
                                run.style = styles[sname]
                            except Exception:
                                try:
                                    run.style = sname
                                except Exception:
                                    pass
                        else:
                            run.style = sname
                    except Exception:
                        pass


# -------------------------
# The hybrid-find function (main logic for each reference)
# -------------------------
# -------------------------
# Helper: Detect Type Heuristics
# -------------------------
def is_likely_book_input(text: str) -> bool:
    """Return True if text likely represents a Book citation."""
    t = text.lower()
    # Indicators: "Press", "Publisher", "Sons", "Wiley", "Ed.", "Edition", locations like "NY:", "DC:"
    book_indicators = [
        'press', 'publisher', 'published', 'publishing', 'wiley', 'sons', 'ltd', 'inc.', 'co.', 
        'edition', '(ed.)', '(eds.)', 'volume', ' vol. ', ' pp. ', 'books', 'associates', 'group',
        'routledge', 'sage', 'springer', 'elsevier', 'pearson', 'mcgraw', 'chapman', 'hall',
        'basic books', 'guilford', 'addison-wesley' , 'Little Brown', 'chemical toxicology institute', 'Houghton Mifflin Harcourt', 'Mariner Books', 'New Harbinger'
    ]
    if any(k in t for k in book_indicators):
        return True
    # Pattern: City: Publisher (e.g., "New York: Springer")
    if re.search(r'[A-Z][a-zA-Z\s]+:\s*[A-Z]', text):
        return True
    return False

def is_likely_web_input(text: str) -> bool:
    """Return True if text likely represents a Website/Online Report."""
    t = text.lower()
    return 'http' in t or 'www.' in t or 'retrieved' in t or 'accessed' in t or 'about the ' in t or 'home page' in t

# -------------------------
# The hybrid-find function (main logic for each reference)
# -------------------------
# -------------------------
# The hybrid-find function (main logic for each reference)
# -------------------------
def pubmed_get_by_doi(doi: str) -> Optional[Dict[str, Any]]:
    """
    Attempts to fetch metadata from PubMed using DOI.
    Uses ESearch with [aid] or [doi] tag, then EFetch.
    """
    # 1. Search for ID
    # Try [aid] (Article Identifier) first, then [doi]
    # PubMed often indexes DOIs as Location IDs [lid] or Article IDs [aid]
    
    # Check cache via simulated search key
    cache_key = f"doi|{doi}"
    if cache_key in REF_CACHE['pubmed_search']:
         ids = REF_CACHE['pubmed_search'][cache_key]
    else:
        # Strategy: term="10.xxx/yyy[aid]" OR term="10.xxx/yyy[doi]"
        # We'll try generic search which usually resolves DOI
        try:
            # quote the DOI?
            q = f'{doi}[aid] OR {doi}[doi]'
            params = {'db': 'pubmed', 'term': q, 'retmax': 1, 'retmode': 'json'}
            r = SESSION.get(f"{NCBI_BASE}/esearch.fcgi", params=params, timeout=PUBMED_TIMEOUT)
            if r.status_code == 200:
                ids = r.json().get('esearchresult', {}).get('idlist', [])
                with CACHE_LOCK:
                    REF_CACHE['pubmed_search'][cache_key] = ids
            else:
                ids = []
        except RequestException:
            ids = []

    if not ids:
        return None

    # 2. Fetch Metadata
    pid = ids[0]
    root = pubmed_fetch_xml(pid)
    if root is not None:
        return pubmed_parse_article_from_xml(root)
    return None

def preprocess_reference_text(text: str) -> str:
    """Cleans up common OCR/formatting issues in references."""
    s = text
    # Fix spaced parens: ( 2020 ) -> (2020)
    s = re.sub(r'\(\s+(\d{4})\s+\)', r'(\1)', s)
    # Fix double dots: .. -> .
    s = re.sub(r'\.\.+', '.', s)
    return s.strip()

def detect_reference_style(raw_text: str) -> str:
    """Detects if a reference is likely AMA, APA, or CMS."""
    s = normalize_whitespace(raw_text)
    
    # APA: Look for (Year), (n.d.), or (in press) pattern
    # Handles "Author (Year)." and "Org (Abbrev) (Year)."
    # Updated to handle spaces inside parens like (2020 )
    if re.search(r'\(\s*(?:\d{4}[a-z]?|n\.d\.|in press)\s*\)', s, re.IGNORECASE):
        return 'APA'
        
    # AMA: Year;Vol (Journal) or Publisher; Year (Book)
    # Handles cases where numbering is stripped. 
    # Look for "2020;" or "; 2020" or "1. Author" (if numbering exists)
    # Also "Year:Pages" (2011:25-30) and "In: ... eds"
    if re.search(r'\b\d{4};', s) or re.search(r'[:;]\s*\d{4}\b', s) or \
       re.search(r'\b\d{4}:\d+', s) or re.search(r'^\d+\.\s+[A-Z]', s) or \
       re.search(r'\bIn:\s+.*eds?\.?', s, re.IGNORECASE):
        return 'AMA'
        
    # CMS: Quotes or "no. X"
    if re.search(r'\",\s+', s) or re.search(r'\bno\.\s+\d+', s):
        return 'CMS'
        
    return 'Unknown'

def find_best_metadata_for_reference(raw_ref: str, style_name: str) -> Tuple[Optional[Dict[str, Any]], str, float]:
    raw_ref = normalize_whitespace(raw_ref)
    detected_style = detect_reference_style(raw_ref)
    
    # Input Type Detection
    def detect_manual_type(t: str) -> Optional[str]:
        tl = t.lower()
        
        # Check for Journal patterns to EXCLUDE from manual book detection
        if re.search(r'\b\d+\s*\(\d+\)\s*:\s*\d+', t) or re.search(r'\b\d+\s*:\s*\d+[-–]\d+', t):
            if 'in:' not in tl and '(eds' not in tl:
                return None # Defer to API / fallback as Journal

        # Explicit Journal Check (Safety)
        if 'journal' in tl or 'doi.org' in tl:
            return None # Defer to API / fallback as Journal

        # Thesis
        if 'dissertation' in tl or 'master\'s thesis' in tl or 'phd thesis' in tl or 'doctoral thesis' in tl:
            return 'thesis'
        # Web
        if 'http' in tl or 'www.' in tl or 'retrieved' in tl or 'accessed' in tl or 'about the ' in tl or 'home page' in tl:
            return 'web'
        # Edited Book / Chapter
        if re.search(r'\bIn:\s', t, re.IGNORECASE) or '(eds' in tl or '(ed.' in tl:
            return 'edited_book'
            
        # Book indicators
        book_inds = [
            'press', 'publisher', 'wiley', 'sons', 'ltd', 'inc.', 'co.', 
            'edition', ' vol. ', 'books', 'associates', 'group',
            'routledge', 'sage', 'springer', 'elsevier', 'pearson', 'mcgraw', 'chapman', 'hall',
            'basic books', 'guilford', 'addison-wesley', 'williams & wilkins', 'williams and wilkins', 
            'chemical toxicology institute', 'Little Brown', 'Houghton Mifflin Harcourt', 'Mariner Books', 'New Harbinger'
        ]
        
        for k in book_inds:
             k_lower = k.lower()
             if ' ' in k_lower:
                  if k_lower in tl: return 'book'
             elif re.search(r'\b' + re.escape(k_lower) + r'\b', tl):
                  return 'book'

        if re.search(r'\b[A-Za-z\s]+:\s*[A-Za-z\s&]+(?:[;.]\s*\d{4}.*)?$', t): 
            return 'book'
            
        # If it has a year and period-separated structure, could be a journal - defer to API
        if re.search(r'\b(19|20)\d{2}\b', t) and '.' in t:
            return None
            
        return 'other'

    manual_type = detect_manual_type(raw_ref)
    
    if manual_type:
         if manual_type in ('book', 'edited_book'):
             # Try Google Books (logic remains)
             if style_name == 'REF-N': p_tmp = parse_ama_reference_raw(raw_ref)
             else: p_tmp = parse_apa_reference_raw(raw_ref)
             
             t_query = (p_tmp.get('title')[0] if isinstance(p_tmp.get('title'), list) and p_tmp.get('title') else p_tmp.get('title')) or raw_ref
             a_query = None
             if p_tmp.get('authors'):
                 a_parts = p_tmp['authors'].split(',')
                 if a_parts: a_query = a_parts[0].strip().split()[-1] 
             
             gb_res = search_google_books(t_query, a_query)
             if gb_res:
                 found_title = gb_res['title'][0]
                 sim = similarity(t_query, found_title)
                 if sim > 0.7:
                     # If Google Books didn't return a publisher, try extracting from raw text
                     if not gb_res.get('publisher'):
                         # For APA: publisher is typically the last segment after the last period
                         # e.g. "... (Year). Title. Publisher."
                         raw_parts = raw_ref.rsplit('.', 2)
                         if len(raw_parts) >= 2:
                             possible_pub = raw_parts[-2].strip()
                             # Only use if it looks like a publisher (not a title or year)
                             if possible_pub and not re.match(r'^\(\d{4}\)$', possible_pub) and len(possible_pub) > 2:
                                 gb_res['publisher'] = possible_pub

                     # Replace books.google.com URL with a useful isbn.org link if possible,
                     # otherwise keep it as a bib_url fallback (better than blanking it).
                     gb_url = gb_res.get('URL', '')
                     if gb_url and 'books.google' in gb_url:
                         # Try to extract the Google Books volume ID as a DOI-less identifier
                         # and prefer the raw ISBN from infoLink if available
                         m_isbn = re.search(r'isbn=(\d[\dX]+)', gb_url, re.IGNORECASE)
                         m_id   = re.search(r'[?&]id=([^&]+)', gb_url)
                         if m_isbn:
                             isbn_val = m_isbn.group(1)
                             gb_res['URL'] = f"https://www.isbn.org/{isbn_val}"
                         elif m_id:
                             # No ISBN, but keep a clean Open Library link as fallback
                             gb_res['URL'] = f"https://openlibrary.org/search?q={requests.utils.quote(found_title)}"
                         # else: keep original URL so citation has some link
                     return gb_res, 'google_books', sim
             
             # If Google Books fails, return manual_skip BUT with parsed metadata
             if style_name == 'REF-N': p_tmp = parse_ama_reference_raw(raw_ref)
             else: p_tmp = parse_apa_reference_raw(raw_ref)
             p_tmp['manual_type'] = manual_type
             p_tmp['type'] = manual_type
             p_tmp['detected_style'] = detected_style
             return p_tmp, 'manual_skip', 1.0
             

             
         # For other manual types (web, thesis, other), also return parsed metadata
         if style_name == 'REF-N': p_tmp = parse_ama_reference_raw(raw_ref)
         else: p_tmp = parse_apa_reference_raw(raw_ref)
         p_tmp['manual_type'] = manual_type
         p_tmp['type'] = manual_type
         p_tmp['detected_style'] = detected_style
         return p_tmp, 'manual_skip', 1.0

    # API Validation Path
    is_book_input = False
    is_web_input = False

    # 1. DOI Validation (Priority: PubMed -> CrossRef)
    doi = extract_doi_from_text(raw_ref)
    if doi:
        # Try PubMed first
        pm_doi_item = pubmed_get_by_doi(doi)
        if pm_doi_item:
            # Convert to unified/Crossref-like structure
            unified = pubmed_to_crossref_like(pm_doi_item)
            return unified, 'doi_pubmed', 1.0
        
        # Try CrossRef second
        cr = crossref_get_by_doi(doi)
        if cr:
            return cr, 'doi_crossref', 1.0

    # 2. Search Fallback (Priority: PubMed -> CrossRef)
    if style_name == 'REF-N':
        parsed = parse_ama_reference_raw(raw_ref)
    else:
        parsed = parse_apa_reference_raw(raw_ref)

    title_raw = parsed.get('title')
    title = (title_raw[0] if isinstance(title_raw, list) and title_raw else title_raw) or raw_ref
    journal_raw = parsed.get('container-title')
    journal = (journal_raw[0] if isinstance(journal_raw, list) and journal_raw else journal_raw) or None
    year = parsed.get('year') or None
    authors_str = parsed.get('authors') or ''

    # A. PubMed Search
    pm_ids = pubmed_search_ids(title, journal, year, max_results=4)
    pubmed_items = []
    for pid in pm_ids:
        root = pubmed_fetch_xml(pid)
        if root is None: continue
        pm_unified = pubmed_parse_article_from_xml(root)
        if pm_unified:
            pubmed_items.append(pm_unified)
    
    # Pick best PubMed
    best_pm, best_pm_score = None, 0.0
    tnorm = normalize_whitespace(title).lower()
    for pm in pubmed_items:
        ititle = (pm.get('title') or [''])[0]
        sc = similarity(tnorm, ititle or '')
        if sc > best_pm_score:
            best_pm_score = sc
            best_pm = pm
    
    # If PubMed is good enough (>= 0.8), stop here?
    # User said "first pubmed then cross ref site".
    # Assuming "strict priority", if PubMed finds a high confidence match, we use it.
    if best_pm and best_pm_score >= 0.85:
        return pubmed_to_crossref_like(best_pm), 'pubmed', best_pm_score

    # B. CrossRef Search (if PubMed failed or low score)
    cr_candidates = crossref_search(title, journal, year, rows=CROSSREF_ROWS)
    best_cr, best_cr_score = crossref_pick_best(title, cr_candidates) # Use helper directly
    
    # Compare strategies
    # If PubMed score is decent but < 0.85, and CrossRef is better?
    # We compare them now.
    
    final_best = None
    final_source = ''
    final_score = 0.0
    
    if best_pm and best_pm_score >= SIMILARITY_MIN:
        final_best = pubmed_to_crossref_like(best_pm)
        final_source = 'pubmed'
        final_score = best_pm_score
        
    if best_cr and best_cr_score >= SIMILARITY_MIN:
        # If CrossRef is significantly better, switch.
        # But prefer PubMed if close.
        threshold = 0.10 # CrossRef must be 10% better to override decent PubMed
        if not final_best:
            final_best = best_cr
            final_source = 'crossref'
            final_score = best_cr_score
        elif best_cr_score > (final_score + threshold):
             final_best = best_cr
             final_source = 'crossref'
             final_score = best_cr_score
             
    # --- FILTERING LOGIC ---
    if final_best:
        res_type = final_best.get('type', '').lower()
        if is_book_input and 'journal' in res_type:
            chosen_title = (final_best.get('title') or [''])[0]
            sim_pure = similarity(title, chosen_title)
            reject = False
            if sim_pure < 0.98: reject = True
            if len(chosen_title) > len(title) * 1.5: reject = True
            if reject: return None, 'filtered_book_mismatch', 0.0
        
        if is_web_input and 'journal' in res_type:
             if final_score < 0.98: return None, 'filtered_web_mismatch', 0.0

    # Fallback to web link detection if nothing found
    if not final_best:
        url_match = re.search(r'https?://\S+', raw_ref)
        if url_match:
            url = url_match.group(0).rstrip('.,;)')
            web_item = {
                'type': 'web',
                'title': [title],
                'author': parse_authors_string(authors_str),
                'container-title': [journal] if journal else [],
                'year': year,
                'URL': url,
                'DOI': doi
            }
            return web_item, 'web', 0.5
            
        return None, '', 0.0
    
    return final_best, final_source, final_score


# -------------------------
# DOCX processing
# -------------------------
def process_docx_file(input_docx: Path, output_dir: Optional[Path] = None, target_style: str = "Auto") -> Dict[str, Path]:
    if output_dir is None:
        output_dir = input_docx.parent
    
    if not output_dir.exists():
        output_dir.mkdir(parents=True, exist_ok=True)

    # Define output paths
    # We don't necessarily need a backup if we are working on a copy or if output is separate
    # But let's keep logic similar: Input -> Output
    
    output_docx = output_dir / f"{input_docx.stem}_fixed{input_docx.suffix}"
    log_file = output_dir / f"{input_docx.stem}_fix_log.txt"

    # We don't need to copy to backup if we are writing to a new output file usually
    # But if output_dir is same as input, we might overwrite? 
    # The original logic created `_fixed`, so it didn't overwrite original.
    # Let's Skip backup creation for web-app flow to save space/time, or make it optional.
    # For now, I'll drop the backup copy since we are generating a new file.

    doc = Document(input_docx)

    # ensure granular character styles exist
    style_names = [
        'bib_alt-year', 'bib_article', 'bib_base', 'bib_book', 'bib_chapterno', 'bib_chaptertitle', 
        'bib_comment', 'bib_confacronym', 'bib_confdate', 'bib_conference', 'bib_conflocation', 
        'bib_confpaper', 'bib_confproceedings', 'bib_day', 'bib_deg', 'bib_doi', 'bib_ed-etal', 
        'bib_ed-fname', 'bib_editionno', 'bib_ed-organization', 'bib_ed-suffix', 'bib_ed-surname', 
        'bib_etal', 'bib_extlink', 'bib_fname', 'bib_fpage', 'bib_institution', 'bib_isbn', 
        'bib_issue', 'bib_journal', 'bib_location', 'bib_lpage', 'bib_medline', 'bib_month', 
        'bib_number', 'bib_organization', 'bib_pagecount', 'bib_papernumber', 'bib_patent', 
        'bib_publisher', 'bib_reportnum', 'bib_school', 'bib_season', 'bib_series', 'bib_seriesno', 
        'bib_suffix', 'bib_suppl', 'bib_surname', 'bib_title', 'bib_trans', 'bib_unpubl', 'bib_url', 
        'bib_volcount', 'bib_volume', 'bib_year'
    ]
    for sname in style_names:
        try:
            if sname not in doc.styles:
                doc.styles.add_style(sname, WD_STYLE_TYPE.CHARACTER)
        except Exception:
            pass

    log_lines = []
    total = 0
    changed = 0
    unresolved = []
    all_ref_texts = [] # Track for duplicate detection

    in_ref_section = False
    
    
    # --- PROCESSING: PHASE 1 (Identify & Submit) ---
    tasks = []
    
    print("Phase 1: Scanning document and submitting API tasks...")
    for para in doc.paragraphs:
        raw = (para.text or '').strip()
        if not raw:
            continue
            
        # Check for section markers
        raw_lower = raw.lower()
        if '<ref-open>' in raw:
            in_ref_section = True
            continue
        if '<ref-close>' in raw:
            in_ref_section = False
            continue
            
        # Automatic Section Detection (Robustness) - REMOVED per user request
        # Strict <ref-open> / <ref-close> logic only.
        
        style = None
        try:
            style = para.style.name
        except Exception:
            style = None

        # Helper to recover from missing close tag if we hit a Heading
        if in_ref_section and style and style.startswith('Heading'):
             # Optional safety: close section if we hit a chapter heading
             if not any(k in raw_lower for k in ('references', 'bibliography', 'works cited', 'literature cited')):
                 in_ref_section = False
                 continue

        # STRICT CHECK: Only process if inside explicit reference section
        if not in_ref_section:
            continue
            
        if not style or style not in ('REF-N', 'REF-U', 'REF'):
            if re.match(r'^\[?\d+\]?\.?', raw):
                style = 'REF-N'
            elif raw.startswith('REF'):
                style = 'REF' # Heuristic if raw starts with REF? unlikely, but safety
            else:
                style = 'REF-U'
        
        # Debugging: Trace "suspicious" large captures
        if in_ref_section:
            if len(raw) < 40 and raw_lower in ('appendix', 'tables', 'figures', 'index'):
                in_ref_section = False
                continue

        # cleanup leading DOI urls (DOIs that appear at the BEGINNING of a reference, e.g. from copy-paste errors)
        # Only match if the DOI appears within the first 60 characters (truly leading DOI)
        m_head_doi = re.match(r'(?i)^(https?://doi\.org/10\.\S+)\s*', raw)
        if m_head_doi:
             raw_clean = raw[m_head_doi.end():].strip().lstrip('.,;)')
        else:
             raw_clean = raw
             
        # Apply robust preprocessing (fix spaced parens, double dots, etc.)
        raw_clean = preprocess_reference_text(raw_clean)
        
        # PRESERVE NUMBERING Logic
        current_numbering_prefix = ""
        m_num = re.match(r'^(\[?\d+\]?[\.\)]?)\s+', raw_clean)
        if m_num:
            current_numbering_prefix = m_num.group(0) 
            raw_for_search = raw_clean[len(current_numbering_prefix):]
        else:
            raw_for_search = raw_clean
            
        # Do NOT force style here; we need detected style for correct parsing.
        # We will apply target_style during generation.
            
        tasks.append({
            'para': para,
            'raw': raw,
            'style': style,
            'raw_clean': raw_clean,
            'raw_for_search': raw_for_search,
            'prefix': current_numbering_prefix
        })
    # ============================================================
    # STEP 1: Detect citation format for each reference (AMA/APA/CMS)
    # ============================================================
    for t in tasks:
        t['detected_format'] = detect_reference_style(t['raw_for_search'])

    format_counts = {}
    for t in tasks:
        fmt = t['detected_format']
        format_counts[fmt] = format_counts.get(fmt, 0) + 1
    print(f"Format detection: {format_counts}")

    # Majority-vote: fill 'Unknown' refs with the document's dominant style
    # so ambiguous references use the correct parser (AMA vs APA).
    known_counts = {k: v for k, v in format_counts.items() if k != 'Unknown'}
    dominant_style = max(known_counts, key=known_counts.get) if known_counts else 'AMA'
    unknowns = [t for t in tasks if t['detected_format'] == 'Unknown']
    if unknowns:
        print(f"\n=== Unknown References: {len(unknowns)} refs will inherit dominant style '{dominant_style}' ===")
        for t in unknowns:
            print(f"  UNKNOWN: {t['raw_clean'][:80]}")
            t['detected_format'] = dominant_style  # inherit majority-vote style
        print("=" * 60 + "\n")
    
    # ============================================================
    # STEP 2 & 3: Type classification + Journal API validation
    # Non-journal types (book, web, thesis, other) get granular styles only
    # Journal types get validated with PubMed/CrossRef
    # ============================================================
    print(f"Submitting {len(tasks)} references for parallel validation...")
    # PubMed rate-limit: 3 req/s (no API key) / 10 req/s (with key).
    # 30 workers causes cascading 429s even with retries.
    # 8 workers gives good IO concurrency while staying within limits.
    _n_workers = min(8, max(1, len(tasks)))
    with concurrent.futures.ThreadPoolExecutor(max_workers=_n_workers) as executor:
        for t in tasks:
            t['future'] = executor.submit(find_best_metadata_for_reference, t['raw_for_search'], t['style'])
            
        # --- PROCESSING: PHASE 2 (Apply) ---
        print("Phase 2: Applying changes...")
        for i, t in enumerate(tasks):
            para = t['para']
            raw = t['raw']
            style = t['style']
            raw_clean = t['raw_clean'] # needed?
            current_numbering_prefix = t['prefix']
            raw_for_search = t['raw_for_search']
            
            all_ref_texts.append({'id': len(all_ref_texts) + 1, 'text': raw_clean})
            total += 1
            log_lines.append(f"---\nOriginal ({style}): {raw}\n")
            
            comment_text_to_add = None

            try:
                # GET RESULT
                try:
                    item, source, score = t['future'].result() # Blocks if not ready
                except Exception as e:
                    logger.error(f"Task failed for ref '{raw[:30]}...': {e}")
                    raise e
                
                # --- STEP 2: GRANULAR TYPE CLASSIFICATION ---
                # (Already done in find_best_metadata_for_reference)
                
                if source == 'manual_skip':
                    source = 'manual_fallback'
                    log_lines.append(f"Detected type: {item.get('manual_type')}, Style: {item.get('detected_style', 'Unknown')}. Bypassing API validation, applying granular styles only.\n")

                # --- STEP 4: FINAL FORMAT VALIDATION ---
                # Check for missing elements based on reference type
                validation_errors = []
                if item:
                    itype = item.get('type')
                    # normalize type
                    if itype in ('book', 'edited_book'):
                         # For manual_fallback, the parsed metadata may not include publisher,
                         # but the raw text often does. Check raw text for publisher indicators.
                         has_publisher = bool(item.get('publisher'))
                         if not has_publisher and source == 'manual_fallback':
                             # Check raw text for common publisher names/patterns
                             publisher_indicators = [
                                 'press', 'publisher', 'wiley', 'sons', 'inc.', 'co.', 
                                 'routledge', 'sage', 'springer', 'elsevier', 'pearson',
                                 'basic books', 'guilford', 'appleton', 'grune', 'stratton',
                                 'oxford university', 'cambridge university', 'publishing',
                                 'saunders', 'mosby', 'lippincott', 'associates'
                             ]
                             raw_lower = raw.lower()
                             has_publisher = any(ind in raw_lower for ind in publisher_indicators)
                         if not has_publisher and not item.get('URL') and not item.get('DOI'):
                             validation_errors.append("Publisher missing")
                         if not item.get('year'):
                             validation_errors.append("Year missing")
                         if not item.get('author') and not item.get('editor'):
                             validation_errors.append("Author/Editor missing")
                    elif itype == 'web':
                         if not item.get('URL'):
                             validation_errors.append("URL missing")
                    elif itype == 'thesis':
                         if not item.get('publisher'): # Institution
                             validation_errors.append("Institution missing")
                    elif itype in ('journal-article', 'article-journal'):
                        # Journal format validation
                        if not item.get('volume') and not item.get('DOI'):
                            validation_errors.append("Volume missing")
                        if not item.get('page') and not item.get('DOI'):
                            validation_errors.append("Pages missing")
                        if not item.get('DOI') and not item.get('URL'):
                            validation_errors.append("DOI/URL missing")

                if validation_errors:
                    comment_text_to_add = f"Missing elements: {', '.join(validation_errors)}"
                
                # Explicit handling for filtered Book/Web inputs:
                if source in ('filtered_book_mismatch', 'filtered_web_mismatch'):
                    log_lines.append(f"Skipped API match due to {source}. Reverting to fallback styling.\n")
                    if not comment_text_to_add: 
                        comment_text_to_add = "Reference validation mismatch: Used standard formatting on original text."
                    
                    source = 'fallback'
                    item = None 
                
                is_original_web = bool(re.search(r'https?://', raw))
                
                # Check for web conservative fallback
                if is_original_web and item and source != 'web' and score < 0.9 and source != 'fallback':
                     log_lines.append("Low-confidence cross-type match for web reference; Reverting to fallback styling.\n")
                     if not comment_text_to_add:
                        comment_text_to_add = "Low confidence API match. Applied standard formatting; please verify manually."
                     source = 'fallback'
                     item = None
                
                # Mix policy
                final_choice_use_api = False
                if item:
                    candidate_doi = item.get('DOI') or item.get('doi') or ''
                    candidate_doi = candidate_doi.strip()
                    if candidate_doi:
                        final_choice_use_api = (score >= 0.60)
                    else:
                        final_choice_use_api = (score >= 0.75)

                if not item or not final_choice_use_api or score < SIMILARITY_MIN:
                    # Use style-parsed data from current para if available.
                    # Note: only safe to call parse_reference_from_styles if no writes
                    # have occurred to para yet in this iteration (which is the case here).
                    parsed_style_data = parse_reference_from_styles(para)
                    if parsed_style_data and parsed_style_data.get('title') and parsed_style_data['title'][0]:
                        if item and score >= 0.60:
                            pass
                        else:
                            item = parsed_style_data
                            source = 'style_parsing'
                            score = 1.0
                            log_lines.append("Using style-based parsing as fallback.\n")
                    else:
                        # FALLBACK — use raw_for_search (already stripped of numbering prefix)
                        # Using plain `raw` would embed the number in the parsed title.
                        comment_text_to_add = "No confident API match found. Applied standard formatting; please verify manually."
                        if source == 'skip_validation':
                            comment_text_to_add = "Validation skipped for Book/Web source. Applied standard formatting."

                        if style == 'REF-N':
                            fallback_parsed = parse_ama_reference_raw(raw_for_search)
                        else:
                            fallback_parsed = parse_apa_reference_raw(raw_for_search)

                        item = fallback_parsed
                        source = 'fallback'
                        score = 0.0
                        
                        if 'title' not in item: item['title'] = fallback_parsed.get('title', [''])
                        if 'URL' not in item: item['URL'] = ''
                        if 'DOI' not in item: item['DOI'] = ''

                # Safety Check
                if source not in ('style_parsing', 'fallback', 'manual_fallback') and item:
                    if style == 'REF-N': 
                         orig_parsed = parse_ama_reference_raw(raw)
                    else:
                         orig_parsed = parse_apa_reference_raw(raw)
                    
                    orig_auth_str = (orig_parsed.get('authors') or '').lower()
                    _ot = orig_parsed.get('title', '')
                    orig_title = (_ot[0] if isinstance(_ot, list) and _ot else (_ot or '')).lower()
                    
                    cand_title = (item.get('title') or [''])[0]
                    cand_authors = item.get('author', [])
                    
                    real_sim = similarity(orig_title, cand_title)

                    # BYPASS SIMILARITY CHECK FOR GEMINI
                    # Gemini might rewrite titles (e.g. expanding abbrevs), leading to mismatches.
                    # We trust Gemini's intent if source is 'gemini_ai'.
                    if source == 'gemini_ai':
                        real_sim = 1.0
                        author_match = True  # skip author check — Gemini result is trusted
                    else:
                        # Normal API path: check author match independently
                        author_match = False
                        if cand_authors:
                            first_family = cand_authors[0].get('family', '').lower()
                            if first_family and first_family in orig_auth_str:
                                author_match = True
                            if not author_match:
                                for a in cand_authors[:3]:  # check first 3 authors
                                    if a.get('family', '').lower() in orig_auth_str:
                                        author_match = True
                                        break
                    
                    # Logic:
                    # - If Title Sim < 0.5 -> Reject (Too different)
                    # - If Title Sim < 0.8 AND Author Match is False -> Reject (Different paper)
                    # - If Score > 0.95 (High Confidence) -> be more lenient?
                    
                    reject_candidate = False
                    
                    if real_sim < 0.4: 
                        reject_candidate = True # Very wildly different title
                        log_lines.append(f"Rejected API result due to Title Mismatch (Sim: {real_sim:.2f}).\n")
                        
                    elif real_sim < 0.8 and not author_match:
                        # Title is somewhat different AND Author doesn't match -> Likely wrong paper
                        reject_candidate = True
                        log_lines.append(f"Rejected API result due to Author Mismatch + Low Title Sim ({real_sim:.2f}).\n")
                        
                    elif score < 0.9 and not author_match:
                         # Even if title is decent (0.8+), if author is completely wrong and score isn't super high -> Reject
                         # E.g. "Similar title" but different authors
                         reject_candidate = True
                         log_lines.append(f"Rejected API result: Good title ({real_sim:.2f}) but Author Mismatch & Score < 0.9.\n")

                    if reject_candidate:
                        # Revert to original / annotate
                        log_lines.append("Safety Check Failed. Reverting to fallback styling.\n")
                        # unresolved.append(raw) # User requested to not report these as validation errors if styling is applied
                        comment_text_to_add = "Validation Failed: Matched metadata differs significantly from original text. Applied standard formatting."
                        
                        # Force fallback
                        source = 'fallback'
                        
                        # We need to re-parse 'raw' to get the fallback structure
                        if style == 'REF-N': 
                             item = parse_ama_reference_raw(raw)
                        else:
                             item = parse_apa_reference_raw(raw)
                        
                        # Ensure basic fields exist for flow compatibility
                        if 'title' not in item: item['title'] = [raw] # fallback title
                        if 'URL' not in item: item['URL'] = ''
                        if 'DOI' not in item: item['DOI'] = ''
                        
                        # Proceed to write...
                        reject_candidate = False # handled style parsing fallback handling below checks checks check checks check
                        # We don't continue; we let it flow down to Generation

                # -----------------------------------------------------

                # ensure item has title; fallback to parsed title if missing
                cr_item = item
                current_title = (cr_item.get('title') or [''])[0] if isinstance(cr_item.get('title'), list) else (cr_item.get('title') or '')
                if not current_title or current_title == 'No title available' or current_title.lower().startswith('published '):
                    # attempt to fill from original parsing
                    if style == 'REF-N':
                        fallback_parsed = parse_ama_reference_raw(raw)
                    else:
                        fallback_parsed = parse_apa_reference_raw(raw)
                    _ft = fallback_parsed.get('title', [''])
                    _ft_str = _ft[0] if isinstance(_ft, list) and _ft else (_ft or '')
                    if _ft_str:
                        cr_item['title'] = [_ft_str]
                        log_lines.append("Filled missing title from original parsing.\n")
                    else:
                        # cannot fill title -> leave original (annotate)
                        log_lines.append("Could not extract title from API or styles; leaving original and annotating.\n")
                        unresolved.append(raw)
                        comment_text = "Reference not updated — missing or invalid title in metadata."
                        added = try_add_word_comment(doc, para, comment_text, author="RefFix", initials="RF")
                        if not added:
                            para.add_run(f" [COMMENT: {comment_text}]").style = 'bib_comment' if 'bib_comment' in doc.styles else None
                        continue

                # generate citation segments
                if source in ('fallback', 'manual_fallback'):
                     # Use raw_for_search to prevent numbering issues
                     # For non-journal types (book, web, thesis, edited_book, other),
                     # generate_fallback_citation handles them correctly using raw text
                     segments = generate_fallback_citation(item, raw_for_search, style)
                else:
                    # API-validated references (journals) — use structured generators
                    # Determine Output Style
                    start_style = style
                    output_style = start_style
                    
                    if target_style == 'AMA':
                        output_style = 'REF-N'
                    elif target_style == 'APA':
                        output_style = 'REF-U'

                    if output_style == 'REF-N':
                        segments = generate_ama_citation(cr_item)
                    elif output_style == 'REF-U':
                        segments = generate_apa_citation(cr_item)
                    elif output_style == 'REF':
                        segments = generate_chicago_citation(cr_item)
                    else:
                        segments = generate_ama_citation(cr_item)
                
                # --- RE-ATTACH NUMBERING ---
                # --- RE-ATTACH NUMBERING ---
                if current_numbering_prefix:
                    # Only number (and period) should have bib_number style, not the tab
                    m_digits = re.search(r'\d+', current_numbering_prefix)
                    if m_digits:
                        number_part = f"{m_digits.group(0)}."
                        # Insert number with bib_number style
                        segments.insert(0, (number_part, 'bib_number'))
                        # Insert tab without style (after the number)
                        segments.insert(1, ("\t", None))
                    else:
                        # Fallback if no digits found
                        segments.insert(0, (current_numbering_prefix, 'bib_number'))

                full_text = "".join(s[0] for s in segments)
                final_url = cr_item.get('URL')
                if not final_url and cr_item.get('DOI'):
                    doi_val = cr_item['DOI']
                    if isinstance(doi_val, str) and doi_val.startswith('http'):
                        final_url = doi_val
                    else:
                        final_url = f"https://doi.org/{doi_val}"

                if not final_url:
                    m = re.search(r'https?://\S+', full_text)
                    if m:
                        final_url = m.group(0).rstrip('.,;)')


                validation_msg = ""
                if final_url:
                    # OPTIMIZATION: Skip rigorous URL validation to speed up processing
                    # validation_msg is only for logging.
                    # is_valid = validate_url(final_url, timeout=3) 
                    # For now, just assume valid if syntax is okay
                    is_valid = True 
                    if not is_valid:
                        validation_msg = f" [WARNING: URL validation failed for {final_url}]"
                    else:
                        validation_msg = " [URL Syntax Valid]"

                # Write back
                preserve_styles = (source == 'style_parsing')
                write_citation_with_styles(para, segments, preserve_original_styles=preserve_styles, styles=doc.styles, doc=doc)
                
                # Add Comment if flagged (AFTER writing, so anchors attach to new runs)
                if comment_text_to_add:
                     try_add_word_comment(doc, para, comment_text_to_add, author="RefFix", initials="RF")
                     log_lines.append(f"Added comment: {comment_text_to_add}\n")

                log_lines.append(f"Source: {source}, Score: {score:.3f}{validation_msg}\nNew: {full_text}\n")
                changed += 1

            except Exception as e:
                log_lines.append(f"ERROR processing reference: {repr(e)}\n")
                unresolved.append(raw)


    try:
        doc.save(output_docx)
        logger.info(f"Saved output to: {output_docx}")
    except Exception as e:
        logger.error(f"Failed to save document: {repr(e)}")
        raise

    # --- DUPLICATE DETECTION ---
    try:
        if all_ref_texts:
            log_lines.append("\n" + "="*30 + "\nDUPLICATE REFERENCES CHECK\n" + "="*30 + "\n")
            duplicates = find_duplicates(all_ref_texts)
            if duplicates:
                log_lines.append(f"Found {len(duplicates)} potential duplicates:\n")
                for d in duplicates:
                    log_lines.append(f"  Ref #{d['id']} is duplicate of #{d['duplicate_of']} (Score: {d['score']}%)\n")
                    log_lines.append(f"    Text: {d['text']}\n")
            else:
                log_lines.append("No duplicates found.\n")
    except Exception as e:
        log_lines.append(f"\nError checking duplicates: {e}\n")

    # write log
    header = [
        f"Reference Fix Log",
        f"Input: {input_docx}",
        f"Output: {output_docx}",
        f"Timestamp: {datetime.now().isoformat()}",
        f"Total styled refs encountered: {total}",
        f"Total changed: {changed}",
        "-" * 60,
        ""
    ]
    with open(log_file, 'w', encoding='utf-8') as fh:
        fh.write("\n".join(header))
        fh.write("\n".join(log_lines))
        if unresolved:
            fh.write("\nUnresolved references (no confident match):\n")
            for u in unresolved:
                fh.write(u + "\n")

    logger.info("Done.")
    logger.info("Input : %s", input_docx)
    logger.info("Output: %s", output_docx)
    logger.info("Log   : %s", log_file)
    logger.info("Total refs: %d, Changed: %d, Unresolved: %d", total, changed, len(unresolved))

    return {
        'output_docx': output_docx,
        'log_file': log_file
    }

# -------------------------
# CLI runner
# -------------------------
if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Fix references in a DOCX file using PubMed/CrossRef.")
    parser.add_argument("input_file", nargs='?', help="Path to the input .docx file")
    args = parser.parse_args()

    if args.input_file:
        input_path = Path(args.input_file)
        if not input_path.exists():
            print(f"Error: File not found: {input_path}")
            exit(1)
        process_docx_file(input_path)
    else:
        print("Usage: python ReferencesStructing.py <input_docx>")
