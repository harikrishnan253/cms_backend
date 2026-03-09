
import re
import os
import uuid
from datetime import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# Namespace map for OXML
NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
}

class TechnicalEditor:
    def __init__(self):
        self.revision_id_counter = 0
        # Define patterns and their default options
        # Format: key -> { 'label': str, 'pattern': regex, 'options': [], 'type': 'regex'|'func', 'category': str }
        self.rules = {
            # --- Punctuation & Symbols ---
            "unpaired_quotes": {
                "label": "Unpaired Quotes",
                "pattern": None, # Handled by function
                "type": "function",
                "func": self._check_unpaired_quotes,
                "category": "Punctuation & Symbols",
                "options": ["Ignore", "Highlight"] 
            },
            "special_chars": {
                "label": "Special Characters (Non-ASCII)",
                "pattern": r"[^\x00-\x7F]+",
                "category": "Punctuation & Symbols",
                "options": ["Ignore", "Highlight"]
            },
            "trademark": {
                "label": "Trademark Symbols (®©™)",
                "pattern": r"[®©™]",
                "category": "Punctuation & Symbols",
                "options": ["Ignore", "Highlight"]
            },
            "greek": {
                "label": "Greek Text/Symbols",
                "pattern": r"[\u0370-\u03FF]+", # Greek block
                "category": "Punctuation & Symbols",
                "options": ["Ignore", "Highlight"]
            },
            "unpaired_parens": { # Added for completeness
                "label": "Unpaired Parentheses",
                "pattern": None,
                "type": "function",
                "func": self._check_unpaired_parens,
                "category": "Punctuation & Symbols",
                "options": ["Ignore", "Highlight"]
            },

            # --- Numbers & Units ---
            "thousand_separator": {
                "label": "Thousand Separator (e.g., 1000 -> 1,000)",
                "pattern": r"\b\d{4,}\b", # Matches 4+ digits. Need verification logic to avoid years.
                "category": "Numbers & Units",
                "options": ["x,xxx", "xxxx"]
            },
            "spell_out_numbers": {
                "label": "Spell out numbers 0-99",
                "pattern": r"\b([0-9]|[1-9][0-9])\b",
                "category": "Numbers & Units",
                "options": ["Spell Out", "Keep Digits"]
            },
             "degree_symbol": {
                "label": "Degree Symbol",
                "pattern": r"\b\d+\s*(degrees?|°)\b", 
                "category": "Numbers & Units",
                "options": ["°", "degrees"] 
            },
            "percent_symbol": {
                "label": "Percent Symbol",
                "pattern": r"\b(percent|per cent|%)\b",
                "category": "Numbers & Units",
                "options": ["%", "percent", "per cent"]
            },
            "units_mmHg": {
               "label": "mmHg Spacing",
               "pattern": r"\bmm\s*Hg\b",
               "category": "Numbers & Units",
               "options": ["mmHg", "mm Hg"]
            },
             "micro_units": {
               "label": "Micro Units (mcg/mcm/µL)",
               "pattern": r"\b(mcg|mcm|mcl|micron|micro)\b",
               "category": "Numbers & Units",
               "options": ["µg", "µm", "µL", "mcg", "mcm"]
            },
            "per_usage": {
                "label": "'Per' Usage (per dL vs /dL)",
                "pattern": r"\b(per|/)\s*(dL|L|min|mL|µL)\b",
                "category": "Numbers & Units",
                "options": ["per unit", "/unit"]
            },

            # --- Abbreviation & Terminology ---
            "latin_abbrev": {
                "label": "Latin Abbreviations (e.g., i.e.)",
                "pattern": r"\b(e\.g\.|eg|i\.e\.|ie)\b",
                "category": "Abbreviation & Terminology",
                "options": ["e.g.", "i.e.", "eg", "ie", "Highlight"]
            },
             "xray": {
                "label": "X-Ray Terminology",
                "pattern": r"\b[xX][- ]?ray\b",
                "category": "Abbreviation & Terminology",
                "options": ["X-ray", "X Ray", "x ray"]
            },
             "vs": {
                "label": "Versus",
                "pattern": r"\b(vs\.|versus|v\.|vs)\b",
                "category": "Abbreviation & Terminology",
                "options": ["vs.", "versus", "v.", "vs"]
            },
             "chemical_prefixes": {
                "label": "Chemical Prefixes (D-, L-, N-)",
                "pattern": r"\b[DLN]-\w+", # Matches D-Glucose etc. Should check for Small Caps?
                "category": "Abbreviation & Terminology",
                "options": ["Small Caps Prefix", "Regular Prefix"]
            },
             "chemical_formulas": {
                "label": "Chemical Formulas (O2, CO2)",
                "pattern": r"\b(CO2|O2|H2O)\b", # Basic list
                "category": "Abbreviation & Terminology",
                "options": ["Subscript Numbers", "Small Caps (User Request)", "Regular"]
            },
             "time_format": {
                "label": "Time Format (AM/PM)",
                "pattern": r"\b(AM|PM|a\.m\.|p\.m\.|am|pm)\b",
                "category": "Abbreviation & Terminology",
                "options": ["SMALL CAPS (AM/PM)", "a.m./p.m.", "AM/PM"]
            },
             "dsm_covid": {
                 "label": "Terms (DSM, COVID-19)",
                 "pattern": r"\b(DSM|COVID[- ]?19)\b",
                 "category": "Abbreviation & Terminology",
                 "options": ["DSM", "COVID-19"]
             },

            # --- Citations & References ---
            "fig_table_cite": {
                "label": "Figure/Table Citations",
                "pattern": r"\b(Figure|Fig\.|Table|Tab\.)\s*\d+",
                "category": "Citations & References",
                "options": ["Figure X", "Fig. X", "Highlight"]
            },
            "cross_ref": {
                "label": "Cross References (Section/Chapter)",
                "pattern": r"\b(Section|section|Chapter|chapter)\s*\d+",
                "category": "Citations & References",
                "options": ["Capitalized", "lowercase"]
            },
             "p_value": {
                "label": "P Value Formatting",
                "pattern": r"\b[pP][- ]?value\b",
                "category": "Citations & References",
                "options": ["P value", "p value", "p-value"]
            },
            
            # --- Comparisons ---
            "comparison": {
                "label": "Comparison Operators",
                "pattern": r"(>|<|≥|≤|=|less than|greater than|equal to)",
                "category": "Comparisons",
                "options": ["<", ">", "=", "less than", "greater than"]
            }
        }

    def scan(self, file_path: str) -> dict:
        """
        Scans values in the document and returns found patterns.
        """
        doc = Document(file_path)
        full_text = []
        for p in doc.paragraphs:
            full_text.append(p.text)
        text = "\n".join(full_text)

        results = {}
        
        # 1. Regex Checks
        for key, rule in self.rules.items():
            if rule.get('type') == 'function':
                # Functional check
                check_res = rule['func'](text)
                if check_res:
                    results[key] = {
                        "label": rule['label'],
                        "found": check_res['found'],
                        "count": check_res['count'],
                        "options": rule['options'],
                        "category": rule['category']
                    }
            else:
                # Regex check
                matches = list(set(re.findall(rule['pattern'], text, re.IGNORECASE)))
                if matches:
                    results[key] = {
                        "label": rule['label'],
                        "found": matches,
                        "options": rule['options'],
                        "count": len(re.findall(rule['pattern'], text, re.IGNORECASE)),
                        "category": rule['category']
                    }
        
        return results

    def _check_unpaired_quotes(self, text):
        count_double = text.count('"')
        count_single = text.count("'") # This is tricky due to apostrophes
        
        # Heuristic for detecting smart quotes vs straight quotes
        count_smart_double_open = text.count('“')
        count_smart_double_close = text.count('”')
        
        issues = []
        if count_double % 2 != 0:
            issues.append(f"Unpaired straight quotes (\") found: {count_double}")
        
        if count_smart_double_open != count_smart_double_close:
             issues.append(f"Unpaired smart quotes (open={count_smart_double_open}, close={count_smart_double_close})")

        if issues:
            return {"found": issues, "count": len(issues)}
        return None

    def _check_unpaired_parens(self, text):
        open_p = text.count('(')
        close_p = text.count(')')
        if open_p != close_p:
            return {"found": [f"Open: {open_p}, Close: {close_p}"], "count": abs(open_p - close_p)}
        return None

    def process(self, input_path: str, output_path: str, user_choices: dict, author="System"):
        """
        Applies replacements with Track Changes.
        """
        doc = Document(input_path)
        
        # Logic to separate Regex replacements vs Formatting/Functional changes
        
        regex_replacements = {}
        formatting_rules = {}

        for key, choice in user_choices.items():
            if choice == "Ignore": continue
            
            if key in self.rules:
                rule = self.rules[key]
                if rule.get('type') == 'function':
                    # Functional rules (e.g. highlight unpaired matches)
                    # For V1, we might just SKIP functional rules in "Apply" unless they imply a fix.
                    # "Highlight" choice would mean adding a comment or highlight style?
                    pass
                elif "Small Caps" in choice:
                     # Handle formatting rules separately
                     formatting_rules[key] = {
                         'pattern': re.compile(rule['pattern'], re.IGNORECASE),
                         'action': 'small_caps'
                     }
                elif "Subscript" in choice:
                     formatting_rules[key] = {
                         'pattern': re.compile(rule['pattern'], re.IGNORECASE),
                         'action': 'subscript'
                     }
                if choice == "Highlight":
                    # Handle Highlight as a formatting rule
                    # Only add if we have a pattern to highlight!
                    if not rule.get('pattern'):
                         # Functional rule without pattern -> cannot highlight easily with current logic
                         continue
                         
                    formatting_rules[key] = {
                        'pattern': re.compile(rule['pattern'], re.IGNORECASE),
                        'action': 'highlight'
                    }
                    continue

                if choice in ["Spell Out", "Keep Digits", "x,xxx", "xxxx"]:
                     continue
                
                t_pattern = rule['pattern']
                regex_replacements[key] = {
                    'pattern': re.compile(t_pattern, re.IGNORECASE),
                    'replacement': choice
                }
        
        if not regex_replacements and not formatting_rules:
            doc.save(output_path)
            return

        for para in doc.paragraphs:
            self._process_paragraph(para, regex_replacements, formatting_rules, author)
            
        doc.save(output_path)

    def _process_paragraph(self, para, regex_replacements, formatting_rules, author):
        """
        Replaces text and applies formatting in a paragraph with Track Changes.
        """
        
        # 1. Apply Formatting Rules first (non-destructive to text, usually)
        # Actually changing formatting with track changes is tricky (requires modifying rPr with w:rPrChange).
        # For simplicity in this version, we will focus on Text Replacement.
        # Formatting updates will use the same "Delete + Insert new run with formatting" approach.
        
        combined_actions = {}
        combined_actions.update(regex_replacements)
        combined_actions.update(formatting_rules) # These have 'action': 'small_caps' instead of 'replacement'

        for key, rule in combined_actions.items():
            pattern = rule['pattern']
            
            # Prepare actions
            target_text = rule.get('replacement')
            special_action = rule.get('action')

            # We must iterate a COPY of runs because we might modify them
            for run in list(para.runs):
                text = run.text
                if not text: continue
                    
                matches = list(pattern.finditer(text))
                if not matches: continue
                
                new_runs = []
                cursor = 0
                
                for m in matches:
                    start, end = m.span()
                    original_substr = text[start:end]
                    
                    # Determine replacement Text
                    if special_action:
                        # For formatting changes, text stays same (or modified slightly)
                        final_text = original_substr 
                        if "small_caps" in special_action and "AM/PM" in str(key): 
                            final_text = original_substr.upper() # Ensure AM/PM
                    else:
                        final_text = target_text
                    
                    if original_substr == final_text and not special_action:
                        continue

                    # Pre-match
                    if start > cursor:
                        self._create_run_copy(para, run, text[cursor:start], new_runs)
                        
                    if special_action == 'highlight':
                        # For Highlight, DO NOT use Track Changes (Delete/Insert) to avoid text duplication in "All Markup" view.
                        # Just create a new run with the highlight property.
                        from docx.enum.text import WD_COLOR_INDEX
                        
                        # Create the run but DO NOT wrap in w:ins or w:del
                        hl_run = self._create_run_copy(para, run, original_substr, new_runs, add_to_list=False)
                        hl_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        new_runs.append(hl_run)
                    
                    else:
                        # Standard Track Changes (Delete + Insert) for replacements
                        
                        # DELETION
                        del_run = self._create_run_copy(para, run, original_substr, new_runs, add_to_list=False)
                        self._mark_delete(del_run, author, datetime.now())
                        new_runs.append(del_run)
                        
                        # INSERTION (with potentially new formatting)
                        ins_run = self._create_run_copy(para, run, final_text, new_runs, add_to_list=False)
                        
                        # Apply Special Formatting
                        if special_action == 'small_caps':
                            ins_run.font.small_caps = True
                        elif special_action == 'subscript':
                            pass 
                            
                        self._mark_insert(ins_run, author, datetime.now())
                        new_runs.append(ins_run)
                    
                    cursor = end
                
                if cursor < len(text):
                     self._create_run_copy(para, run, text[cursor:], new_runs)
                     
                if new_runs:
                    p_element = para._element
                    try:
                        index = p_element.index(run._element)
                        p_element.remove(run._element)
                        for nr in reversed(new_runs):
                            p_element.insert(index, nr._element)
                    except ValueError:
                        pass

    def _create_run_copy(self, para, source_run, text, list_to_append, add_to_list=True):
        """Creates a new run with same formatting as source."""
        new_run = para.add_run(text)
        # Copy basic formatting
        if source_run.style: new_run.style = source_run.style
        if source_run.bold: new_run.bold = source_run.bold
        if source_run.italic: new_run.italic = source_run.italic
        if source_run.underline: new_run.underline = source_run.underline
        if source_run.font.color.rgb: new_run.font.color.rgb = source_run.font.color.rgb
        if source_run.font.size: new_run.font.size = source_run.font.size
        
        if add_to_list:
            list_to_append.append(new_run)     
        return new_run

    def _mark_delete(self, run, author, date):
        """Wraps run in <w:del>"""
        self.revision_id_counter += 1
        del_node = OxmlElement('w:del')
        del_node.set(qn('w:id'), str(self.revision_id_counter))
        del_node.set(qn('w:author'), author)
        del_node.set(qn('w:date'), date.replace(microsecond=0).isoformat())
        
        # In a deleted run, text should be in <w:delText> not <w:t>
        # Iterate over children and rename w:t to w:delText
        for child in run._element:
            if child.tag == qn('w:t'):
                child.tag = qn('w:delText')
                
        del_node.append(run._element)
        
        class XmlWrapper:
            def __init__(self, element): self._element = element
        return XmlWrapper(del_node)

    def _mark_insert(self, run, author, date):
        """Wraps run in <w:ins>"""
        self.revision_id_counter += 1
        ins_node = OxmlElement('w:ins')
        ins_node.set(qn('w:id'), str(self.revision_id_counter))
        ins_node.set(qn('w:author'), author)
        ins_node.set(qn('w:date'), date.replace(microsecond=0).isoformat())
        ins_node.append(run._element)
        
        class XmlWrapper:
            def __init__(self, element): self._element = element
        return XmlWrapper(ins_node)

