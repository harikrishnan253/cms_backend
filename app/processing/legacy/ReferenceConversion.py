import os
import json
import logging
import re
from typing import Optional, Dict, List, Tuple
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

from gemini_ref_converter import convert_reference, CitationStyle, BIB_FIELDS

# ─────────────────────────────────────────────
# LOGGING
# ─────────────────────────────────────────────
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s: %(message)s")
logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────
# REFERENCE TYPE DETECTION (fallback if Gemini fails)
# ─────────────────────────────────────────────

def detect_source_style(raw_text: str) -> CitationStyle:
    """
    Heuristically detect whether a reference is AMA or APA formatted.
    AMA: typically starts with number or surname without comma-year pattern.
    APA: typically has (Year) pattern near start, author "Surname, I." format.
    """
    stripped = raw_text.strip()

    # AMA indicators
    if re.match(r'^\[?\d+\]?\.?\s+', stripped):            # numbered reference
        return CitationStyle.AMA
    if re.search(r'\.\s+\d{4};', stripped):                 # Year;Volume pattern
        return CitationStyle.AMA
    if re.search(r'\bpp?\.\s*\d+[-–]\d+', stripped):        # pp. 12-34 rare in AMA
        return CitationStyle.APA

    # APA indicators
    if re.search(r'\(\d{4}[a-z]?\)', stripped):             # (Year) or (2022a)
        return CitationStyle.APA
    if re.search(r'https://doi\.org/', stripped):           # APA uses full DOI URL
        return CitationStyle.APA
    if re.search(r',\s+\d+\(\d+\),\s+\d+', stripped):     # volume(issue), page
        return CitationStyle.APA

    # Default to APA
    return CitationStyle.APA


def detect_ref_type_from_metadata(metadata: Dict) -> str:
    """Derive a human-readable ref type label from bib_ metadata."""
    return metadata.get("bib_reftype") or "unknown"


# ─────────────────────────────────────────────
# FORMATTING FROM METADATA
# (Replaces the broken item-mapping + generate_*_citation approach)
# ─────────────────────────────────────────────

def format_apa_from_metadata(meta: Dict) -> str:
    """
    Build APA 7th edition string directly from bib_ metadata fields.
    Used as fallback if Gemini's formatted_output is empty.
    """
    ref_type = meta.get("bib_reftype", "journal")
    parts = []

    # Authors
    surnames = [s.strip() for s in (meta.get("bib_surname") or "").split("|") if s.strip()]
    fnames   = [f.strip() for f in (meta.get("bib_fname")   or "").split("|") if f.strip()]
    authors  = []
    for i, surname in enumerate(surnames):
        initial = fnames[i] if i < len(fnames) else ""
        # Format initials: "John A" → "J. A."
        initials_fmt = " ".join(f"{p[0]}." for p in initial.split() if p) if initial else ""
        authors.append(f"{surname}, {initials_fmt}".strip(", "))

    if authors:
        if len(authors) > 20:
            author_str = ", ".join(authors[:19]) + ", ... " + authors[-1]
        elif len(authors) > 1:
            author_str = ", ".join(authors[:-1]) + ", & " + authors[-1]
        else:
            author_str = authors[0]
        parts.append(author_str + ".")

    # Year
    year = meta.get("bib_year", "n.d.")
    parts.append(f"({year}).")

    if ref_type == "journal":
        title   = meta.get("bib_title", "")
        journal = meta.get("bib_journal", "")
        volume  = meta.get("bib_volume", "")
        issue   = meta.get("bib_issue", "")
        fpage   = meta.get("bib_fpage", "")
        lpage   = meta.get("bib_lpage", "")
        doi     = meta.get("bib_doi", "")

        if title:   parts.append(f"{title}.")
        vol_issue   = f"*{journal}*" if journal else ""
        if volume:  vol_issue += f", *{volume}*"
        if issue:   vol_issue += f"({issue})"
        pages = f"{fpage}–{lpage}" if fpage and lpage else fpage or lpage
        if pages:   vol_issue += f", {pages}"
        if vol_issue: parts.append(vol_issue + ".")
        if doi:     parts.append(f"https://doi.org/{doi}")

    elif ref_type in ("book", "edited_book"):
        book_title = meta.get("bib_book") or meta.get("bib_title", "")
        edition    = meta.get("bib_editionno", "")
        publisher  = meta.get("bib_publisher", "")
        doi        = meta.get("bib_doi", "")
        url        = meta.get("bib_url", "")
        ed_surnames = [s.strip() for s in (meta.get("bib_ed_surname") or "").split("|") if s.strip()]
        ed_fnames   = [f.strip() for f in (meta.get("bib_ed_fname")   or "").split("|") if f.strip()]

        if ref_type == "edited_book" and ed_surnames and not authors:
            eds = []
            for i, s in enumerate(ed_surnames):
                ini = ed_fnames[i] if i < len(ed_fnames) else ""
                eds.append(f"{s}, {ini}".strip(", "))
            ed_label = "(Ed.)" if len(eds) == 1 else "(Eds.)"
            parts.insert(0, f"{', '.join(eds)} {ed_label}.")

        title_str = f"*{book_title}*" if book_title else ""
        if edition: title_str += f" ({_ordinal(edition)} ed.)"
        if title_str: parts.append(title_str + ".")
        if publisher: parts.append(publisher + ".")
        if doi:       parts.append(f"https://doi.org/{doi}")
        elif url:     parts.append(url)

    elif ref_type == "book_chapter":
        chapter   = meta.get("bib_chaptertitle") or meta.get("bib_title", "")
        book      = meta.get("bib_book", "")
        edition   = meta.get("bib_editionno", "")
        fpage     = meta.get("bib_fpage", "")
        lpage     = meta.get("bib_lpage", "")
        publisher = meta.get("bib_publisher", "")
        doi       = meta.get("bib_doi", "")
        ed_surnames = [s.strip() for s in (meta.get("bib_ed_surname") or "").split("|") if s.strip()]
        ed_fnames   = [f.strip() for f in (meta.get("bib_ed_fname")   or "").split("|") if f.strip()]

        if chapter: parts.append(f"{chapter}.")
        editors = []
        for i, s in enumerate(ed_surnames):
            ini = ed_fnames[i] if i < len(ed_fnames) else ""
            ini_fmt = " ".join(f"{p[0]}." for p in ini.split() if p) if ini else ""
            editors.append(f"{ini_fmt} {s}".strip())
        ed_label = "Ed." if len(editors) == 1 else "Eds."
        in_str = "In " + ", ".join(editors) + f" ({ed_label}.), " if editors else "In "
        book_str = f"*{book}*" if book else ""
        if edition: book_str += f" ({_ordinal(edition)} ed.)"
        pages = f"pp. {fpage}–{lpage}" if fpage and lpage else (f"pp. {fpage}" if fpage else "")
        chunk = in_str + book_str
        if pages: chunk += f" ({pages})"
        parts.append(chunk + ".")
        if publisher: parts.append(publisher + ".")
        if doi:       parts.append(f"https://doi.org/{doi}")

    elif ref_type == "thesis":
        title   = meta.get("bib_title", "")
        deg     = meta.get("bib_deg", "Doctoral dissertation")
        school  = meta.get("bib_school", "")
        url     = meta.get("bib_url", "")
        if title:  parts.append(f"*{title}* [{deg}, {school}]." if school else f"*{title}* [{deg}].")
        if url:    parts.append(url)

    elif ref_type == "conference":
        title       = meta.get("bib_title", "")
        conf        = meta.get("bib_conference", "")
        confloc     = meta.get("bib_conflocation", "")
        confdate    = meta.get("bib_confdate", "")
        doi         = meta.get("bib_doi", "")
        if title:   parts.append(f"*{title}* [Conference session].")
        conf_str = conf
        if confdate:   conf_str += f", {confdate}"
        if confloc:    conf_str += f", {confloc}"
        if conf_str:   parts.append(conf_str + ".")
        if doi:        parts.append(f"https://doi.org/{doi}")

    elif ref_type in ("website", "ereference"):
        title    = meta.get("bib_title", "")
        site     = meta.get("bib_journal") or meta.get("bib_book", "")
        accessed = meta.get("bib_accessed", "")
        url      = meta.get("bib_url", "")
        if title:    parts.append(f"{title}.")
        if site:     parts.append(f"*{site}*.")
        if accessed: parts.append(f"Retrieved {accessed}, from")
        if url:      parts.append(url)

    elif ref_type == "report":
        title     = meta.get("bib_title", "")
        repnum    = meta.get("bib_reportnum", "")
        inst      = meta.get("bib_institution", "")
        doi       = meta.get("bib_doi", "")
        url       = meta.get("bib_url", "")
        title_str = f"*{title}*" if title else ""
        if repnum: title_str += f" (Report No. {repnum})"
        if title_str: parts.append(title_str + ".")
        if inst:      parts.append(inst + ".")
        if doi:       parts.append(f"https://doi.org/{doi}")
        elif url:     parts.append(url)

    return " ".join(parts)


def format_ama_from_metadata(meta: Dict) -> str:
    """
    Build AMA 11th edition string directly from bib_ metadata fields.
    Used as fallback if Gemini's formatted_output is empty.
    """
    ref_type = meta.get("bib_reftype", "journal")
    parts = []

    # Authors
    surnames = [s.strip() for s in (meta.get("bib_surname") or "").split("|") if s.strip()]
    fnames   = [f.strip() for f in (meta.get("bib_fname")   or "").split("|") if f.strip()]
    authors  = []
    for i, surname in enumerate(surnames):
        initial = fnames[i] if i < len(fnames) else ""
        initials_fmt = "".join(p[0] for p in initial.split() if p) if initial else ""
        authors.append(f"{surname} {initials_fmt}".strip())

    if authors:
        if len(authors) > 6:
            author_str = ", ".join(authors[:6]) + ", et al"
        else:
            author_str = ", ".join(authors)
        parts.append(author_str + ".")

    if ref_type == "journal":
        title   = meta.get("bib_title", "")
        journal = meta.get("bib_journal", "")
        year    = meta.get("bib_year", "")
        volume  = meta.get("bib_volume", "")
        issue   = meta.get("bib_issue", "")
        fpage   = meta.get("bib_fpage", "")
        lpage   = meta.get("bib_lpage", "")
        doi     = meta.get("bib_doi", "")

        if title:   parts.append(f"{title}.")
        vol_str = journal or ""
        if year:    vol_str += f". {year}"
        if volume:  vol_str += f";{volume}"
        if issue:   vol_str += f"({issue})"
        pages = f"{fpage}-{lpage}" if fpage and lpage else fpage or lpage
        if pages:   vol_str += f":{pages}"
        if vol_str: parts.append(vol_str + ".")
        if doi:     parts.append(f"doi:{doi}")

    elif ref_type in ("book", "edited_book"):
        book_title = meta.get("bib_book") or meta.get("bib_title", "")
        edition    = meta.get("bib_editionno", "")
        publisher  = meta.get("bib_publisher", "")
        year       = meta.get("bib_year", "")
        doi        = meta.get("bib_doi", "")
        url        = meta.get("bib_url", "")

        ed_surnames = [s.strip() for s in (meta.get("bib_ed_surname") or "").split("|") if s.strip()]
        ed_fnames   = [f.strip() for f in (meta.get("bib_ed_fname")   or "").split("|") if f.strip()]

        if ref_type == "edited_book" and ed_surnames and not authors:
            eds = []
            for i, s in enumerate(ed_surnames):
                ini = ed_fnames[i] if i < len(ed_fnames) else ""
                initials_fmt = "".join(p[0] for p in ini.split() if p) if ini else ""
                eds.append(f"{s} {initials_fmt}".strip())
            ed_label = "ed." if len(eds) == 1 else "eds."
            parts.append(", ".join(eds) + f", {ed_label}.")

        title_str = book_title or ""
        if edition and edition != "1": title_str += f". {_ordinal(edition)} ed."
        if title_str: parts.append(title_str + ".")
        if publisher: parts.append(publisher + ";")
        if year:      parts.append(year + ".")
        if doi:       parts.append(f"doi:{doi}")
        elif url:     parts.append(url)

    elif ref_type == "book_chapter":
        chapter   = meta.get("bib_chaptertitle") or meta.get("bib_title", "")
        book      = meta.get("bib_book", "")
        edition   = meta.get("bib_editionno", "")
        fpage     = meta.get("bib_fpage", "")
        lpage     = meta.get("bib_lpage", "")
        publisher = meta.get("bib_publisher", "")
        year      = meta.get("bib_year", "")
        doi       = meta.get("bib_doi", "")
        ed_surnames = [s.strip() for s in (meta.get("bib_ed_surname") or "").split("|") if s.strip()]
        ed_fnames   = [f.strip() for f in (meta.get("bib_ed_fname")   or "").split("|") if f.strip()]

        if chapter: parts.append(f"{chapter}.")
        editors = []
        for i, s in enumerate(ed_surnames):
            ini = ed_fnames[i] if i < len(ed_fnames) else ""
            initials_fmt = "".join(p[0] for p in ini.split() if p) if ini else ""
            editors.append(f"{s} {initials_fmt}".strip())
        ed_label = "ed." if len(editors) == 1 else "eds."
        in_str = "In: " + ", ".join(editors) + f", {ed_label}. " if editors else "In: "
        book_str = book or ""
        if edition and edition != "1": book_str += f". {_ordinal(edition)} ed."
        parts.append(in_str + book_str + ".")
        if publisher: parts.append(publisher + ";")
        if year:      parts.append(year + ".")
        pages = f"{fpage}-{lpage}" if fpage and lpage else fpage or lpage
        if pages:     parts[-1] = parts[-1].rstrip(".") + f":{pages}."
        if doi:       parts.append(f"doi:{doi}")

    elif ref_type == "thesis":
        title  = meta.get("bib_title", "")
        deg    = meta.get("bib_deg", "doctoral dissertation")
        school = meta.get("bib_school", "")
        year   = meta.get("bib_year", "")
        url    = meta.get("bib_url", "")
        if title:  parts.append(f"{title} [{deg}].")
        if school: parts.append(school + ";")
        if year:   parts.append(year + ".")
        if url:    parts.append(url)

    elif ref_type == "conference":
        title    = meta.get("bib_title", "")
        conf     = meta.get("bib_conference", "")
        confloc  = meta.get("bib_conflocation", "")
        confdate = meta.get("bib_confdate", "")
        doi      = meta.get("bib_doi", "")
        if title: parts.append(f"{title}.")
        conf_str = f"Paper presented at: {conf}" if conf else ""
        if confdate: conf_str += f"; {confdate}"
        if confloc:  conf_str += f"; {confloc}"
        if conf_str: parts.append(conf_str + ".")
        if doi:      parts.append(f"doi:{doi}")

    elif ref_type in ("website", "ereference"):
        title    = meta.get("bib_title", "")
        site     = meta.get("bib_journal") or meta.get("bib_book", "")
        year     = meta.get("bib_year", "")
        accessed = meta.get("bib_accessed", "")
        url      = meta.get("bib_url", "")
        if title:    parts.append(f"{title}.")
        if site:     parts.append(f"{site}.")
        if year:     parts.append(f"Published {year}.")
        if accessed: parts.append(f"Accessed {accessed}.")
        if url:      parts.append(url)

    elif ref_type == "report":
        title   = meta.get("bib_title", "")
        repnum  = meta.get("bib_reportnum", "")
        inst    = meta.get("bib_institution", "")
        year    = meta.get("bib_year", "")
        doi     = meta.get("bib_doi", "")
        url     = meta.get("bib_url", "")
        if title:  parts.append(f"{title}.")
        if inst:   parts.append(inst + ";")
        if year:   parts.append(year + ".")
        if repnum: parts.append(f"Report No. {repnum}.")
        if doi:    parts.append(f"doi:{doi}")
        elif url:  parts.append(url)

    return " ".join(parts)


def _ordinal(n: str) -> str:
    """Convert "2" → "2nd", "3" → "3rd", etc."""
    try:
        n_int = int(n)
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(n_int % 10 if n_int % 100 not in (11, 12, 13) else 0, "th")
        return f"{n_int}{suffix}"
    except (ValueError, TypeError):
        return str(n)


# ─────────────────────────────────────────────
# PARAGRAPH FORMATTING HELPERS
# ─────────────────────────────────────────────

def _clear_paragraph_text(para) -> None:
    """Remove all runs from a paragraph without disturbing its style."""
    for run in para.runs:
        run.text = ""
    # Remove any extra runs
    p_elem = para._p
    for r in p_elem.findall(qn("w:r")):
        p_elem.remove(r)


def _write_styled_runs(para, segments: List[Tuple[str, Optional[str]]], doc=None) -> None:
    """
    Clear the paragraph and write segments as individual Word character-style runs.

    segments: list of (text, style_name_or_None)
                  style_name should be a bib_* character style name.
    """
    _clear_paragraph_text(para)
    styles = doc.styles if doc is not None else None

    for text, style_name in segments:
        if not text:
            continue
        run = para.add_run(text)
        if style_name:
            try:
                if styles is not None:
                    # Create the character style on-the-fly if it doesn't exist yet
                    try:
                        from docx.enum.style import WD_STYLE_TYPE
                        if style_name not in styles:
                            styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
                        run.style = styles[style_name]
                    except Exception:
                        try:
                            run.style = style_name
                        except Exception:
                            pass
                else:
                    run.style = style_name
            except Exception:
                pass  # Best-effort: leave unstyled if style application fails


def _set_paragraph_text(para, text: str) -> None:
    """
    Fallback: set plain text on a paragraph as a single unstyled run.
    Prefer _write_styled_runs() for production use.
    """
    _clear_paragraph_text(para)
    run = para.add_run(text)
    return run


# ─────────────────────────────────────────────
# SEGMENT BUILDERS  (bib_ metadata → styled segments)
# ─────────────────────────────────────────────


def _split_pipe(value: Optional[str]) -> List[str]:
    """Split a pipe-delimited metadata string, skip blanks."""
    if not value:
        return []
    return [v.strip() for v in value.split("|") if v.strip()]


def build_segments_ama(meta: Dict) -> List[Tuple[str, Optional[str]]]:
    """
    Build AMA 11th-edition styled segments from bib_ metadata.
    Returns list of (text, bib_style_name_or_None) tuples.
    """
    segs: List[Tuple[str, Optional[str]]] = []
    ref_type = (meta.get("bib_reftype") or "journal").lower()

    # ── Authors ──────────────────────────────────────────────────
    surnames = _split_pipe(meta.get("bib_surname"))
    fnames   = _split_pipe(meta.get("bib_fname"))
    n_auth   = len(surnames)

    if n_auth == 0:
        # Organisational / no author
        org = meta.get("bib_organization") or meta.get("bib_institution") or ""
        if org:
            segs.append((org, "bib_organization"))
    else:
        # AMA: ≤6 list all; >6 → first 6 + et al  (strict rule: >6 → first 3, but Gemini already handles this)
        subset = surnames if n_auth <= 6 else surnames[:6]
        for i, surname in enumerate(subset):
            if i > 0:
                segs.append((", ", None))
            segs.append((surname, "bib_surname"))
            initial = fnames[i] if i < len(fnames) else ""
            # AMA initials: "John A" → "JA" (no dots)
            initials_str = "".join(p[0].upper() for p in initial.split() if p)
            if initials_str:
                segs.append((" ", None))
                segs.append((initials_str, "bib_fname"))
        if n_auth > 6:
            segs.append((", ", None))
            segs.append(("et al", "bib_etal"))
    segs.append((". ", None))

    # ── Title ─────────────────────────────────────────────────────
    chapter_title = meta.get("bib_chaptertitle") or ""
    main_title    = meta.get("bib_title") or ""
    book_title    = meta.get("bib_book") or ""

    if ref_type == "book_chapter" and chapter_title:
        segs.append((chapter_title, "bib_chaptertitle"))
        segs.append((". ", None))
    elif main_title:
        segs.append((main_title, "bib_article" if ref_type == "journal" else "bib_title"))
        segs.append((". ", None))

    # ── In: editors (book chapter) ────────────────────────────────
    if ref_type == "book_chapter":
        ed_surnames = _split_pipe(meta.get("bib_ed_surname") or meta.get("bib_ed-surname"))
        ed_fnames   = _split_pipe(meta.get("bib_ed_fname")  or meta.get("bib_ed-fname"))
        segs.append(("In: ", None))
        if ed_surnames:
            for i, es in enumerate(ed_surnames):
                if i > 0:
                    segs.append((", ", None))
                segs.append((es, "bib_ed-surname"))
                ei = ed_fnames[i] if i < len(ed_fnames) else ""
                ei_str = "".join(p[0].upper() for p in ei.split() if p)
                if ei_str:
                    segs.append((" ", None))
                    segs.append((ei_str, "bib_ed-fname"))
            ed_label = "ed." if len(ed_surnames) == 1 else "eds."
            segs.append((", " + ed_label + " ", None))
        if book_title:
            segs.append((book_title, "bib_book"))
            segs.append((". ", None))

    # ── Journal / publisher section ────────────────────────────────
    if ref_type == "journal":
        journal = meta.get("bib_journal") or ""
        year    = meta.get("bib_year") or ""
        volume  = meta.get("bib_volume") or ""
        issue   = meta.get("bib_issue") or ""
        fpage   = meta.get("bib_fpage") or ""
        lpage   = meta.get("bib_lpage") or ""

        if journal:
            segs.append((journal, "bib_journal"))
            segs.append((".", None))
        if year:
            segs.append((" ", None))
            segs.append((year, "bib_year"))
        if volume:
            segs.append((";", None))
            segs.append((volume, "bib_volume"))
        if issue:
            segs.append(("(", None))
            segs.append((issue, "bib_issue"))
            segs.append((") ", None))
        pages_str = f"{fpage}-{lpage}" if fpage and lpage else (fpage or lpage)
        if pages_str:
            segs.append((":", None))
            segs.append((pages_str, "bib_fpage"))
        segs.append((".", None))

    elif ref_type in ("book", "edited_book", "book_chapter"):
        edition   = meta.get("bib_editionno") or ""
        publisher = meta.get("bib_publisher") or ""
        year      = meta.get("bib_year") or ""

        # For plain book, include book title here (chapter already done above)
        if ref_type != "book_chapter" and book_title:
            segs.append((book_title, "bib_book"))
            segs.append((". ", None))
        if edition and edition not in ("1", "1st"):
            segs.append((_ordinal(edition) + " ed. ", "bib_editionno"))
        if publisher:
            segs.append((publisher, "bib_publisher"))
            segs.append(("; ", None))
        if year:
            segs.append((year, "bib_year"))
        # Append page range for chapters
        if ref_type == "book_chapter":
            fpage = meta.get("bib_fpage") or ""
            lpage = meta.get("bib_lpage") or ""
            pages_str = f"{fpage}-{lpage}" if fpage and lpage else (fpage or lpage)
            if pages_str:
                segs.append((":", None))
                segs.append((pages_str, "bib_fpage"))
        segs.append((".", None))

    elif ref_type == "conference":
        conf     = meta.get("bib_conference") or ""
        confloc  = meta.get("bib_conflocation") or ""
        confdate = meta.get("bib_confdate")    or meta.get("bib_year") or ""

        if title := meta.get("bib_title") or "":
            segs.append((title, "bib_confpaper"))
            segs.append((". ", None))
        segs.append(("Paper presented at: ", None))
        if conf:
            segs.append((conf, "bib_conference"))
        if confdate:
            segs.append(("; ", None))
            segs.append((confdate, "bib_confdate"))
        if confloc:
            segs.append(("; ", None))
            segs.append((confloc, "bib_conflocation"))
        segs.append((".", None))

    elif ref_type in ("website", "ereference"):
        title    = meta.get("bib_title") or ""
        year     = meta.get("bib_year") or ""
        accessed = meta.get("bib_accessed") or ""
        url      = meta.get("bib_url") or ""
        if title:
            segs.append((title, "bib_title"))
            segs.append((". ", None))
        if year:
            segs.append(("Published ", None))
            segs.append((year, "bib_year"))
            segs.append((". ", None))
        if accessed:
            segs.append(("Accessed ", None))
            segs.append((accessed, "bib_accessed"))
            segs.append((". ", None))
        if url:
            segs.append((url, "bib_url"))

    # ── DOI ────────────────────────────────────────────────────────
    doi = (meta.get("bib_doi") or "").strip().lstrip("doi:").lstrip()
    if doi and ref_type not in ("website", "ereference"):
        segs.append((" doi:", None))
        segs.append((doi, "bib_doi"))

    return segs


def build_segments_apa(meta: Dict) -> List[Tuple[str, Optional[str]]]:
    """
    Build APA 7th-edition styled segments from bib_ metadata.
    Returns list of (text, bib_style_name_or_None) tuples.
    """
    segs: List[Tuple[str, Optional[str]]] = []
    ref_type = (meta.get("bib_reftype") or "journal").lower()

    # ── Authors ──────────────────────────────────────────────────
    surnames = _split_pipe(meta.get("bib_surname"))
    fnames   = _split_pipe(meta.get("bib_fname"))
    n_auth   = len(surnames)

    if n_auth == 0:
        org = meta.get("bib_organization") or meta.get("bib_institution") or ""
        if org:
            segs.append((org, "bib_organization"))
    else:
        subset = surnames if n_auth <= 20 else surnames[:19]
        for i, surname in enumerate(subset):
            if i > 0:
                segs.append((", ", None))
                if i == n_auth - 1 and n_auth <= 20:
                    segs.append(("& ", None))
            segs.append((surname, "bib_surname"))
            initial = fnames[i] if i < len(fnames) else ""
            # APA initials: "John A" → "J. A."
            parts_i = [p[0].upper() + "." for p in initial.split() if p]
            if parts_i:
                segs.append((", ", None))
                segs.append((" ".join(parts_i), "bib_fname"))
        if n_auth > 20:
            segs.append((", … ", None))
            segs.append((surnames[-1], "bib_surname"))
            last_initial = fnames[-1] if len(fnames) >= n_auth else ""
            parts_i = [p[0].upper() + "." for p in last_initial.split() if p]
            if parts_i:
                segs.append((", ", None))
                segs.append((" ".join(parts_i), "bib_fname"))
    segs.append((" (", None))
    segs.append((meta.get("bib_year") or "n.d.", "bib_year"))
    segs.append(("). ", None))

    # ── Titles ────────────────────────────────────────────────────
    chapter_title = meta.get("bib_chaptertitle") or ""
    main_title    = meta.get("bib_title") or ""
    book_title    = meta.get("bib_book") or ""

    if ref_type == "book_chapter" and chapter_title:
        segs.append((chapter_title, "bib_chaptertitle"))
        segs.append((". ", None))
    elif ref_type in ("book", "edited_book") and book_title:
        segs.append((book_title, "bib_book"))
        segs.append((". ", None))
    elif main_title:
        style = "bib_article" if ref_type == "journal" else "bib_title"
        segs.append((main_title, style))
        segs.append((". ", None))

    # ── In: editors (book chapter) ────────────────────────────────
    if ref_type == "book_chapter":
        ed_surnames = _split_pipe(meta.get("bib_ed_surname") or meta.get("bib_ed-surname"))
        ed_fnames   = _split_pipe(meta.get("bib_ed_fname")   or meta.get("bib_ed-fname"))
        segs.append(("In ", None))
        if ed_surnames:
            for i, es in enumerate(ed_surnames):
                if i > 0:
                    segs.append((", ", None))
                ei = ed_fnames[i] if i < len(ed_fnames) else ""
                parts_i = [p[0].upper() + "." for p in ei.split() if p]
                if parts_i:
                    segs.append((" ".join(parts_i) + " ", "bib_ed-fname"))
                segs.append((es, "bib_ed-surname"))
            ed_label = "(Ed.)," if len(ed_surnames) == 1 else "(Eds.),"
            segs.append((" " + ed_label + " ", None))
        if book_title:
            segs.append((book_title, "bib_book"))
        edition = meta.get("bib_editionno") or ""
        fpage   = meta.get("bib_fpage") or ""
        lpage   = meta.get("bib_lpage") or ""
        if edition and edition not in ("1", "1st"):
            segs.append((f" ({_ordinal(edition)} ed.,", None))
            pages_str = f"pp. {fpage}–{lpage}" if fpage and lpage else (f"pp. {fpage}" if fpage else "")
            segs.append((" " + pages_str + ")", None))
        elif fpage:
            pages_str = f"pp. {fpage}–{lpage}" if lpage else f"pp. {fpage}"
            segs.append((f" ({pages_str})", None))
        segs.append((". ", None))

    # ── Journal section ───────────────────────────────────────────
    if ref_type == "journal":
        journal = meta.get("bib_journal") or ""
        volume  = meta.get("bib_volume") or ""
        issue   = meta.get("bib_issue") or ""
        fpage   = meta.get("bib_fpage") or ""
        lpage   = meta.get("bib_lpage") or ""

        if journal:
            segs.append((journal, "bib_journal"))
        if volume:
            segs.append((", ", None))
            segs.append((volume, "bib_volume"))
        if issue:
            segs.append(("(", None))
            segs.append((issue, "bib_issue"))
            segs.append((") ", None))
        pages_str = f"{fpage}–{lpage}" if fpage and lpage else (fpage or lpage)
        if pages_str:
            pages_str = pages_str.replace("-", "–")
            segs.append((", ", None))
            segs.append((pages_str, "bib_fpage"))
        segs.append((".", None))

    elif ref_type in ("book", "edited_book"):
        edition   = meta.get("bib_editionno") or ""
        publisher = meta.get("bib_publisher") or ""
        if edition and edition not in ("1", "1st"):
            segs.append((f"({_ordinal(edition)} ed.). ", "bib_editionno"))
        if publisher:
            segs.append((publisher, "bib_publisher"))
            segs.append((". ", None))

    elif ref_type in ("website", "ereference"):
        site     = meta.get("bib_journal") or meta.get("bib_book") or ""
        accessed = meta.get("bib_accessed") or ""
        url      = meta.get("bib_url") or ""
        if site:
            segs.append((site, "bib_journal"))
            segs.append((". ", None))
        if accessed:
            segs.append(("Retrieved " + accessed + ", from ", None))
        if url:
            segs.append((url, "bib_url"))

    elif ref_type == "conference":
        conf    = meta.get("bib_conference") or ""
        confloc = meta.get("bib_conflocation") or ""
        confdate = meta.get("bib_confdate") or ""
        segs.append(("[Conference session]. ", None))
        if conf:
            segs.append((conf, "bib_conference"))
        if confdate:
            segs.append((", " + confdate, None))
        if confloc:
            segs.append((", " + confloc, None))
        segs.append((".", None))

    # ── DOI / URL ─────────────────────────────────────────────────
    doi = (meta.get("bib_doi") or "").strip().lstrip("doi:").lstrip()
    url = meta.get("bib_url") or ""
    if doi:
        segs.append((" https://doi.org/", "bib_doi"))
        segs.append((doi, "bib_doi"))
    elif url and ref_type not in ("website", "ereference"):
        segs.append((" ", None))
        segs.append((url, "bib_url"))

    return segs


# ─────────────────────────────────────────────
# CONVERSION LOG ENTRY
# ─────────────────────────────────────────────

class ConversionLogEntry:
    def __init__(self, original: str, converted: str, ref_type: str,
                 source_style: str, target_style: str, notes: Optional[str] = None,
                 error: Optional[str] = None):
        self.original     = original
        self.converted    = converted
        self.ref_type     = ref_type
        self.source_style = source_style
        self.target_style = target_style
        self.notes        = notes
        self.error        = error

    def to_log_line(self) -> str:
        lines = [
            f"  TYPE:    {self.ref_type}",
            f"  FROM:    [{self.source_style}] {self.original}",
            f"  TO:      [{self.target_style}] {self.converted}",
        ]
        if self.notes:
            lines.append(f"  NOTES:   {self.notes}")
        if self.error:
            lines.append(f"  ERROR:   {self.error}")
        return "\n".join(lines)


# ─────────────────────────────────────────────
# MAIN PROCESSOR
# ─────────────────────────────────────────────

def process_conversion(
    input_docx: Path,
    output_dir: Optional[Path] = None,
    source_style: str = "Auto",       # "AMA", "APA", or "Auto"
    target_style: str = "APA",        # "AMA" or "APA"
    model_name: str = "gemini-2.0-flash",
    prefer_gemini_output: bool = True, # If True, use Gemini's formatted_output directly
) -> Dict[str, Path]:
    """
    Convert all references in a Word document between AMA 11th and APA 7th.

    The document must contain <ref-open> and <ref-close> tags to delimit
    the reference section.

    Args:
        input_docx:           Path to the input .docx file.
        output_dir:           Directory for output files (defaults to input dir).
        source_style:         "AMA", "APA", or "Auto" (auto-detect per reference).
        target_style:         "AMA" or "APA".
        model_name:           Gemini model to use.
        prefer_gemini_output: Use Gemini's formatted_output directly vs rebuilding
                              from metadata. Recommended: True.

    Returns:
        Dict with paths: output_docx, log_file, json_dump
    """
    # ── Validate inputs ───────────────────────────────────────────
    input_docx = Path(input_docx)
    if not input_docx.exists():
        raise FileNotFoundError(f"Input file not found: {input_docx}")

    target_style = target_style.upper().strip()
    if target_style not in ("AMA", "APA"):
        raise ValueError(f"target_style must be 'AMA' or 'APA', got: {target_style}")

    if output_dir is None:
        output_dir = input_docx.parent
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    target_enum = CitationStyle.APA if target_style == "APA" else CitationStyle.AMA

    # ── Output paths ──────────────────────────────────────────────
    stem = input_docx.stem
    output_docx_path = output_dir / f"{stem}_Converted.docx"
    log_file_path    = output_dir / f"{stem}_conversion_log.txt"
    json_dump_path   = output_dir / f"{stem}_metadata_dump.json"

    # ── Load document ─────────────────────────────────────────────
    doc = Document(input_docx)

    log_entries: List[ConversionLogEntry] = []
    json_records: List[Dict] = []
    log_header: List[str] = [
        f"Reference Conversion Log",
        f"Input:         {input_docx.name}",
        f"Source Style:  {source_style}",
        f"Target Style:  {target_style}",
        f"Model:         {model_name}",
        "=" * 60,
        ""
    ]

    total_count     = 0
    converted_count = 0
    error_count     = 0
    in_ref_section  = False

    # ── Process paragraphs ────────────────────────────────────────
    for para in doc.paragraphs:
        raw_text = para.text.strip()
        if not raw_text:
            continue

        raw_lower = raw_text.lower()

        # Section boundary detection
        if "<ref-open>" in raw_lower:
            in_ref_section = True
            logger.info("Entering reference section.")
            continue
        if "<ref-close>" in raw_lower:
            in_ref_section = False
            logger.info("Exiting reference section.")
            continue
        if not in_ref_section:
            continue

        # Skip very short lines (headers, labels, etc.)
        if len(raw_text) < 15:
            continue

        total_count += 1
        logger.info(f"[{total_count}] Processing: {raw_text[:80]}...")

        # ── Auto-detect source style ──────────────────────────────
        if source_style.upper() == "AUTO":
            detected_source = detect_source_style(raw_text)
        else:
            detected_source = CitationStyle.AMA if source_style.upper() == "AMA" else CitationStyle.APA

        # Skip if source == target (nothing to convert)
        if detected_source == target_enum:
            logger.info(f"  Skipping: already in {target_style}")
            continue

        # ── Call Gemini ───────────────────────────────────────────
        result = convert_reference(
            raw_text=raw_text,
            source_style=detected_source,
            target_style=target_enum,
            model_name=model_name,
        )

        if not result:
            error_count += 1
            entry = ConversionLogEntry(
                original=raw_text, converted="[FAILED]",
                ref_type="unknown",
                source_style=detected_source.value,
                target_style=target_style,
                error="Gemini returned no result"
            )
            log_entries.append(entry)
            logger.warning(f"  Gemini failed for reference {total_count}")
            continue

        metadata   = result.get("metadata", {})
        ref_type   = detect_ref_type_from_metadata(metadata)
        gemini_out = result.get("formatted_output", "").strip()
        notes      = result.get("conversion_notes")

        # ── Choose formatted text ─────────────────────────────────
        if prefer_gemini_output and gemini_out:
            final_text = gemini_out
        else:
            # Rebuild from metadata as fallback
            if target_style == "APA":
                final_text = format_apa_from_metadata(metadata)
            else:
                final_text = format_ama_from_metadata(metadata)

        if not final_text.strip():
            error_count += 1
            entry = ConversionLogEntry(
                original=raw_text, converted="[EMPTY OUTPUT]",
                ref_type=ref_type,
                source_style=detected_source.value,
                target_style=target_style,
                error="Both Gemini output and metadata fallback produced empty string"
            )
            log_entries.append(entry)
            continue

        # ── Update paragraph in-place with granular bib_ styles ─────
        try:
            if target_style == "AMA":
                segs = build_segments_ama(metadata)
            else:
                segs = build_segments_apa(metadata)

            if segs:
                _write_styled_runs(para, segs, doc=doc)
            else:
                # Fallback to plain text if metadata produced no segments
                _set_paragraph_text(para, final_text)
        except Exception as _seg_err:
            logger.warning(f"  Segment build failed ({_seg_err}); falling back to plain text.")
            _set_paragraph_text(para, final_text)
        converted_count += 1

        entry = ConversionLogEntry(
            original=raw_text,
            converted=final_text,
            ref_type=ref_type,
            source_style=detected_source.value,
            target_style=target_style,
            notes=notes
        )
        log_entries.append(entry)

        # Collect metadata for JSON dump
        json_records.append({
            "index":        total_count,
            "ref_type":     ref_type,
            "source_style": detected_source.value,
            "target_style": target_style,
            "original":     raw_text,
            "converted":    final_text,
            "notes":        notes,
            "metadata":     metadata,
        })

        logger.info(f"  ✓ [{ref_type}] → {final_text[:80]}...")

    # ── Save document ─────────────────────────────────────────────
    doc.save(output_docx_path)
    logger.info(f"Saved converted document: {output_docx_path}")

    # ── Write log ─────────────────────────────────────────────────
    summary = [
        "",
        "=" * 60,
        f"SUMMARY",
        f"  Total references found:  {total_count}",
        f"  Successfully converted:  {converted_count}",
        f"  Errors:                  {error_count}",
        f"  Skipped (same style):    {total_count - converted_count - error_count}",
    ]

    with open(log_file_path, "w", encoding="utf-8") as f:
        f.write("\n".join(log_header) + "\n")
        for i, entry in enumerate(log_entries, 1):
            f.write(f"[{i}]\n{entry.to_log_line()}\n\n")
        f.write("\n".join(summary) + "\n")

    logger.info(f"Log written: {log_file_path}")

    # ── Write JSON dump ───────────────────────────────────────────
    with open(json_dump_path, "w", encoding="utf-8") as f:
        json.dump(json_records, f, indent=2, ensure_ascii=False)

    logger.info(f"Metadata dump: {json_dump_path}")

    return {
        "output_docx": output_docx_path,
        "log_file":    log_file_path,
        "json_dump":   json_dump_path,
    }


# ─────────────────────────────────────────────
# CLI ENTRY POINT
# ─────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Convert references in a Word document between AMA and APA styles.")
    parser.add_argument("input",          type=str,            help="Path to input .docx file")
    parser.add_argument("--output-dir",   type=str,            help="Output directory (default: same as input)")
    parser.add_argument("--source-style", type=str, default="Auto",  choices=["AMA", "APA", "Auto"], help="Source citation style")
    parser.add_argument("--target-style", type=str, default="APA",   choices=["AMA", "APA"],         help="Target citation style")
    parser.add_argument("--model",        type=str, default="gemini-2.0-flash",                       help="Gemini model name")
    parser.add_argument("--no-gemini-output", action="store_true", help="Rebuild from metadata instead of using Gemini's formatted output")
    args = parser.parse_args()

    paths = process_conversion(
        input_docx=Path(args.input),
        output_dir=Path(args.output_dir) if args.output_dir else None,
        source_style=args.source_style,
        target_style=args.target_style,
        model_name=args.model,
        prefer_gemini_output=not args.no_gemini_output,
    )

    print("\nConversion complete:")
    for k, v in paths.items():
        print(f"  {k}: {v}")