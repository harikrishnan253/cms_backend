import re
import os
import io
import zipfile
from collections import defaultdict

try:
    from flask import Flask, request, send_file, render_template, redirect, url_for, session
except ImportError:
    class MockApp:
        def route(self, *args, **kwargs): return lambda f: f
        def __setitem__(self, key, value): pass
    Flask = lambda *args, **kwargs: MockApp()
    request = session = send_file = render_template = redirect = url_for = None
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import RGBColor
try:
    from utils import track_changes
    TRACK_CHANGES_ENABLED = True
except ImportError:
    track_changes = None
    TRACK_CHANGES_ENABLED = False
import logging

app = Flask(__name__)
app.secret_key = "secret_key_for_session_encryption"
UPLOAD_DIR = "temp_reports"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# =====================================================
# Helpers & Core Logic
# =====================================================

def iter_document_paragraphs(doc):
    """
    Iterate through all paragraphs in the document body in order,
    including those inside tables.
    """
    body = doc._element.body
    for child in body:
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p


def get_numbers(text):
    """
    Extract numbers from text like '1', '2-5', '1, 3, 5'.
    Handles ranges "1-5" -> [1, 2, 3, 4, 5].
    """
    nums = []
    # Matches: (start)-(end) OR (single)
    # Allows hyphen, en dash, em dash
    pattern = re.compile(r'(\d+)\s*[-–—]\s*(\d+)|(\d+)')
    
    for start, end, single in pattern.findall(text):
        if start and end:
            try:
                s, e = int(start), int(end)
                if s <= e:
                    nums.extend(range(s, e + 1))
            except ValueError:
                pass
        elif single:
            try:
                nums.append(int(single))
            except ValueError:
                pass
    return nums


def format_numbers(nums):
    """
    Format a list of numbers into a string like '1-3, 5'.
    Collapses ranges of 3 or more (e.g. 1,2,3 -> 1-3).
    """
    nums = sorted(set(nums))
    if not nums:
        return ""

    parts = []
    if not nums:
        return ""

    start = prev = nums[0]

    for n in nums[1:]:
        if n == prev + 1:
            prev = n
        else:
            length = prev - start + 1
            if length >= 3:
                parts.append(f"{start}-{prev}")
            elif length == 2:
                parts.append(f"{start},{prev}")
            else:
                parts.append(str(start))
            start = prev = n

    length = prev - start + 1
    if length >= 3:
        parts.append(f"{start}-{prev}")
    elif length == 2:
        parts.append(f"{start},{prev}")
    else:
        parts.append(str(start))

    return ", ".join(parts)


def is_citation_run(run):
    """
    Determine if a run is part of a citation.
    Strictly checks for 'cite_bib' styles.
    """
    if run.style and run.style.name in ["cite_bib"]:
        return True
    return False


class ReferenceProcessor:
    def __init__(self, doc):
        self.doc = doc
        
    def get_references_in_bibliography(self):
        """
        Returns a Set of IDs found in the bibliography sections (REF-N style).
        Also returns a list of objects for reordering later.
        """
        refs_found = set()
        ref_objects = [] # list of dicts: {'id': int, 'para': p, 'run': r}

        for para in self.doc.paragraphs:
            if para.style and para.style.name == "REF-N":
                found_id = None
                bib_run = None
                
                # Try finding styled run
                for run in para.runs:
                    if run.style and run.style.name == "bib_number":
                        nums = get_numbers(run.text)
                        if nums:
                            found_id = nums[0]
                            bib_run = run
                            break
                
                if found_id is not None:
                    refs_found.add(found_id)
                    ref_objects.append({
                        'id': found_id,
                        'para': para,
                        'run': bib_run
                    })
                    
        return refs_found, ref_objects

    def get_citations_in_text(self):
        """
        Scans document for citations.
        Returns:
            all_cited_ids: list of all IDs in order of appearance (with duplicates)
            appearance_order: list of unique IDs in order of first appearance
        """
        all_cited_ids = []
        appearance_order = []
        seen = set()
        
        for para in iter_document_paragraphs(self.doc):
            # 1. Process runs
            current_group = []
            
            for run in para.runs:
                if is_citation_run(run):
                    current_group.append(run)
                else:
                    if current_group:
                        # Flush group
                        text = "".join(r.text for r in current_group)
                        nums = get_numbers(text)
                        all_cited_ids.extend(nums)
                        for n in nums:
                            if n not in seen:
                                seen.add(n)
                                appearance_order.append(n)
                        current_group = []
            
            # Flush trailing group
            if current_group:
                text = "".join(r.text for r in current_group)
                nums = get_numbers(text)
                all_cited_ids.extend(nums)
                for n in nums:
                    if n not in seen:
                        seen.add(n)
                        appearance_order.append(n)
                        
        return all_cited_ids, appearance_order

    def find_duplicates(self, ref_objects):
        """
        Finds duplicate references using fuzzy matching (difflib).
        Returns a list of dicts: {'id': int, 'text': str, 'duplicate_of': int, 'score': float}
        """
        import difflib
        
        duplicates = []
        processed_refs = [] # list of (id, clean_text)
        
        # 1. Pre-process all candidates
        for obj in ref_objects:
            full_text = obj['para'].text.strip()
            # Remove leading numbering like "1. ", "[1] "
            clean_text = re.sub(r'^\[?\d+\]?[\.\s]*', '', full_text)
            processed_refs.append({'id': obj['id'], 'text': clean_text})
            
        # 2. Compare O(N^2)
        # We only check forward to avoid double reporting (A=B, B=A)
        # We assume the *earlier* ID is the "original" and later is "duplicate"
        n = len(processed_refs)
        matcher = difflib.SequenceMatcher(None, "", "")
        
        for i in range(n):
            ref_a = processed_refs[i]
            text_a = ref_a['text']
            len_a = len(text_a)
            
            if len_a == 0:
                continue
                
            matcher.set_seq1(text_a)
            
            for j in range(i + 1, n):
                ref_b = processed_refs[j]
                text_b = ref_b['text']
                len_b = len(text_b)
                
                if len_b == 0: 
                    continue
                    
                # Optimization: Length ratio check
                # If lengths differ significantly, they can't be high matches
                # If ratio > 0.85, then min_len / max_len must be roughly > 0.85
                # We use 0.6 as a conservative safety net, but 0.8 is probably safe if threshold is 0.85.
                if min(len_a, len_b) / max(len_a, len_b) < 0.6:
                    continue
                
                matcher.set_seq2(text_b)
                
                # Performance Optimization: Check cheap upper bounds first
                if matcher.real_quick_ratio() < 0.85:
                    continue
                if matcher.quick_ratio() < 0.85:
                    continue
                    
                ratio = matcher.ratio()
                
                # Threshold: 0.85 (85% similar)
                if ratio > 0.85:
                    duplicates.append({
                        'id': ref_b['id'], # The later one is the duplicate
                        'text': ref_b['text'][:100] + "...",
                        'duplicate_of': ref_a['id'],
                        'score': round(ratio * 100, 1)
                    })
                    
        return duplicates

    def get_validation_stats(self):
        bib_refs, ref_objects = self.get_references_in_bibliography()
        all_cited, _ = self.get_citations_in_text()
        
        unique_cited = set(all_cited)
        
        # Missing: Cited but not in Bib
        missing = sorted(unique_cited - bib_refs)
        
        # Unused: In Bib but not Cited
        unused = sorted(bib_refs - unique_cited)
        
        # Duplicates
        duplicates = self.find_duplicates(ref_objects)
        
        # Sequence Issues
        sequence_issues = []
        seen_in_seq = []
        previous_max = 0
        
        for n in all_cited:
            if n not in seen_in_seq:
                if n < previous_max:
                     pass
                
                if n != len(seen_in_seq) + 1:
                     sequence_issues.append({
                         "position": len(seen_in_seq) + 1,
                         "current": n,
                         "expected": len(seen_in_seq) + 1
                     })
                
                seen_in_seq.append(n)
                previous_max = max(previous_max, n)
                
        return {
            "total_references": len(bib_refs),
            "total_citations": len(all_cited),
            "missing_references": missing,
            "unused_references": unused,
            "duplicate_references": duplicates,
            "sequence_issues": sequence_issues,
            "is_perfect": (not missing and not unused and not sequence_issues and not duplicates)
        }

    def renumber(self):
        """
        Renumber citations and reorder bibliography.
        Returns: mapping (Old -> New)
        """
        _, appearance_order = self.get_citations_in_text()
        
        # Ensure 'cite_bib' style exists
        from docx.enum.style import WD_STYLE_TYPE
        styles = self.doc.styles
        try:
            styles['cite_bib']
        except KeyError:
            s = styles.add_style('cite_bib', WD_STYLE_TYPE.CHARACTER)
            s.font.superscript = True

        # Create Mapping
        mapping = {} 
        new_id = 1
        for old_id in appearance_order:
            mapping[old_id] = new_id
            new_id += 1
            
        for para in iter_document_paragraphs(self.doc):
            i = 0
            while i < len(para.runs):
                run = para.runs[i]
                
                if is_citation_run(run):
                    txt = run.text
                    nums = get_numbers(txt)
                    if nums:
                         new_nums = [mapping.get(n, n) for n in nums]
                         new_text = format_numbers(new_nums)
                         
                         is_renumbered = (nums != new_nums)
                         highlight_color = "008000" if is_renumbered else None
                         
                         style_name = run.style.name if run.style else "cite_bib"
                         
                         if TRACK_CHANGES_ENABLED:
                             # Must replace the whole run
                             track_changes.delete_tracked_run(para, run)
                             
                             run_del = run._element.getparent()
                             anchor = run_del if run_del.tag == track_changes.qn('w:del') else run._element
                             
                             ins_new = track_changes.add_tracked_text(para, new_text, style=style_name, color=highlight_color)
                             anchor.addnext(ins_new)
                         else:
                             run.text = new_text
                             if is_renumbered:
                                 run.font.color.rgb = RGBColor(0, 128, 0)
                
                i += 1
                
                i += 1

        # 2. Reorder Bibliography
        _, ref_objects = self.get_references_in_bibliography()
        
        # Sort objects into Cited and Uncited
        cited_refs = []
        uncited_refs = []
        
        for obj in ref_objects:
            if obj['id'] in mapping:
                obj['new_id'] = mapping[obj['id']]
                cited_refs.append(obj)
            else:
                uncited_refs.append(obj)
        
        if not ref_objects:
            return mapping

        # Find anchor (min index)
        body = self.doc._element.body
        
        indices = []
        for obj in ref_objects:
            try:
                idx = body.index(obj['para']._element)
                indices.append(idx)
            except ValueError:
                pass 
        
        if not indices:
            return mapping
            
        anchor = min(indices)
        
        # Remove all
        for obj in ref_objects:
             p = obj['para']._element
             if p.getparent() == body:
                 body.remove(p)
                 
        # Insert Cited (Sorted)
        cited_refs.sort(key=lambda x: x['new_id'])
        
        insert_idx = anchor
        for obj in cited_refs:
            # Update ID text
            if obj['run']:
                old_text = obj['run'].text
                new_text = str(obj['new_id'])
                
                if old_text != new_text:
                    if TRACK_CHANGES_ENABLED:
                        style_name = obj['run'].style.name if obj['run'].style else None
                        
                        track_changes.delete_tracked_run(obj['para'], obj['run'])
                        run_del = obj['run']._element.getparent()
                        anchor = run_del if run_del.tag == track_changes.qn('w:del') else obj['run']._element
                        
                        ins_new = track_changes.add_tracked_text(obj['para'], new_text, style=style_name)
                        anchor.addnext(ins_new)
                    else:
                        obj['run'].text = new_text
            
            body.insert(insert_idx, obj['para']._element)
            insert_idx += 1
            
        # Insert Uncited (Appended after cited)
        for obj in uncited_refs:
            body.insert(insert_idx, obj['para']._element)
            insert_idx += 1
            
        return mapping


def process_document(file):
    doc = Document(file)
    processor = ReferenceProcessor(doc)
    
    # Check BEFORE
    before_stats = processor.get_validation_stats()
    
    # DECISION:
    # 1. If Unused References exist -> ABORT renumbering.
    if before_stats["unused_references"]:
        return doc, before_stats, before_stats, {}, "Aborted: Document validation failed due to unused references."

    # 2. If Perfect -> No need.
    if before_stats["is_perfect"]:
        return doc, before_stats, before_stats, {}, "Validation completed."
        
    # 3. If Missing Refs -> Can't safely renumber usually
    if before_stats["missing_references"]:
         return doc, before_stats, before_stats, {}, "Aborted: Missing references detected."

    # DO RENUMBER
    mapping = processor.renumber()
    
    # Check AFTER (Validate result)
    after_stats = processor.get_validation_stats()
    
    # Determine status message
    changes_made = False
    if mapping:
        for k, v in mapping.items():
            if k != v:
                changes_made = True
                break

    if before_stats["duplicate_references"]:
        count = len(before_stats['duplicate_references'])
        prefix = "Renumbering" if changes_made else "Validation"
        status_msg = f"{prefix} completed with {count} duplicate{'s' if count > 1 else ''}."
    elif changes_made:
        status_msg = "Renumbering completed successfully."
    else:
        status_msg = "Validation completed."

    return doc, before_stats, after_stats, mapping, status_msg


# =====================================================
# Flask Routes
# =====================================================
@app.route("/")
def upload_file():
    return render_template("upload.html")


@app.route("/process", methods=["GET", "POST"])
def process():
    if request.method == "POST":
        file = request.files.get("file")
        if not file or not file.filename.endswith(".docx"):
            return "Invalid file", 400

        doc, before, after, mapping, status_msg = process_document(file)

        base = os.path.splitext(file.filename)[0]
        doc_path = os.path.join(UPLOAD_DIR, f"{base}_renumbered.docx")
        report_path = os.path.join(UPLOAD_DIR, f"{base}_validation.txt")

        doc.save(doc_path)

        with open(report_path, "w", encoding="utf-8") as f:
            f.write(f"STATUS: {status_msg}\n")
            f.write("VALIDATION BEFORE\n")
            f.write(str(before) + "\n\n")
            f.write("VALIDATION AFTER\n")
            f.write(str(after) + "\n\n")
            if mapping:
                f.write("RENUMBERING MAPPING (Old -> New)\n")
                for old, new in sorted(mapping.items(), key=lambda x: x[1]):
                    f.write(f"{old} -> {new}\n")

        # Create ZIP package
        zip_filename = f"{base}_results.zip"
        zip_path = os.path.join(UPLOAD_DIR, zip_filename)
        
        # Validation HTML Report (Offline)
        html_report_filename = f"{base}_results.html"
        html_report_path = os.path.join(UPLOAD_DIR, html_report_filename)
        
        # Render the template for offline use
        # Note: We pass offline_mode=True to make links relative
        html_content = render_template(
            "validation_results.html",
            filename=file.filename,
            results=after,
            before=before,
            mapping=mapping,
            status_msg=status_msg,
            report_file=os.path.basename(report_path),
            doc_file=os.path.basename(doc_path),
            zip_file=None, # No zip button in offline report
            offline_mode=True 
        )
        
        with open(html_report_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        with zipfile.ZipFile(zip_path, 'w') as zf:
             # Add Doc
             zf.write(doc_path, arcname=os.path.basename(doc_path))
             # Add Text Report
             zf.write(report_path, arcname=os.path.basename(report_path))
             # Add HTML Report
             zf.write(html_report_path, arcname=os.path.basename(html_report_path))

        # Store data in session for GET request
        session['processing_result'] = {
            'filename': file.filename,
            'before': before,
            'after': after,
            'mapping': mapping,
            'status_msg': status_msg,
            'report_file': os.path.basename(report_path),
            'doc_file': os.path.basename(doc_path),
            'zip_file': zip_filename
        }
        
        return redirect(url_for('process'))

    # GET request - retrieve from session
    result = session.get('processing_result')
    if not result:
        return redirect(url_for('upload_file'))
        
    return render_template(
        "validation_results.html",
        filename=result['filename'],
        results=result['after'],
        before=result['before'],
        mapping=result['mapping'],
        status_msg=result['status_msg'],
        report_file=result['report_file'],
        doc_file=result['doc_file'],
        zip_file=result.get('zip_file')
    )


@app.route("/download/<path:filename>")
def download_file(filename):
    # Security: Ensure filename is in UPLOAD_DIR
    return send_file(os.path.join(UPLOAD_DIR, filename), as_attachment=True)


if __name__ == "__main__":
    app.run(debug=True, port=5000)
