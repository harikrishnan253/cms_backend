import os
import re
import logging
try:
    from flask import Flask, render_template, request, send_file, redirect, url_for, flash, make_response, session
    from werkzeug.utils import secure_filename
except ImportError:
    class MockApp:
        config = {}
        secret_key = ""
        def route(self, *args, **kwargs): return lambda f: f
        def __setitem__(self, key, value): pass
    Flask = lambda *args, **kwargs: MockApp()
    secure_filename = lambda x: x
    request = session = render_template = redirect = url_for = flash = make_response = send_file = None
from docx import Document
import io
import zipfile
from citation_parsers import get_parser, auto_detect_style
from validation_core import CitationProcessor, ValidationReport
try:
    from utils import track_changes
    TRACK_CHANGES_ENABLED = True
except ImportError:
    track_changes = None
    TRACK_CHANGES_ENABLED = False

# Configure logging (import logging only once — the second import at line 11 was removed)
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.secret_key = os.environ.get('SESSION_SECRET', 'dev-secret-key')

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'docx'}

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append(para.text)
    return paragraphs


def normalize_citation_key(author_part, year):
    author_clean = re.sub(r'[^\w\s&]', '', author_part).strip()
    author_clean = re.sub(r'\s+', ' ', author_clean)
    return f"{author_clean}|{year}"


# Ported from Referencenumvalidation.py
def find_duplicates(references, reference_details):
    """
    Finds duplicate references using fuzzy matching (difflib) on FULL TEXT.
    Args:
        references: Dictionary of reference objects.
        reference_details: Dictionary containing 'text' for each reference key.
    Returns:
        List of dicts: {'id': str, 'text': str, 'duplicate_of': str, 'score': float}
    """
    import difflib
    
    duplicates = []
    processed_refs = [] 
    for key, data in references.items():
        # Use full original text for comparison to avoid false positives on "Author (Year)"
        text = reference_details.get(key, {}).get('text', data.get('display', ''))
        processed_refs.append({'id': key, 'text': text})
        
    n = len(processed_refs)
    for i in range(n):
        ref_a = processed_refs[i]
        for j in range(i + 1, n):
            ref_b = processed_refs[j]
            
            # Simple fuzzy match on the display string
            len_a = len(ref_a['text'])
            len_b = len(ref_b['text'])
            if len_a == 0 or len_b == 0: continue
            
            if min(len_a, len_b) / max(len_a, len_b) < 0.6:
                continue
            
            ratio = difflib.SequenceMatcher(None, ref_a['text'], ref_b['text']).ratio()
            
            if ratio > 0.85:
                duplicates.append({
                    'id': ref_b['id'], 
                    'text': ref_b['text'][:100],
                    'duplicate_of': ref_a['id'],
                    'score': round(ratio * 100, 1)
                })
        

    return duplicates


def parse_single_citation(cite_text):
    cite_text = cite_text.strip()
    cite_text = re.sub(r'\[[^\]]+\]', '', cite_text).strip()
    cite_text = re.sub(r'^(see|cf\.?|e\.g\.?,?|i\.e\.?,?)\s+', '', cite_text, flags=re.IGNORECASE).strip()
    
    if re.search(r'\bp\.?\s*\d+', cite_text, re.IGNORECASE):
        return [(None, None)]
    
    if re.search(r'(January|February|March|April|May|June|July|August|September|October|November|December)', cite_text, re.IGNORECASE):
        return [(None, None)]
    
    years = re.findall(r'\b((?:19|20)\d{2})[a-z]?\b', cite_text)
    if years:
        author_part = re.sub(r',?\s*(?:19|20)\d{2}[a-z]?,?\s*', '', cite_text).strip()
        author_part = re.sub(r',\s*$', '', author_part).strip()
        author_part = re.sub(r'^\d+,?\s*', '', author_part).strip()
        
        if not author_part or len(author_part) < 2:
            return [(None, None)]
        
        results = []
        for year in years:
            results.append((author_part, year))
        return results if results else [(None, None)]
    return [(None, None)]


def find_citations_in_text(paragraphs, parser):
    """
    Find citations using the provided parser.
    """
    citations = {}
    citation_locations = {}
    
    for i, para in enumerate(paragraphs):
        # Stop if we hit the bibliography section start
        if '<ref-open>' in para:
            break
            
        # Pass the whole paragraph text to the parser
        # The parser is now responsible for finding parenthetical AND narrative citations
        found_citations = parser.parse_citation(para)
        
        for cite in found_citations:
            author = cite['author']
            year = cite['year']
            
            # Create a unique key
            citation_key = normalize_citation_key(author, year)
            
            # Construct display string
            if cite['type'] == 'parenthetical':
                display_key = f"({author}, {year})"
            else:
                display_key = f"{author} ({year})"
            
            if citation_key not in citations:
                citations[citation_key] = {
                    'display': display_key,
                    'author': author,
                    'year': year,
                    'type': cite['type'],
                    'warnings': cite.get('warnings', []),
                    'raw': cite.get('raw', '')
                }
            else:
                # Merge warnings if new ones found
                citations[citation_key]['warnings'].extend([w for w in cite.get('warnings', []) if w not in citations[citation_key]['warnings']])

            if citation_key not in citation_locations:
                citation_locations[citation_key] = []
            citation_locations[citation_key].append(i + 1)
    
    return citations, citation_locations


def find_references_in_bibliography(paragraphs, parser):
    """
    Find references strictly between <ref-open> and <ref-close> tags.
    """
    references = {}
    reference_details = {}
    abbreviation_map = {}
    
    in_references_section = False
    
    for i, para in enumerate(paragraphs):
        para_stripped = para.strip()
        
        # Check for tags
        if '<ref-open>' in para_stripped:
            in_references_section = True
            # If the tag is on its own line, continue. If inline, we might need to parse the rest?
            # Assuming tag is a section delimiter.
            continue
        
        if '<ref-close>' in para_stripped:
            in_references_section = False
            continue
            
        if in_references_section and para_stripped:
            # Parse reference using parser
            ref_data = parser.parse_reference(para_stripped)
            
            if ref_data:
                author_display = ref_data['author']
                year = ref_data['year']
                full_author = ref_data['full_author']
                abbreviations = ref_data['abbreviations']
                
                ref_key = normalize_citation_key(author_display, year)
                
                if ref_key not in references:
                    references[ref_key] = {
                        'display': f"{author_display} ({year})",
                        'author': author_display,
                        'year': year,
                        'full_author': full_author,
                        'abbreviations': abbreviations
                    }
                    reference_details[ref_key] = {
                        'line': i + 1,
                        'text': para_stripped[:150] + ('...' if len(para_stripped) > 150 else '')
                    }
                    
                    for abbr in abbreviations:
                        abbr_key = f"{abbr}|{year}"
                        abbreviation_map[abbr_key] = ref_key
    
    return references, reference_details, abbreviation_map


def extract_first_surname(author_str):
    author_str = author_str.strip()
    author_str = re.sub(r'\s*et\s+al\.?\s*', '', author_str, flags=re.IGNORECASE)
    
    if ',' in author_str:
        return author_str.split(',')[0].strip()
    
    parts = author_str.split()
    return parts[0] if parts else ''


def match_citation_to_reference(citation, references):
    cite_author = citation['author'].strip()
    cite_year = citation['year']
    cite_author_lower = cite_author.lower()
    
    for ref_key, ref_data in references.items():
        ref_year = ref_data['year']
        ref_full_author = ref_data.get('full_author', ref_data['author'])
        ref_full_lower = ref_full_author.lower()
        
        if cite_year != ref_year:
            continue
        
        cite_first = extract_first_surname(cite_author)
        ref_first = extract_first_surname(ref_full_author)
        
        if cite_first and ref_first and cite_first == ref_first:
            return ref_key
        
        if cite_author_lower in ref_full_lower or ref_full_lower.startswith(cite_author_lower):
            return ref_key
        
        # Word subset matching for basic cases
        cite_words = set(re.findall(r'\b[a-z]{2,}\b', cite_author_lower))
        cite_words -= {'et', 'al', 'and', 'the'}
        ref_words = set(re.findall(r'\b[a-z]{2,}\b', ref_full_lower))
        
        if cite_words and cite_words.issubset(ref_words):
            return ref_key
    
    return None


import difflib
from difflib import SequenceMatcher

def normalize_text_for_comparison(text):
    """Normalize text for flexible matching."""
    # Replace 'and' with '&'
    text = re.sub(r'\band\b', '&', text, flags=re.IGNORECASE)
    # Remove dots, commas, extra spaces
    text = re.sub(r'[.,]', '', text)
    text = text.strip().lower()
    # Remove leading 'the'
    text = re.sub(r'^the\s+', '', text)
    return text

def check_smart_match(cite_data, references):
    """
    Advanced matching logic for:
    1. Introduction of abbreviations: "Organization [Org]" -> Match "Organization"
    2. Et al: "Smith et al" -> Match "Smith, Jones..."
    3. Proper subsets/variations: "Smith & Jones" -> Match "Smith, ... & Jones"
    4. Narrative conjunctions: "Smith and Jones" -> Match "Smith & Jones"
    """
    cite_author = cite_data['author']
    cite_year = cite_data['year']
    
    cite_norm = normalize_text_for_comparison(cite_author)
    
    # 1. Handle Abbreviation Introduction: "Name [Abbr]"
    # Extract "Name" part
    prefix_match = re.match(r'^(.*?)\s*\[.*?\]', cite_author)
    citation_prefix_norm = None
    if prefix_match:
        citation_prefix_norm = normalize_text_for_comparison(prefix_match.group(1))

    # 2. Handle "et al"
    is_etal = 'et al' in cite_norm
    cite_first_surname = cite_norm.split()[0] if cite_norm else ""
    
    # 3. Handle list of names (e.g. "Smith & Jones")
    cite_names = [n.strip() for n in re.split(r'[&]', cite_norm)]
    
    for ref_key, ref_data in references.items():
        if cite_year and ref_data['year'] != cite_year:
            continue
            
        ref_author = ref_data.get('full_author', ref_data['author'])
        ref_norm = normalize_text_for_comparison(ref_author)
        
        # A. Direct Normalized Match (covers "Smith and Jones" vs "Smith & Jones")
        if cite_norm == ref_norm:
            return ref_key
            
        # B. Abbreviation Definition Match
        # Citation: "National Org [NO]" vs Ref: "National Org"
        if citation_prefix_norm and citation_prefix_norm == ref_norm:
            return ref_key
            
        # C. Et Al Match
        # Citation: "Smith et al" vs Ref: "Smith, Jones..."
        # Rule: First surnames match
        if is_etal:
            ref_first_surname = ref_norm.split()[0]
            if cite_first_surname == ref_first_surname:
                return ref_key
                
        # D. Word Subset Match (Robust for "Smith, Jones" vs "Smith, A., Jones, B.")
        # Only check if citation has multiple words (potential authors)
        cite_words = set(re.findall(r'\b[a-z]{2,}\b', cite_norm))
        # Remove common stopwords from citation side to avoid false positives matching "and" to "and"
        cite_words -= {'and', 'the', 'et', 'al'}
        
        if len(cite_words) > 1:
            ref_words = set(re.findall(r'\b[a-z]{2,}\b', ref_norm))
            # Check if All significant citation words are in reference
            if cite_words.issubset(ref_words):
                return ref_key

    return None

def check_spelling_mismatch(cite_author, references):
    """
    Check for spelling mismatches using difflib.
    Returns reference key if a close match is found.
    """
    cite_author_norm = normalize_text_for_comparison(cite_author)
    
    best_match = None
    max_ratio = 0.0
    
    for ref_key, ref_data in references.items():
        ref_author = ref_data['author']
        ref_norm = normalize_text_for_comparison(ref_author)
        
        # simple ratio check
        ratio = difflib.SequenceMatcher(None, cite_author_norm, ref_norm).ratio()
        
        if ratio > 0.8: # >80% similarity
            if ratio > max_ratio:
                max_ratio = ratio
                best_match = ref_key
                
    return best_match


def check_et_al_misuse(cite_data, ref_data):
    """
    Check if 'et al.' is used incorrectly in citation based on reference author count.
    
    APA 7th Edition Rules:
    - 1-2 authors: Always cite all authors
    - 3+ authors: Use 'et al.' after first author
    
    Args:
        cite_data: Citation data dict with 'author' key
        ref_data: Reference data dict with 'full_author' key
        
    Returns:
        Dict with 'has_error' (bool) and 'message' (str) if error found, None otherwise
    """
    cite_author = cite_data['author'].strip()
    ref_full_author = ref_data.get('full_author', ref_data['author'])
    
    # Check if citation uses 'et al.'
    has_et_al = 'et al' in cite_author.lower()
    
    # Count authors in reference
    # Method: Count commas and ampersands
    # "Smith, J." = 1 author (1 comma)
    # "Smith, J., & Jones, M." = 2 authors (3 commas, 1 ampersand)
    # "Smith, J., Jones, M., & Brown, K." = 3 authors (5 commas, 1 ampersand)
    
    if '&' in ref_full_author:
        # Count author segments by splitting on '&'
        parts = ref_full_author.split('&')
        # First part has N-1 authors (separated by commas)
        # Last part has 1 author
        first_part_authors = len([p for p in parts[0].split(',') if p.strip() and not re.match(r'^[A-Z]\.?$', p.strip())])
        author_count = first_part_authors + 1
    else:
        # Single author (has comma for "Last, F.")
        author_count = 1
    
    # Check for misuse
    if has_et_al and author_count <= 2:
        # Error: Using et al. with 1-2 authors
        if author_count == 1:
            correct_form = ref_data['author']
        else:
            # Extract both author surnames for 2-author case
            ref_display = ref_data['author']
            correct_form = ref_display
        
        return {
            'has_error': True,
            'message': f"Change author per reference - use '{correct_form}', not 'et al.' (reference has only {author_count} author{'s' if author_count > 1 else ''})",
            'author_count': author_count,
            'correct_form': correct_form
        }
    
    elif not has_et_al and author_count >= 3:
        # Warning: Should use et al. with 3+ authors (but this is less critical)
        first_author = extract_first_surname(ref_full_author)
        correct_form = f"{first_author} et al."
        
        return {
            'has_error': True,
            'message': f"Consider using 'et al.' - reference has {author_count} authors, APA allows '{correct_form}'",
            'author_count': author_count,
            'correct_form': correct_form,
            'severity': 'warning'  # Less severe than incorrect et al. usage
        }
    
    return None

def get_citation_matches(citations, references, abbreviation_map):
    """
    Match citations to references using Exact, Abbreviation, and Smart/Fuzzy matching.
    Returns: matched_citations (set), matched_references (set), matched_pairs (dict: cite_key -> ref_key)
    """
    matched_citations = set()
    matched_references = set()
    matched_pairs = {}
    
    for cite_key, cite_data in citations.items():
        cite_author = cite_data['author'].strip()
        cite_year = cite_data['year']
        abbr_key = f"{cite_author}|{cite_year}"
        
        matched_ref_key = None
        
        # 1. Exact Match via Key
        if cite_key in references:
            matched_ref_key = cite_key
            
        # 2. Abbreviation Match
        elif abbr_key in abbreviation_map:
            matched_ref_key = abbreviation_map[abbr_key]
            
        # 3. Smart Match / Fuzzy Match
        else:
            smart_match_key = check_smart_match(cite_data, references)
            if smart_match_key:
                matched_ref_key = smart_match_key
        
        if matched_ref_key:
            matched_citations.add(cite_key)
            matched_references.add(matched_ref_key)
            matched_pairs[cite_key] = matched_ref_key
            
    return matched_citations, matched_references, matched_pairs


def check_abbreviation_usage(matched_pairs, citations, references, citation_locations):
    """
    Validate proper usage of abbreviations (First vs Subsequent usage).
    """
    abbreviation_errors = []
    from collections import defaultdict
    ref_usage_map = defaultdict(list)
    
    # 1. Build Usage Map
    for cite_key, ref_key in matched_pairs.items():
        cite_data = citations[cite_key]
        ref_data = references.get(ref_key)
        
        if not ref_data: continue

        ref_abbrs = ref_data.get('abbreviations', [])
        if not ref_abbrs: continue

        is_intro = '[' in cite_data['author'] 
        
        # Check if citation IS the abbreviation
        is_abbr = False
        cite_author_clean = cite_data['author'].strip()
        for abbr in ref_abbrs:
             if abbr == cite_author_clean: 
                 is_abbr = True
                 break
        
        # Check if full name
        ref_full_text = ref_data.get('full_author', '').split('(')[0].strip()
        is_full = False
        if cite_author_clean.lower().startswith(ref_full_text.lower()):
            is_full = True
            
        if is_intro: 
            is_full = True 
            is_abbr = False 

        if citation_locations.get(cite_key):
             for loc in citation_locations[cite_key]:
                 ref_usage_map[ref_key].append({
                     'loc': loc,
                     'text': cite_data['display'],
                     'is_intro': is_intro,
                     'is_abbr': is_abbr,
                     'is_full': is_full,
                     'ref_abbr': ref_abbrs[0]
                 })

    # 2. Validate Usage Order
    for ref_key, usages in ref_usage_map.items():
        usages.sort(key=lambda x: x['loc'])
        
        if not usages: continue
            
        # Check FIRST usage
        first_usage = usages[0]
        if first_usage['is_abbr'] and not first_usage['is_full']:
             abbreviation_errors.append({
                'citation': first_usage['text'],
                'message': f"First confirmation of abbreviation should define it. Use 'Full Name [Abbr]' instead of '{first_usage['text']}'.",
                'locations': [first_usage['loc']],
                'ref_abbr': first_usage['ref_abbr']
            })
            
        # Check SUBSEQUENT usages
        ref_abbr = usages[0]['ref_abbr']
        for i in range(1, len(usages)):
            usage = usages[i]
            
            if usage['is_intro']: 
                 abbreviation_errors.append({
                    'citation': usage['text'],
                    'message': f"Abbreviation already introduced. Use '{ref_abbr}' instead.",
                    'locations': [usage['loc']],
                    'ref_abbr': ref_abbr
                })
            elif usage['is_full'] and not usage['is_abbr']: 
                 abbreviation_errors.append({
                    'citation': usage['text'],
                    'message': f"Abbreviation previously introduced. Consider using '{ref_abbr}' instead.",
                    'locations': [usage['loc']],
                    'ref_abbr': ref_abbr,
                    'severity': 'warning'
                })
                
    return abbreviation_errors



# ─────────────────────────────────────────────────────────────────────────────
# LOW-LEVEL HELPERS
# ─────────────────────────────────────────────────────────────────────────────

from lxml import etree as _lxml_etree
from docx.oxml.ns import qn as _qn
from docx.oxml import OxmlElement as _OxmlElement
from docx.opc.part import Part as _Part
from docx.opc.packuri import PackURI as _PackURI

_W_NS_AP   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_MC_NS_AP  = "http://schemas.openxmlformats.org/markup-compatibility/2006"
_W14_NS_AP = "http://schemas.microsoft.com/office/word/2010/wordml"
_NSMAP_AP  = {"w": _W_NS_AP, "mc": _MC_NS_AP, "w14": _W14_NS_AP}
_COMMENTS_REL_AP = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
_COMMENTS_CT_AP  = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"


def _ensure_cite_bib_style(doc):
    """Create the cite_bib character style if it doesn't already exist."""
    from docx.enum.style import WD_STYLE_TYPE
    if "cite_bib" not in [s.name for s in doc.styles]:
        doc.styles.add_style("cite_bib", WD_STYLE_TYPE.CHARACTER)


def _get_comments_part(doc):
    """Resolve (or create) word/comments.xml Part, returning the Part object."""
    doc_part = doc.part
    for rel in doc_part.rels.values():
        if rel.reltype == _COMMENTS_REL_AP:
            return rel.target_part
    for p in doc_part.package.parts:
        if p.partname == "/word/comments.xml":
            doc_part.relate_to(p, _COMMENTS_REL_AP)
            return p
    root = _lxml_etree.Element(f"{{{_W_NS_AP}}}comments", nsmap=_NSMAP_AP)
    root.set(f"{{{_MC_NS_AP}}}Ignorable", "w14 wp14")
    blob = _lxml_etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    part = _Part(_PackURI("/word/comments.xml"), _COMMENTS_CT_AP, blob, doc_part.package)
    doc_part.package.add_part(part)
    doc_part.relate_to(part, _COMMENTS_REL_AP)
    return part


def _add_word_comment(doc, para, text, author="APA Checker", initials="AC"):
    """
    Attach a Word comment to the first runs of *para* using lxml so the
    namespace prefix stays as w: (not ns0:) — avoids the OOXML schema error.
    """
    if not para.runs:
        return False
    try:
        cp = _get_comments_part(doc)
        try:
            tree = _lxml_etree.fromstring(cp._blob)
        except Exception:
            tree = _lxml_etree.Element(f"{{{_W_NS_AP}}}comments", nsmap=_NSMAP_AP)

        existing = [
            int(c.get(f"{{{_W_NS_AP}}}id"))
            for c in tree.findall(f"{{{_W_NS_AP}}}comment")
            if c.get(f"{{{_W_NS_AP}}}id")
        ]
        cid = max(existing) + 1 if existing else 0

        cel = _lxml_etree.SubElement(tree, f"{{{_W_NS_AP}}}comment")
        cel.set(f"{{{_W_NS_AP}}}id",       str(cid))
        cel.set(f"{{{_W_NS_AP}}}author",   author)
        cel.set(f"{{{_W_NS_AP}}}initials", initials)
        pel = _lxml_etree.SubElement(cel, f"{{{_W_NS_AP}}}p")
        rel = _lxml_etree.SubElement(pel, f"{{{_W_NS_AP}}}r")
        tel = _lxml_etree.SubElement(rel,  f"{{{_W_NS_AP}}}t")
        tel.text = text

        cp._blob = _lxml_etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)

        p_el = para._element
        start = _OxmlElement("w:commentRangeStart"); start.set(_qn("w:id"), str(cid))
        end   = _OxmlElement("w:commentRangeEnd");   end.set(_qn("w:id"),   str(cid))
        ref   = _OxmlElement("w:commentReference");  ref.set(_qn("w:id"),   str(cid))
        ref_r = _OxmlElement("w:r"); ref_r.append(ref)

        pPr = p_el.find(_qn("w:pPr"))
        if pPr is not None:
            pPr.addnext(start)
        else:
            p_el.insert(0, start)
        p_el.append(end)
        p_el.append(ref_r)
        return True
    except Exception as exc:
        logger.debug("Comment failed: %s", exc)
        return False


def _find_runs_for_text(para, needle):
    """
    Return the list of runs in *para* whose combined text contains *needle*.
    Tries exact match first, then case-insensitive.
    """
    if not needle:
        return []
    full = "".join(r.text for r in para.runs)
    pos = full.find(needle)
    if pos == -1:
        pos = full.lower().find(needle.lower())
    if pos == -1:
        return []
    end = pos + len(needle)
    result, cur = [], 0
    for r in para.runs:
        rlen = len(r.text)
        if cur < end and cur + rlen > pos:
            result.append(r)
        cur += rlen
    return result


def _apply_style_to_runs(runs, style_name):
    """Apply a character style to a list of runs (best-effort)."""
    for r in runs:
        try:
            r.style = style_name
        except Exception:
            pass


# ─────────────────────────────────────────────────────────────────────────────
# CORE: SINGLE-PASS PROCESS & ANNOTATE
# ─────────────────────────────────────────────────────────────────────────────

def process_and_annotate_document(input_path, output_path, citation_style=None):
    """
    Single-pass pipeline:
      1. Auto-detect APA style (or use provided).
      2. Find every in-text citation (body paragraphs only, before <ref-open>).
      3. Find every reference entry (between <ref-open> … <ref-close>).
      4. Match citations → references:
           a. Exact key match
           b. Smart / normalised match (handles «and» vs «&», et al, abbreviations)
           c. Fuzzy spelling match (SequenceMatcher ≥ 0.80)
      5. For each body paragraph:
           • Locate citation run(s) and apply  cite_bib  character style
             (green highlight = matched, yellow = unmatched)
           • Add Word comment for MISSING REFERENCE (no match at all)
           • Add Word comment for SPELLING MISMATCH (close but different spelling)
      6. For unused references add Word comment UNUSED REFERENCE.
      7. Save annotated doc to output_path.
    Returns a summary dict.
    """
    import difflib
    from docx.enum.text import WD_COLOR_INDEX

    # ── 0. Setup ────────────────────────────────────────────────────────────
    if citation_style is None:
        citation_style = "apa"

    try:
        parser = get_parser(citation_style)
    except Exception:
        parser = get_parser("apa")
        citation_style = "apa"

    doc = Document(input_path)
    _ensure_cite_bib_style(doc)

    paragraphs_text = [p.text for p in doc.paragraphs]

    # ── 1. Find citations (body only, stop at <ref-open>) ──────────────────
    citations = {}          # cite_key -> {author, year, display, type, raw}
    cite_para_idx = {}      # cite_key -> list[paragraph_index (0-based)]

    for idx, para_text in enumerate(paragraphs_text):
        if "<ref-open>" in para_text:
            break
        found = parser.parse_citation(para_text)
        for cite in found:
            author = cite.get("author", "").strip()
            year   = cite.get("year",   "").strip()
            if not author or not year:
                continue
            key = normalize_citation_key(author, year)
            if key not in citations:
                citations[key] = {
                    "author":  author,
                    "year":    year,
                    "display": f"({author}, {year})" if cite.get("type") == "parenthetical"
                               else f"{author} ({year})",
                    "type":    cite.get("type", "parenthetical"),
                    "raw":     cite.get("raw", ""),
                }
            cite_para_idx.setdefault(key, [])
            if idx not in cite_para_idx[key]:
                cite_para_idx[key].append(idx)

    # ── 2. Find references (between <ref-open> … <ref-close>) ──────────────
    references      = {}  # ref_key -> {author, year, full_author, display}
    ref_para_idx    = {}  # ref_key -> paragraph_index (0-based)
    in_refs         = False

    for idx, para_text in enumerate(paragraphs_text):
        stripped = para_text.strip()
        if "<ref-open>"  in stripped: in_refs = True;  continue
        if "<ref-close>" in stripped: in_refs = False; continue
        if not in_refs or not stripped:
            continue
        ref_data = parser.parse_reference(stripped)
        if not ref_data:
            continue
        author_disp = ref_data.get("author", "")
        year        = ref_data.get("year",   "")
        full_author = ref_data.get("full_author", author_disp)
        key = normalize_citation_key(author_disp, year)
        if key not in references:
            references[key] = {
                "author":      author_disp,
                "year":        year,
                "full_author": full_author,
                "display":     f"{author_disp} ({year})",
                "text":        stripped[:150],
            }
            ref_para_idx[key] = idx

    # ── 3. Match every citation → a reference ──────────────────────────────
    # Returns (ref_key | None, match_type)
    def _match(cite_key, cite_data):
        # a. Exact key
        if cite_key in references:
            return cite_key, "exact"
        # b. Smart / normalised
        smart = check_smart_match(cite_data, references)
        if smart:
            return smart, "smart"
        # c. Fuzzy spelling (author only, same year)
        ca_norm = normalize_text_for_comparison(cite_data["author"])
        best_k, best_r = None, 0.0
        for rk, rd in references.items():
            if rd["year"] != cite_data["year"]:
                continue
            ra_norm = normalize_text_for_comparison(rd["author"])
            ratio = difflib.SequenceMatcher(None, ca_norm, ra_norm).ratio()
            if ratio > best_r:
                best_r, best_k = ratio, rk
        if best_k and best_r >= 0.80:
            return best_k, "spelling" if best_r < 0.97 else "smart"
        return None, "none"

    matched_pairs  = {}   # cite_key -> ref_key
    match_types    = {}   # cite_key -> "exact"|"smart"|"spelling"|"none"
    matched_refs   = set()

    for ck, cd in citations.items():
        rk, mt = _match(ck, cd)
        match_types[ck] = mt
        if rk:
            matched_pairs[ck] = rk
            matched_refs.add(rk)

    # ── 4. Build summary counts ─────────────────────────────────────────────
    missing_citations  = [ck for ck in citations if ck not in matched_pairs]
    spelling_mismatches = [ck for ck, mt in match_types.items() if mt == "spelling"]
    unused_references  = [rk for rk in references if rk not in matched_refs]

    # ── 5. Annotate body paragraphs ─────────────────────────────────────────
    comments_added = 0
    styles_applied = 0

    for cite_key, para_indices in cite_para_idx.items():
        cd        = citations[cite_key]
        is_matched = cite_key in matched_pairs
        is_spelling = match_types.get(cite_key) == "spelling"
        highlight  = WD_COLOR_INDEX.BRIGHT_GREEN if is_matched and not is_spelling \
                     else WD_COLOR_INDEX.YELLOW

        # Search candidates: raw text first, then display form
        candidates = []
        if cd.get("raw"):
            candidates.append(cd["raw"])
        if cd["display"] not in candidates:
            candidates.append(cd["display"])

        for pidx in para_indices:
            if pidx >= len(doc.paragraphs):
                continue
            para = doc.paragraphs[pidx]

            # Find runs
            target_runs = []
            for cand in candidates:
                target_runs = _find_runs_for_text(para, cand)
                if target_runs:
                    break
            # Fallback: search by year string
            if not target_runs and cd["year"]:
                target_runs = _find_runs_for_text(para, cd["year"])

            # Apply cite_bib style + highlight
            if target_runs:
                _apply_style_to_runs(target_runs, "cite_bib")
                for r in target_runs:
                    try:
                        r.font.highlight_color = highlight
                    except Exception:
                        pass
                styles_applied += 1

            # Add comment if needed
            if not is_matched:
                msg = (f"MISSING REFERENCE: '{cd['display']}' has no matching entry "
                       f"in the bibliography.")
                if _add_word_comment(doc, para, msg):
                    comments_added += 1

            elif is_spelling:
                ref_data  = references[matched_pairs[cite_key]]
                msg = (f"SPELLING MISMATCH: Cited as '{cd['author']}' but bibliography "
                       f"has '{ref_data['author']}'. Please verify spelling.")
                if _add_word_comment(doc, para, msg):
                    comments_added += 1

    # ── 6. Annotate unused references ───────────────────────────────────────
    for rk in unused_references:
        pidx = ref_para_idx.get(rk)
        if pidx is None or pidx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[pidx]
        msg  = (f"UNUSED REFERENCE: '{references[rk]['display']}' is in the bibliography "
                f"but is never cited in the text.")
        if _add_word_comment(doc, para, msg):
            comments_added += 1

    # ── 7. Save ─────────────────────────────────────────────────────────────
    doc.save(output_path)

    return {
        "total_citations":    len(citations),
        "total_references":   len(references),
        "matched_count":      len(matched_pairs),
        "missing_citations":  len(missing_citations),
        "spelling_mismatches":len(spelling_mismatches),
        "unused_references":  len(unused_references),
        "styles_applied":     styles_applied,
        "comments_added":     comments_added,
        "citation_style":     citation_style.upper(),
        # detail lists for report
        "_missing":  [citations[k]["display"]           for k in missing_citations],
        "_spelling": [(citations[k]["display"],
                       references[matched_pairs[k]]["author"])
                      for k in spelling_mismatches if k in matched_pairs],
        "_unused":   [references[k]["display"]          for k in unused_references],
    }


def generate_report(results, filename):
    """Build a plain-text validation report from process_and_annotate_document results."""
    total_issues = (results["missing_citations"]
                    + results["spelling_mismatches"]
                    + results["unused_references"])
    lines = []
    lines.append(f"STATUS: APA Name/Year: {total_issues} issue(s)")
    lines.append("")
    lines.append("=" * 60)
    lines.append("APA CITATION VALIDATION REPORT")
    lines.append("=" * 60)
    lines.append(f"\nDocument : {filename}")
    lines.append(f"Style    : {results.get('citation_style', 'APA')}")
    lines.append("-" * 60)
    lines.append("\nSUMMARY:")
    lines.append(f"  In-text citations found   : {results['total_citations']}")
    lines.append(f"  Bibliography entries found : {results['total_references']}")
    lines.append(f"  Matched (valid)           : {results['matched_count']}")
    lines.append(f"  Missing references        : {results['missing_citations']}")
    lines.append(f"  Spelling mismatches       : {results['spelling_mismatches']}")
    lines.append(f"  Unused references         : {results['unused_references']}")
    lines.append(f"  cite_bib styles applied   : {results['styles_applied']}")
    lines.append(f"  Word comments added       : {results['comments_added']}")

    if results["_missing"]:
        lines.append("\n" + "-" * 60)
        lines.append("MISSING REFERENCES (cited but not in bibliography):")
        lines.append("-" * 60)
        for item in results["_missing"]:
            lines.append(f"  {item}")

    if results["_spelling"]:
        lines.append("\n" + "-" * 60)
        lines.append("SPELLING MISMATCHES:")
        lines.append("-" * 60)
        for cited, ref_auth in results["_spelling"]:
            lines.append(f"  Cited:  {cited}")
            lines.append(f"  Ref:    {ref_auth}")

    if results["_unused"]:
        lines.append("\n" + "-" * 60)
        lines.append("UNUSED REFERENCES (in bibliography but never cited):")
        lines.append("-" * 60)
        for item in results["_unused"]:
            lines.append(f"  {item}")

    lines.append("\n" + "=" * 60)
    lines.append("END OF REPORT")
    lines.append("=" * 60)
    return "\n".join(lines)


# ─────────────────────────────────────────────────────────────────────────────
# ReferencesStructing optional dependency
# ─────────────────────────────────────────────────────────────────────────────
import shutil
import uuid
from pathlib import Path

try:
    import ReferencesStructing as RS
    RS_AVAILABLE = True
except Exception as _rs_err:
    RS = None  # type: ignore
    RS_AVAILABLE = False
    logger.warning(f"ReferencesStructing not available: {_rs_err}")


    
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file selected', 'error')
            return redirect(request.url)
        
        files = request.files.getlist('file')
        if not files or files[0].filename == '':
            flash('No file selected', 'error')
            return redirect(request.url)

        # Options
        check_validation = request.form.get('check_validation') == 'yes'
        check_structuring = request.form.get('check_structuring') == 'yes'
        citation_style = request.form.get('citation_style', 'auto')
        if citation_style == 'auto': citation_style = None

        if not check_validation and not check_structuring:
            flash('Please select at least one option (Reference Check or Structuring)', 'warning')
            return redirect(request.url)

        # Create Batch Directory
        batch_id = str(uuid.uuid4())[:8]
        batch_dir = os.path.join(app.config['UPLOAD_FOLDER'], f"batch_{batch_id}")
        os.makedirs(batch_dir, exist_ok=True)
        
        results_map = [] # To store (file_path, arcname)

        try:
            for file in files:
                if not (file.filename and allowed_file(file.filename)):
                    continue
                
                original_filename = secure_filename(file.filename)
                file_path = os.path.join(batch_dir, original_filename)
                file.save(file_path)

                # Track current working document
                current_doc_path = file_path
                current_doc_name = original_filename

                # 1. STRUCTURING PHASE
                if check_structuring:
                    if not RS_AVAILABLE:
                        logger.warning("Structuring skipped — ReferencesStructing module not available.")
                        flash("Structuring module not available. Validation only.", "warning")
                    else:
                        try:
                            logger.info(f"Structuring references for {original_filename}")
                            struct_res = RS.process_docx_file(Path(file_path), Path(batch_dir))

                            structured_path = str(struct_res['output_docx'])
                            log_path = str(struct_res['log_file'])

                            if os.path.exists(structured_path):
                                results_map.append((structured_path, f"Structured/{current_doc_name.replace('.docx', '_Structured.docx')}"))
                                results_map.append((log_path, f"Logs/{current_doc_name.replace('.docx', '_Structuring_Log.txt')}"))
                                # Hand off structured doc to validation phase
                                current_doc_path = structured_path
                                current_doc_name = os.path.basename(structured_path)
                        except Exception as e:
                            logger.error(f"Error structuring {original_filename}: {e}")
                            results_map.append((file_path, f"Errors/{original_filename}_Structuring_Failed.docx"))

                # 2. VALIDATION PHASE
                if check_validation:
                    try:
                        logger.info(f"Validating {current_doc_name}")

                        # Output path for the annotated DOCX
                        annotated_filename = current_doc_name.replace('.docx', '_Annotated.docx')
                        annotated_path = os.path.join(batch_dir, annotated_filename)

                        # Single-pass: APA check/fix, cite_bib style, match, comments
                        processor = CitationProcessor(current_doc_path)
                        report    = processor.process(annotated_path)

                        # Save plain-text report
                        report_text = (
                            f"Document: {current_doc_name}\n\n"
                            + report.summary()
                        )
                        report_filename = f"{current_doc_name}_Report.txt"
                        report_path = os.path.join(batch_dir, report_filename)
                        with open(report_path, 'w', encoding='utf-8') as f:
                            f.write(report_text)

                        results_map.append((report_path,    f"Reports/{report_filename}"))
                        results_map.append((annotated_path, f"Annotated/{annotated_filename}"))

                        s = report.stats
                        logger.info(
                            f"Validation done — matched:{s.get('matched',0)} "
                            f"missing:{s.get('missing',0)} "
                            f"year:{s.get('year_mismatch',0)} "
                            f"spelling:{s.get('spelling_mismatch',0)} "
                            f"unused:{s.get('unused',0)}"
                        )
                    except Exception as e:
                        logger.error(f"Error validating {current_doc_name}: {e}", exc_info=True)

            # GENERATE ZIP
            if not results_map:
                flash('No results generated. Please check files and try again.', 'error')
                return redirect(request.url)

            memory_file = io.BytesIO()
            with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zf:
                for src_path, arc_name in results_map:
                    if os.path.exists(src_path):
                        zf.write(src_path, arc_name)
            
            memory_file.seek(0)
            
            # Use send_file with a callback to clean up? 
            # Flask send_file doesn't support cleanup callback easily in older versions, 
            # but we can trust OS temp cleaning or do it in a finally block if we weren't returning
            # Since we are returning a stream, we can't delete immediately.
            # Best practice: schedule a cleanup or trust the unique dir is small enough until standard purge.
            # OR read bytes and delete.
            
            response = make_response(send_file(
                memory_file,
                mimetype='application/zip',
                as_attachment=True,
                download_name=f"Processed_Results_{batch_id}.zip"
            ))
            
            # Set cookie for frontend to detect download completion
            token = request.form.get('download_token')
            if token:
                response.set_cookie('download_token', token, max_age=10, path='/')
                
            return response

        except Exception as e:
            logger.error(f"Batch processing error: {e}", exc_info=True)
            flash(f'Error processing files: {str(e)}', 'error')
            return redirect(request.url)

    return render_template('index.html')

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
