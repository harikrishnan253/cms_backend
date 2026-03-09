"""
validation_core.py

Responsibilities:
  1. Detect whether in-text citations are APA-formatted.
  2. Fix non-APA citations to correct APA 7th format.
  3. Apply 'cite_bib' character style to every in-text citation.
  4. Highlight:
       GREEN  — citation matched to a reference
       YELLOW — citation not matched / year or spelling mismatch
  5. Insert Word comments for:
       - Missing reference  (citation has no bibliography entry)
       - Spelling mismatch  (close match found but author spelling differs)
       - Year mismatch      (author matches but year differs)
       - Unused reference   (bibliography entry never cited)
       - Format fixed       (citation was not APA, was auto-corrected)

Fixed from original proposal:
  - BUG 1: xml.etree.ElementTree → lxml (was producing ns0: namespace prefix)
  - BUG 2: style membership check uses [s.name for s in doc.styles]
  - BUG 3: replace_text_in_para uses direct element insertion, not add_run
  - BUG 4: initials filter tightened to catch "A. B." patterns
"""

import re
import logging
import difflib
from typing import Optional, List, Dict, Tuple, Any
from collections import defaultdict

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from lxml import etree as _ET      # FIX BUG 1 — lxml keeps w: prefix

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# XML / NAMESPACE CONSTANTS
# ─────────────────────────────────────────────────────────────────────────────

_W_NS    = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_MC_NS   = "http://schemas.openxmlformats.org/markup-compatibility/2006"
_W14_NS  = "http://schemas.microsoft.com/office/word/2010/wordml"
_NSMAP   = {"w": _W_NS, "mc": _MC_NS, "w14": _W14_NS}
_COMMENTS_REL = ("http://schemas.openxmlformats.org/officeDocument/"
                 "2006/relationships/comments")
_COMMENTS_CT  = ("application/vnd.openxmlformats-officedocument"
                 ".wordprocessingml.comments+xml")


# ─────────────────────────────────────────────────────────────────────────────
# APA FORMAT CHECKER & FIXER
# ─────────────────────────────────────────────────────────────────────────────

# APA parenthetical:  (Author, Year)  or  (Author et al., Year)
_RE_APA_PAREN = re.compile(
    r'\(([A-Z][^()]{1,80}?,\s*(?:19|20)\d{2}[a-z]?)\)',
    re.UNICODE
)
# APA narrative:  Author (Year)
_RE_APA_NARRATIVE = re.compile(
    r'([A-Z][A-Za-zÀ-ÖØ-öø-ÿ\-\' ]{1,60}?(?:\s+et\s+al\.)?)\s*\((\d{4}[a-z]?)\)',
    re.UNICODE
)
# AMA-style narrative: "Smith 2020" (no comma, no parens)
_RE_AMA_NARRATIVE = re.compile(
    r'\b([A-Z][a-z]+(?:\s+et\s+al\.)?)\s+((?:19|20)\d{2})\b(?![,\)])'
)


def is_apa_citation(text: str) -> bool:
    """Return True if *text* contains at least one APA-formatted citation."""
    return bool(_RE_APA_PAREN.search(text) or _RE_APA_NARRATIVE.search(text))


def has_non_apa_citation(text: str) -> bool:
    """Return True if *text* appears to contain AMA-style citations."""
    return bool(_RE_AMA_NARRATIVE.search(text))


def fix_citation_to_apa(text: str) -> Tuple[str, List[Dict]]:
    """
    Attempt to convert non-APA in-text citations to APA format.

    Returns:
        (fixed_text, list_of_changes)
    Each change dict has keys: original, fixed, type.
    """
    changes: List[Dict] = []
    result = text

    # Fix AMA narrative: "Smith 2020" → "Smith (2020)"
    def _fix_ama(m):
        orig  = m.group(0)
        fixed = f"{m.group(1)} ({m.group(2)})"
        if orig != fixed:
            changes.append({"original": orig, "fixed": fixed, "type": "ama_narrative"})
        return fixed

    result = _RE_AMA_NARRATIVE.sub(_fix_ama, result)

    return result, changes


# ─────────────────────────────────────────────────────────────────────────────
# AUTHOR / YEAR HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def normalize_author(s: str) -> str:
    """Lowercase, strip 'et al.', punctuation, extra whitespace."""
    s = re.sub(r"\s*et\s+al\.?", "", s, flags=re.IGNORECASE)
    s = re.sub(r"[.,&]", "", s)
    return re.sub(r"\s+", " ", s).strip().lower()


def extract_first_surname(author_str: str) -> str:
    tokens = normalize_author(author_str).split()
    return tokens[0] if tokens else ""


def fuzzy_author_match(a: str, b: str) -> float:
    return difflib.SequenceMatcher(
        None, normalize_author(a), normalize_author(b)
    ).ratio()

def check_acronym_match(cite_author: str, ref_full: str) -> bool:
    """Check if cyt_author is an acronym (like 'WHO') for the ref_full organization."""
    acronym = re.sub(r'[^A-Za-z]', '', cite_author)
    if not acronym or not acronym.isupper() or len(acronym) < 2:
        return False
    words = [w for w in re.findall(r'[A-Za-z]+', ref_full) if w.lower() not in ('and', 'of', 'for', 'the', 'in', 'on', 'to')]
    if not words:
        return False
    first_letters = "".join(w[0] for w in words).upper()
    it = iter(first_letters)
    return all(c in it for c in acronym)


# ─────────────────────────────────────────────────────────────────────────────
# CITATION → REFERENCE MATCHING
# ─────────────────────────────────────────────────────────────────────────────

def match_citation(cite_author: str, cite_year: str,
                   references: Dict[str, dict]) -> Tuple[Optional[str], str]:
    """
    Returns (ref_key | None, match_type).
    match_type: 'exact' | 'smart' | 'spelling_mismatch' | 'year_mismatch' | 'not_found'
    """
    cite_norm  = normalize_author(cite_author)
    cite_first = extract_first_surname(cite_author)
    cite_words = (set(re.findall(r"\b[a-z]{2,}\b", cite_norm))
                  - {"and", "et", "al"})

    year_candidates: List[str] = []
    close_spelling:  List[Tuple[float, str, str]] = []   # (ratio, key, ref_year)

    for ref_key, ref in references.items():
        ref_full  = ref.get("full_author", ref["author"])
        ref_norm  = normalize_author(ref_full)
        ref_year  = ref["year"]
        ref_first = extract_first_surname(ref_full)
        year_ok   = (cite_year == ref_year)

        # 1. Exact normalised
        if cite_norm == ref_norm:
            if year_ok:
                return ref_key, "exact"
            year_candidates.append(ref_key)
            continue

        # 2. First-surname + year
        if cite_first and ref_first and cite_first == ref_first and year_ok:
            return ref_key, "smart"

        # 3. Subset ("Smith & Jones" ⊆ "Smith, A., Jones, B., Brown")
        ref_words = set(re.findall(r"\b[a-z]{2,}\b", ref_norm))
        if cite_words and cite_words.issubset(ref_words) and year_ok:
            return ref_key, "smart"

        # 4. Acronym match (e.g. 'APA' -> 'American Psychiatric Association')
        if check_acronym_match(cite_author, ref_full):
            if year_ok:
                return ref_key, "smart"
            year_candidates.append(ref_key)
            continue

        # 4. Fuzzy spelling
        ratio = fuzzy_author_match(cite_author, ref_full)
        if ratio >= 0.80:
            close_spelling.append((ratio, ref_key, ref_year))

    # Year mismatch wins over spelling mismatch
    if year_candidates:
        return year_candidates[0], "year_mismatch"

    if close_spelling:
        close_spelling.sort(reverse=True)
        _, ref_key, _ = close_spelling[0]
        return ref_key, "spelling_mismatch"

    return None, "not_found"


# ─────────────────────────────────────────────────────────────────────────────
# WORD DOCUMENT STYLE / RUN UTILITIES
# ─────────────────────────────────────────────────────────────────────────────

def ensure_style(doc: Document, style_name: str) -> None:
    """Create *style_name* character style if it doesn't exist. (FIX BUG 2)"""
    try:
        existing_names = [s.name for s in doc.styles]   # was: "style_name not in doc.styles"
        if style_name not in existing_names:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
    except Exception:
        pass


def _get_run_map(paragraph) -> List[Tuple[int, int, Any]]:
    """Return [(start, end, run), …] for every run in *paragraph*."""
    result, pos = [], 0
    for run in paragraph.runs:
        n = len(run.text)
        result.append((pos, pos + n, run))
        pos += n
    return result


def apply_style_to_citation(paragraph, search_text: str,
                             style_name: str, highlight: WD_COLOR_INDEX,
                             doc: Document) -> bool:
    """
    Find *search_text* in *paragraph*, apply *style_name* + *highlight* to
    the matching runs and return True.  Falls back to case-insensitive search.

    FIX BUG 3: uses element-level insertion instead of paragraph.add_run()
    to avoid corrupting run order in paragraphs with many runs.
    """
    if not search_text:
        return False

    run_map = _get_run_map(paragraph)
    if not run_map:
        return False
    full = "".join(r.text for _, _, r in run_map)

    start = full.find(search_text)
    if start == -1:
        start = full.lower().find(search_text.lower())
    if start == -1:
        return False

    end = start + len(search_text)

    # Runs that overlap [start, end)
    affected = [(s, e, r) for s, e, r in run_map if s < end and e > start]
    if not affected:
        return False

    s0, _, r0 = affected[0]
    sn, _, rn = affected[-1]

    pre  = r0.text[:start - s0]
    post = rn.text[end - sn:]

    # Clear every affected run
    for _, _, r in affected:
        r.text = ""

    first_run = r0
    first_run.text = pre

    # Build the styled citation run using lxml so we control placement
    first_el = first_run._element
    p_el     = paragraph._element

    # Create a new <w:r> with the right style and highlight
    new_r = OxmlElement("w:r")

    # Copy rPr from the first run if available
    rPr_src = first_el.find(qn("w:rPr"))
    if rPr_src is not None:
        import copy
        new_r.append(copy.deepcopy(rPr_src))

    # Apply character style in rPr
    rPr = new_r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        new_r.insert(0, rPr)

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), style_name)
    rPr.insert(0, rStyle)

    # Apply highlight
    hl_map = {
        WD_COLOR_INDEX.BRIGHT_GREEN: "green",
        WD_COLOR_INDEX.YELLOW:       "yellow",
    }
    if highlight in hl_map:
        hl_el = OxmlElement("w:highlight")
        hl_el.set(qn("w:val"), hl_map[highlight])
        rPr.append(hl_el)

    # Text element
    t_el = OxmlElement("w:t")
    t_el.text = search_text
    if search_text != search_text.strip():
        t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(t_el)

    # Insert right after first_el
    first_el.addnext(new_r)

    # Insert post-text run if needed
    if post:
        post_r = OxmlElement("w:r")
        post_t = OxmlElement("w:t")
        post_t.text = post
        if post != post.strip():
            post_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        post_r.append(post_t)
        new_r.addnext(post_r)

    return True


# ─────────────────────────────────────────────────────────────────────────────
# WORD COMMENT INSERTION  (FIX BUG 1 — uses lxml, not xml.etree)
# ─────────────────────────────────────────────────────────────────────────────

def _get_or_create_comments_part(doc: Document):
    """Return (or create) the /word/comments.xml Part."""
    doc_part = doc.part
    for rel in doc_part.rels.values():
        if rel.reltype == _COMMENTS_REL:
            return rel.target_part
    root = _ET.Element(f"{{{_W_NS}}}comments", nsmap=_NSMAP)
    root.set(f"{{{_MC_NS}}}Ignorable", "w14 wp14")
    blob = _ET.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    part = Part(PackURI("/word/comments.xml"), _COMMENTS_CT, blob, doc_part.package)
    doc_part.package.add_part(part)
    doc_part.relate_to(part, _COMMENTS_REL)
    return part


def add_word_comment(doc: Document, paragraph, comment_text: str,
                     author: str = "Ref Validator",
                     initials: str = "RV") -> bool:
    """
    Attach a Word comment to *paragraph* using lxml — keeps w: namespace prefix.
    """
    if not paragraph.runs:
        return False
    try:
        cp = _get_or_create_comments_part(doc)

        try:
            tree = _ET.fromstring(cp._blob)
        except Exception:
            tree = _ET.Element(f"{{{_W_NS}}}comments", nsmap=_NSMAP)

        existing_ids = [
            int(c.get(f"{{{_W_NS}}}id") or 0)
            for c in tree.findall(f"{{{_W_NS}}}comment")
        ]
        new_id = max(existing_ids, default=-1) + 1

        # <w:comment>
        cel = _ET.SubElement(tree, f"{{{_W_NS}}}comment")
        cel.set(f"{{{_W_NS}}}id",       str(new_id))
        cel.set(f"{{{_W_NS}}}author",   author)
        cel.set(f"{{{_W_NS}}}initials", initials)
        pel = _ET.SubElement(cel, f"{{{_W_NS}}}p")
        rel = _ET.SubElement(pel, f"{{{_W_NS}}}r")
        tel = _ET.SubElement(rel,  f"{{{_W_NS}}}t")
        tel.text = comment_text

        cp._blob = _ET.tostring(tree, xml_declaration=True,
                                encoding="UTF-8", standalone=True)

        # Anchor markers
        p_elem = paragraph._element
        start_el = OxmlElement("w:commentRangeStart")
        start_el.set(qn("w:id"), str(new_id))
        end_el = OxmlElement("w:commentRangeEnd")
        end_el.set(qn("w:id"), str(new_id))
        ref_run = OxmlElement("w:r")
        ref_ref = OxmlElement("w:commentReference")
        ref_ref.set(qn("w:id"), str(new_id))
        ref_run.append(ref_ref)

        pPr = p_elem.find(qn("w:pPr"))
        if pPr is not None:
            pPr.addnext(start_el)
        else:
            p_elem.insert(0, start_el)
        p_elem.append(end_el)
        p_elem.append(ref_run)
        return True

    except Exception as exc:
        logger.debug("Comment insertion failed: %s", exc)
        return False


# ─────────────────────────────────────────────────────────────────────────────
# BIBLIOGRAPHY PARSER HELPER
# ─────────────────────────────────────────────────────────────────────────────

# FIX BUG 4: tightened initials filter — matches "A.", "A. B.", "AB." etc.
_RE_INITIALS = re.compile(r"^(?:[A-Z]\.?\s*)+$")


def _format_author_display(raw: str) -> str:
    """
    Produce a short citation-style author display from a bibliography author string:
      1 author  → "Smith"
      2 authors → "Smith & Jones"
      3+        → "Smith et al."
    """
    parts    = re.split(r"\s*[,&]\s*", raw)
    surnames = [
        p.strip() for p in parts
        if p.strip()
        and len(p.strip()) > 1
        and not _RE_INITIALS.match(p.strip())   # FIX BUG 4
    ]
    if not surnames:
        return raw.strip()
    if len(surnames) == 1:
        return surnames[0]
    if len(surnames) == 2:
        return f"{surnames[0]} & {surnames[1]}"
    return f"{surnames[0]} et al."


def extract_ref_author_year(text: str) -> Tuple[str, str]:
    """
    Extract (author_display_string, year) from a bibliography entry line.
    Handles APA: "Author, A. A. (Year)." and AMA: "Author A. Title. J. Year;…"
    """
    # APA: "...author block... (Year[a-z]?)"
    m = re.match(
        r"^(?P<authors>[^(]{4,150}?)\s*\((?P<year>(?:19|20)\d{2}[a-z]?|n\.d\.)\)",
        text
    )
    if m:
        raw_authors  = m.group("authors").strip().rstrip(",").strip()
        year         = m.group("year")
        return _format_author_display(raw_authors), year

    # AMA / plain: find first 4-digit year, take text-before-first-period as authors
    m2 = re.search(r"\b((?:19|20)\d{2})\b", text)
    if m2:
        year  = m2.group(1)
        parts = text.split(".")
        if parts:
            return _format_author_display(parts[0].strip()), year

    return "", ""


# ─────────────────────────────────────────────────────────────────────────────
# CITATION EXTRACTOR
# ─────────────────────────────────────────────────────────────────────────────

def extract_citations_from_text(text: str) -> List[Dict]:
    """
    Return list of citation dicts found in *text*:
      {raw, author, year, type, start, end}
    Detects both parenthetical and narrative APA citations.
    """
    found: List[Dict] = []

    # 1. Parenthetical: (Author, Year) or (Author, Year; Author2, Year2)
    for m in re.finditer(r'\(([^()]+)\)', text):
        content = m.group(1)
        for sub_m in re.finditer(r'([^;()0-9]+?),\s*((?:19|20)\d{2}[a-z]?|n\.d\.)', content):
            raw_sub = sub_m.group(0).strip()
            author = sub_m.group(1).strip()
            year = sub_m.group(2).strip()
            
            # Avoid matching prefixes like "e.g., "
            author = re.sub(r'^(?:e\.g\.,?|i\.e\.,?|see|cf\.)\s+', '', author, flags=re.IGNORECASE).strip()
            
            # Allow well-known prefixes or acronyms containing numbers/brackets/hyphens. 
            if not author or len(author) < 2:
                continue

            found.append({
                "raw":    raw_sub,
                "author": author,
                "year":   year,
                "type":   "parenthetical",
                "start":  m.start() + sub_m.start() + 1,
                "end":    m.start() + sub_m.end() + 1,
            })

    # 2. Narrative: Author (Year)  — skip ranges already captured above
    occupied = {(d["start"], d["end"]) for d in found}
    for m in re.finditer(
        r'([A-Za-zÀ-ÖØ-öø-ÿ\-\' \[\]&]+?(?:\s+et\s+al\.?)?)\s*\(((?:19|20)\d{2}[a-z]?|n\.d\.)\)',
        text
    ):
        if any(m.start() >= s and m.end() <= e for s, e in occupied):
            continue
            
        raw_author = m.group(1).strip()
        year = m.group(2).strip()
        
        # Strip narrative intro text
        raw_author = re.sub(r'^(?:according\s+to|suggested\s+by|offered\s+by|idea\s+is|by)\s+', '', raw_author, flags=re.IGNORECASE).strip()
        
        # Target end to capture just author names, discarding previous sentence parts.
        # Check for APA structure: [Capitalized words] and maybe van/der + et al.
        m2 = re.search(r'\b(?:(?:van|der|de|la|von|da|di)\s+)*[A-ZÀ-Ö][A-Za-zÀ-ÖØ-öø-ÿ\-\']+(?:,?\s*(?:and|&)\s*(?:(?:van|der|de|la|von|da|di)\s+)*[A-ZÀ-Ö]?[A-Za-zÀ-ÖØ-öø-ÿ\-\']+)*(?:\s+et\s+al\.?)?$', raw_author)
        
        if m2:
            author = m2.group(0)
            raw = f"{author} ({year})"
            start_off = raw_author.rfind(author)
        else:
            # simple fallback
            author = raw_author.split()[-1] if raw_author else ""
            raw = f"{author} ({year})"
            start_off = raw_author.rfind(author) if author else 0
            
        if not author or len(author) < 2 or (not author[0].isupper() and not author.startswith(('van', 'der', 'de'))):
            continue

        found.append({
            "raw":    raw,
            "author": author,
            "year":   year,
            "type":   "narrative",
            "start":  m.start() + start_off,
            "end":    m.end(),
        })

    return found


# ─────────────────────────────────────────────────────────────────────────────
# MAIN PROCESSOR
# ─────────────────────────────────────────────────────────────────────────────

class CitationProcessor:
    """
    Single-pass processor that:
      1. Detects & fixes non-APA citations.
      2. Applies cite_bib character style + colour highlight.
      3. Validates every citation against the bibliography.
      4. Inserts Word comments for every issue found.

    Usage:
        processor = CitationProcessor("input.docx")
        report    = processor.run()
        processor.save("output.docx")
        print(report.summary())
    """

    STYLE_NAME = "cite_bib"
    GREEN  = WD_COLOR_INDEX.BRIGHT_GREEN
    YELLOW = WD_COLOR_INDEX.YELLOW

    def __init__(self, doc_path: str):
        self.doc_path = doc_path
        self.doc      = Document(doc_path)
        ensure_style(self.doc, self.STYLE_NAME)

        self.references:     Dict[str, dict] = {}
        self.ref_para_idx:   Dict[str, int]  = {}
        self.issues:         List[dict]      = []
        self._stats:         dict            = defaultdict(int)

    # ── Step 1: Build reference index ───────────────────────────────────────

    def _parse_bibliography(self) -> None:
        in_refs = False
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if "<ref-open>" in text.lower():
                in_refs = True
                continue
            if "<ref-close>" in text.lower():
                in_refs = False
                continue
            if not (in_refs and text):
                continue

            author, year = extract_ref_author_year(text)
            if not author or not year:
                continue

            key = f"{normalize_author(author)}|{year}"
            if key not in self.references:
                self.references[key] = {
                    "author":      author,
                    "full_author": author,
                    "year":        year,
                    "display":     f"{author} ({year})",
                    "text":        text,
                    "para_idx":    i,
                    "cited":       False,
                }
                self.ref_para_idx[key] = i

    # ── Main run ─────────────────────────────────────────────────────────────

    def run(self) -> "ValidationReport":
        self._parse_bibliography()
        cited_ref_keys: set = set()

        for para_idx, para in enumerate(self.doc.paragraphs):
            text = para.text
            if not text.strip():
                continue
            if "<ref-open>" in text.lower():
                break  # entered bibliography — stop scanning

            # ── A. Detect & flag non-APA format ─────────────────────────────
            if has_non_apa_citation(text) and not is_apa_citation(text):
                fixed_text, fix_changes = fix_citation_to_apa(text)
                for ch in fix_changes:
                    self.issues.append({
                        "type":      "format_fixed",
                        "para_idx":  para_idx,
                        "para":      para,
                        "original":  ch["original"],
                        "fixed":     ch["fixed"],
                        "message": (
                            f"FORMAT FIXED: '{ch['original']}' "
                            f"→ '{ch['fixed']}' (APA style applied)."
                        ),
                    })
                    self._stats["format_fixed"] += 1
                # Use fixed text for citation extraction below
                text = fixed_text

            # ── B. Extract APA citations and validate ────────────────────────
            for cite in extract_citations_from_text(text):
                raw    = cite["raw"]
                author = cite["author"]
                year   = cite["year"]

                ref_key, match_type = match_citation(author, year, self.references)

                if match_type in ("exact", "smart"):
                    # ✅ Matched — green
                    self.references[ref_key]["cited"] = True
                    cited_ref_keys.add(ref_key)
                    apply_style_to_citation(para, raw, self.STYLE_NAME,
                                            self.GREEN, self.doc)
                    self._stats["matched"] += 1

                elif match_type == "year_mismatch":
                    ref = self.references[ref_key]
                    apply_style_to_citation(para, raw, self.STYLE_NAME,
                                            self.YELLOW, self.doc)
                    self.issues.append({
                        "type":     "year_mismatch",
                        "para_idx": para_idx,
                        "para":     para,
                        "raw":      raw,
                        "message": (
                            f"YEAR MISMATCH: Cited year '{year}' but bibliography "
                            f"has '{ref['year']}' for author '{ref['author']}'."
                        ),
                    })
                    self._stats["year_mismatch"] += 1

                elif match_type == "spelling_mismatch":
                    ref = self.references[ref_key]
                    apply_style_to_citation(para, raw, self.STYLE_NAME,
                                            self.YELLOW, self.doc)
                    self.issues.append({
                        "type":     "spelling_mismatch",
                        "para_idx": para_idx,
                        "para":     para,
                        "raw":      raw,
                        "message": (
                            f"SPELLING MISMATCH: Cited as '{author}' but "
                            f"bibliography has '{ref['author']}'. "
                            f"Please check spelling."
                        ),
                    })
                    self._stats["spelling_mismatch"] += 1

                else:
                    # ❌ Not found — yellow
                    apply_style_to_citation(para, raw, self.STYLE_NAME,
                                            self.YELLOW, self.doc)
                    self.issues.append({
                        "type":     "missing_reference",
                        "para_idx": para_idx,
                        "para":     para,
                        "raw":      raw,
                        "message": (
                            f"MISSING REFERENCE: '{raw}' has no matching "
                            f"entry in the bibliography."
                        ),
                    })
                    self._stats["missing"] += 1

        # ── C. Flag unused references ────────────────────────────────────────
        for key, ref in self.references.items():
            if not ref["cited"] and key not in cited_ref_keys:
                pidx = ref.get("para_idx")
                self.issues.append({
                    "type":     "unused_reference",
                    "para_idx": pidx,
                    "para":     (self.doc.paragraphs[pidx]
                                 if pidx is not None else None),
                    "message": (
                        f"UNUSED REFERENCE: '{ref['display']}' is in the "
                        f"bibliography but never cited in the text."
                    ),
                })
                self._stats["unused"] += 1

        # ── D. Insert Word comments for all issues ───────────────────────────
        for issue in self.issues:
            para = issue.get("para")
            if para is None:
                continue
            try:
                add_word_comment(self.doc, para, issue["message"])
            except Exception as exc:
                logger.debug("Could not add comment: %s", exc)

        return ValidationReport(self.issues, self._stats, len(self.references))

    def save(self, output_path: str) -> None:
        self.doc.save(output_path)
        logger.info("Saved annotated document: %s", output_path)

    # Convenience: run + save in one call
    def process(self, output_path: str) -> "ValidationReport":
        report = self.run()
        self.save(output_path)
        return report


# ─────────────────────────────────────────────────────────────────────────────
# REPORT
# ─────────────────────────────────────────────────────────────────────────────

class ValidationReport:
    def __init__(self, issues: List[dict], stats: dict, total_refs: int):
        self.issues     = issues
        self.stats      = dict(stats)
        self.total_refs = total_refs

    def summary(self) -> str:
        s = self.stats
        lines = [
            "=" * 60,
            "CITATION VALIDATION SUMMARY",
            "=" * 60,
            f"  Matched citations:     {s.get('matched', 0)}",
            f"  Missing references:    {s.get('missing', 0)}",
            f"  Year mismatches:       {s.get('year_mismatch', 0)}",
            f"  Spelling mismatches:   {s.get('spelling_mismatch', 0)}",
            f"  Format fixes applied:  {s.get('format_fixed', 0)}",
            f"  Unused references:     {s.get('unused', 0)}",
            f"  Total bibliography:    {self.total_refs}",
            "-" * 60,
        ]
        type_order = [
            ("missing_reference",   "MISSING REFERENCES"),
            ("year_mismatch",       "YEAR MISMATCHES"),
            ("spelling_mismatch",   "SPELLING MISMATCHES"),
            ("format_fixed",        "FORMAT FIXES"),
            ("unused_reference",    "UNUSED REFERENCES"),
        ]
        for issue_type, heading in type_order:
            items = [i for i in self.issues if i["type"] == issue_type]
            if items:
                lines += ["", heading, "-" * 40]
                for item in items:
                    lines.append(
                        f"  Para {item.get('para_idx', '?'):>4}: {item['message']}"
                    )
        lines += ["", "=" * 60, "END OF REPORT", "=" * 60]
        return "\n".join(lines)

    def to_dict(self) -> dict:
        return {
            "stats":  self.stats,
            "issues": [
                {k: v for k, v in i.items() if k != "para"}
                for i in self.issues
            ],
        }

    @property
    def has_issues(self) -> bool:
        return bool(self.issues)
