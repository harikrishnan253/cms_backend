"""
Core document styling and tagging module.
Handles applying styles and tags to DOCX documents.
"""

import logging
import re
from typing import Literal, Dict, Any
from docx import Document
from docx.oxml.ns import qn
from .annotator import annotate_document, is_list_paragraph
from .logger_config import get_logger
from .hierarchy_manager import enforce_hierarchy

logger = get_logger(__name__)


from .rules_loader import get_rules_loader

def tag_tables(doc: Document, mode: Literal["style", "tag"] = "style") -> None:
    """
    Apply styles or tags to table cells.
    
    Args:
        doc: python-docx Document object
        mode: "style" to apply Word styles, "tag" to prefix text with [TAG]
    
    Raises:
        ValueError: If mode is invalid
    """
    if mode not in ("style", "tag"):
        raise ValueError(f"Invalid mode: {mode}. Must be 'style' or 'tag'")
    
    rules_loader = get_rules_loader()
    table_config = rules_loader.get_table_config()
    list_config = rules_loader.get_list_patterns()
    
    header_style_name = table_config.get("header_style", "Table Header")
    body_style_name = table_config.get("body_style", "Table Body")
    
    number_pattern = list_config.get("number_pattern", r"^\d+\.\s+")
    bullet_pattern = list_config.get("bullet_pattern", r"^[•\-\–]\s+")

    for table_idx, table in enumerate(doc.tables):
        try:
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if not text:
                            continue
                        
                        # Determine if header or body
                        is_header = text.isupper() and len(text) > 2
                        
                        if is_header:
                            tag = "TBCH"
                            style = header_style_name
                        else:
                            # Check for lists
                            is_list = is_list_paragraph(para)
                            
                            if re.match(number_pattern, text):
                                tag = "TNL-MID"
                                style = "TNL-MID"
                            elif re.match(bullet_pattern, text):
                                tag = "TBL-MID"
                                style = "TBL-MID"
                            elif is_list:
                                # Fallback for structural lists without text markers
                                tag = "TBL-MID"
                                style = "TBL-MID"
                            else:
                                tag = "TB"
                                style = body_style_name
                        
                        if mode == "style":
                            try:
                                para.style = style
                                logger.debug(f"Applied style '{style}' to table[{table_idx}].row[{row_idx}].cell[{cell_idx}]")
                            except KeyError:
                                logger.warning(f"Style '{style}' not found in template. Creating it as default Paragraph Style.")
                                try:
                                    # Create the style if missing
                                    styles = doc.styles
                                    new_style = styles.add_style(style, 1) # 1 is WD_STYLE_TYPE.PARAGRAPH
                                    # Optional: Set base style
                                    new_style.base_style = styles['Normal']
                                    para.style = style
                                except Exception as e2:
                                    logger.error(f"Failed to create style '{style}': {e2}")
                        else:  # tag mode
                            if para.runs:
                                try:
                                    para.runs[0].text = f"[{tag}] " + para.runs[0].text
                                    logger.debug(f"Added tag '{tag}' to table[{table_idx}].row[{row_idx}].cell[{cell_idx}]")
                                except Exception as e:
                                    logger.error(f"Failed to tag table cell: {e}")
        except Exception as e:
            logger.error(f"Error processing table {table_idx}: {e}")


from docx.enum.style import WD_STYLE_TYPE

def process_cross_references(doc: Document) -> None:
    """
    Process cross-references: wrap in brackets and apply character style.
    Warning: This recreates paragraph runs, identifying plain text matches. 
    Existing run-level formatting (bold/italic) inside the specific paragraph *might* be reset 
    if strictly text-based splitting is used. 
    """
    rules_loader = get_rules_loader()
    xref_config = rules_loader.get_cross_references()
    
    if not xref_config:
        return

    # Collect all paragraphs
    all_paras = list(doc.paragraphs)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                all_paras.extend(c.paragraphs)

    for rule_name, rule in xref_config.items():
        pattern = rule.get("pattern")
        style_name = rule.get("style", "Hyperlink")
        
        # Validate style exists or create it (Character style)
        try:
            doc.styles[style_name]
        except KeyError:
            # Create character style
            logger.info(f"Creating missing character style: {style_name}")
            try:
                styles = doc.styles
                char_style = styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
                # Dictionary check for base style?
                # char_style.base_style = styles['Default Paragraph Font'] 
            except Exception as e:
                logger.error(f"Failed to create character style {style_name}: {e}")

        # Compile regex once
        regex = re.compile(pattern)
        
        count = 0
        total_paras = len(all_paras)
        logger.info(f"Scanning {total_paras} paragraphs for pattern: {pattern[:20]}...")
        
        for para in all_paras:
            count += 1
            if count % 500 == 0:
                logger.info(f"Scanned {count}/{total_paras} paragraphs...")
                
            text = para.text
            if not text:
                continue
            
            # Skip checking if no match at all to save time
            if not regex.search(text):
                continue
                
            # Skip checking if no match at all to save time
            if not regex.search(text):
                continue
                
            # Iterate over runs to preserve formatting
            # We must iterate a copy because we might modify the list of runs
            
            # --- Better Implementation below ---
            # 1. Capture existing state
            original_runs = []
            for r in para.runs:
                original_runs.append({
                    "text": r.text,
                    "style": r.style,
                    "bold": r.bold,
                    "italic": r.italic,
                    "underline": r.underline,
                    "superscript": r.font.superscript if r.font else None,
                    "subscript": r.font.subscript if r.font else None,
                    "strike": r.font.strike if r.font else None,
                    "font": r.font.name if r.font else None,
                    "size": r.font.size if r.font else None,
                    "color": r.font.color.rgb if r.font and r.font.color else None,
                    "highlight": r.font.highlight_color if r.font else None
                })
            
            # 2. Clear
            para.clear()
            
            # 3. Rebuild with Splitting
            for run_data in original_runs:
                text_chunk = run_data["text"]
                matches = list(regex.finditer(text_chunk))
                
                if not matches:
                    # No cross-ref in this run, restore as is
                    new_run = para.add_run(text_chunk)
                    # Restore props
                    _restore_run_props(new_run, run_data)
                else:
                    # Split this run's text
                    cursor = 0
                    for m in matches:
                        m_start, m_end = m.span()
                        
                        # Pre-match text (inherit origin props)
                        if m_start > cursor:
                            r = para.add_run(text_chunk[cursor:m_start])
                            _restore_run_props(r, run_data)
                        
                        # Match text (Apply XREF style, but keep boldness if it was bold?)
                        # Usually Xref style overrides, but let's keep bold/italic valid if implicit.
                        r_match = para.add_run(text_chunk[m_start:m_end])
                        try:
                            r_match.style = style_name
                        except: pass
                        # Optional: Force keep bold if original was bold? 
                        # Usually character style handles format. User asked to "retain".
                        # If original was bold, we should probably keep it bold even if Hyperlinked.
                        if run_data["bold"]: r_match.bold = True
                        if run_data["italic"]: r_match.italic = True
                        
                        cursor = m_end
                    
                    # Post-match text
                    if cursor < len(text_chunk):
                        r = para.add_run(text_chunk[cursor:])
                        _restore_run_props(r, run_data)

def _restore_run_props(run, data):
    # Helper to restore ALL formatting
    if data["style"]: 
        try: run.style = data["style"]
        except: pass
    
    # Boolean properties
    if data["bold"] is not None: run.bold = data["bold"]
    if data["italic"] is not None: run.italic = data["italic"]
    if data["underline"] is not None: run.underline = data["underline"]
    if data.get("superscript") is not None: run.font.superscript = data["superscript"]
    if data.get("subscript") is not None: run.font.subscript = data["subscript"]
    if data.get("strike") is not None: run.font.strike = data["strike"]
    # Double strike is rare but possible
    # if data.get("double_strike") is not None: run.font.double_strike = data["double_strike"]
    
    # Advanced Font properties
    if data["font"]: run.font.name = data["font"]
    if data["size"]: run.font.size = data["size"]
    if data["color"]: run.font.color.rgb = data["color"]
    if data.get("highlight"): run.font.highlight_color = data["highlight"]

def process_docx(
    input_path: str,
    output_path: str,
    mode: Literal["style", "tag"] = "style"
) -> Dict[str, Any]:
    """
    Process a DOCX document by annotating and styling paragraphs and tables.
    
    Args:
        input_path: Path to input DOCX file
        output_path: Path to save output DOCX file
        mode: "style" to apply Word styles, "tag" to prefix text with [TAG]
    
    Returns:
        Dictionary with processing results:
        - success: bool
        - paragraphs_processed: int
        - tables_processed: int
        - errors: list of error messages
    
    Raises:
        FileNotFoundError: If input file doesn't exist
        ValueError: If mode is invalid
    """
    if mode not in ("style", "tag"):
        raise ValueError(f"Invalid mode: {mode}. Must be 'style' or 'tag'")
    
    result = {
        "success": False,
        "paragraphs_processed": 0,
        "tables_processed": 0,
        "errors": []
    }
    
    try:
        logger.info(f"Opening document: {input_path}")
        doc = Document(input_path)
        
        logger.info(f"Annotating document with mode: {mode}")
        annotations = annotate_document(doc)
        
        if mode == "style":
            logger.info("Enforcing heading hierarchy and validation")
            annotations = enforce_hierarchy(annotations)
        
        # Process annotations
        for idx, item in enumerate(annotations):
            try:
                para = item["para"]
                tag = item["tag"]
                style = item["style"]
                
                if mode == "style":
                    try:
                        para.style = style
                        result["paragraphs_processed"] += 1
                        logger.debug(f"Paragraph {idx}: Applied style '{style}' (tag: {tag})")
                    except KeyError:
                        logger.warning(f"Paragraph {idx}: Style '{style}' not found in template. Creating it as default Paragraph Style.")
                        try:
                            # Create the style if missing
                            styles = doc.styles
                            new_style = styles.add_style(style, 1) # 1 is WD_STYLE_TYPE.PARAGRAPH
                            new_style.base_style = styles['Normal']
                            para.style = style
                            logger.info(f"Created and applied new style '{style}'")
                            result["paragraphs_processed"] += 1
                        except Exception as e2:
                            logger.error(f"Failed to create style '{style}': {e2}")
                            result["errors"].append(f"Style '{style}' missing and creation failed: {str(e2)}")
                else:  # tag mode
                    if para.runs:
                        try:
                            para.runs[0].text = f"[{tag}] " + para.runs[0].text
                            result["paragraphs_processed"] += 1
                            logger.debug(f"Paragraph {idx}: Added tag '{tag}'")
                        except Exception as e:
                            logger.error(f"Paragraph {idx}: Failed to add tag: {e}")
                            result["errors"].append(f"Paragraph {idx} tag error: {str(e)}")
                    else:
                        logger.warning(f"Paragraph {idx}: No runs found (empty paragraph)")
            
            except KeyError as e:
                logger.error(f"Annotation {idx} missing key: {e}")
                result["errors"].append(f"Invalid annotation format: {str(e)}")
            except Exception as e:
                logger.error(f"Error processing paragraph {idx}: {e}")
                result["errors"].append(f"Paragraph {idx} error: {str(e)}")
        
        # Process tables
        try:
            logger.info(f"Processing tables (mode: {mode})")
            tag_tables(doc, mode)
            result["tables_processed"] = len(doc.tables)
        except Exception as e:
            logger.error(f"Error processing tables: {e}")
            result["errors"].append(f"Table processing error: {str(e)}")

        # Process cross-references
        if mode == "style":
            try:
                logger.info("Processing cross-references")
                process_cross_references(doc)
            except Exception as e:
                logger.error(f"Error processing cross-references: {e}")
                result["errors"].append(f"Xref error: {str(e)}")
        
        # Save document
        logger.info(f"Saving document to: {output_path}")
        doc.save(output_path)
        
        result["success"] = True
        logger.info(f"Document processed successfully. Paragraphs: {result['paragraphs_processed']}, Tables: {result['tables_processed']}")
        
    except FileNotFoundError:
        msg = f"Input file not found: {input_path}"
        logger.error(msg)
        result["errors"].append(msg)
    except Exception as e:
        msg = f"Unexpected error processing document: {e}"
        logger.error(msg, exc_info=True)
        result["errors"].append(msg)
    
    return result


if __name__ == "__main__":
    # Example usage
    result = process_docx(
        input_path="input.docx",
        output_path="processed_input.docx",
        mode="style"
    )
    print(f"Processing result: {result}")

