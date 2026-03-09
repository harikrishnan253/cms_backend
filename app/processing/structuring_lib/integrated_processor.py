"""
Integration module connecting annotation, styling, and enhanced processing.
Combines all components into a cohesive pipeline for document processing.
"""

import logging
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from docx.document import Document
from docx.oxml.ns import qn

from .logger_config import get_logger
from .rules_loader import get_rules_loader
from .enhanced_processor import DocumentProcessor, BlockState

logger = get_logger(__name__)


@dataclass
class ProcessingResult:
    """Result from document processing"""
    success: bool
    paragraphs_processed: int
    styles_applied: int
    warnings: List[str]
    errors: List[str]
    state_transitions: List[Tuple[BlockState, BlockState]]
    cross_references_found: int
    validation_issues: List[str]


class IntegratedDocumentProcessor:
    """
    Unified processor combining annotation, styling, and enhancement.
    Orchestrates the document processing pipeline.
    """
    
    def __init__(self):
        """Initialize integrated processor"""
        self.logger = logger
        self.rules_loader = get_rules_loader()
        self.doc_processor = DocumentProcessor()
        self.result: Optional[ProcessingResult] = None
    
    def process_document_integrated(self, document: Document) -> ProcessingResult:
        """
        Process document through integrated pipeline.
        
        Pipeline steps:
        1. Validate document structure
        2. Validate template styles
        3. Process paragraphs with state machine
        4. Detect tables and headers
        5. Preserve formatting when applying styles
        6. Validate heading hierarchy
        7. Find cross-references
        8. Handle empty paragraphs
        
        Args:
            document: python-docx Document object
            
        Returns:
            ProcessingResult with detailed processing information
        """
        self.logger.info(f"Starting integrated processing for document with {len(document.paragraphs)} paragraphs")
        
        errors: List[str] = []
        warnings: List[str] = []
        styles_applied = 0
        state_transitions: List[Tuple[BlockState, BlockState]] = []
        
        try:
            # Step 1: Validate template styles
            self.logger.debug("Validating template styles")
            style_warnings = self.doc_processor.validate_styles_in_template(document)
            warnings.extend(style_warnings)
            self.logger.debug(f"Template validation found {len(style_warnings)} warnings")
            
            # Step 2: Process each paragraph
            self.logger.debug(f"Processing {len(document.paragraphs)} paragraphs")
            for para_idx, para in enumerate(document.paragraphs):
                try:
                    # Update block state
                    old_state = self.doc_processor.state
                    self.doc_processor.update_block_state(para.text)
                    new_state = self.doc_processor.state
                    
                    if old_state != new_state:
                        state_transitions.append((old_state, new_state))
                        self.logger.debug(f"State transition: {old_state} -> {new_state} at paragraph {para_idx}")
                    
                    # Handle empty paragraphs
                    if not para.text.strip():
                        handling = self.doc_processor.handle_empty_paragraph(para)
                        if handling:
                            self.logger.debug(f"Handled empty paragraph at index {para_idx}: {handling}")
                            continue
                    
                    # Apply style based on patterns
                    style = self._get_style_for_paragraph(para.text)
                    if style:
                        # Preserve formatting while applying
                        if para.runs:
                            para.style = style
                            styles_applied += 1
                        self.logger.debug(f"Applied style '{style}' to paragraph {para_idx}")
                    
                    # Find cross-references
                    refs = self.doc_processor.find_cross_references(para.text)
                    if refs:
                        for ref_type, numbers in refs.items():
                            self.logger.debug(f"Found {ref_type} references: {numbers}")
                
                except Exception as e:
                    error_msg = f"Error processing paragraph {para_idx}: {str(e)}"
                    self.logger.error(error_msg)
                    errors.append(error_msg)
            
            # Step 3: Process tables
            self.logger.debug(f"Processing {len(document.tables)} tables")
            for table_idx, table in enumerate(document.tables):
                try:
                    for row_idx, row in enumerate(table.rows):
                        for cell_idx, cell in enumerate(row.cells):
                            cell_text = cell.text.strip()
                            if cell_text:
                                # Smart table header detection
                                score = self.doc_processor.detect_table_header_smart(
                                    cell_text, 
                                    row_idx, 
                                    cell_idx, 
                                    len(table.rows)
                                )
                                
                                if score > 0.6:  # Threshold
                                    try:
                                        cell.paragraphs[0].style = "Table Header"
                                        styles_applied += 1
                                    except Exception as e:
                                        self.logger.debug(f"Could not apply table header style: {e}")
                
                except Exception as e:
                    error_msg = f"Error processing table {table_idx}: {str(e)}"
                    self.logger.warning(error_msg)
                    warnings.append(error_msg)
            
            # Step 4: Validate heading hierarchy
            self.logger.debug("Validating heading hierarchy")
            hierarchy_issues = self.doc_processor.validate_heading_hierarchy(document)
            if hierarchy_issues:
                warnings.extend(hierarchy_issues)
                self.logger.debug(f"Heading hierarchy validation found {len(hierarchy_issues)} issues")
            
            # Compile result
            self.result = ProcessingResult(
                success=len(errors) == 0,
                paragraphs_processed=len(document.paragraphs),
                styles_applied=styles_applied,
                warnings=warnings,
                errors=errors,
                state_transitions=state_transitions,
                cross_references_found=sum(len(v) for v in self.doc_processor.cross_references.values()),
                validation_issues=self.doc_processor.validation_warnings
            )
            
            self.logger.info(
                f"Processing complete: {styles_applied} styles applied, "
                f"{len(errors)} errors, {len(warnings)} warnings"
            )
            
            return self.result
        
        except Exception as e:
            self.logger.error(f"Critical error during processing: {str(e)}")
            return ProcessingResult(
                success=False,
                paragraphs_processed=0,
                styles_applied=styles_applied,
                warnings=warnings,
                errors=[f"Critical error: {str(e)}"] + errors,
                state_transitions=state_transitions,
                cross_references_found=0,
                validation_issues=[]
            )
    
    def _get_style_for_paragraph(self, text: str) -> Optional[str]:
        """
        Determine appropriate style for paragraph text.
        
        Checks rules in priority order and returns first matching style.
        
        Args:
            text: Paragraph text to analyze
            
        Returns:
            Style name if match found, None otherwise
        """
        try:
            paragraphs = self.rules_loader.get_paragraphs()
            
            # Sort by priority (descending)
            for para_rule in sorted(paragraphs, key=lambda x: x.get("priority", 0), reverse=True):
                pattern = para_rule.get("pattern")
                style = para_rule.get("style")
                
                if pattern and style:
                    try:
                        if __import__('re').match(pattern, text):
                            return style
                    except Exception as e:
                        self.logger.debug(f"Pattern matching error: {e}")
            
            return None
        
        except Exception as e:
            self.logger.error(f"Error determining style: {e}")
            return None
    
    def get_processing_summary(self) -> Dict[str, Any]:
        """
        Get summary of last processing result.
        
        Returns:
            Dictionary with processing statistics and details
        """
        if not self.result:
            return {"status": "No processing completed yet"}
        
        return {
            "status": "SUCCESS" if self.result.success else "FAILED",
            "paragraphs_processed": self.result.paragraphs_processed,
            "styles_applied": self.result.styles_applied,
            "cross_references_found": self.result.cross_references_found,
            "state_transitions_count": len(self.result.state_transitions),
            "warnings_count": len(self.result.warnings),
            "errors_count": len(self.result.errors),
            "warnings": self.result.warnings[:5],  # First 5
            "errors": self.result.errors[:5],  # First 5
        }


# Global instance
_integrated_processor: Optional[IntegratedDocumentProcessor] = None


def get_integrated_processor() -> IntegratedDocumentProcessor:
    """Get or create global integrated processor instance."""
    global _integrated_processor
    if _integrated_processor is None:
        _integrated_processor = IntegratedDocumentProcessor()
    return _integrated_processor


def reset_processor():
    """Reset global processor instance (useful for testing)."""
    global _integrated_processor
    _integrated_processor = None
