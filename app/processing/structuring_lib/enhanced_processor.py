"""
Enhanced document processor with state machine, validation, and advanced features.
Handles empty paragraphs, run preservation, table detection, heading hierarchy, etc.
"""

import re
import logging
from typing import List, Dict, Any, Optional, Tuple, Set
from enum import Enum
from docx.text.paragraph import Paragraph
from docx.oxml.shared import OxmlElement

from .logger_config import get_logger
from .rules_loader import get_rules_loader

logger = get_logger(__name__)


class BlockState(Enum):
    """Document block states for state machine"""
    NORMAL = "normal"
    LEARNING_OBJECTIVES = "learning_objectives"
    REFERENCES = "references"
    FIGURES_SECTION = "figures_section"


class DocumentProcessor:
    """Enhanced document processor with advanced features"""
    
    def __init__(self):
        """Initialize processor with rules"""
        self.rules_loader = get_rules_loader()
        self.state = BlockState.NORMAL
        self.state_history: List[BlockState] = []
        self.cross_references: Dict[str, List[int]] = {}  # figure/table num -> para indices
        self.heading_levels: Set[str] = set()  # Track used heading levels
        self.validation_warnings: List[str] = []
    
    def detect_block_end(self, text: str) -> bool:
        """
        Check if paragraph marks end of current block.
        
        Args:
            text: Paragraph text to check
        
        Returns:
            True if this paragraph ends the block
        """
        patterns = self.rules_loader.get_block_end_patterns()
        
        for pattern_def in patterns:
            pattern = pattern_def.get("pattern", "")
            try:
                if re.match(pattern, text.strip()):
                    logger.debug(f"Detected block end pattern: {pattern_def.get('description')}")
                    return True
            except re.error as e:
                logger.error(f"Invalid block end pattern: {e}")
        
        return False
    
    def update_block_state(self, text: str) -> None:
        """
        Update state machine based on paragraph text.
        
        Args:
            text: Paragraph text
        """
        markers = self.rules_loader.get_block_start_markers()
        
        # Check for block start markers
        for marker, block_type in markers.items():
            if text.strip() == marker:
                old_state = self.state
                
                if block_type == "LEARNING_OBJECTIVES_BLOCK":
                    self.state = BlockState.LEARNING_OBJECTIVES
                elif block_type == "REFERENCES_BLOCK":
                    self.state = BlockState.REFERENCES
                
                self.state_history.append(self.state)
                logger.debug(f"State transition: {old_state.value} → {self.state.value}")
                return
        
        # Check for block end
        if self.detect_block_end(text):
            old_state = self.state
            self.state = BlockState.NORMAL
            self.state_history.append(self.state)
            logger.debug(f"State transition: {old_state.value} → {self.state.value}")
    
    def detect_table_header_smart(self, text: str, row_index: int, cell_index: int, 
                                  total_rows: int) -> float:
        """
        Smart table header detection using multiple heuristics.
        
        Args:
            text: Cell text
            row_index: Row position (0-indexed)
            cell_index: Column position
            total_rows: Total rows in table
        
        Returns:
            Score from 0 to 1 (higher = more likely to be header)
        """
        if not text or not text.strip():
            return 0.0
        
        table_config = self.rules_loader.get_table_config()
        header_methods = table_config.get("header_detection", [])
        
        score = 0.0
        
        for method in header_methods:
            method_type = method.get("method")
            weight = method.get("weight", 0.5)
            
            if method_type == "uppercase_only":
                # Check if text is all uppercase
                if text.isupper() and len(text) > 2:
                    score += weight
                    logger.debug(f"Header method 'uppercase': +{weight}")
            
            elif method_type == "short_text":
                # Short text is likely a header
                if len(text) < 50:
                    score += weight * 0.5
                    logger.debug(f"Header method 'short_text': +{weight * 0.5}")
            
            elif method_type == "row_position":
                # First row is likely header
                if row_index == 0:
                    score += weight
                    logger.debug(f"Header method 'row_position': +{weight}")
            
            elif method_type == "content_pattern":
                # Common header patterns
                if any(keyword in text.lower() for keyword in 
                       ['name', 'id', 'description', 'date', 'value', 'count', 'label']):
                    score += weight
                    logger.debug(f"Header method 'content_pattern': +{weight}")
        
        return min(score, 1.0)  # Cap at 1.0
    
    def preserve_run_formatting(self, para: Paragraph, tag: str) -> None:
        """
        Add tag prefix while preserving formatting (bold, italic, hyperlinks).
        
        Args:
            para: Paragraph to modify
            tag: Tag to add as prefix
        """
        if not para.runs:
            logger.warning("Paragraph has no runs, cannot add tag")
            return
        
        # Get the first run's formatting
        first_run = para.runs[0]
        
        # Create new run for tag
        tag_run = first_run.insert_paragraph_before(f"[{tag}] ").runs[0]
        
        # Copy formatting from original to tag run
        try:
            if first_run.bold:
                tag_run.bold = True
            if first_run.italic:
                tag_run.italic = True
            if first_run.underline:
                tag_run.underline = True
            if first_run.font.color.rgb:
                tag_run.font.color.rgb = first_run.font.color.rgb
            if first_run.font.size:
                tag_run.font.size = first_run.font.size
            
            logger.debug(f"Added tag [{tag}] with preserved formatting")
        
        except Exception as e:
            logger.warning(f"Failed to copy formatting: {e}")
            # Fallback: just add the tag
            first_run.text = f"[{tag}] " + first_run.text
    
    def validate_heading_hierarchy(self, style: str, para_index: int) -> bool:
        """
        Validate proper heading hierarchy (Heading 1 → Heading 2 → Heading 3).
        
        Args:
            style: Style name (e.g., "Heading 1")
            para_index: Paragraph index for context
        
        Returns:
            True if valid, False if invalid
        """
        hierarchy_config = self.rules_loader.get_heading_hierarchy()
        
        if not hierarchy_config:
            return True  # No hierarchy validation configured
        
        valid_transitions = hierarchy_config.get("valid_transitions", [])
        enforcement = hierarchy_config.get("enforcement", "warn")
        
        # Get last heading level used
        last_level = self._get_last_heading_level()
        
        # Find valid transitions from current level
        for transition in valid_transitions:
            if transition.get("from") == last_level:
                allowed = transition.get("to", [])
                if style in allowed:
                    self.heading_levels.add(style)
                    return True
        
        # Invalid transition
        message = f"Para {para_index}: Invalid heading transition from {last_level} to {style}"
        
        if enforcement == "error":
            raise ValueError(message)
        elif enforcement == "warn":
            logger.warning(message)
            self.validation_warnings.append(message)
        elif enforcement == "auto-fix":
            logger.info(f"Auto-fixing: {message}")
            # Could auto-fix by adjusting style
        
        return enforcement != "error"
    
    def _get_last_heading_level(self) -> Optional[str]:
        """Get the most recently used heading level"""
        for heading in ["Heading 1", "Heading 2", "Heading 3", "Title"]:
            if heading in self.heading_levels:
                return heading
        return None
    
    def find_cross_references(self, doc) -> Dict[str, List[int]]:
        """
        Find cross-references to figures and tables.
        
        Args:
            doc: python-docx Document
        
        Returns:
            Dict mapping reference type to paragraph indices
        """
        cross_refs = {
            "figure_references": [],
            "table_references": [],
            "box_references": []
        }
        
        for idx, para in enumerate(doc.paragraphs):
            text = para.text
            
            # Look for figure references
            if re.search(r"(?:Figure|Fig\.?|see Fig)\s+\d+", text, re.IGNORECASE):
                cross_refs["figure_references"].append(idx)
            
            # Look for table references
            if re.search(r"(?:Table|Tbl\.?)\s+\d+", text, re.IGNORECASE):
                cross_refs["table_references"].append(idx)
            
            # Look for box references
            if re.search(r"(?:Box)\s+\d+", text, re.IGNORECASE):
                cross_refs["box_references"].append(idx)
        
        self.cross_references = cross_refs
        logger.info(f"Found {len(cross_refs['figure_references'])} figure refs, "
                   f"{len(cross_refs['table_references'])} table refs")
        
        return cross_refs
    
    def handle_empty_paragraph(self, para: Paragraph) -> Tuple[str, str]:
        """
        Handle empty paragraphs with configurable behavior.
        
        Args:
            para: Empty paragraph
        
        Returns:
            Tuple of (tag, style)
        """
        defaults = self.rules_loader.rules.get("defaults", {})
        empty_config = defaults.get("empty_paragraph", {})
        
        tag = empty_config.get("tag", "EMPTY")
        style = empty_config.get("style", "Normal")
        
        logger.debug(f"Handling empty paragraph: tag={tag}, style={style}")
        
        return tag, style
    
    def validate_styles_in_template(self, doc) -> List[str]:
        """
        Validate that all referenced styles exist in the document template.
        
        Args:
            doc: python-docx Document
        
        Returns:
            List of missing style names
        """
        available_styles = {style.name for style in doc.styles}
        referenced_styles = set()
        
        # Collect all referenced styles from rules
        for para_rule in self.rules_loader.get_paragraphs():
            style = para_rule.get("style")
            if style:
                referenced_styles.add(style)
        
        # Check table styles
        table_config = self.rules_loader.get_table_config()
        if "header_style" in table_config:
            referenced_styles.add(table_config["header_style"])
        if "body_style" in table_config:
            referenced_styles.add(table_config["body_style"])
        
        # Find missing styles
        missing = referenced_styles - available_styles
        
        if missing:
            logger.warning(f"Missing styles in template: {missing}")
            for style in missing:
                self.validation_warnings.append(f"Style not in template: {style}")
        
        return sorted(list(missing))
    
    def process_document_enhanced(self, doc, mode: str = "style") -> Dict[str, Any]:
        """
        Process document with all enhancements.
        
        Args:
            doc: python-docx Document
            mode: "style" or "tag"
        
        Returns:
            Processing result dictionary
        """
        result = {
            "success": True,
            "paragraphs_processed": 0,
            "tables_processed": 0,
            "cross_references_found": 0,
            "validation_warnings": [],
            "missing_styles": []
        }
        
        try:
            # Validate styles
            missing = self.validate_styles_in_template(doc)
            result["missing_styles"] = missing
            
            # Find cross-references
            cross_refs = self.find_cross_references(doc)
            result["cross_references_found"] = sum(len(v) for v in cross_refs.values())
            
            # Process paragraphs
            for para_idx, para in enumerate(doc.paragraphs):
                try:
                    text = para.text.strip()
                    
                    # Update state machine
                    self.update_block_state(text)
                    
                    # Handle empty paragraphs
                    if not text:
                        tag, style = self.handle_empty_paragraph(para)
                        if mode == "style":
                            try:
                                para.style = style
                            except KeyError:
                                pass
                        result["paragraphs_processed"] += 1
                        continue
                    
                    # Process normally (would use enhanced annotator)
                    result["paragraphs_processed"] += 1
                
                except Exception as e:
                    logger.error(f"Error processing paragraph {para_idx}: {e}")
            
            # Process tables
            for table_idx, table in enumerate(doc.tables):
                try:
                    table_config = self.rules_loader.get_table_config()
                    header_threshold = table_config.get("header_threshold", 1.0)
                    header_style = table_config.get("header_style", "Table Header")
                    body_style = table_config.get("body_style", "Table Body")
                    
                    for row_idx, row in enumerate(table.rows):
                        for cell_idx, cell in enumerate(row.cells):
                            for para in cell.paragraphs:
                                text = para.text.strip()
                                if not text:
                                    continue
                                
                                # Smart header detection
                                score = self.detect_table_header_smart(
                                    text, row_idx, cell_idx, len(table.rows)
                                )
                                
                                is_header = score >= header_threshold
                                style = header_style if is_header else body_style
                                
                                if mode == "style":
                                    try:
                                        para.style = style
                                    except KeyError:
                                        pass
                    
                    result["tables_processed"] += 1
                
                except Exception as e:
                    logger.error(f"Error processing table {table_idx}: {e}")
            
            result["validation_warnings"] = self.validation_warnings
        
        except Exception as e:
            logger.error(f"Error in enhanced processing: {e}", exc_info=True)
            result["success"] = False
        
        return result


if __name__ == "__main__":
    print("Enhanced document processor module loaded")
    processor = DocumentProcessor()
    print(f"✓ Processor initialized with {len(processor.rules_loader.get_paragraphs())} rules")
