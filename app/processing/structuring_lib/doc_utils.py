from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
import logging
from typing import List, Dict, Any

logger = logging.getLogger(__name__)

def iter_block_items(doc):
    """
    Generate a stream of document body items (paragraphs and tables) in proper order.
    """
    parent = doc
    if hasattr(parent, 'element'):
        body = parent.element.body
    else:
        return

    # We maintain indices to map back to the high-level lists
    # doc.paragraphs only includes immediate children of body that are P
    # doc.tables only includes immediate children of body that are Tbl
    p_counter = 0
    t_counter = 0

    for child in body.iterchildren():
        if isinstance(child, CT_P):
            if p_counter < len(doc.paragraphs):
                yield "paragraph", doc.paragraphs[p_counter], p_counter
                p_counter += 1
        elif isinstance(child, CT_Tbl):
            if t_counter < len(doc.tables):
                yield "table", doc.tables[t_counter], t_counter
                t_counter += 1

def extract_document_structure(docx_path: str) -> List[Dict[str, Any]]:
    """
    Reads a DOCX and returns a list of items (paras and tables) in order.
    IDs: 
      - Main Body Para: p-{index}
      - Table Para: t-{t_idx}-r{r_idx}-c{c_idx}-p{p_idx}
    """
    doc = Document(docx_path)
    structure = []
    
    for item_type, item, idx in iter_block_items(doc):
        if item_type == "paragraph":
            text = item.text.strip()
            if text:
                structure.append({
                    "id": f"p-{idx}",
                    "text": text,
                    "style": item.style.name if item.style else "Normal",
                    "type": "paragraph"
                })
        
        elif item_type == "table":
            table_data = {
                "id": f"t-{idx}",
                "type": "table",
                "rows": []
            }
            
            table = item
            t_idx = idx
            for r_idx, row in enumerate(table.rows):
                row_data = {"cells": []}
                for c_idx, cell in enumerate(row.cells):
                    cell_data = {"paragraphs": []}
                    for p_idx, para in enumerate(cell.paragraphs):
                        text = para.text.strip()
                        if text:
                            cell_data["paragraphs"].append({
                                "id": f"t-{t_idx}-r{r_idx}-c{c_idx}-p{p_idx}",
                                "text": text, 
                                "style": para.style.name if para.style else "Normal",
                                "type": "table_paragraph"
                            })
                    row_data["cells"].append(cell_data)
                table_data["rows"].append(row_data)
            
            # Only add if table has content
            has_content = any(any(c["paragraphs"] for c in r["cells"]) for r in table_data["rows"])
            if has_content:
                structure.append(table_data)

    return structure

def update_document_structure(docx_path: str, output_path: str, updates: Dict[str, str]) -> bool:
    """
    Apply style updates using complex IDs.
    """
    try:
        doc = Document(docx_path)
        count = 0
        
        for item_id, new_style_name in updates.items():
            target_para = None
            
            # Parse ID
            parts = item_id.split('-')
            
            if item_id.startswith('p-'):
                # Body Paragraph: p-{index}
                try:
                    p_idx = int(parts[1])
                    if p_idx < len(doc.paragraphs):
                        target_para = doc.paragraphs[p_idx]
                except (IndexError, ValueError):
                    pass
                    
            elif item_id.startswith('t-'):
                # Table Paragraph: t-{t}-r{r}-c{c}-p{p}
                # Format: t, 0, r0, c1, p0 -> parts: ['t', '0', 'r0', 'c1', 'p0']
                # Need to strip chars
                try:
                    t_idx = int(parts[1])
                    r_idx = int(parts[2][1:]) # strip 'r'
                    c_idx = int(parts[3][1:]) # strip 'c'
                    p_idx = int(parts[4][1:]) # strip 'p'
                    
                    if t_idx < len(doc.tables):
                        table = doc.tables[t_idx]
                        if r_idx < len(table.rows):
                            row = table.rows[r_idx]
                            if c_idx < len(row.cells):
                                cell = row.cells[c_idx]
                                if p_idx < len(cell.paragraphs):
                                    target_para = cell.paragraphs[p_idx]
                except (IndexError, ValueError, AttributeError) as e:
                    logger.warning(f"Failed to parse or find ID {item_id}: {e}")

            # Apply Style
            if target_para:
                try:
                    target_para.style = new_style_name
                    count += 1
                except KeyError:
                    # Style Creation Logic
                    logger.info(f"Style '{new_style_name}' not found. Creating it.")
                    try:
                        styles = doc.styles
                        new_style = styles.add_style(new_style_name, 1) # Paragraph style
                        if 'Normal' in styles:
                            new_style.base_style = styles['Normal']
                        target_para.style = new_style_name
                        count += 1
                    except Exception as e:
                        logger.error(f"Failed to create style '{new_style_name}': {e}")

        doc.save(output_path)
        logger.info(f"Updated {count} items in {output_path}")
        return True
    
    except Exception as e:
        logger.error(f"Failed to update doc: {e}", exc_info=True)
        return False

def load_document(docx_path: str) -> Document:
    """Load a DOCX document."""
    return Document(docx_path)

def save_document(doc: Document, output_path: str) -> None:
    """Save a DOCX document."""
    doc.save(output_path)
